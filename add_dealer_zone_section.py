#!/usr/bin/env python3
"""
Add Dealer Zone Section with Screenshots to Koleksiyon Document
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import os

print("="*80)
print("ADDING DEALER ZONE SECTION TO KOLEKSIYON DOCUMENT")
print("="*80 + "\n")

# Load existing document
doc = Document('/Users/caglarozyildirim/WebstormProjects/Deneme/Koleksiyon_Mobilya_Bayi_Portali_Kapsam_Dokumani.docx')

# Find the Web Sitesi Entegrasyonu section and add content after "Dealer Zone" section
insert_position = None
for idx, para in enumerate(doc.paragraphs):
    if "Dealer Zone" in para.text or "Web Sitesi Navigasyon Yapısı" in para.text:
        # Insert after this paragraph
        insert_position = idx + 15  # After existing navigation structure
        break

if not insert_position:
    insert_position = len(doc.paragraphs) - 20

print(f"📍 Insert position: {insert_position}")

# Add new content before screenshots section
# Find EK: MEVCUT SİSTEM section
screenshots_idx = None
for idx, para in enumerate(doc.paragraphs):
    if "EK: MEVCUT SİSTEM" in para.text:
        screenshots_idx = idx
        break

if screenshots_idx:
    insert_position = screenshots_idx

print(f"📍 Final insert position: {insert_position}")

# Insert new section
print("📝 Adding Dealer Zone Implementation Section...")

# Add page break before new section
doc.paragraphs[insert_position - 1]._element.addnext(
    doc.add_paragraph()._element
)

# Add heading
heading = doc.add_heading('DEALER ZONE İMPLEMENTASYONU', level=2)
heading_para = heading._element
doc.paragraphs[insert_position]._element.addnext(heading_para)

# Add intro paragraph
intro = doc.add_paragraph(
    "Mevcut Koleksiyon web sitesinde (https://www.koleksiyondesign.com) "
    "'Biz kimiz' menüsü altında 'Bayilik başvurusu' sayfası bulunmaktadır. "
    "Bu sayfa yeniden yapılandırılarak 'Dealer Zone' bölümü eklenecektir."
)
intro_element = intro._element
heading_para.addnext(intro_element)

# Current navigation structure
nav_heading = doc.add_heading('Mevcut Web Sitesi Navigasyonu', level=3)
nav_element = nav_heading._element
intro_element.addnext(nav_element)

nav_desc = doc.add_paragraph(
    "Koleksiyon web sitesinde şu navigasyon yapısı mevcuttur:"
)
nav_desc_element = nav_desc._element
nav_element.addnext(nav_desc_element)

# Navigation list
nav_structure = [
    "Ana Menü > Biz kimiz (dropdown)",
    "  → Hakkımızda",
    "  → Tarihçe",
    "  → Ortak tasarım",
    "  → Sürdürülebilirlik",
    "  → Ödüller",
    "  → Sertifikalar",
    "  → Haberler ve Etkinlikler",
    "  → Kariyer",
    "  → Yatırımcı İlişkileri",
    "İletişim Alt Bölümü:",
    "  → Bize ulaşın",
    "  → Mağazalar",
    "  → Bayilik başvurusu ← MEVCUT SAYFA"
]

last_element = nav_desc_element
for item in nav_structure:
    nav_item = doc.add_paragraph(f"  {item}")
    nav_item.paragraph_format.left_indent = Inches(0.25)
    if "←" in item:
        for run in nav_item.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 0, 0)
    nav_item_element = nav_item._element
    last_element.addnext(nav_item_element)
    last_element = nav_item_element

# Add screenshot 1
screenshot1_heading = doc.add_heading('Ekran Görüntüsü 1: Mevcut Navigasyon Menüsü', level=3)
screenshot1_heading_element = screenshot1_heading._element
last_element.addnext(screenshot1_heading_element)
last_element = screenshot1_heading_element

screenshot1_desc = doc.add_paragraph(
    "Aşağıdaki görselde 'Biz kimiz' menüsü altında 'İletişim' bölümünde "
    "'Bayilik başvurusu' linki görülmektedir:"
)
screenshot1_desc_element = screenshot1_desc._element
last_element.addnext(screenshot1_desc_element)
last_element = screenshot1_desc_element

# Add space
space1 = doc.add_paragraph()
space1_element = space1._element
last_element.addnext(space1_element)
last_element = space1_element

# Add first image
try:
    img1_path = "/Users/caglarozyildirim/Downloads/image-20251015-060827.png"
    pic1 = doc.add_picture(img1_path, width=Inches(4.5))
    pic1_para = doc.paragraphs[-1]
    pic1_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic1_element = pic1_para._element
    last_element.addnext(pic1_element)
    last_element = pic1_element
    print("  ✓ Screenshot 1 added")
except Exception as e:
    print(f"  ✗ Error adding screenshot 1: {e}")

# Add caption
caption1 = doc.add_paragraph("Şekil: 'Biz kimiz > İletişim' menüsünde 'Bayilik başvurusu' linki")
caption1.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in caption1.runs:
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = RGBColor(128, 128, 128)
caption1_element = caption1._element
last_element.addnext(caption1_element)
last_element = caption1_element

# Add space
space2 = doc.add_paragraph()
space2_element = space2._element
last_element.addnext(space2_element)
last_element = space2_element

# Current page structure
current_page_heading = doc.add_heading('Mevcut Bayilik Başvuru Sayfası', level=3)
current_page_element = current_page_heading._element
last_element.addnext(current_page_element)
last_element = current_page_element

current_page_desc = doc.add_paragraph(
    "Mevcut bayilik başvuru sayfasında kullanıcılara şu mesaj sunulmaktadır:"
)
current_page_desc_element = current_page_desc._element
last_element.addnext(current_page_desc_element)
last_element = current_page_desc_element

quote = doc.add_paragraph(
    "\"Koleksiyon ailesine katılmak ister misiniz? Bu seçkin markanın bayilik "
    "fırsatlarıyla ilgileniyorsanız, aşağıdaki formu doldurmanız yeterli. Bayi "
    "Yönetimi ekibimiz en kısa sürede sizinle iletişime geçerek detayları paylaşacaktır. "
    "Gelin, birlikte ilham veren mekânlara imza atalım.\""
)
for run in quote.runs:
    run.font.italic = True
    run.font.color.rgb = RGBColor(64, 64, 64)
quote.paragraph_format.left_indent = Inches(0.5)
quote.paragraph_format.right_indent = Inches(0.5)
quote_element = quote._element
last_element.addnext(quote_element)
last_element = quote_element

# Add space
space3 = doc.add_paragraph()
space3_element = space3._element
last_element.addnext(space3_element)
last_element = space3_element

# Add screenshot 2
screenshot2_heading = doc.add_heading('Ekran Görüntüsü 2: Dealer Zone Ekleme Alanı', level=3)
screenshot2_heading_element = screenshot2_heading._element
last_element.addnext(screenshot2_heading_element)
last_element = screenshot2_heading_element

screenshot2_desc = doc.add_paragraph(
    "Kırmızı dikdörtgen ile işaretlenmiş alan, 'Dealer Zone' bölümünün "
    "ekleneceği yeri göstermektedir. Bu bölüm 'Başvuru formu' butonunun "
    "hemen altına yerleştirilecektir:"
)
screenshot2_desc_element = screenshot2_desc._element
last_element.addnext(screenshot2_desc_element)
last_element = screenshot2_desc_element

# Add space
space4 = doc.add_paragraph()
space4_element = space4._element
last_element.addnext(space4_element)
last_element = space4_element

# Add second image
try:
    img2_path = "/Users/caglarozyildirim/Downloads/image-20251015-060917.png"
    pic2 = doc.add_picture(img2_path, width=Inches(5.5))
    pic2_para = doc.paragraphs[-1]
    pic2_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic2_element = pic2_para._element
    last_element.addnext(pic2_element)
    last_element = pic2_element
    print("  ✓ Screenshot 2 added")
except Exception as e:
    print(f"  ✗ Error adding screenshot 2: {e}")

# Add caption
caption2 = doc.add_paragraph("Şekil: 'Dealer Zone' bölümünün ekleneceği alan (kırmızı ile işaretli)")
caption2.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in caption2.runs:
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = RGBColor(128, 128, 128)
caption2_element = caption2._element
last_element.addnext(caption2_element)
last_element = caption2_element

# Add space
space5 = doc.add_paragraph()
space5_element = space5._element
last_element.addnext(space5_element)
last_element = space5_element

# Proposed Dealer Zone Section
proposed_heading = doc.add_heading('Önerilen Dealer Zone Yapısı', level=3)
proposed_element = proposed_heading._element
last_element.addnext(proposed_element)
last_element = proposed_element

proposed_desc = doc.add_paragraph(
    "'Başvuru formu' butonunun altına, 3 kartlı bir 'Dealer Zone' bölümü eklenecektir:"
)
proposed_desc_element = proposed_desc._element
last_element.addnext(proposed_desc_element)
last_element = proposed_desc_element

# Card 1
card1_para = doc.add_paragraph()
card1_run = card1_para.add_run("Kart 1: Bayilik Başvurusu")
card1_run.font.bold = True
card1_run.font.color.rgb = RGBColor(0, 112, 192)
card1_element = card1_para._element
last_element.addnext(card1_element)
last_element = card1_element

card1_details = [
    "İkon: 📋 Başvuru formu ikonu",
    "Başlık: 'Bayilik Başvurusu'",
    "Açıklama: 'Fiziksel mağaza açmak için genel başvuru'",
    "Buton: 'Başvur' → Başvuru formu sayfasına yönlendirir",
    "Renk Teması: Mavi"
]

for detail in card1_details:
    detail_para = doc.add_paragraph(f"  • {detail}")
    detail_para.paragraph_format.left_indent = Inches(0.5)
    detail_element = detail_para._element
    last_element.addnext(detail_element)
    last_element = detail_element

# Card 2
card2_para = doc.add_paragraph()
card2_run = card2_para.add_run("Kart 2: Panel Girişi")
card2_run.font.bold = True
card2_run.font.color.rgb = RGBColor(0, 176, 80)
card2_element = card2_para._element
last_element.addnext(card2_element)
last_element = card2_element

card2_details = [
    "İkon: 🔐 Kilit ikonu",
    "Başlık: 'Panel Girişi'",
    "Açıklama: 'Onaylı bayiler için fiyat listesi erişimi'",
    "Buton: 'Giriş Yap' → Login sayfasına yönlendirir",
    "Renk Teması: Yeşil"
]

for detail in card2_details:
    detail_para = doc.add_paragraph(f"  • {detail}")
    detail_para.paragraph_format.left_indent = Inches(0.5)
    detail_element = detail_para._element
    last_element.addnext(detail_element)
    last_element = detail_element

# Card 3
card3_para = doc.add_paragraph()
card3_run = card3_para.add_run("Kart 3: Panele Üye Ol")
card3_run.font.bold = True
card3_run.font.color.rgb = RGBColor(255, 192, 0)
card3_element = card3_para._element
last_element.addnext(card3_element)
last_element = card3_element

card3_details = [
    "İkon: ✍️ Kayıt ikonu",
    "Başlık: 'Panele Üye Ol'",
    "Açıklama: 'Online fiyat erişimi için kayıt olun'",
    "Buton: 'Kayıt Ol' → Self-registration formuna yönlendirir",
    "Renk Teması: Turuncu"
]

for detail in card3_details:
    detail_para = doc.add_paragraph(f"  • {detail}")
    detail_para.paragraph_format.left_indent = Inches(0.5)
    detail_element = detail_para._element
    last_element.addnext(detail_element)
    last_element = detail_element

# Design specifications
design_heading = doc.add_heading('Tasarım Spesifikasyonları', level=3)
design_element = design_heading._element
last_element.addnext(design_element)
last_element = design_element

design_desc = doc.add_paragraph(
    "Dealer Zone bölümü aşağıdaki tasarım özelliklerine sahip olacaktır:"
)
design_desc_element = design_desc._element
last_element.addnext(design_desc_element)
last_element = design_desc_element

design_specs = [
    "Layout: 3 kart yan yana (Desktop), dikey (Mobil)",
    "Kart Boyutu: Her kart eşit genişlikte, hover efekti ile",
    "Boşluk: Kartlar arası 20px margin",
    "Gölge: Hafif drop-shadow efekti",
    "Typography: Başlıklar Calibri Bold 18pt, Açıklama Calibri Regular 12pt",
    "Buton Stili: Primary button, border-radius 4px, hover efekti",
    "Responsive: Mobile-first yaklaşım, breakpoint 768px",
    "Bölüm Başlığı: 'Dealer Zone' veya 'Bayi Portalı' (büyük başlık, merkezi)",
    "Alt Açıklama: 'Fiyat listelerine erişim ve bayilik işlemleri için'",
    "Arka Plan: Beyaz veya açık gri, sayfa ile uyumlu"
]

for spec in design_specs:
    spec_para = doc.add_paragraph(f"• {spec}")
    spec_para.paragraph_format.left_indent = Inches(0.5)
    spec_element = spec_para._element
    last_element.addnext(spec_element)
    last_element = spec_element

# Implementation notes
impl_heading = doc.add_heading('İmplementasyon Notları', level=3)
impl_element = impl_heading._element
last_element.addnext(impl_element)
last_element = impl_element

impl_notes = [
    "Mevcut 'Başvuru formu' butonu değiştirilmeyecek, Dealer Zone ek olarak eklenecek",
    "Dealer Zone bölümü sayfa içinde scroll ile görünür olacak (hemen alt bölümde)",
    "SEO için 'Dealer Zone' bölümü sayfa başlığına eklenecek",
    "Google Analytics event tracking her kart için ayrı ayrı kurulacak",
    "3 farklı landing page hazırlanacak (Başvuru, Login, Kayıt)",
    "URL yapısı: /bayilik-basvurusu (mevcut), /dealer/login, /dealer/register",
    "Dil desteği: TR/EN switch butonu header'da mevcut yapı kullanılacak",
    "Koleksiyon kurumsal renkleri kullanılacak (Mavi: #003366, Turuncu: #FF6600)"
]

for note in impl_notes:
    note_para = doc.add_paragraph(f"• {note}")
    note_para.paragraph_format.left_indent = Inches(0.5)
    note_element = note_para._element
    last_element.addnext(note_element)
    last_element = note_element

# Add page break
page_break = doc.add_page_break()
page_break_element = page_break._element
last_element.addnext(page_break_element)

print("✅ Dealer Zone section added successfully!")

# Update version table
for table in doc.tables:
    if len(table.rows) >= 2 and len(table.columns) >= 4:
        if "Version" in table.rows[0].cells[0].text:
            # Update version info
            table.rows[1].cells[0].text = "1.2"
            table.rows[1].cells[1].text = datetime.now().strftime('%d.%m.%Y')
            table.rows[1].cells[3].text = "Dealer Zone implementasyon detayları ve ekran görüntüleri eklendi"
            print("✅ Version table updated to 1.2")
            break

# Save document
output_file = "/Users/caglarozyildirim/WebstormProjects/Deneme/Koleksiyon_Mobilya_Bayi_Portali_Kapsam_Dokumani.docx"
doc.save(output_file)

print("\n" + "="*80)
print("✅ DOCUMENT UPDATED SUCCESSFULLY!")
print("="*80)
print(f"📁 File: {output_file}")
print(f"✨ Version: 1.2")
print(f"📸 Screenshots added: 2")
print(f"🎨 Dealer Zone design specs included")
print(f"📋 3-card structure detailed")
print("="*80 + "\n")
