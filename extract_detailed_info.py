#!/usr/bin/env python3
"""
Desktop/new dokümanlarından detaylı bilgi çıkaran script
"""
import json
from pathlib import Path
from collections import defaultdict

def extract_detailed_info():
    desktop_path = Path.home() / "Desktop" / "new"
    result = {
        'screens': {},
        'data_structures': {},
        'workflows': [],
        'requirements': {}
    }

    # Python kütüphaneleri yüklü mü kontrol et
    try:
        from docx import Document
        import openpyxl

        # 1. Screen Designs Excel'den ekran bilgilerini çıkar
        screen_designs_path = desktop_path / "Screen Designs.xlsx"
        if screen_designs_path.exists():
            print("=== Screen Designs Analizi ===\n")
            wb = openpyxl.load_workbook(str(screen_designs_path), read_only=True, data_only=True)

            # Activities X Screens sayfasını oku
            if 'Activities X Screens' in wb.sheetnames:
                sheet = wb['Activities X Screens']
                print("Activities X Screens Sayfası:")

                # Başlık satırını bul (genellikle 4-6 arası)
                headers = []
                screen_names = []

                for row_idx in range(1, 10):
                    row = list(sheet.iter_rows(min_row=row_idx, max_row=row_idx, values_only=True))[0]
                    if 'Screen Name' in str(row):
                        # Ekran isimlerini al
                        next_row = list(sheet.iter_rows(min_row=row_idx+1, max_row=row_idx+1, values_only=True))[0]
                        for idx, cell in enumerate(next_row):
                            if cell and idx > 3:
                                screen_names.append((idx, str(cell)))
                        break

                print(f"Bulunan ekranlar ({len(screen_names)} adet):")
                for idx, name in screen_names[:20]:
                    print(f"  Sütun {idx}: {name}")
                    result['screens'][name] = {'column': idx, 'activities': []}

            # Info sayfasındaki ekran açıklamalarını oku
            if 'Info' in wb.sheetnames:
                sheet = wb['Info']
                print("\n\nInfo Sayfası - Ekran Açıklamaları:")

                for row in sheet.iter_rows(min_row=7, max_row=50, values_only=True):
                    if row[1]:  # Ekran adı
                        screen_name = str(row[1])
                        purpose = str(row[2]) if row[2] else ""
                        print(f"  {screen_name}: {purpose[:100]}")

        # 2. Data Structure Excel'den tablo yapılarını çıkar
        data_struct_path = desktop_path / "Data Structure.xlsx"
        if data_struct_path.exists():
            print("\n\n=== Data Structure Analizi ===\n")
            wb = openpyxl.load_workbook(str(data_struct_path), read_only=True, data_only=True)

            tables_info = {}

            for sheet_name in wb.sheetnames[:15]:
                sheet = wb[sheet_name]

                # Tablo adını ve alan bilgilerini çıkar
                table_name = sheet_name
                fields = []

                # İlk 50 satırı tara
                for row in sheet.iter_rows(min_row=2, max_row=50, values_only=True):
                    # Field name, Type, Length gibi sütunlar ara
                    if row[1] and row[2]:  # Alan adı ve tipi varsa
                        field_name = str(row[1])
                        field_type = str(row[2])
                        field_length = str(row[3]) if len(row) > 3 and row[3] else ""

                        if field_name not in ['Field name', 'Fields', ''] and not field_name.startswith('Table'):
                            fields.append({
                                'name': field_name,
                                'type': field_type,
                                'length': field_length
                            })

                if fields:
                    tables_info[table_name] = fields
                    print(f"{table_name} Tablosu: {len(fields)} alan")
                    result['data_structures'][table_name] = fields

        # 3. Requirement Analysis dokümanından önemli bölümleri çıkar
        req_doc_path = desktop_path / "Maintenance Management Application Requirement Analysis (Version1).docx"
        if req_doc_path.exists():
            print("\n\n=== Requirement Analysis İçeriği ===\n")
            doc = Document(str(req_doc_path))

            current_section = ""
            requirements_text = []

            for para in doc.paragraphs[:500]:
                if para.style.name.startswith('Heading'):
                    current_section = para.text
                    print(f"{para.style.name}: {para.text}")

                # İlk birkaç paragrafın metnini topla
                if para.text and len(para.text) > 20:
                    requirements_text.append({
                        'section': current_section,
                        'text': para.text[:200]
                    })

            # Tabloları incele
            print(f"\nToplam tablo sayısı: {len(doc.tables)}")
            for idx, table in enumerate(doc.tables[:3]):
                print(f"\nTablo {idx+1}:")
                for row_idx, row in enumerate(table.rows[:5]):
                    cells = [cell.text[:50] for cell in row.cells[:5]]
                    print(f"  Satır {row_idx+1}: {cells}")

        # 4. Workflows listesi
        workflows_path = desktop_path / "Workflows"
        if workflows_path.exists():
            print("\n\n=== Workflows ===\n")
            for f in workflows_path.glob("*.vsdx"):
                workflow_name = f.stem
                result['workflows'].append(workflow_name)
                print(f"  - {workflow_name}")

        # 5. Sonuçları JSON olarak kaydet
        output_path = Path("desktop_new_analysis.json")
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(result, f, ensure_ascii=False, indent=2)

        print(f"\n\nAnaliz sonuçları {output_path} dosyasına kaydedildi.")

    except ImportError as e:
        print(f"Gerekli kütüphane yüklü değil: {e}")
        print("Lütfen şu komutları çalıştırın:")
        print("  pip install python-docx openpyxl")
    except Exception as e:
        print(f"Hata: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    extract_detailed_info()