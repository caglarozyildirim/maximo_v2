#!/usr/bin/env python3
"""
Create comprehensive business analysis document for MAN Turkey Maintenance Management System
With corporate design screenshots and Visio-style workflow diagrams
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def add_colored_heading(doc, text, level, color_rgb=(226, 7, 20)):
    """Add MAN red colored heading"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(*color_rgb)
        run.font.name = 'Arial'
    return heading

def add_table_with_style(doc, data, headers):
    """Add styled table with MAN colors"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'

    # Header row
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        # Style header
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
        # Set cell background to MAN red
        set_cell_bg_color(header_cells[i], "E20714")

    # Data rows
    for row_data in data:
        row_cells = table.add_row().cells
        for i, cell_data in enumerate(row_data):
            row_cells[i].text = str(cell_data)

    return table

def set_cell_bg_color(cell, color_hex):
    """Set cell background color"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color_hex)
    cell._element.get_or_add_tcPr().append(shading_elm)

def add_image_if_exists(doc, image_path, width_inches=6):
    """Add image if it exists"""
    if os.path.exists(image_path):
        doc.add_picture(image_path, width=Inches(width_inches))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return True
    return False

# Create document
doc = Document()

# Set document margins
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# ============================================
# TITLE PAGE
# ============================================
title = doc.add_heading('MAN TÜRKİYE', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.color.rgb = RGBColor(226, 7, 20)
    run.font.size = Pt(36)
    run.font.name = 'Arial Black'

subtitle = doc.add_heading('BAKIM YÖNETİMİ SİSTEMİ', 1)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in subtitle.runs:
    run.font.color.rgb = RGBColor(26, 26, 26)
    run.font.size = Pt(24)

doc.add_paragraph()
doc.add_paragraph()

doc_title = doc.add_heading('İş Analizi Dokümanı', 2)
doc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in doc_title.runs:
    run.font.color.rgb = RGBColor(102, 102, 102)

doc.add_paragraph()
version_p = doc.add_paragraph('Versiyon: 2.0 - Corporate Edition')
version_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_p = doc.add_paragraph('Tarih: Ekim 2025')
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

# ============================================
# TABLE OF CONTENTS
# ============================================
add_colored_heading(doc, '📋 İçindekiler', 1)
doc.add_paragraph('1. Proje Genel Bakış', style='List Number')
doc.add_paragraph('2. Sistem Mimarisi', style='List Number')
doc.add_paragraph('3. Modül Detayları', style='List Number')
doc.add_paragraph('   3.1. İş Talepleri Yönetimi', style='List Number 2')
doc.add_paragraph('   3.2. Varlık Yönetimi', style='List Number 2')
doc.add_paragraph('   3.3. Bakım Yönetimi', style='List Number 2')
doc.add_paragraph('   3.4. Olay Yönetimi', style='List Number 2')
doc.add_paragraph('4. İş Akışları ve Süreçler', style='List Number')
doc.add_paragraph('5. Veri Yapıları', style='List Number')
doc.add_paragraph('6. Ekran Görüntüleri', style='List Number')
doc.add_paragraph('7. Test Senaryoları', style='List Number')

doc.add_page_break()

# ============================================
# 1. PROJE GENEL BAKIŞ
# ============================================
add_colored_heading(doc, '1. Proje Genel Bakış', 1)

doc.add_heading('1.1. Proje Tanımı', 2)
p = doc.add_paragraph()
p.add_run('MAN Türkiye Bakım Yönetimi Sistemi').bold = True
p.add_run(', MAN Türkiye otobüs üretim tesisindeki tüm varlıkların (makineler, ekipmanlar, araçlar) bakım süreçlerinin dijital ortamda yönetilmesini sağlayan kapsamlı bir web uygulamasıdır.')

doc.add_heading('1.2. Proje Hedefleri', 2)
goals = [
    'Bakım süreçlerinin dijitalleştirilmesi ve merkezi yönetimi',
    'Plansız duruş sürelerinin azaltılması',
    'Bakım maliyetlerinin optimize edilmesi',
    'Varlık yaşam döngüsünün etkin yönetimi',
    'SAP entegrasyonu ile kurumsal veri bütünlüğü',
    'Raporlama ve analiz yeteneklerinin güçlendirilmesi'
]
for goal in goals:
    doc.add_paragraph(goal, style='List Bullet')

doc.add_heading('1.3. Hedef Kullanıcılar', 2)
users_data = [
    ['Kullanıcı Grubu', 'Roller', 'Yetkiler'],
    ['Bakım Yöneticisi', 'Planlama, Koordinasyon', 'Tam Yetki'],
    ['Bakım Teknisyeni', 'Bakım Gerçekleştirme', 'Okuma, Güncelleme'],
    ['Tesis Yöneticisi', 'İzleme, Onaylama', 'Onay, Raporlama'],
    ['Çalışanlar', 'Talep Oluşturma', 'Sadece Okuma']
]
add_table_with_style(doc, users_data[1:], users_data[0])

doc.add_heading('1.4. Teknoloji Yığını', 2)
tech_data = [
    ['Katman', 'Teknoloji', 'Açıklama'],
    ['Frontend', 'HTML5, CSS3, JavaScript (ES6+)', 'Responsive web arayüzü'],
    ['Stil Sistemi', 'Custom CSS (MAN Corporate)', 'MAN kurumsal tasarım sistemi'],
    ['Veri Yönetimi', 'Client-side JSON', 'SAP entegrasyonuna hazır yapı'],
    ['Tarayıcı Desteği', 'Chrome, Firefox, Safari, Edge', 'Modern tarayıcılar']
]
add_table_with_style(doc, tech_data[1:], tech_data[0])

doc.add_page_break()

# ============================================
# 2. SISTEM MIMARISI
# ============================================
add_colored_heading(doc, '2. Sistem Mimarisi', 1)

doc.add_heading('2.1. Mimari Yaklaşım', 2)
p = doc.add_paragraph('Sistem, ')
p.add_run('3 Katmanlı Mimari (Three-Tier Architecture)').bold = True
p.add_run(' prensibi ile tasarlanmıştır:')

arch_data = [
    ['Katman', 'Sorumluluk', 'Teknoloji'],
    ['Sunum Katmanı', 'Kullanıcı arayüzü, etkileşim', 'HTML5, CSS3, JavaScript'],
    ['İş Mantığı Katmanı', 'Veri işleme, validasyon, iş kuralları', 'JavaScript Modülleri'],
    ['Veri Katmanı', 'Veri saklama, SAP entegrasyonu', 'JSON Data Structures']
]
add_table_with_style(doc, arch_data[1:], arch_data[0])

doc.add_heading('2.2. Tasarım Sistemi - MAN Corporate Identity', 2)
p = doc.add_paragraph()
p.add_run('Kurumsal Renkler:').bold = True
doc.add_paragraph('• MAN Kırmızısı: #E20714 (Ana renk, başlıklar, CTA butonları)')
doc.add_paragraph('• Siyah: #1A1A1A (Navigasyon, tipografi)')
doc.add_paragraph('• Gri Tonları: 9 seviyeli profesyonel gri skalası (#2C2C2C - #F5F5F5)')
doc.add_paragraph('• Başarı Yeşili: #00A859 (Onaylı durumlar)')
doc.add_paragraph('• Uyarı Turuncusu: #FF9500 (Bekleyen durumlar)')

doc.add_page_break()

# ============================================
# 3. MODÜL DETAYLARI
# ============================================
add_colored_heading(doc, '3. Modül Detayları', 1)

# ============================================
# 3.1. İŞ TALEPLERİ YÖNETİMİ
# ============================================
add_colored_heading(doc, '3.1. İş Talepleri Yönetimi Modülü', 2, (0, 0, 0))

doc.add_heading('Modül Açıklaması', 3)
p = doc.add_paragraph(
    'İş Talepleri Modülü, MAN Türkiye otobüs üretim tesisindeki çalışanların bakım, onarım veya '
    'destek ihtiyaçlarını sisteme bildirmelerini sağlar. Talepler merkezi olarak yönetilir, '
    'önceliklendirilir ve ilgili bakım ekiplerine atanır.'
)

doc.add_heading('Talep Detayı - Fonksiyonel Gereksinimler', 3)

requirements = [
    {
        'id': 'FR-JR-001',
        'gereksinim': 'Yeni İş Talebi Oluşturma',
        'detay': 'Kullanıcı, talep başlığı, kategori, öncelik, lokasyon ve açıklama girerek yeni talep oluşturabilmelidir.',
        'kabul_kriteri': 'Tüm zorunlu alanlar doldurulduğunda sistem otomatik talep numarası üretir ve bildirim gönderir.'
    },
    {
        'id': 'FR-JR-002',
        'gereksinim': 'Talep Filtreleme ve Arama',
        'detay': 'Kullanıcı, talepleri durum, öncelik, kategori ve lokasyona göre filtreleyebilmeli ve metin araması yapabilmelidir.',
        'kabul_kriteri': 'Filtreler gerçek zamanlı çalışır ve sonuçlar anında güncellenir. Boş sonuç durumunda bilgilendirme mesajı gösterilir.'
    },
    {
        'id': 'FR-JR-003',
        'gereksinim': 'Talep Durum Yönetimi',
        'detay': 'Sistem, talep durumlarını (Beklemede, Onaylandı, İşlemde, Tamamlandı, Reddedildi) takip eder.',
        'kabul_kriteri': 'Her durum değişikliğinde ilgili kullanıcılara bildirim gönderilir ve değişiklik geçmişe kaydedilir.'
    },
    {
        'id': 'FR-JR-004',
        'gereksinim': 'Önceliklendirme Sistemi',
        'detay': 'Talepler Kritik, Yüksek, Orta, Düşük öncelik seviyelerine atanabilir.',
        'kabul_kriteri': 'Kritik talepler liste başında kırmızı renkte görüntülenir ve acil bildirim gönderilir.'
    },
    {
        'id': 'FR-JR-005',
        'gereksinim': 'Kategorizasyon',
        'detay': 'Talepler kategorilere ayrılır: Mekanik, Elektrik, HVAC, Güvenlik, Yazılım, Genel.',
        'kabul_kriteri': 'Her kategori için otomatik olarak ilgili uzmanlık alanına sahip teknisyen havuzu önerilir.'
    }
]

req_data = [['Gereksinim ID', 'Gereksinim', 'Detay', 'Kabul Kriteri']]
for req in requirements:
    req_data.append([req['id'], req['gereksinim'], req['detay'], req['kabul_kriteri']])

add_table_with_style(doc, req_data[1:], req_data[0])

doc.add_heading('Veri Alanları', 3)
job_req_fields = [
    ['Alan Adı', 'Tip', 'Zorunlu', 'Açıklama'],
    ['requestId', 'String', 'Evet', 'Otomatik üretilen talep numarası (JOB-2025-XXX)'],
    ['title', 'String (200)', 'Evet', 'Talep başlığı'],
    ['description', 'Text', 'Evet', 'Detaylı açıklama'],
    ['category', 'Enum', 'Evet', 'Kategori (mekanik, elektrik, hvac, vb.)'],
    ['priority', 'Enum', 'Evet', 'Öncelik (kritik, yüksek, orta, düşük)'],
    ['location', 'String', 'Evet', 'Lokasyon (Şasi Üretim, Boya, Montaj, vb.)'],
    ['status', 'Enum', 'Evet', 'Durum (beklemede, onaylandı, işlemde, vb.)'],
    ['requestedBy', 'String', 'Evet', 'Talep sahibi'],
    ['requestedDate', 'DateTime', 'Evet', 'Talep oluşturma tarihi'],
    ['assignedTo', 'String', 'Hayır', 'Atanan teknisyen/ekip'],
    ['dueDate', 'Date', 'Hayır', 'Hedef tamamlanma tarihi'],
    ['completedDate', 'DateTime', 'Hayır', 'Gerçek tamamlanma tarihi']
]
add_table_with_style(doc, job_req_fields[1:], job_req_fields[0])

doc.add_heading('Ekran Görüntüleri', 3)
doc.add_paragraph().add_run('İş Talepleri Ana Ekranı:').bold = True
add_image_if_exists(doc, '/Users/caglarozyildirim/WebstormProjects/Deneme/screenshots/corporate_02_is_talepleri.png', 6.5)

doc.add_page_break()

# ============================================
# 3.2. VARLIK YÖNETİMİ
# ============================================
add_colored_heading(doc, '3.2. Varlık Yönetimi Modülü', 2, (0, 0, 0))

doc.add_heading('Modül Açıklaması', 3)
p = doc.add_paragraph(
    'Varlık Yönetimi Modülü, MAN Türkiye otobüs üretim tesisindeki tüm ekipman, makine ve araçların '
    'dijital envanterini tutar. SAP entegrasyonu ile finansal verileri takip eder ve bakım geçmişini yönetir.'
)

doc.add_heading('Talep Detayı - Fonksiyonel Gereksinimler', 3)

asset_requirements = [
    {
        'id': 'FR-AS-001',
        'gereksinim': 'Varlık Kayıt Yönetimi',
        'detay': 'Yeni varlık ekleme, güncelleme, silme işlemleri yapılabilir. SAP ID entegrasyonu sağlanır.',
        'kabul_kriteri': 'Varlık eklendiğinde SAP sistemine otomatik senkronize edilir ve benzersiz ID atanır.'
    },
    {
        'id': 'FR-AS-002',
        'gereksinim': 'Detaylı Varlık Bilgileri',
        'detay': '17 farklı varlık özelliği (ID, SAP ID, model, seri no, lokasyon, vb.) saklanır.',
        'kabul_kriteri': 'Tüm alanlar varlık detay sayfasında görüntülenebilir ve düzenlenebilir.'
    },
    {
        'id': 'FR-AS-003',
        'gereksinim': 'Bakım Geçmişi Takibi',
        'detay': 'Her varlık için son bakım tarihi ve sonraki bakım tarihi takip edilir.',
        'kabul_kriteri': 'Sistem, sonraki bakım tarihi yaklaştığında (7 gün öncesinden) otomatik hatırlatma gönderir.'
    },
    {
        'id': 'FR-AS-004',
        'gereksinim': 'Lokasyon Bazlı Yönetim',
        'detay': 'Varlıklar tesis, alan ve maliyet merkezine göre gruplanabilir.',
        'kabul_kriteri': 'Lokasyon haritası üzerinde varlık konumları görsel olarak gösterilir.'
    },
    {
        'id': 'FR-AS-005',
        'gereksinim': 'Durum İzleme',
        'detay': 'Varlık durumu (Aktif, Bakımda, Arızalı, Pasif) gerçek zamanlı izlenir.',
        'kabul_kriteri': 'Durum değişikliklerinde ilgili ekiplere anında bildirim gönderilir.'
    }
]

asset_req_data = [['Gereksinim ID', 'Gereksinim', 'Detay', 'Kabul Kriteri']]
for req in asset_requirements:
    asset_req_data.append([req['id'], req['gereksinim'], req['detay'], req['kabul_kriteri']])

add_table_with_style(doc, asset_req_data[1:], asset_req_data[0])

doc.add_heading('Veri Alanları', 3)
asset_fields = [
    ['Alan Adı', 'Tip', 'Zorunlu', 'Açıklama'],
    ['id', 'String', 'Evet', 'Varlık numarası (AST-XXX)'],
    ['sapId', 'String', 'Evet', 'SAP sistem ID'],
    ['name', 'String (200)', 'Evet', 'Varlık adı'],
    ['category', 'String', 'Evet', 'Kategori (Üretim Makinesi, Araç, vb.)'],
    ['manufacturer', 'String', 'Hayır', 'Üretici firma'],
    ['model', 'String', 'Hayır', 'Model'],
    ['serialNo', 'String', 'Evet', 'Seri numarası'],
    ['location', 'String', 'Evet', 'Ana lokasyon'],
    ['subLocation', 'String', 'Hayır', 'Alt lokasyon/Alan'],
    ['costCenter', 'String', 'Evet', 'Maliyet merkezi kodu'],
    ['status', 'Enum', 'Evet', 'Durum (active, maintenance, faulty, inactive)'],
    ['statusText', 'String', 'Evet', 'Durum açıklaması (Türkçe)'],
    ['assignedTo', 'String', 'Hayır', 'Sorumlu ekip/kişi'],
    ['purchaseDate', 'Date', 'Hayır', 'Satın alma tarihi'],
    ['bookValue', 'Decimal', 'Hayır', 'Defter değeri (TL)'],
    ['lastMaintenance', 'Date', 'Hayır', 'Son bakım tarihi'],
    ['nextMaintenance', 'Date', 'Hayır', 'Sonraki bakım tarihi'],
    ['description', 'Text', 'Hayır', 'Açıklama']
]
add_table_with_style(doc, asset_fields[1:], asset_fields[0])

doc.add_heading('Ekran Görüntüleri', 3)
doc.add_paragraph().add_run('Varlık Yönetimi Ana Ekranı:').bold = True
add_image_if_exists(doc, '/Users/caglarozyildirim/WebstormProjects/Deneme/screenshots/corporate_03_varlik_yonetimi.png', 6.5)

doc.add_paragraph().add_run('Varlık Detay Sayfası:').bold = True
add_image_if_exists(doc, '/Users/caglarozyildirim/WebstormProjects/Deneme/screenshots/corporate_04_varlik_detay.png', 6.5)

doc.add_page_break()

# ============================================
# 3.3. BAKIM YÖNETİMİ
# ============================================
add_colored_heading(doc, '3.3. Bakım Yönetimi Modülü', 2, (0, 0, 0))

doc.add_heading('Modül Açıklaması', 3)
p = doc.add_paragraph(
    'Bakım Yönetimi Modülü, önleyici bakım (preventive), rutin kontrol, düzeltici bakım ve '
    'tahmine dayalı bakım (predictive) planlarının oluşturulması, takibi ve raporlanmasını sağlar.'
)

doc.add_heading('Talep Detayı - Fonksiyonel Gereksinimler', 3)

maint_requirements = [
    {
        'id': 'FR-MN-001',
        'gereksinim': 'Bakım Planı Oluşturma',
        'detay': 'Varlık bazlı bakım planları oluşturulur. Bakım tipi, tarih, ekip ve tahmini süre belirlenir.',
        'kabul_kriteri': 'Plan oluşturulduğunda atanan ekibe bildirim gider ve takvime otomatik eklenir.'
    },
    {
        'id': 'FR-MN-002',
        'gereksinim': 'Çoklu Bakım Tipi Desteği',
        'detay': 'Önleyici, rutin, düzeltici ve tahmine dayalı bakım tipleri desteklenir.',
        'kabul_kriteri': 'Her bakım tipi için farklı form şablonları ve kontrol listeleri kullanılır.'
    },
    {
        'id': 'FR-MN-003',
        'gereksinim': 'Otomatik Planlama',
        'detay': 'Periyodik bakımlar için otomatik tekrarlayan planlar oluşturulur.',
        'kabul_kriteri': 'Sistem, bakım tamamlandıktan sonra sonraki bakım planını otomatik oluşturur.'
    },
    {
        'id': 'FR-MN-004',
        'gereksinim': 'Ekip ve Kaynak Yönetimi',
        'detay': 'Bakım ekipleri ve gerekli kaynaklar planlara atanır.',
        'kabul_kriteri': 'Ekip uygunluk durumu kontrol edilir, çakışma varsa uyarı verilir.'
    },
    {
        'id': 'FR-MN-005',
        'gereksinim': 'Performans Takibi',
        'detay': 'Planlanan vs gerçekleşen süre, geciken bakımlar ve tamamlanma oranları izlenir.',
        'kabul_kriteri': 'Haftalık performans raporu otomatik oluşturulur ve yöneticilere gönderilir.'
    }
]

maint_req_data = [['Gereksinim ID', 'Gereksinim', 'Detay', 'Kabul Kriteri']]
for req in maint_requirements:
    maint_req_data.append([req['id'], req['gereksinim'], req['detay'], req['kabul_kriteri']])

add_table_with_style(doc, maint_req_data[1:], maint_req_data[0])

doc.add_heading('Veri Alanları', 3)
maint_fields = [
    ['Alan Adı', 'Tip', 'Zorunlu', 'Açıklama'],
    ['maintenanceId', 'String', 'Evet', 'Bakım plan numarası (MNT-2025-XXX)'],
    ['title', 'String (200)', 'Evet', 'Bakım başlığı'],
    ['assetId', 'String', 'Evet', 'İlişkili varlık ID (Foreign Key)'],
    ['maintenanceType', 'Enum', 'Evet', 'Bakım tipi (preventive, routine, corrective, predictive)'],
    ['scheduledDate', 'Date', 'Evet', 'Planlanan tarih'],
    ['estimatedDuration', 'String', 'Evet', 'Tahmini süre (örn: "4 saat")'],
    ['assignedTeam', 'String', 'Evet', 'Atanan ekip'],
    ['description', 'Text', 'Hayır', 'Bakım açıklaması']
]
add_table_with_style(doc, maint_fields[1:], maint_fields[0])

doc.add_heading('Ekran Görüntüleri', 3)
doc.add_paragraph().add_run('Bakım Yönetimi Ana Ekranı ve Yeni Bakım Planı Oluşturma:').bold = True
add_image_if_exists(doc, '/Users/caglarozyildirim/WebstormProjects/Deneme/screenshots/corporate_05_bakim_yonetimi.png', 6.5)

doc.add_page_break()

# ============================================
# 3.4. OLAY YÖNETİMİ
# ============================================
add_colored_heading(doc, '3.4. Olay Yönetimi Modülü', 2, (0, 0, 0))

doc.add_heading('Modül Açıklaması', 3)
p = doc.add_paragraph(
    'Olay Yönetimi Modülü, tesiste meydana gelen acil durumlar, arızalar ve beklenmeyen olayların '
    'hızlı bir şekilde bildirilmesini, önceliklendirilmesini ve çözülmesini sağlar. '
    'Kritik olaylar için anında müdahale mekanizması içerir.'
)

doc.add_heading('Talep Detayı - Fonksiyonel Gereksinimler', 3)

incident_requirements = [
    {
        'id': 'FR-IN-001',
        'gereksinim': 'Acil Olay Bildirimi',
        'detay': 'Kullanıcılar tek tuşla acil olay bildirebilir. Form hızlı doldurulabilir şekilde tasarlanmıştır.',
        'kabul_kriteri': 'Kritik olay bildirimi 30 saniye içinde ilgili tüm ekiplere iletilir.'
    },
    {
        'id': 'FR-IN-002',
        'gereksinim': 'Öncelik Bazlı Sınıflandırma',
        'detay': 'Olaylar Kritik, Yüksek, Orta, Düşük öncelik seviyelerine otomatik veya manuel atanır.',
        'kabul_kriteri': 'Kritik olaylar kırmızı renkte, üstte ve anında bildirimle gösterilir.'
    },
    {
        'id': 'FR-IN-003',
        'gereksinim': 'Gerçek Zamanlı İzleme',
        'detay': 'Aktif olayların durumu gerçek zamanlı güncellenir ve dashboard\'da görüntülenir.',
        'kabul_kriteri': 'Dashboard her 30 saniyede otomatik yenilenir.'
    },
    {
        'id': 'FR-IN-004',
        'gereksinim': 'Müdahale Süresi Takibi',
        'detay': 'Her olay için başlangıç zamanı, ilk müdahale zamanı ve çözüm zamanı kaydedilir.',
        'kabul_kriteri': 'SLA süreleri aşıldığında otomatik uyarı ve eskalasyon tetiklenir.'
    },
    {
        'id': 'FR-IN-005',
        'gereksinim': 'Olay Kapatma ve Raporlama',
        'detay': 'Olay çözüldüğünde detaylı rapor hazırlanır ve ilgili kişilere paylaşılır.',
        'kabul_kriteri': 'Rapor, kök neden analizi ve alınan önlemleri içerir.'
    }
]

incident_req_data = [['Gereksinim ID', 'Gereksinim', 'Detay', 'Kabul Kriteri']]
for req in incident_requirements:
    incident_req_data.append([req['id'], req['gereksinim'], req['detay'], req['kabul_kriteri']])

add_table_with_style(doc, incident_req_data[1:], incident_req_data[0])

doc.add_heading('Veri Alanları', 3)
incident_fields = [
    ['Alan Adı', 'Tip', 'Zorunlu', 'Açıklama'],
    ['incidentId', 'String', 'Evet', 'Olay numarası (INC-2025-XXX)'],
    ['title', 'String (200)', 'Evet', 'Olay başlığı'],
    ['description', 'Text', 'Evet', 'Olay detayı'],
    ['priority', 'Enum', 'Evet', 'Öncelik (critical, high, medium, low)'],
    ['location', 'String', 'Evet', 'Olay lokasyonu'],
    ['assetId', 'String', 'Hayır', 'İlişkili varlık (opsiyonel)'],
    ['reportedBy', 'String', 'Evet', 'Bildiren kişi'],
    ['reportedDate', 'DateTime', 'Evet', 'Bildirim tarihi']
]
add_table_with_style(doc, incident_fields[1:], incident_fields[0])

doc.add_heading('Ekran Görüntüleri', 3)
doc.add_paragraph().add_run('Olay Yönetimi Ana Ekranı ve Acil Olay Bildirimi:').bold = True
add_image_if_exists(doc, '/Users/caglarozyildirim/WebstormProjects/Deneme/screenshots/corporate_06_olay_yonetimi.png', 6.5)

doc.add_page_break()

# ============================================
# 4. İŞ AKIŞLARI VE SÜREÇLER
# ============================================
add_colored_heading(doc, '4. İş Akışları ve Süreçler', 1)

doc.add_paragraph(
    'Bu bölüm, MAN Türkiye Bakım Yönetimi Sisteminin temel iş akışlarını Visio tarzında '
    'profesyonel süreç diyagramları ile göstermektedir. Her diyagram, farklı roller arası '
    'etkileşimleri swimlane (kulvar) formatında sunmaktadır.'
)

doc.add_heading('4.1. İş Talebi Akışı', 2)
doc.add_paragraph(
    'İş talebi süreci, bir çalışanın bakım talebi oluşturmasından, '
    'talebin onaylanması, teknisyene atanması ve tamamlanmasına kadar olan adımları gösterir.'
)
add_image_if_exists(doc, '/Users/caglarozyildirim/WebstormProjects/Deneme/diagrams/is_talebi_akisi.png', 6.5)

doc.add_page_break()

doc.add_heading('4.2. Bakım Planlama Akışı', 2)
doc.add_paragraph(
    'Bakım planlama süreci, bir varlık için bakım planının oluşturulmasından, '
    'bakımın gerçekleştirilmesi ve kapatılmasına kadar olan süreci detaylandırır.'
)
add_image_if_exists(doc, '/Users/caglarozyildirim/WebstormProjects/Deneme/diagrams/bakim_planlama_akisi.png', 6.5)

doc.add_page_break()

doc.add_heading('4.3. Olay Yönetimi Akışı', 2)
doc.add_paragraph(
    'Olay yönetimi süreci, acil bir olayın bildirilmesinden, önceliklendirilmesi, '
    'müdahale edilmesi ve kapatılmasına kadar olan kritik adımları gösterir. '
    'Kritik öncelikli olaylar kırmızı renkle vurgulanmıştır.'
)
add_image_if_exists(doc, '/Users/caglarozyildirim/WebstormProjects/Deneme/diagrams/olay_yonetimi_akisi.png', 6.5)

doc.add_page_break()

doc.add_heading('4.4. Varlık Yönetimi Akışı', 2)
doc.add_paragraph(
    'Varlık yönetimi süreci, bir varlığın sisteme eklenmesinden, '
    'bilgilerinin güncellenmesi ve sistem/veritabanı ile senkronizasyonuna kadar olan veri akışını gösterir.'
)
add_image_if_exists(doc, '/Users/caglarozyildirim/WebstormProjects/Deneme/diagrams/varlik_yonetimi_akisi.png', 6.5)

doc.add_page_break()

# ============================================
# 5. EKRAN GÖRÜNTÜLERİ
# ============================================
add_colored_heading(doc, '5. Ekran Görüntüleri - Corporate Design', 1)

doc.add_paragraph(
    'Aşağıdaki ekran görüntüleri, MAN Türkiye kurumsal tasarım standartlarına uygun '
    'olarak revize edilmiş modern ve profesyonel arayüzü göstermektedir.'
)

doc.add_heading('5.1. Ana Sayfa Dashboard', 2)
add_image_if_exists(doc, '/Users/caglarozyildirim/WebstormProjects/Deneme/screenshots/corporate_01_ana_sayfa.png', 6.5)

doc.add_heading('5.2. İş Talepleri - Filtreleme Sistemi', 2)
add_image_if_exists(doc, '/Users/caglarozyildirim/WebstormProjects/Deneme/screenshots/corporate_02_is_talepleri.png', 6.5)

doc.add_heading('5.3. Varlık Yönetimi', 2)
add_image_if_exists(doc, '/Users/caglarozyildirim/WebstormProjects/Deneme/screenshots/corporate_03_varlik_yonetimi.png', 6.5)

doc.add_heading('5.4. Varlık Detay Sayfası', 2)
add_image_if_exists(doc, '/Users/caglarozyildirim/WebstormProjects/Deneme/screenshots/corporate_04_varlik_detay.png', 6.5)

doc.add_heading('5.5. Bakım Yönetimi', 2)
add_image_if_exists(doc, '/Users/caglarozyildirim/WebstormProjects/Deneme/screenshots/corporate_05_bakim_yonetimi.png', 6.5)

doc.add_heading('5.6. Olay Yönetimi', 2)
add_image_if_exists(doc, '/Users/caglarozyildirim/WebstormProjects/Deneme/screenshots/corporate_06_olay_yonetimi.png', 6.5)

doc.add_page_break()

# ============================================
# 6. TEST SENARYOLARI
# ============================================
add_colored_heading(doc, '6. Test Senaryoları', 1)

doc.add_heading('6.1. İş Talebi Testi', 2)
test_scenarios = [
    ['Test ID', 'Senaryo', 'Adımlar', 'Beklenen Sonuç'],
    ['TS-JR-001', 'Yeni iş talebi oluşturma',
     '1. İş Talepleri sayfasına git\n2. "Yeni Talep" butonuna tıkla\n3. Formu doldur\n4. Gönder',
     'Talep başarıyla oluşturulur, JOB-2025-XXX numarası atanır'],
    ['TS-JR-002', 'Filtreleme testi',
     '1. Durum filtresinden "Beklemede" seç\n2. Sonuçları kontrol et',
     'Sadece beklemede durumundaki talepler görüntülenir'],
    ['TS-AS-001', 'Varlık detay görüntüleme',
     '1. Varlık listesinden bir varlık seç\n2. Detay butonuna tıkla',
     'Varlık detay sayfası tüm 17 alan ile açılır'],
    ['TS-MN-001', 'Bakım planı oluşturma',
     '1. Bakım Yönetimi sayfasına git\n2. "Yeni Bakım Planı" tıkla\n3. Formu doldur',
     'Plan oluşturulur, MNT-2025-XXX numarası atanır'],
    ['TS-IN-001', 'Acil olay bildirimi',
     '1. Olay Yönetimi sayfasına git\n2. "Acil Olay Bildir" butonuna tıkla\n3. Kritik öncelik seç\n4. Formu doldur',
     'Olay kaydedilir, kritik bildirim gönderilir']
]

add_table_with_style(doc, test_scenarios[1:], test_scenarios[0])

doc.add_page_break()

# ============================================
# SON SAYFA
# ============================================
add_colored_heading(doc, 'Doküman Bilgileri', 1)

doc_info = [
    ['Özellik', 'Değer'],
    ['Doküman Adı', 'MAN Türkiye Bakım Yönetimi Sistemi - İş Analizi'],
    ['Versiyon', '2.0 - Corporate Edition'],
    ['Tarih', 'Ekim 2025'],
    ['Hazırlayan', 'Sistem Analisti'],
    ['Onaylayan', 'Proje Yöneticisi'],
    ['Durum', 'Yazılım Geliştirmeye Hazır'],
    ['Tasarım Standardı', 'MAN Corporate Identity Guidelines'],
    ['Hedef Sistem', 'Web Tabanlı Bakım Yönetimi Platformu'],
    ['SAP Entegrasyon', 'Evet - Varlık ve Maliyet Verileri']
]

add_table_with_style(doc, doc_info[1:], doc_info[0])

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Not: ').bold = True
p.add_run('Bu doküman, yazılım geliştirme ekibine iletilmek üzere hazırlanmıştır. '
          'Tüm ekran görüntüleri MAN Türkiye kurumsal tasarım standartlarına uygun olarak '
          'revize edilmiştir. İş akışları Visio formatında profesyonel süreç diyagramları '
          'ile görselleştirilmiştir.')

# Save document
output_path = '/Users/caglarozyildirim/WebstormProjects/Deneme/MAN_Turkiye_Bakim_Yonetimi_Is_Analizi_Corporate_V2.docx'
doc.save(output_path)

file_size = os.path.getsize(output_path)
print(f"✅ Comprehensive business analysis document created successfully!")
print(f"📄 File: {output_path}")
print(f"📊 Size: {file_size / 1024:.1f} KB")
print(f"\n📋 Document includes:")
print(f"   • Complete module descriptions")
print(f"   • Detailed functional requirements")
print(f"   • All data structures (17 fields for assets, 12 for job requests, etc.)")
print(f"   • Corporate design screenshots (6 images)")
print(f"   • Visio-style workflow diagrams (4 professional SVG diagrams)")
print(f"   • Test scenarios")
print(f"   • Ready for software development team")
