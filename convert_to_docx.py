#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re

def convert_markdown_to_docx(md_file, docx_file):
    """Convert Markdown file to DOCX with formatting"""

    # Read markdown file
    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    # Create document
    doc = Document()

    # Set up styles
    styles = doc.styles

    # Title style
    try:
        title_style = styles['Title']
    except:
        title_style = styles.add_style('Title', WD_STYLE_TYPE.PARAGRAPH)
    title_style.font.size = Pt(24)
    title_style.font.bold = True
    title_style.font.color.rgb = RGBColor(0, 51, 102)

    # Process each line
    in_table = False
    table_data = []
    in_code_block = False

    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        # Skip empty lines (but add paragraph for spacing)
        if not line.strip():
            if not in_table:
                doc.add_paragraph()
            i += 1
            continue

        # Handle H1 (# )
        if line.startswith('# ') and not line.startswith('## '):
            text = line[2:].strip()
            p = doc.add_heading(text, level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue

        # Handle H2 (## )
        if line.startswith('## ') and not line.startswith('### '):
            text = line[3:].strip()
            doc.add_heading(text, level=2)
            i += 1
            continue

        # Handle H3 (### )
        if line.startswith('### ') and not line.startswith('#### '):
            text = line[4:].strip()
            doc.add_heading(text, level=3)
            i += 1
            continue

        # Handle H4 (#### )
        if line.startswith('#### '):
            text = line[5:].strip()
            doc.add_heading(text, level=4)
            i += 1
            continue

        # Handle bold text (**text**)
        if line.startswith('**') and line.endswith('**') and len(line) > 4:
            text = line[2:-2]
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = True
            i += 1
            continue

        # Handle horizontal rule (---)
        if line.strip() == '---':
            doc.add_paragraph('_' * 80)
            i += 1
            continue

        # Handle bullet lists (- )
        if line.startswith('- ') or line.startswith('  - ') or line.startswith('    - '):
            indent_level = (len(line) - len(line.lstrip())) // 2
            text = line.strip()[2:]

            # Remove markdown bold
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)

            p = doc.add_paragraph(text, style='List Bullet')
            if indent_level > 0:
                p.paragraph_format.left_indent = Inches(0.5 * indent_level)
            i += 1
            continue

        # Handle numbered lists (1. )
        if re.match(r'^\d+\.\s', line):
            text = re.sub(r'^\d+\.\s', '', line)
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            doc.add_paragraph(text, style='List Number')
            i += 1
            continue

        # Handle tables
        if '|' in line and not in_table:
            # Start of table
            table_data = []
            in_table = True

        if in_table:
            if '|' in line:
                # Parse table row
                cells = [cell.strip() for cell in line.split('|')]
                cells = [c for c in cells if c]  # Remove empty cells

                # Skip separator line (|---|---|)
                if all(re.match(r'^-+$', cell) for cell in cells):
                    i += 1
                    continue

                table_data.append(cells)
                i += 1
                continue
            else:
                # End of table
                if table_data:
                    # Create table
                    num_rows = len(table_data)
                    num_cols = max(len(row) for row in table_data)

                    table = doc.add_table(rows=num_rows, cols=num_cols)
                    table.style = 'Light Grid Accent 1'

                    for row_idx, row_data in enumerate(table_data):
                        for col_idx, cell_data in enumerate(row_data):
                            if col_idx < num_cols:
                                cell = table.rows[row_idx].cells[col_idx]
                                cell.text = cell_data
                                # Bold first row (header)
                                if row_idx == 0:
                                    for paragraph in cell.paragraphs:
                                        for run in paragraph.runs:
                                            run.bold = True

                    doc.add_paragraph()  # Add spacing after table

                in_table = False
                table_data = []
                continue

        # Handle regular paragraphs
        text = line

        # Process inline bold
        if '**' in text:
            p = doc.add_paragraph()
            parts = re.split(r'(\*\*.*?\*\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                elif part:
                    p.add_run(part)
        else:
            doc.add_paragraph(text)

        i += 1

    # Handle remaining table if file ends with table
    if in_table and table_data:
        num_rows = len(table_data)
        num_cols = max(len(row) for row in table_data)

        table = doc.add_table(rows=num_rows, cols=num_cols)
        table.style = 'Light Grid Accent 1'

        for row_idx, row_data in enumerate(table_data):
            for col_idx, cell_data in enumerate(row_data):
                if col_idx < num_cols:
                    cell = table.rows[row_idx].cells[col_idx]
                    cell.text = cell_data
                    if row_idx == 0:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.bold = True

    # Save document
    doc.save(docx_file)
    print(f"✅ DOCX dokümanı başarıyla oluşturuldu!")
    print(f"📄 Dosya yolu: {docx_file}")

    return docx_file

if __name__ == "__main__":
    md_file = "/Users/caglarozyildirim/WebstormProjects/Deneme/Bakim_Yonetim_Uygulamasi_Is_Analizi.md"
    docx_file = "/Users/caglarozyildirim/WebstormProjects/Deneme/Bakim_Yonetim_Uygulamasi_Is_Analizi.docx"

    convert_markdown_to_docx(md_file, docx_file)
