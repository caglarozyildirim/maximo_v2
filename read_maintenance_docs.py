#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import json
from docx import Document
import openpyxl
from pathlib import Path

def read_docx(file_path):
    """Read a .docx file and return its content"""
    try:
        doc = Document(file_path)
        content = {
            'paragraphs': [],
            'tables': []
        }

        # Read paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                content['paragraphs'].append({
                    'text': para.text,
                    'style': para.style.name if para.style else 'Normal'
                })

        # Read tables
        for table_idx, table in enumerate(doc.tables):
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)
            content['tables'].append({
                'index': table_idx,
                'data': table_data
            })

        return content
    except Exception as e:
        return {'error': str(e)}

def read_xlsx(file_path):
    """Read an .xlsx file and return its content"""
    try:
        wb = openpyxl.load_workbook(file_path, data_only=True)
        content = {}

        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            sheet_data = []

            for row in sheet.iter_rows(values_only=True):
                # Convert None to empty string
                row_data = [str(cell) if cell is not None else '' for cell in row]
                sheet_data.append(row_data)

            content[sheet_name] = sheet_data

        return content
    except Exception as e:
        return {'error': str(e)}

def main():
    base_path = Path("/Users/caglarozyildirim/Desktop/Şirketler/MAN Türkiye/Maintenance Management Application")

    results = {}

    # Read main requirement document
    print("Reading Requirement Analysis...")
    req_doc = base_path / "Maintenance Management Application Requirement Analysis (draft).docx"
    results['requirement_analysis'] = read_docx(req_doc)

    # Read Asset Assignment Form
    print("Reading Asset Assignment Form...")
    asset_form = base_path / "Asset Assignment Form.docx"
    results['asset_assignment_form'] = read_docx(asset_form)

    # Read Asset Retirement Printout
    print("Reading Asset Retirement Printout...")
    retirement_doc = base_path / "Asset Retirement Printout.docx"
    results['asset_retirement_printout'] = read_docx(retirement_doc)

    # Read Data Structure
    print("Reading Data Structure...")
    data_structure = base_path / "Data Structure.xlsx"
    results['data_structure'] = read_xlsx(data_structure)

    # Read Screen Designs
    print("Reading Screen Designs...")
    screen_designs = base_path / "Screen Designs.xlsx"
    results['screen_designs'] = read_xlsx(screen_designs)

    # Read Overall Status
    print("Reading Overall Status...")
    overall_status = base_path / "Overall Status.xlsx"
    results['overall_status'] = read_xlsx(overall_status)

    # Read Use Cases
    print("Reading Use Cases...")
    use_cases_path = base_path / "Use Cases"
    results['use_cases'] = {}

    for uc_file in use_cases_path.glob("*.docx"):
        if not uc_file.name.startswith('~'):  # Skip temp files
            uc_name = uc_file.stem
            results['use_cases'][uc_name] = read_docx(uc_file)

    for uc_file in use_cases_path.glob("*.xlsx"):
        if not uc_file.name.startswith('~'):
            uc_name = uc_file.stem
            results['use_cases'][uc_name] = read_xlsx(uc_file)

    # Save results to JSON
    output_file = "/Users/caglarozyildirim/WebstormProjects/Deneme/maintenance_docs_content.json"
    with open(output_file, 'w', encoding='utf-8') as f:
        json.dump(results, f, ensure_ascii=False, indent=2)

    print(f"\nAll documents read successfully!")
    print(f"Results saved to: {output_file}")
    print(f"\nSummary:")
    print(f"- Requirement Analysis: {len(results['requirement_analysis'].get('paragraphs', []))} paragraphs, {len(results['requirement_analysis'].get('tables', []))} tables")
    print(f"- Data Structure: {len(results['data_structure'])} sheets")
    print(f"- Screen Designs: {len(results['screen_designs'])} sheets")
    print(f"- Overall Status: {len(results['overall_status'])} sheets")
    print(f"- Use Cases: {len(results['use_cases'])} documents")

if __name__ == "__main__":
    main()