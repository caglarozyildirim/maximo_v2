#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
Magdeburger doküman formatını analiz eder
"""

from docx import Document
import json

def read_magdeburger_format():
    """Magdeburger doküman formatını okur"""

    doc_path = '/Users/caglarozyildirim/Desktop/Teklifler/Magdeburger - Fraud Süreç Teklifi.docx'
    doc = Document(doc_path)

    content = {
        'paragraphs': [],
        'tables': [],
        'headings': []
    }

    for para in doc.paragraphs:
        if para.text.strip():
            style = para.style.name if para.style else 'Normal'
            content['paragraphs'].append({
                'text': para.text,
                'style': style
            })

            if 'Heading' in style:
                content['headings'].append(para.text)

    for i, table in enumerate(doc.tables):
        table_data = []
        for row in table.rows:
            row_data = [cell.text for cell in row.cells]
            table_data.append(row_data)
        content['tables'].append({
            'table_no': i + 1,
            'data': table_data
        })

    # Save to JSON
    output_path = '/tmp/magdeburger_format.json'
    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(content, f, ensure_ascii=False, indent=2)

    print(f"✓ Magdeburger format analizi tamamlandı: {output_path}")
    print(f"\nBaşlıklar ({len(content['headings'])}):")
    for heading in content['headings'][:15]:
        print(f"  - {heading}")

    print(f"\nToplam paragraf: {len(content['paragraphs'])}")
    print(f"Toplam tablo: {len(content['tables'])}")

if __name__ == '__main__':
    read_magdeburger_format()
