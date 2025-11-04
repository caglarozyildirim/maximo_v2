#!/usr/bin/env python3
"""
Create Koleksiyon Mobilya Dealer Portal - Business Analysis & Scope Document
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

# Koleksiyon Corporate Colors
KOLEKSIYON_BLUE = RGBColor(0, 51, 102)
KOLEKSIYON_ORANGE = RGBColor(255, 102, 0)
DARK_GRAY = RGBColor(51, 51, 51)
LIGHT_GRAY = RGBColor(128, 128, 128)
SUCCESS_GREEN = RGBColor(34, 139, 34)

doc = Document()

# Set default font
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

print("="*80)
print("CREATING KOLEKSIYON MOBILYA DEALER PORTAL")
print("BUSINESS ANALYSIS & SCOPE DOCUMENT")
print("="*80 + "\n")

# ============================================================================
# TITLE PAGE
# ============================================================================
print("📄 Creating title page...")

title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title_para.add_run("KOLEKSİYON MOBİLYA\nBAYİ PORTALI PROJESİ")
title_run.font.size = Pt(24)
title_run.font.bold = True
title_run.font.color.rgb = KOLEKSIYON_BLUE

subtitle_para = doc.add_paragraph()
subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_run = subtitle_para.add_run("\nİş Analizi ve Kapsam Dokümanı")
subtitle_run.font.size = Pt(16)
subtitle_run.font.color.rgb = KOLEKSIYON_ORANGE

doc.add_paragraph("\n" * 3)

# Document info table
info_para = doc.add_paragraph()
info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
info_run = info_para.add_run(f"Versiyon: 1.0\n")
info_run.font.size = Pt(12)
info_run = info_para.add_run(f"Tarih: {datetime.now().strftime('%d.%m.%Y')}\n")
info_run.font.size = Pt(12)
info_run = info_para.add_run("Hazırlayan: İş Analisti")
info_run.font.size = Pt(12)

doc.add_page_break()

# ============================================================================
# TABLE OF CONTENTS
# ============================================================================
print("📋 Creating table of contents...")

toc_heading = doc.add_heading('İÇİNDEKİLER', 1)
toc_heading.runs[0].font.color.rgb = KOLEKSIYON_BLUE

toc_items = [
    "1. YÖNETİCİ ÖZETİ",
    "2. PROJE TANIMI VE AMAÇ",
    "3. MEVCUT DURUM ANALİZİ",
    "4. İŞ PROBLEMİ VE ÇÖZÜM",
    "5. PROJE KAPSAMI",
    "6. KAPSAM DIŞI KONULAR",
    "7. FONKSİYONEL GEREKSİNİMLER",
    "8. TEKNİK GEREKSİNİMLER",
    "9. SİSTEM MİMARİSİ VE ENTEGRASYONLAR",
    "10. VERİ YAPILARI",
    "11. KULLANICI ROLLERİ VE YETKİLENDİRME",
    "12. İŞ KURALLARI",
    "13. KULLANICI DENEYİMİ (UX) VE ARAYÜZ TASARIMI",
    "14. GÜVENLİK VE ERİŞİM YÖNETİMİ",
    "15. PROJE TAKVİMİ VE KİLOMETRE TAŞLARI",
    "16. RİSKLER VE VARSAYIMLAR",
    "17. BAŞARI KRİTERLERİ",
    "18. GELECEK AŞAMALAR VE ÖNERİLER"
]

for item in toc_items:
    toc_para = doc.add_paragraph(item, style='List Number')
    toc_para.runs[0].font.size = Pt(11)

doc.add_page_break()

# ============================================================================
# 1. YÖNETİCİ ÖZETİ
# ============================================================================
print("📝 Section 1: Yönetici Özeti...")

heading = doc.add_heading('1. YÖNETİCİ ÖZETİ', 1)
heading.runs[0].font.color.rgb = KOLEKSIYON_BLUE

doc.add_paragraph(
    "Koleksiyon Mobilya, ana ERP sistemi olarak SAP kullanan ve katalog yönetimi için "
    "mevcut CMS (Butterfly/PIM) sistemini kullanan lider bir mobilya üreticisidir. "
    "Bu proje kapsamında, bayilerin SAP'de yer alan ürün fiyat listelerine erişebileceği, "
    "web tabanlı güvenli bir Bayi Portalı geliştirilecektir."
)

doc.add_paragraph(
    "Projenin temel amacı, bayilere ürün bilgileri ve fiyatlandırma kırılımlarını "
    "çevrimiçi olarak sunmak, manuel süreçleri azaltmak ve bayi memnuniyetini artırmaktır. "
    "Portal, mevcut CMS altyapısı üzerine entegre edilecek ve SAP'den gerçek zamanlı veri "
    "akışı sağlanacaktır."
)

# Key highlights
doc.add_heading('Proje Özellikleri', 2)
highlights_table = doc.add_table(rows=1, cols=2)
highlights_table.style = 'Light Grid Accent 1'

highlights_data = [
    ("Proje Adı", "Koleksiyon Mobilya Bayi Portalı"),
    ("Proje Tipi", "Web Tabanlı Portal Geliştirme"),
    ("Ana Entegrasyon", "SAP ERP & Butterfly CMS"),
    ("Hedef Kullanıcılar", "Bayiler, Distribütörler, Yetkili Müşteriler"),
    ("Tahmini Süre", "3-4 Ay (Faz 1)"),
    ("Proje Durumu", "Planlama ve Gereksinim Analizi"),
]

for label, value in highlights_data:
    row_cells = highlights_table.add_row().cells
    row_cells[0].text = label
    row_cells[1].text = value
    # Bold first column
    for para in row_cells[0].paragraphs:
        for run in para.runs:
            run.font.bold = True
            run.font.color.rgb = KOLEKSIYON_BLUE

doc.add_paragraph()

# ============================================================================
# 2. PROJE TANIMI VE AMAÇ
# ============================================================================
print("📝 Section 2: Proje Tanımı ve Amaç...")

doc.add_page_break()
heading = doc.add_heading('2. PROJE TANIMI VE AMAÇ', 1)
heading.runs[0].font.color.rgb = KOLEKSIYON_BLUE

doc.add_heading('2.1. Proje Tanımı', 2)
doc.add_paragraph(
    "Koleksiyon Mobilya Bayi Portalı, bayilerin ürün fiyat listelerine, teknik "
    "dokümanlara ve ürün özelliklerine güvenli bir şekilde erişebilecekleri, web tabanlı "
    "bir self-servis platformdur. Portal, SAP ERP sisteminden gelen fiyatlandırma verilerini "
    "mevcut CMS (Butterfly PIM) altyapısı üzerinden bayilere sunacaktır."
)

doc.add_heading('2.2. Proje Amaçları', 2)

objectives = [
    ("İş Amaçları", [
        "Bayi memnuniyetini ve deneyimini iyileştirmek",
        "Fiyat listesi paylaşım süreçlerini dijitalleştirmek",
        "Manuel e-posta ve Excel gönderimlerini azaltmak",
        "Bayilere 7/24 erişim imkanı sağlamak",
        "Müşteri ilişkilerini güçlendirmek"
    ]),
    ("Operasyonel Amaçlar", [
        "Fiyat güncellemelerini otomatikleştirmek",
        "SAP ve CMS arasında gerçek zamanlı veri senkronizasyonu sağlamak",
        "Para birimi bazlı fiyatlandırma kontrolü yapmak",
        "Bayi onay süreçlerini sistematikleştirmek",
        "Erişim yetkilendirmesini merkezi yönetmek"
    ]),
    ("Teknik Amaçlar", [
        "Mevcut CMS altyapısını genişletmek",
        "SAP ile güvenli API entegrasyonu kurmak",
        "Çok dilli (TR/EN) destek sağlamak",
        "Responsive ve kullanıcı dostu arayüz geliştirmek",
        "Güvenli kimlik doğrulama ve yetkilendirme uygulamak"
    ])
]

for category, items in objectives:
    para = doc.add_paragraph()
    run = para.add_run(f"• {category}:")
    run.font.bold = True
    run.font.color.rgb = KOLEKSIYON_ORANGE

    for item in items:
        doc.add_paragraph(f"  - {item}", style='List Bullet 2')

doc.add_paragraph()

# ============================================================================
# 3. MEVCUT DURUM ANALİZİ
# ============================================================================
print("📝 Section 3: Mevcut Durum Analizi...")

doc.add_page_break()
heading = doc.add_heading('3. MEVCUT DURUM ANALİZİ', 1)
heading.runs[0].font.color.rgb = KOLEKSIYON_BLUE

doc.add_heading('3.1. Mevcut Sistem Mimarisi', 2)

current_systems = doc.add_table(rows=1, cols=3)
current_systems.style = 'Light Grid Accent 1'

# Headers
headers = current_systems.rows[0].cells
headers[0].text = "Sistem"
headers[1].text = "Kullanım Amacı"
headers[2].text = "Durum"

for cell in headers:
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.bold = True
            run.font.color.rgb = DARK_GRAY

# Data
systems_data = [
    ("SAP ERP", "Ana iş süreçleri, stok, fiyatlandırma, envanter yönetimi", "Aktif - Ana Sistem"),
    ("Butterfly CMS (PIM)", "Katalog yönetimi, ürün bilgileri, içerik yönetimi", "Aktif - Katalog"),
    ("Web Sitesi", "Ürün tanıtımı, kurumsal iletişim, bayilik başvuruları", "Aktif - Public"),
    ("E-posta / Excel", "Fiyat listesi paylaşımı (Manuel)", "Aktif - Manuel Süreç")
]

for system, purpose, status in systems_data:
    row_cells = current_systems.add_row().cells
    row_cells[0].text = system
    row_cells[1].text = purpose
    row_cells[2].text = status

doc.add_paragraph()

doc.add_heading('3.2. Mevcut Süreç Akışı', 2)

current_process_steps = [
    "Bayi, fiyat listesi talebini e-posta veya telefon ile iletir",
    "Koleksiyon yetkili personeli talebi alır ve değerlendirir",
    "Para birimi ve yetkilendirme kontrolü manuel yapılır",
    "SAP'den ilgili fiyat listesi Excel formatında export edilir",
    "Fiyat listesi e-posta ile bayiye gönderilir",
    "Bayi, Excel dosyasını indirir ve kullanır",
    "Fiyat güncellemelerinde süreç tekrar eder"
]

for i, step in enumerate(current_process_steps, 1):
    para = doc.add_paragraph(f"{i}. {step}")
    para.style = 'List Number'

doc.add_paragraph()

doc.add_heading('3.3. Mevcut Durumun Sorunları', 2)

problems_table = doc.add_table(rows=1, cols=2)
problems_table.style = 'Light Grid Accent 1'

problems_headers = problems_table.rows[0].cells
problems_headers[0].text = "Sorun"
problems_headers[1].text = "Etki"

for cell in problems_headers:
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(139, 0, 0)  # Dark red

problems_data = [
    ("Manuel süreçler", "Zaman kaybı, insan hatası riski, düşük verimlilik"),
    ("Gecikmeli bilgi akışı", "Güncel olmayan fiyat bilgileri, bayi memnuniyetsizliği"),
    ("Takip zorluğu", "Kimin hangi fiyat listesini gördüğü takip edilemiyor"),
    ("Güvenlik riskleri", "E-posta ile gönderilen Excel dosyaları kontrol dışı paylaşılabiliyor"),
    ("Ölçeklenebilirlik sorunu", "Bayi sayısı arttıkça manuel süreç yönetilemez hale geliyor"),
    ("Çok dilli destek eksikliği", "Uluslararası bayiler için iletişim zorluğu"),
    ("Para birimi yönetimi", "Her bayi için ayrı Excel hazırlanması gerekiyor")
]

for problem, impact in problems_data:
    row_cells = problems_table.add_row().cells
    row_cells[0].text = problem
    row_cells[1].text = impact

doc.add_paragraph()

# ============================================================================
# 4. İŞ PROBLEMİ VE ÇÖZÜM
# ============================================================================
print("📝 Section 4: İş Problemi ve Çözüm...")

doc.add_page_break()
heading = doc.add_heading('4. İŞ PROBLEMİ VE ÇÖZÜM', 1)
heading.runs[0].font.color.rgb = KOLEKSIYON_BLUE

doc.add_heading('4.1. İş Problemi', 2)

problem_para = doc.add_paragraph()
problem_run = problem_para.add_run(
    "Koleksiyon Mobilya'nın mevcut bayi fiyat listesi paylaşım süreci manuel, yavaş "
    "ve hata yapmaya açıktır. Bayiler güncel fiyat bilgilerine anında erişememekte, "
    "Koleksiyon personeli her talep için manuel işlem yapmak zorunda kalmaktadır. "
    "Bu durum hem operasyonel maliyetleri artırmakta hem de bayi deneyimini olumsuz "
    "etkilemektedir."
)

doc.add_paragraph()

doc.add_heading('4.2. Önerilen Çözüm', 2)

solution_para = doc.add_paragraph()
solution_run = solution_para.add_run(
    "Web tabanlı bir Bayi Portalı geliştirerek, onaylı bayilerin SAP'deki güncel "
    "fiyat listelerine 7/24 self-servis olarak erişmeleri sağlanacaktır. Portal, "
    "mevcut CMS (Butterfly) altyapısı üzerine entegre edilecek ve SAP ile gerçek "
    "zamanlı veri alışverişi yapacaktır."
)

doc.add_paragraph()

doc.add_heading('4.3. Çözümün Faydaları', 2)

benefits_table = doc.add_table(rows=1, cols=3)
benefits_table.style = 'Light Grid Accent 1'

benefits_headers = benefits_table.rows[0].cells
benefits_headers[0].text = "Paydaş"
benefits_headers[1].text = "Fayda"
benefits_headers[2].text = "Ölçülebilir Sonuç"

for cell in benefits_headers:
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.bold = True
            run.font.color.rgb = SUCCESS_GREEN

benefits_data = [
    ("Bayiler", "7/24 güncel fiyat erişimi, self-servis", "Talep süresinde %90 azalma"),
    ("Koleksiyon Satış Ekibi", "Manuel işlem yükünden kurtulma", "Operasyonel maliyet %70 azalma"),
    ("Koleksiyon Yönetimi", "Merkezi kontrol ve raporlama", "Gerçek zamanlı kullanım metrikleri"),
    ("IT Ekibi", "Otomatik entegrasyon", "Manuel export-import %100 kaldırma"),
    ("Müşteriler", "Daha hızlı teklif alma", "Satış döngüsünde hızlanma")
]

for stakeholder, benefit, result in benefits_data:
    row_cells = benefits_table.add_row().cells
    row_cells[0].text = stakeholder
    row_cells[1].text = benefit
    row_cells[2].text = result

doc.add_paragraph()

# ============================================================================
# 5. PROJE KAPSAMI
# ============================================================================
print("📝 Section 5: Proje Kapsamı...")

doc.add_page_break()
heading = doc.add_heading('5. PROJE KAPSAMI', 1)
heading.runs[0].font.color.rgb = KOLEKSIYON_BLUE

doc.add_heading('5.1. Kapsam İçi Özellikler (In-Scope)', 2)

in_scope_items = {
    "Kullanıcı Yönetimi": [
        "Bayi kayıt (register) sistemi",
        "Login/Logout işlemleri",
        "Şifre sıfırlama ve yenileme",
        "Profil yönetimi (İletişim bilgileri, şirket bilgileri)",
        "Dil seçimi (TR/EN)"
    ],
    "Onay ve Yetkilendirme": [
        "Yeni bayi kayıt taleplerinin Koleksiyon'a bildirilmesi (e-posta)",
        "Admin panelinden bayi onaylama/reddetme",
        "Para birimi bazlı yetkilendirme (TRY, USD, EUR, vb.)",
        "Rol tabanlı erişim kontrolü"
    ],
    "Fiyat Listesi Görüntüleme": [
        "SAP fiyatlandırma kodları görüntüleme (35/36 kodları)",
        "3000'lu son kırılım kodları listeleme",
        "Ürün bazlı fiyat detayları",
        "Para birimine göre fiyat gösterimi",
        "Fiyat listesi filtreleme ve arama",
        "Excel export özelliği"
    ],
    "Ürün Bilgileri": [
        "Ürün tip kodu bilgileri",
        "Ürün teknik özellikler (mevcut PIM'den)",
        "PDF spec dosyaları indirme",
        "Ürün görselleri (CMS'den)"
    ],
    "SAP Entegrasyonu": [
        "SAP'den gerçek zamanlı fiyat çekimi (API)",
        "Fiyat güncelleme senkronizasyonu",
        "Ürün kodu eşleştirmesi",
        "Para birimi dönüşümleri"
    ],
    "CMS Entegrasyonu": [
        "Butterfly PIM ile entegrasyon",
        "Ürün katalog bilgilerinin gösterimi",
        "Teknik doküman erişimi",
        "Görsel ve içerik yönetimi"
    ],
    "Web Sitesi Entegrasyonu": [
        "Ana web sitesinden portal yönlendirmesi",
        "\"Bayi Portalı\" erişim linki (Biz Kimiz > İletişim > Bayilik Başvurusu)",
        "Gizli/dolaylı erişim noktası (SEO dışı)",
        "Single Sign-On (SSO) hazırlığı (gelecek için altyapı)"
    ],
    "Güvenlik": [
        "HTTPS/SSL sertifikası",
        "Güvenli kimlik doğrulama (JWT/OAuth2)",
        "Parola şifreleme",
        "Oturum yönetimi ve timeout",
        "Brute-force koruması",
        "GDPR uyumlu veri yönetimi"
    ],
    "Raporlama (Admin)": [
        "Aktif bayi listesi",
        "Bayi erişim logları",
        "En çok görüntülenen ürünler",
        "Fiyat listesi export istatistikleri"
    ]
}

for category, items in in_scope_items.items():
    para = doc.add_paragraph()
    run = para.add_run(f"📌 {category}")
    run.font.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = KOLEKSIYON_ORANGE

    for item in items:
        doc.add_paragraph(f"✓ {item}", style='List Bullet 2')

doc.add_paragraph()

# ============================================================================
# 6. KAPSAM DIŞI KONULAR
# ============================================================================
print("📝 Section 6: Kapsam Dışı Konular...")

doc.add_page_break()
heading = doc.add_heading('6. KAPSAM DIŞI KONULAR (Out-of-Scope)', 1)
heading.runs[0].font.color.rgb = KOLEKSIYON_BLUE

doc.add_paragraph(
    "Aşağıdaki özellikler Faz 1 kapsamında yer almamaktadır. Bu özellikler gelecek "
    "fazlarda değerlendirilebilir veya ayrı projeler olarak ele alınabilir:"
)

out_of_scope_items = {
    "Sipariş Yönetimi": [
        "Online sipariş verme",
        "Sepet ve checkout süreci",
        "Ödeme entegrasyonu",
        "Sipariş takibi"
    ],
    "Stok Yönetimi": [
        "Gerçek zamanlı stok sorgulama",
        "Stok rezervasyonu",
        "Teslimat tarihi tahmini"
    ],
    "Gelişmiş Ürün Özellikleri (Faz 2)": [
        "360° ürün görselleri",
        "B, T, U alanları (Boyut, Teknik resim, Uzun metin)",
        "Ürün karşılaştırma",
        "Kişiselleştirilmiş kataloglar",
        "Bakım talimatları otomatik üretimi"
    ],
    "Lojistik ve Kargo": [
        "Kargo takibi",
        "Ambalaj bilgileri",
        "Nakliye hesaplamaları"
    ],
    "CRM Entegrasyonu": [
        "Müşteri ilişkileri yönetimi",
        "Lead tracking",
        "Satış raporlamaları"
    ],
    "Mobil Uygulama": [
        "iOS/Android native app",
        "Push notification"
    ],
    "Toplu İşlemler": [
        "Excel'den toplu ürün yükleme",
        "Toplu fiyat güncelleme"
    ],
    "Diğer": [
        "SAP'e geri besleme (PIM'den SAP'e veri yazma)",
        "Multi-currency otomatik döviz çevrimi (manuel seçim olacak)",
        "Kullanım senaryoları otomatik yazımı",
        "Eski ürün arşivleme (manuel olarak yapılacak)"
    ]
}

for category, items in out_of_scope_items.items():
    para = doc.add_paragraph()
    run = para.add_run(f"❌ {category}")
    run.font.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(139, 0, 0)  # Dark red

    for item in items:
        doc.add_paragraph(f"  • {item}", style='List Bullet 2')

doc.add_paragraph()

# ============================================================================
# 7. FONKSİYONEL GEREKSİNİMLER
# ============================================================================
print("📝 Section 7: Fonksiyonel Gereksinimler...")

doc.add_page_break()
heading = doc.add_heading('7. FONKSİYONEL GEREKSİNİMLER', 1)
heading.runs[0].font.color.rgb = KOLEKSIYON_BLUE

doc.add_heading('7.1. Kullanıcı Kayıt ve Giriş', 2)

requirements_table = doc.add_table(rows=1, cols=4)
requirements_table.style = 'Light Grid Accent 1'

req_headers = requirements_table.rows[0].cells
req_headers[0].text = "REQ ID"
req_headers[1].text = "Gereksinim"
req_headers[2].text = "Öncelik"
req_headers[3].text = "Detay"

for cell in req_headers:
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.bold = True
            run.font.size = Pt(9)

requirements_data = [
    ("FR-001", "Bayi Kayıt Formu", "Yüksek", "Şirket adı, vergi no, yetkili kişi, e-posta, telefon, adres, dil seçimi"),
    ("FR-002", "E-posta Doğrulama", "Yüksek", "Kayıt sonrası doğrulama e-postası gönderilmeli"),
    ("FR-003", "Admin Onay Süreci", "Yüksek", "Yeni kayıtlar için Koleksiyon'a bildirim gitmeli, admin onaylamalı"),
    ("FR-004", "Para Birimi Seçimi", "Yüksek", "Admin, bayi için para birimi ataması yapmalı (TRY, USD, EUR, vb.)"),
    ("FR-005", "Login Sistemi", "Yüksek", "E-posta ve şifre ile giriş, 'Beni Hatırla' özelliği"),
    ("FR-006", "Şifre Sıfırlama", "Orta", "E-posta ile şifre sıfırlama linki gönderimi"),
    ("FR-007", "Profil Güncelleme", "Orta", "Bayi kendi profilini güncelleyebilmeli"),
]

for req_id, req, priority, detail in requirements_data:
    row_cells = requirements_table.add_row().cells
    row_cells[0].text = req_id
    row_cells[1].text = req
    row_cells[2].text = priority
    row_cells[3].text = detail

    # Color code priority
    if priority == "Yüksek":
        for para in row_cells[2].paragraphs:
            for run in para.runs:
                run.font.color.rgb = RGBColor(255, 0, 0)
                run.font.bold = True

doc.add_paragraph()

doc.add_heading('7.2. Fiyat Listesi ve Ürün Görüntüleme', 2)

price_req_table = doc.add_table(rows=1, cols=4)
price_req_table.style = 'Light Grid Accent 1'

# Headers (reuse structure)
headers = price_req_table.rows[0].cells
headers[0].text = "REQ ID"
headers[1].text = "Gereksinim"
headers[2].text = "Öncelik"
headers[3].text = "Detay"

for cell in headers:
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.bold = True
            run.font.size = Pt(9)

price_requirements = [
    ("FR-101", "Fiyat Listesi Görüntüleme", "Yüksek", "SAP fiyatlandırma kodları (35/36), 3000'lu kırılımlar"),
    ("FR-102", "Para Birimine Göre Gösterim", "Yüksek", "Bayinin atanmış para birimi ile fiyat gösterimi"),
    ("FR-103", "Ürün Filtreleme", "Yüksek", "Ürün kodu, kategori, fiyat aralığı filtresi"),
    ("FR-104", "Arama İşlevi", "Yüksek", "Ürün adı, kod, özellik bazlı arama"),
    ("FR-105", "Excel Export", "Orta", "Görüntülenen liste Excel formatında indirilebilmeli"),
    ("FR-106", "PDF Spec İndirme", "Orta", "Ürün tip koduna ait PDF teknik döküman indirme"),
    ("FR-107", "Dil Desteği", "Yüksek", "TR/EN dil seçimine göre ürün bilgileri gösterimi"),
    ("FR-108", "Ürün Görselleri", "Orta", "CMS'den ürün görselleri entegrasyonu"),
]

for req_id, req, priority, detail in price_requirements:
    row_cells = price_req_table.add_row().cells
    row_cells[0].text = req_id
    row_cells[1].text = req
    row_cells[2].text = priority
    row_cells[3].text = detail

    if priority == "Yüksek":
        for para in row_cells[2].paragraphs:
            for run in para.runs:
                run.font.color.rgb = RGBColor(255, 0, 0)
                run.font.bold = True

doc.add_paragraph()

doc.add_heading('7.3. Admin Paneli', 2)

admin_req_table = doc.add_table(rows=1, cols=4)
admin_req_table.style = 'Light Grid Accent 1'

headers = admin_req_table.rows[0].cells
headers[0].text = "REQ ID"
headers[1].text = "Gereksinim"
headers[2].text = "Öncelik"
headers[3].text = "Detay"

for cell in headers:
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.bold = True
            run.font.size = Pt(9)

admin_requirements = [
    ("FR-201", "Bayi Onay/Red", "Yüksek", "Bekleyen başvuruları onaylama veya reddetme"),
    ("FR-202", "Para Birimi Atama", "Yüksek", "Bayi bazlı para birimi seçimi ve güncelleme"),
    ("FR-203", "Bayi Yönetimi", "Yüksek", "Bayi listesi, detayları, aktivasyon/deaktivasyon"),
    ("FR-204", "Erişim Logları", "Orta", "Bayi login, görüntüleme, indirme logları"),
    ("FR-205", "Raporlama", "Orta", "Kullanım istatistikleri, en çok görüntülenen ürünler"),
    ("FR-206", "E-posta Bildirimleri", "Orta", "Yeni başvuru, onay/red bildirimleri"),
]

for req_id, req, priority, detail in admin_requirements:
    row_cells = admin_req_table.add_row().cells
    row_cells[0].text = req_id
    row_cells[1].text = req
    row_cells[2].text = priority
    row_cells[3].text = detail

    if priority == "Yüksek":
        for para in row_cells[2].paragraphs:
            for run in para.runs:
                run.font.color.rgb = RGBColor(255, 0, 0)
                run.font.bold = True

doc.add_paragraph()

# ============================================================================
# 8. TEKNİK GEREKSİNİMLER
# ============================================================================
print("📝 Section 8: Teknik Gereksinimler...")

doc.add_page_break()
heading = doc.add_heading('8. TEKNİK GEREKSİNİMLER', 1)
heading.runs[0].font.color.rgb = KOLEKSIYON_BLUE

tech_req_table = doc.add_table(rows=1, cols=3)
tech_req_table.style = 'Light Grid Accent 1'

headers = tech_req_table.rows[0].cells
headers[0].text = "Kategori"
headers[1].text = "Gereksinim"
headers[2].text = "Teknik Detay"

for cell in headers:
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.bold = True

tech_requirements = [
    ("Frontend", "Responsive Web Tasarımı", "Desktop, Tablet, Mobile uyumlu (Bootstrap/Tailwind)"),
    ("Frontend", "Modern Framework", "React.js / Vue.js / Angular önerisi"),
    ("Frontend", "Çok Dilli Destek", "i18n kütüphanesi (TR/EN)"),
    ("Backend", "API Geliştirme", "RESTful API (Node.js/Python/PHP)"),
    ("Backend", "Kimlik Doğrulama", "JWT veya OAuth2 token bazlı auth"),
    ("Backend", "Session Yönetimi", "Redis ile session storage"),
    ("Veritabanı", "İlişkisel Veritabanı", "PostgreSQL veya MySQL"),
    ("Veritabanı", "Veri Modeli", "Kullanıcılar, Bayiler, Fiyat Listeleri, Loglar"),
    ("Entegrasyon", "SAP API", "SAP OData/RFC protokolü ile fiyat çekimi"),
    ("Entegrasyon", "CMS API", "Butterfly PIM API entegrasyonu"),
    ("Entegrasyon", "E-posta Servisi", "SMTP / SendGrid / AWS SES"),
    ("Güvenlik", "SSL/TLS", "HTTPS zorunlu, sertifika yönetimi"),
    ("Güvenlik", "Parola Hash", "bcrypt/Argon2 ile şifreleme"),
    ("Güvenlik", "Rate Limiting", "API request limitleri"),
    ("Güvenlik", "CORS", "Cross-origin kontrolleri"),
    ("Performans", "Caching", "Redis cache (fiyat listeleri, ürün bilgileri)"),
    ("Performans", "CDN", "Statik dosyalar için CDN kullanımı"),
    ("Performans", "Response Time", "Sayfa yükleme < 2 saniye"),
    ("Hosting", "Cloud Platform", "AWS / Azure / Google Cloud önerisi"),
    ("Hosting", "Konteynerizasyon", "Docker container yapısı"),
    ("Hosting", "CI/CD", "Otomatik deployment pipeline"),
    ("Monitoring", "Log Yönetimi", "Centralized logging (ELK/Splunk)"),
    ("Monitoring", "Error Tracking", "Sentry/Rollbar entegrasyonu"),
    ("Monitoring", "Uptime Monitoring", "Pingdom/UptimeRobot"),
]

for category, requirement, detail in tech_requirements:
    row_cells = tech_req_table.add_row().cells
    row_cells[0].text = category
    row_cells[1].text = requirement
    row_cells[2].text = detail

    # Bold category
    for para in row_cells[0].paragraphs:
        for run in para.runs:
            run.font.bold = True
            run.font.color.rgb = KOLEKSIYON_BLUE

doc.add_paragraph()

# ============================================================================
# 9. SİSTEM MİMARİSİ VE ENTEGRASYONLAR
# ============================================================================
print("📝 Section 9: Sistem Mimarisi...")

doc.add_page_break()
heading = doc.add_heading('9. SİSTEM MİMARİSİ VE ENTEGRASYONLAR', 1)
heading.runs[0].font.color.rgb = KOLEKSIYON_BLUE

doc.add_heading('9.1. Mimari Genel Bakış', 2)

doc.add_paragraph(
    "Bayi Portalı, 3-katmanlı (3-tier) modern web mimarisi kullanacaktır:"
)

arch_layers = [
    ("Presentation Layer (Sunum Katmanı)", [
        "Web Browser (Responsive UI)",
        "React.js / Vue.js Frontend",
        "Çok dilli arayüz (i18n)"
    ]),
    ("Application Layer (Uygulama Katmanı)", [
        "RESTful API Gateway",
        "Authentication Service (JWT/OAuth2)",
        "Business Logic Layer",
        "Integration Services (SAP, CMS)"
    ]),
    ("Data Layer (Veri Katmanı)", [
        "PostgreSQL/MySQL Database",
        "Redis Cache",
        "File Storage (PDF, Images)"
    ])
]

for layer, components in arch_layers:
    para = doc.add_paragraph()
    run = para.add_run(f"🔹 {layer}")
    run.font.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = KOLEKSIYON_BLUE

    for component in components:
        doc.add_paragraph(f"  • {component}", style='List Bullet 2')

doc.add_paragraph()

doc.add_heading('9.2. Entegrasyon Detayları', 2)

integration_table = doc.add_table(rows=1, cols=4)
integration_table.style = 'Light Grid Accent 1'

headers = integration_table.rows[0].cells
headers[0].text = "Entegrasyon"
headers[1].text = "Protokol"
headers[2].text = "Veri Akışı"
headers[3].text = "Sıklık"

for cell in headers:
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.bold = True

integrations = [
    ("SAP ERP → Portal", "OData/RFC API", "Fiyat listeleri, ürün kodları", "Gerçek zamanlı / 15 dk cache"),
    ("CMS (Butterfly) → Portal", "REST API", "Ürün bilgileri, görseller, PDF'ler", "On-demand / 1 saat cache"),
    ("Portal → E-posta", "SMTP/SendGrid", "Kayıt onay bildirimleri", "Event-based"),
    ("Portal → Web Sitesi", "SSO / Deep Link", "Kullanıcı yönlendirme", "On-demand"),
]

for source_target, protocol, data_flow, frequency in integrations:
    row_cells = integration_table.add_row().cells
    row_cells[0].text = source_target
    row_cells[1].text = protocol
    row_cells[2].text = data_flow
    row_cells[3].text = frequency

doc.add_paragraph()

doc.add_heading('9.3. Veri Akış Diyagramı', 2)

doc.add_paragraph("Temel veri akışı:")

flow_steps = [
    "1. Bayi, web sitesinden 'Bayi Portalı' linkine tıklar",
    "2. Portal login sayfasına yönlendirilir",
    "3. Bayi giriş yapar (Authentication Service doğrulama yapar)",
    "4. Portal, SAP API'den bayinin para birimine göre fiyat listesini çeker",
    "5. CMS API'den ürün detayları ve görseller alınır",
    "6. Veri cache'lenir (Redis)",
    "7. Frontend'e JSON formatında döndürülür",
    "8. Kullanıcı arayüzde fiyat listesini görür ve filtreler",
    "9. Gerekirse Excel export işlemi yapılır",
    "10. Tüm işlemler loglanır (audit trail)"
]

for step in flow_steps:
    doc.add_paragraph(step, style='List Number')

doc.add_paragraph()

# ============================================================================
# 10. VERİ YAPILARI
# ============================================================================
print("📝 Section 10: Veri Yapıları...")

doc.add_page_break()
heading = doc.add_heading('10. VERİ YAPILARI', 1)
heading.runs[0].font.color.rgb = KOLEKSIYON_BLUE

doc.add_heading('10.1. Ana Veri Tabloları', 2)

# Users Table
doc.add_paragraph()
table_title = doc.add_paragraph()
run = table_title.add_run("📋 Tablo: users (Kullanıcılar)")
run.font.bold = True
run.font.size = Pt(11)
run.font.color.rgb = KOLEKSIYON_ORANGE

users_table = doc.add_table(rows=1, cols=4)
users_table.style = 'Light Grid Accent 1'

headers = users_table.rows[0].cells
headers[0].text = "Alan Adı"
headers[1].text = "Veri Tipi"
headers[2].text = "Zorunlu"
headers[3].text = "Açıklama"

for cell in headers:
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.bold = True
            run.font.size = Pt(9)

users_fields = [
    ("id", "UUID", "Evet", "Primary Key"),
    ("email", "VARCHAR(255)", "Evet", "Unique, login için"),
    ("password_hash", "VARCHAR(255)", "Evet", "bcrypt hash"),
    ("first_name", "VARCHAR(100)", "Evet", "Adı"),
    ("last_name", "VARCHAR(100)", "Evet", "Soyadı"),
    ("phone", "VARCHAR(20)", "Hayır", "Telefon"),
    ("language", "ENUM('tr','en')", "Evet", "Varsayılan: 'tr'"),
    ("role", "ENUM('dealer','admin')", "Evet", "Kullanıcı rolü"),
    ("is_active", "BOOLEAN", "Evet", "Aktif/pasif durum"),
    ("is_approved", "BOOLEAN", "Evet", "Admin onayı"),
    ("approved_by", "UUID", "Hayır", "Onaylayan admin ID"),
    ("approved_at", "TIMESTAMP", "Hayır", "Onay tarihi"),
    ("created_at", "TIMESTAMP", "Evet", "Kayıt tarihi"),
    ("updated_at", "TIMESTAMP", "Evet", "Güncelleme tarihi"),
]

for field_name, data_type, required, description in users_fields:
    row_cells = users_table.add_row().cells
    row_cells[0].text = field_name
    row_cells[1].text = data_type
    row_cells[2].text = required
    row_cells[3].text = description

    for para in row_cells[0].paragraphs:
        for run in para.runs:
            run.font.name = 'Courier New'

doc.add_paragraph()

# Dealers Table
table_title = doc.add_paragraph()
run = table_title.add_run("📋 Tablo: dealers (Bayiler)")
run.font.bold = True
run.font.size = Pt(11)
run.font.color.rgb = KOLEKSIYON_ORANGE

dealers_table = doc.add_table(rows=1, cols=4)
dealers_table.style = 'Light Grid Accent 1'

headers = dealers_table.rows[0].cells
headers[0].text = "Alan Adı"
headers[1].text = "Veri Tipi"
headers[2].text = "Zorunlu"
headers[3].text = "Açıklama"

for cell in headers:
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.bold = True
            run.font.size = Pt(9)

dealers_fields = [
    ("id", "UUID", "Evet", "Primary Key"),
    ("user_id", "UUID", "Evet", "Foreign Key → users.id"),
    ("company_name", "VARCHAR(255)", "Evet", "Şirket adı"),
    ("tax_number", "VARCHAR(50)", "Evet", "Vergi numarası"),
    ("address", "TEXT", "Hayır", "Adres bilgisi"),
    ("city", "VARCHAR(100)", "Hayır", "Şehir"),
    ("country", "VARCHAR(100)", "Evet", "Ülke"),
    ("currency", "VARCHAR(3)", "Evet", "Para birimi kodu (TRY, USD, EUR)"),
    ("sap_customer_code", "VARCHAR(50)", "Hayır", "SAP müşteri kodu (varsa)"),
    ("created_at", "TIMESTAMP", "Evet", "Kayıt tarihi"),
    ("updated_at", "TIMESTAMP", "Evet", "Güncelleme tarihi"),
]

for field_name, data_type, required, description in dealers_fields:
    row_cells = dealers_table.add_row().cells
    row_cells[0].text = field_name
    row_cells[1].text = data_type
    row_cells[2].text = required
    row_cells[3].text = description

    for para in row_cells[0].paragraphs:
        for run in para.runs:
            run.font.name = 'Courier New'

doc.add_paragraph()

# Price Lists Table
table_title = doc.add_paragraph()
run = table_title.add_run("📋 Tablo: price_lists (Fiyat Listeleri - Cache)")
run.font.bold = True
run.font.size = Pt(11)
run.font.color.rgb = KOLEKSIYON_ORANGE

price_table = doc.add_table(rows=1, cols=4)
price_table.style = 'Light Grid Accent 1'

headers = price_table.rows[0].cells
headers[0].text = "Alan Adı"
headers[1].text = "Veri Tipi"
headers[2].text = "Zorunlu"
headers[3].text = "Açıklama"

for cell in headers:
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.bold = True
            run.font.size = Pt(9)

price_fields = [
    ("id", "UUID", "Evet", "Primary Key"),
    ("product_code", "VARCHAR(50)", "Evet", "SAP ürün kodu (35/36 formatı)"),
    ("product_name", "VARCHAR(255)", "Evet", "Ürün adı"),
    ("category", "VARCHAR(100)", "Hayır", "Ürün kategorisi"),
    ("price_code_3000", "VARCHAR(50)", "Evet", "3000'lu son kırılım kodu"),
    ("price_try", "DECIMAL(10,2)", "Hayır", "TRY fiyat"),
    ("price_usd", "DECIMAL(10,2)", "Hayır", "USD fiyat"),
    ("price_eur", "DECIMAL(10,2)", "Hayır", "EUR fiyat"),
    ("pdf_spec_url", "TEXT", "Hayır", "PDF spec dosya linki"),
    ("sap_sync_date", "TIMESTAMP", "Evet", "SAP'den son senkronizasyon"),
    ("created_at", "TIMESTAMP", "Evet", "Kayıt tarihi"),
    ("updated_at", "TIMESTAMP", "Evet", "Güncelleme tarihi"),
]

for field_name, data_type, required, description in price_fields:
    row_cells = price_table.add_row().cells
    row_cells[0].text = field_name
    row_cells[1].text = data_type
    row_cells[2].text = required
    row_cells[3].text = description

    for para in row_cells[0].paragraphs:
        for run in para.runs:
            run.font.name = 'Courier New'

doc.add_paragraph()

# Access Logs Table
table_title = doc.add_paragraph()
run = table_title.add_run("📋 Tablo: access_logs (Erişim Logları)")
run.font.bold = True
run.font.size = Pt(11)
run.font.color.rgb = KOLEKSIYON_ORANGE

logs_table = doc.add_table(rows=1, cols=4)
logs_table.style = 'Light Grid Accent 1'

headers = logs_table.rows[0].cells
headers[0].text = "Alan Adı"
headers[1].text = "Veri Tipi"
headers[2].text = "Zorunlu"
headers[3].text = "Açıklama"

for cell in headers:
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.bold = True
            run.font.size = Pt(9)

log_fields = [
    ("id", "UUID", "Evet", "Primary Key"),
    ("user_id", "UUID", "Evet", "Foreign Key → users.id"),
    ("action", "VARCHAR(50)", "Evet", "login, view_price, export_excel, vb."),
    ("resource", "VARCHAR(255)", "Hayır", "Erişilen kaynak (ürün kodu, vb.)"),
    ("ip_address", "VARCHAR(45)", "Evet", "Kullanıcı IP"),
    ("user_agent", "TEXT", "Hayır", "Browser bilgisi"),
    ("timestamp", "TIMESTAMP", "Evet", "İşlem zamanı"),
]

for field_name, data_type, required, description in log_fields:
    row_cells = logs_table.add_row().cells
    row_cells[0].text = field_name
    row_cells[1].text = data_type
    row_cells[2].text = required
    row_cells[3].text = description

    for para in row_cells[0].paragraphs:
        for run in para.runs:
            run.font.name = 'Courier New'

doc.add_paragraph()

# ============================================================================
# 11. KULLANICI ROLLERİ VE YETKİLENDİRME
# ============================================================================
print("📝 Section 11: Kullanıcı Rolleri...")

doc.add_page_break()
heading = doc.add_heading('11. KULLANICI ROLLERİ VE YETKİLENDİRME', 1)
heading.runs[0].font.color.rgb = KOLEKSIYON_BLUE

roles_table = doc.add_table(rows=1, cols=3)
roles_table.style = 'Light Grid Accent 1'

headers = roles_table.rows[0].cells
headers[0].text = "Rol"
headers[1].text = "Yetkiler"
headers[2].text = "Kısıtlamalar"

for cell in headers:
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.bold = True

roles_data = [
    ("Bayi (Dealer)",
     "• Fiyat listesi görüntüleme\n• Ürün bilgileri görüntüleme\n• PDF indirme\n• Excel export\n• Profil güncelleme",
     "• Sadece kendi para biriminde fiyat görür\n• Başka bayileri göremez\n• Admin paneline erişemez\n• SAP'e yazma yetkisi yok"),

    ("Admin",
     "• Tüm bayi yönetimi\n• Yeni bayi onaylama/reddetme\n• Para birimi atama\n• Erişim loglarını görme\n• Raporlama\n• Tüm fiyat listelerini görme",
     "• SAP ve CMS verilerini değiştiremez\n• Sadece okuma ve yönetim yetkisi"),

    ("Sistem Admin (IT)",
     "• Tüm admin yetkileri\n• Sistem konfigürasyonu\n• Entegrasyon ayarları\n• Kullanıcı rol yönetimi",
     "• İş verilerini değiştirmez (audit trail için)"),
]

for role, permissions, restrictions in roles_data:
    row_cells = roles_table.add_row().cells
    row_cells[0].text = role
    row_cells[1].text = permissions
    row_cells[2].text = restrictions

    for para in row_cells[0].paragraphs:
        for run in para.runs:
            run.font.bold = True
            run.font.color.rgb = KOLEKSIYON_BLUE

doc.add_paragraph()

# ============================================================================
# 12. İŞ KURALLARI
# ============================================================================
print("📝 Section 12: İş Kuralları...")

doc.add_page_break()
heading = doc.add_heading('12. İŞ KURALLARI', 1)
heading.runs[0].font.color.rgb = KOLEKSIYON_BLUE

business_rules = {
    "Kayıt ve Onay": [
        "BR-001: Yeni bayi kayıtları otomatik olarak 'Onay Bekliyor' durumunda olmalıdır",
        "BR-002: Admin onayı olmadan bayi portale erişememelidir",
        "BR-003: Her bayi için sadece bir para birimi atanmalıdır",
        "BR-004: Para birimi ataması sadece admin tarafından yapılabilir",
        "BR-005: E-posta adresi unique olmalıdır (her e-posta sadece bir hesap)",
    ],
    "Fiyat Görüntüleme": [
        "BR-101: Bayi sadece kendisine atanan para biriminde fiyat görmelidir",
        "BR-102: Fiyatlar SAP'den güncel olarak çekilmelidir (max 15 dk gecikme)",
        "BR-103: Ürün PDF spec dosyaları CMS'den dinamik olarak alınmalıdır",
        "BR-104: Eğer bir ürünün fiyatı SAP'de yoksa, 'Fiyat Bilgisi Yok' gösterilmeli",
        "BR-105: 35/36 fiyatlandırma kodları ve 3000'lu kırılımlar SAP yapısına uygun gösterilmeli",
    ],
    "Güvenlik": [
        "BR-201: Şifreler en az 8 karakter, büyük/küçük harf ve rakam içermelidir",
        "BR-202: 5 başarısız login denemesinden sonra hesap 15 dakika kilitlenmelidir",
        "BR-203: Session timeout 30 dakika olmalıdır",
        "BR-204: Tüm hassas işlemler audit log'a kaydedilmelidir",
        "BR-205: HTTPS zorunlu olmalıdır (HTTP kabul edilmemeli)",
    ],
    "Dil ve Lokalizasyon": [
        "BR-301: Kullanıcı dil seçimi (TR/EN) profile kayıtlı olmalıdır",
        "BR-302: Ürün bilgileri login diline göre gösterilmelidir",
        "BR-303: E-posta bildirimleri kullanıcının diline göre gönderilmelidir",
        "BR-304: Sayısal formatlar para birimine göre ayarlanmalıdır (. vs , kullanımı)",
    ],
    "Cache ve Performans": [
        "BR-401: Fiyat listeleri Redis'te 15 dakika cache'lenmelidir",
        "BR-402: Ürün görselleri CDN üzerinden sunulmalıdır",
        "BR-403: API response time 2 saniyeyi geçmemelidir",
        "BR-404: Aynı anda 100+ concurrent user desteklenmelidir",
    ],
    "Veri Senkronizasyonu": [
        "BR-501: SAP ile senkronizasyon hatası durumunda retry mekanizması olmalıdır (3 deneme)",
        "BR-502: Senkronizasyon hataları admin'e bildirilmelidir",
        "BR-503: CMS'den alınan PDF linklerinin geçerliliği kontrol edilmelidir",
        "BR-504: Veri tutarsızlığı durumunda 'Veri Yüklenemedi' mesajı gösterilmeli",
    ]
}

for category, rules in business_rules.items():
    para = doc.add_paragraph()
    run = para.add_run(f"📌 {category}")
    run.font.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = KOLEKSIYON_ORANGE

    for rule in rules:
        doc.add_paragraph(rule, style='List Bullet 2')

doc.add_paragraph()

# ============================================================================
# 13. KULLANICI DENEYİMİ (UX) VE ARAYÜZ TASARIMI
# ============================================================================
print("📝 Section 13: UX ve Arayüz Tasarımı...")

doc.add_page_break()
heading = doc.add_heading('13. KULLANICI DENEYİMİ (UX) VE ARAYÜZ TASARIMI', 1)
heading.runs[0].font.color.rgb = KOLEKSIYON_BLUE

doc.add_heading('13.1. Tasarım Prensipleri', 2)

design_principles = [
    "Sadelik ve Kullanım Kolaylığı: Bayi kullanıcıları teknik bilgisi olmayabilir, arayüz sezgisel olmalı",
    "Koleksiyon Kurumsal Kimliği: Mevcut web sitesi ile görsel uyum sağlanmalı",
    "Responsive Tasarım: Desktop, tablet ve mobil cihazlarda sorunsuz çalışmalı",
    "Hızlı Erişim: 3 tık içinde fiyat listesine ulaşılabilmeli",
    "Görsel Hiyerarşi: Önemli bilgiler (fiyat, ürün kodu) vurgulanmalı",
    "Hata Yönetimi: Kullanıcı dostu hata mesajları ve yönlendirmeler",
]

for i, principle in enumerate(design_principles, 1):
    para = doc.add_paragraph(f"{i}. {principle}")

doc.add_paragraph()

doc.add_heading('13.2. Ana Sayfalar ve Akışlar', 2)

pages_table = doc.add_table(rows=1, cols=3)
pages_table.style = 'Light Grid Accent 1'

headers = pages_table.rows[0].cells
headers[0].text = "Sayfa"
headers[1].text = "Amaç"
headers[2].text = "Önemli Elementler"

for cell in headers:
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.bold = True

pages_data = [
    ("Karşılama/Landing",
     "İlk izlenim, güven oluşturma",
     "Logo, hoşgeldin mesajı, Login/Register butonları, Steelcase benzeri tasarım"),

    ("Kayıt Formu",
     "Yeni bayi kaydı",
     "Şirket bilgileri, iletişim, dil seçimi, GDPR onayı, Kayıt butonu"),

    ("Onay Bekleme",
     "Kayıt sonrası bilgilendirme",
     "Onay bekliyor mesajı, tahmini süre, iletişim bilgileri"),

    ("Login Sayfası",
     "Giriş yapma",
     "E-posta/Şifre, Beni hatırla, Şifremi unuttum linki"),

    ("Dashboard",
     "Ana sayfa, hızlı erişim",
     "Hoşgeldin mesajı, son görüntülenen ürünler, hızlı arama, menü"),

    ("Fiyat Listesi",
     "Ana iş akışı",
     "Ürün listesi (tablo), filtreler, arama, para birimi, Excel export"),

    ("Ürün Detayı",
     "Detaylı bilgi",
     "Ürün görseli, teknik özellikler, fiyat kırılımı, PDF indirme"),

    ("Profil Sayfası",
     "Kullanıcı ayarları",
     "Kişisel bilgiler, şifre değiştirme, dil seçimi, güncelle butonu"),

    ("Admin Paneli",
     "Bayi yönetimi",
     "Bekleyen onaylar, bayi listesi, loglar, raporlar, ayarlar"),
]

for page, purpose, elements in pages_data:
    row_cells = pages_table.add_row().cells
    row_cells[0].text = page
    row_cells[1].text = purpose
    row_cells[2].text = elements

    for para in row_cells[0].paragraphs:
        for run in para.runs:
            run.font.bold = True
            run.font.color.rgb = KOLEKSIYON_BLUE

doc.add_paragraph()

doc.add_heading('13.3. Referans Tasarım', 2)

doc.add_paragraph(
    "📌 Steelcase Dealer Portal benzeri karşılama sayfası istenmektedir. "
    "Örnek özellikler:"
)

steelcase_features = [
    "Minimalist ve profesyonel görünüm",
    "Büyük hero image veya video arkaplan",
    "Net Login ve Register butonları",
    "Güven oluşturan kurumsal mesajlar",
    "Dil seçici (TR/EN) header'da",
    "Footer'da iletişim ve destek bilgileri"
]

for feature in steelcase_features:
    doc.add_paragraph(f"  • {feature}", style='List Bullet')

doc.add_paragraph()

# ============================================================================
# 14. GÜVENLİK VE ERİŞİM YÖNETİMİ
# ============================================================================
print("📝 Section 14: Güvenlik...")

doc.add_page_break()
heading = doc.add_heading('14. GÜVENLİK VE ERİŞİM YÖNETİMİ', 1)
heading.runs[0].font.color.rgb = KOLEKSIYON_BLUE

security_table = doc.add_table(rows=1, cols=3)
security_table.style = 'Light Grid Accent 1'

headers = security_table.rows[0].cells
headers[0].text = "Güvenlik Konusu"
headers[1].text = "Önlem"
headers[2].text = "Uygulama"

for cell in headers:
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.bold = True

security_measures = [
    ("Kimlik Doğrulama",
     "JWT/OAuth2 token bazlı",
     "Login sonrası access token + refresh token"),

    ("Parola Güvenliği",
     "bcrypt hash + salt",
     "Minimum 8 karakter, karmaşıklık kontrolü"),

    ("Session Yönetimi",
     "30 dakika timeout",
     "Redis ile session storage, otomatik logout"),

    ("Brute-Force Koruması",
     "Rate limiting",
     "5 başarısız denemede 15 dk bekleme"),

    ("HTTPS Zorunluluğu",
     "SSL/TLS sertifikası",
     "HTTP trafiği HTTPS'e yönlendirilir"),

    ("API Güvenliği",
     "API Key + Token",
     "Her request'te authorization header kontrolü"),

    ("CORS Politikası",
     "Whitelist",
     "Sadece izinli domain'lerden request kabul"),

    ("SQL Injection",
     "Prepared statements",
     "ORM kullanımı, input validation"),

    ("XSS Koruması",
     "Input sanitization",
     "HTML escape, CSP headers"),

    ("CSRF Koruması",
     "CSRF token",
     "Her form submission için token kontrolü"),

    ("Veri Şifreleme",
     "Database encryption",
     "Hassas alanlar (şifre, kişisel bilgi) şifreli"),

    ("Audit Logging",
     "Tüm işlemler loglanır",
     "Kim, ne zaman, ne yaptı - değiştirilemez log"),

    ("GDPR Uyumu",
     "Veri sahibi hakları",
     "Veri silme, export, onay mekanizmaları"),
]

for topic, measure, implementation in security_measures:
    row_cells = security_table.add_row().cells
    row_cells[0].text = topic
    row_cells[1].text = measure
    row_cells[2].text = implementation

    for para in row_cells[0].paragraphs:
        for run in para.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(139, 0, 0)

doc.add_paragraph()

# ============================================================================
# 15. PROJE TAKVİMİ VE KİLOMETRE TAŞLARI
# ============================================================================
print("📝 Section 15: Proje Takvimi...")

doc.add_page_break()
heading = doc.add_heading('15. PROJE TAKVİMİ VE KİLOMETRE TAŞLARI', 1)
heading.runs[0].font.color.rgb = KOLEKSIYON_BLUE

doc.add_paragraph(
    "Proje 3-4 aylık bir süreç olarak planlanmıştır. İki faz halinde ilerlenir:"
)

doc.add_heading('15.1. Faz 1: Temel Portal (MVP)', 2)

phase1_table = doc.add_table(rows=1, cols=4)
phase1_table.style = 'Light Grid Accent 1'

headers = phase1_table.rows[0].cells
headers[0].text = "Hafta"
headers[1].text = "Aktivite"
headers[2].text = "Çıktı"
headers[3].text = "Sorumlular"

for cell in headers:
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.bold = True

phase1_data = [
    ("Hafta 1-2", "Gereksinim Analizi ve Tasarım", "Detaylı gereksinim dokümanı, Wireframe'ler, DB şema", "BA, UX Designer, Architect"),
    ("Hafta 3-4", "Backend Geliştirme - Altyapı", "API framework, Auth sistemi, DB kurulum", "Backend Developer, DevOps"),
    ("Hafta 5-6", "SAP ve CMS Entegrasyon", "SAP API bağlantısı, CMS API entegrasyonu, Test verileri", "Integration Engineer, SAP Consultant"),
    ("Hafta 7-8", "Frontend Geliştirme - Temel", "Login, Register, Dashboard, Profil sayfaları", "Frontend Developer, UX Designer"),
    ("Hafta 9-10", "Frontend Geliştirme - Fiyat Listesi", "Fiyat listesi görüntüleme, Filtreleme, Excel export", "Frontend Developer"),
    ("Hafta 11", "Admin Paneli", "Bayi onaylama, Para birimi atama, Loglar", "Full-Stack Developer"),
    ("Hafta 12", "Test ve QA", "Fonksiyonel test, Güvenlik test, UAT", "QA Engineer, Test Ekibi"),
    ("Hafta 13-14", "Deployment ve Go-Live", "Production deployment, Monitoring setup, Bayi eğitimi", "DevOps, Project Manager"),
]

for week, activity, output, responsible in phase1_data:
    row_cells = phase1_table.add_row().cells
    row_cells[0].text = week
    row_cells[1].text = activity
    row_cells[2].text = output
    row_cells[3].text = responsible

doc.add_paragraph()

doc.add_heading('15.2. Faz 2: Gelişmiş Özellikler (Opsiyonel)', 2)

phase2_features = [
    "360° ürün görselleri entegrasyonu",
    "B, T, U alanları (Boyut, Teknik resim, Uzun metin) ekleme",
    "Multi-currency otomatik dönüşüm",
    "Gelişmiş raporlama ve analytics",
    "Mobil uygulama (iOS/Android)",
    "SAP'e geri besleme (PIM → SAP)",
    "Sipariş yönetimi modülü (opsiyonel)"
]

for feature in phase2_features:
    doc.add_paragraph(f"  • {feature}", style='List Bullet')

doc.add_paragraph()
doc.add_paragraph(
    "⏱️ Faz 2 için ayrı bir kapsam toplantısı ve planlama yapılacaktır. "
    "Tahmini süre: 2-3 ay ek geliştirme."
)

doc.add_paragraph()

# ============================================================================
# 16. RİSKLER VE VARSAYIMLAR
# ============================================================================
print("📝 Section 16: Riskler ve Varsayımlar...")

doc.add_page_break()
heading = doc.add_heading('16. RİSKLER VE VARSAYIMLAR', 1)
heading.runs[0].font.color.rgb = KOLEKSIYON_BLUE

doc.add_heading('16.1. Proje Riskleri', 2)

risks_table = doc.add_table(rows=1, cols=4)
risks_table.style = 'Light Grid Accent 1'

headers = risks_table.rows[0].cells
headers[0].text = "Risk"
headers[1].text = "Olasılık"
headers[2].text = "Etki"
headers[3].text = "Azaltma Stratejisi"

for cell in headers:
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.bold = True

risks_data = [
    ("SAP API erişim kısıtlamaları",
     "Orta",
     "Yüksek",
     "Erken SAP ekibi ile koordinasyon, test ortamı hazırlığı"),

    ("Fiyat veri formatı uyumsuzluğu",
     "Orta",
     "Orta",
     "SAP ve PIM veri şemalarını erken analiz, mapping dokümanı"),

    ("Güvenlik açıkları",
     "Düşük",
     "Yüksek",
     "Güvenlik testleri, penetration testing, code review"),

    ("Performans sorunları",
     "Orta",
     "Orta",
     "Load testing, cache stratejisi, CDN kullanımı"),

    ("Kullanıcı adaptasyonu zorluğu",
     "Düşük",
     "Orta",
     "Kullanıcı dostu tasarım, eğitim videoları, destek hattı"),

    ("Kapsam artışı (scope creep)",
     "Yüksek",
     "Orta",
     "Net kapsam tanımı, change request süreci"),
]

for risk, probability, impact, mitigation in risks_data:
    row_cells = risks_table.add_row().cells
    row_cells[0].text = risk
    row_cells[1].text = probability
    row_cells[2].text = impact
    row_cells[3].text = mitigation

    # Color code probability
    if probability == "Yüksek":
        for para in row_cells[1].paragraphs:
            for run in para.runs:
                run.font.color.rgb = RGBColor(255, 0, 0)
                run.font.bold = True

    # Color code impact
    if impact == "Yüksek":
        for para in row_cells[2].paragraphs:
            for run in para.runs:
                run.font.color.rgb = RGBColor(255, 0, 0)
                run.font.bold = True

doc.add_paragraph()

doc.add_heading('16.2. Varsayımlar', 2)

assumptions = [
    "SAP API erişimi mevcut ve dokümante edilmiştir",
    "CMS (Butterfly PIM) API entegrasyonu için gerekli yetkiler mevcuttur",
    "Koleksiyon IT ekibi gerekli altyapı desteğini sağlayacaktır",
    "Bayi listesi ve iletişim bilgileri Koleksiyon tarafından sağlanacaktır",
    "Web hosting ve SSL sertifikası Koleksiyon tarafından sağlanacaktır",
    "Proje ekibi tam zamanlı çalışacaktır (4-5 kişi)",
    "Gereksinim değişiklikleri ilk 2 hafta içinde netleştirilecektir",
    "SAP ve CMS sistemleri %99 uptime sağlayacaktır",
    "Koleksiyon business owner karar süreçlerinde hızlı dönüş sağlayacaktır",
]

for assumption in assumptions:
    doc.add_paragraph(f"  ✓ {assumption}", style='List Bullet')

doc.add_paragraph()

# ============================================================================
# 17. BAŞARI KRİTERLERİ
# ============================================================================
print("📝 Section 17: Başarı Kriterleri...")

doc.add_page_break()
heading = doc.add_heading('17. BAŞARI KRİTERLERİ', 1)
heading.runs[0].font.color.rgb = KOLEKSIYON_BLUE

success_table = doc.add_table(rows=1, cols=3)
success_table.style = 'Light Grid Accent 1'

headers = success_table.rows[0].cells
headers[0].text = "Kategori"
headers[1].text = "Metrik"
headers[2].text = "Hedef"

for cell in headers:
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.bold = True

success_criteria = [
    ("Kullanıcı Deneyimi", "Bayi memnuniyet skoru", "4.5/5 üzeri"),
    ("Kullanıcı Deneyimi", "Ortalama fiyat listesine ulaşma süresi", "< 2 dakika"),
    ("Performans", "Sayfa yükleme süresi", "< 2 saniye"),
    ("Performans", "API response time", "< 500ms"),
    ("Performans", "System uptime", "%99.5 üzeri"),
    ("Güvenlik", "Güvenlik açığı (critical)", "0 adet"),
    ("Güvenlik", "Başarılı login rate", "%98 üzeri"),
    ("İş Etkisi", "Manuel fiyat listesi gönderim azalması", "%90 azalma"),
    ("İş Etkisi", "Aktif bayi kullanıcı sayısı", "İlk 3 ayda 50+ bayi"),
    ("İş Etkisi", "Ortalama günlük kullanım", "20+ session/gün"),
    ("Teknik", "SAP veri senkronizasyon başarı oranı", "%99 üzeri"),
    ("Teknik", "Zero data loss", "100%"),
    ("Operasyonel", "Koleksiyon ekip zaman tasarrufu", "%70 azalma manuel işlerde"),
]

for category, metric, target in success_criteria:
    row_cells = success_table.add_row().cells
    row_cells[0].text = category
    row_cells[1].text = metric
    row_cells[2].text = target

    for para in row_cells[0].paragraphs:
        for run in para.runs:
            run.font.bold = True
            run.font.color.rgb = SUCCESS_GREEN

doc.add_paragraph()

# ============================================================================
# 18. GELECEK AŞAMALAR VE ÖNERİLER
# ============================================================================
print("📝 Section 18: Gelecek Aşamalar...")

doc.add_page_break()
heading = doc.add_heading('18. GELECEK AŞAMALAR VE ÖNERİLER', 1)
heading.runs[0].font.color.rgb = KOLEKSIYON_BLUE

doc.add_heading('18.1. Hemen Sonrası (Immediate Next Steps)', 2)

immediate_steps = [
    "Kapsam dokümanının Koleksiyon yönetimi ile onaylanması",
    "Proje ekibinin kurulması ve rollerin atanması",
    "SAP API erişim yetkilerinin alınması",
    "CMS (Butterfly) API dokümantasyonunun incelenmesi",
    "Hosting ve altyapı hazırlıklarının başlatılması",
    "Detaylı wireframe ve mockup tasarımlarının hazırlanması",
    "Proje kick-off toplantısının yapılması"
]

for i, step in enumerate(immediate_steps, 1):
    para = doc.add_paragraph(f"{i}. {step}")

doc.add_paragraph()

doc.add_heading('18.2. Faz 2 Önerileri (Gelecek Geliştirmeler)', 2)

doc.add_paragraph(
    "Faz 1 tamamlandıktan sonra, aşağıdaki özellikler değerlendirilebilir:"
)

future_recommendations = {
    "Kısa Vade (3-6 ay)": [
        "360° ürün görsellerinin entegrasyonu",
        "B, T, U alanları (teknik resimler, boyutlar, uzun metin) ekleme",
        "Gelişmiş filtreleme ve arama özellikleri",
        "Bayi bazlı özel kataloglar oluşturma",
        "Multi-currency otomatik döviz çevrimi"
    ],
    "Orta Vade (6-12 ay)": [
        "Mobil uygulama (iOS/Android)",
        "Sipariş yönetimi modülü (e-commerce entegrasyonu)",
        "Gerçek zamanlı stok sorgulama",
        "CRM entegrasyonu",
        "Gelişmiş analytics ve BI raporlama"
    ],
    "Uzun Vade (12+ ay)": [
        "AI destekli ürün önerisi",
        "Sanal showroom (AR/VR)",
        "Blockchain bazlı supply chain tracking",
        "IoT entegrasyonu (akıllı mobilya)",
        "Uluslararası pazar genişlemesi (çok dilli, çok para birimli)"
    ]
}

for timeframe, features in future_recommendations.items():
    para = doc.add_paragraph()
    run = para.add_run(f"📅 {timeframe}")
    run.font.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = KOLEKSIYON_ORANGE

    for feature in features:
        doc.add_paragraph(f"  • {feature}", style='List Bullet 2')

doc.add_paragraph()

doc.add_heading('18.3. İş Analisti Önerileri', 2)

analyst_recommendations = [
    ("Sürekli İyileştirme",
     "Portal yayına alındıktan sonra bayi geri bildirimlerini düzenli olarak toplayın ve analiz edin. "
     "Her 3 ayda bir kullanıcı deneyimi anketi yapılmasını öneririm."),

    ("Veri Analitiği",
     "Bayi davranışlarını analiz etmek için Google Analytics veya benzeri bir araç entegre edin. "
     "Hangi ürünlerin en çok görüntülendiği, hangi saatlerde trafik arttığı gibi veriler değerlidir."),

    ("Eğitim ve Dokümantasyon",
     "Bayiler için video eğitim serisi ve kullanım kılavuzu hazırlayın. "
     "İlk kullanıcı deneyiminde onboarding wizard eklenebilir."),

    ("Performans İzleme",
     "Sistem performansını sürekli izleyin. SAP ve CMS entegrasyonlarının uptime'ını takip edin. "
     "Slow query'leri tespit edip optimize edin."),

    ("Güvenlik Denetimi",
     "Yılda en az 2 kez penetration testing yaptırın. "
     "GDPR ve KVKK uyumunu düzenli olarak gözden geçirin."),

    ("Ölçeklenebilirlik",
     "Bayi sayısı arttıkça sistemin ölçeklenebilir olduğundan emin olun. "
     "Auto-scaling altyapısı kurulmasını öneririm."),

    ("Yedekleme ve Disaster Recovery",
     "Düzenli veri yedekleme stratejisi oluşturun (günlük yedek). "
     "Disaster recovery planı hazırlayın ve yılda bir kez test edin."),
]

for topic, recommendation in analyst_recommendations:
    para = doc.add_paragraph()
    run = para.add_run(f"💡 {topic}: ")
    run.font.bold = True
    run.font.color.rgb = KOLEKSIYON_BLUE
    para.add_run(recommendation)

doc.add_paragraph()

# ============================================================================
# DOCUMENT CLOSE
# ============================================================================

doc.add_page_break()

# Final notes
final_para = doc.add_paragraph()
final_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
final_run = final_para.add_run("─" * 60)
final_run.font.color.rgb = LIGHT_GRAY

doc.add_paragraph()

closing_para = doc.add_paragraph()
closing_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
closing_run = closing_para.add_run(
    "Bu doküman, Koleksiyon Mobilya Bayi Portalı projesinin\n"
    "iş analizi ve kapsam tanımını içermektedir.\n\n"
    "Sorularınız ve geri bildirimleriniz için proje ekibi ile iletişime geçebilirsiniz."
)
closing_run.font.size = Pt(10)
closing_run.font.color.rgb = DARK_GRAY

doc.add_paragraph()

footer_para = doc.add_paragraph()
footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_run = footer_para.add_run(f"© {datetime.now().year} Koleksiyon Mobilya - Gizli ve Özel")
footer_run.font.size = Pt(9)
footer_run.font.color.rgb = LIGHT_GRAY

# Save document
output_file = "/Users/caglarozyildirim/WebstormProjects/Deneme/Koleksiyon_Mobilya_Bayi_Portali_Kapsam_Dokumani.docx"
doc.save(output_file)

print("\n" + "="*80)
print("✅ DOCUMENT CREATED SUCCESSFULLY!")
print("="*80)
print(f"📁 File: {output_file}")
print(f"📄 Sections: 18 main sections")
print(f"📋 Tables: 15+ detailed tables")
print(f"📊 Requirements: 30+ functional requirements")
print(f"🔒 Security measures: 13 detailed controls")
print(f"📅 Timeline: 14 weeks (Faz 1)")
print("="*80 + "\n")
