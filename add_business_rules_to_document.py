#!/usr/bin/env python3
"""
Add Business Rules to MAN_Turkiye_Bakim_Yonetimi_FINAL_WITH_DATA_STRUCTURES.docx
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

# Paths
base_path = "/Users/caglarozyildirim/WebstormProjects/Deneme"
input_file = f"{base_path}/MAN_Turkiye_Bakim_Yonetimi_FINAL_WITH_DATA_STRUCTURES.docx"
output_file = f"{base_path}/MAN_Turkiye_Bakim_Yonetimi_FINAL_WITH_BUSINESS_RULES.docx"

# MAN Corporate Colors
MAN_RED = RGBColor(226, 7, 20)
MAN_DARK_GRAY = RGBColor(26, 26, 26)
MAN_GRAY = RGBColor(102, 102, 102)
MAN_GREEN = RGBColor(0, 168, 89)

# Business Rules Data Structure
business_rules = {
    'job_request': {
        'title': '2. İŞ TALEPLERİ MODÜLÜ',
        'rules': [
            {
                'title': 'Onay Süreci',
                'items': [
                    'Talebi oluşturan kişinin ilk yöneticisi talebi kontrol edip onaylar',
                    'En düşük maliyet onaylayıcı GL\'dir (Group Leader)'
                ]
            },
            {
                'title': 'Sorumluluk Değişimi',
                'items': [
                    'Çözüm sorumlusu bakım mühendisi veya SL tarafından her zaman değiştirilebilir',
                    'Mevcut sorumlu ve onaylanan kullanıcılar pozisyon değişiklikleri nedeniyle farklı olabilir, bu nedenle onaylayan kullanıcı ve mevcut sorumlu kullanıcı ID\'leri her zaman kaydedilmelidir'
                ]
            },
            {
                'title': 'İptal İşlemleri',
                'items': [
                    'Talep admin tarafından her zaman iptal edilebilir',
                    'Talep oluşturan kişi, yöneticisi tarafından güncelleme yapılmamışsa iptal edebilir'
                ]
            },
            {
                'title': 'Otomatik Atama',
                'items': [
                    '"Current Assignee" (Mevcut Atanan) alanı, işlem beklenen kişiyi kolayca tanımlamak için otomatik olarak doldurulur',
                    'Süreç tamamlandığında bu alan boşaltılır'
                ]
            }
        ],
        'authorization': [
            {
                'title': 'Erişim ve Oluşturma',
                'items': [
                    'Tüm kullanıcılar iş talebi erişebilir ve oluşturabilir',
                    'İş Talebi Formu talep eden tarafından doldurulmalı veya başkası doldurur ve PKI ile onaylanır'
                ]
            },
            {
                'title': 'Görüntüleme Yetkileri',
                'items': [
                    'Kullanıcılar sadece kendileri tarafından oluşturulan talepleri görebilir',
                    'Yöneticiler kendi taleplerini ve personellerinin taleplerini görebilir',
                    'Çözüm Sorumlusu tüm talepleri görebilir'
                ]
            },
            {
                'title': 'Ortak Kullanıcılar',
                'items': [
                    'Departmana atanan ortak kullanıcılar olmalı ve birden fazla kullanıcı tarafından kullanılabilmeli',
                    'Bir talep üzerinde işlem yapıldığında kullanıcı PKI kartları ile onaylanmalı'
                ]
            },
            {
                'title': 'Log ve Kayıt Yönetimi',
                'items': [
                    'Karar logları tutulmalı ve kayıtların bir rapor sayfası olmalı',
                    'Reddetme işlemi admin tarafından geri alınabilir ve log kaydedilmelidir'
                ]
            }
        ]
    },
    'asset_management': {
        'title': '3. VARLIK YÖNETİMİ MODÜLÜ',
        'rules': [
            {
                'title': 'Bakım Envanter Numarası',
                'items': [
                    'Varlık bakım departmanı tarafından SAP\'de varlık numarası almadan önce alınırsa, bakım operasyonunu sürdürmek için bakım envanter numarası oluşturulur',
                    'Bakım envanter numarası zorunlu ve benzersiz olmalıdır',
                    'Kullanıcı bir bakım envanter numarası girdiğinde, benzersiz olup olmadığı kontrol edilmelidir',
                    'Benzersiz değilse hata mesajı gösterilmeli ve kayıt edilmemelidir'
                ]
            },
            {
                'title': 'Tanımlanamayan Varlıklar',
                'items': [
                    'Bir varlık kayıtlarda tanımlanamıyorsa ve SAP\'de veya herhangi bir uygulamada ilk girişi yoksa, bunun için bir kayıt oluşturulur',
                    'Ancak SAP\'e eşleştirilmez (eşleştirme alanı boş bırakılır)'
                ]
            },
            {
                'title': 'Varlık Tipleri ve Zorunluluklar',
                'items': [
                    'Varlık tipleri zorunlu değildir, gerekirse bakım tarafından doldurulur',
                    'Varlık bakım numarası benzersiz olmalı ve kaydetme sırasında kontrol edilmelidir'
                ]
            },
            {
                'title': 'Lokasyon Değişikliği',
                'items': [
                    'Varlık lokasyonu, Alt Lokasyon 1 ve Alt Lokasyon 2, Maliyet Merkezi Varlık Sorumlusu veya Maliyet Merkezi Sorumlusu tarafından değiştirilebilir'
                ]
            }
        ],
        'authorization': [
            {
                'title': 'Oluşturma ve Değiştirme',
                'items': [
                    'Bakım Adminleri varlık oluşturabilir ve değiştirebilir',
                    'Bakım personeli sadece doküman ekleyebilir'
                ]
            },
            {
                'title': 'Görüntüleme Yetkileri',
                'items': [
                    'Bakım personeli tüm bilgileri görüntüleyebilir',
                    'Kullanıcılar sadece kendilerine atanan varlıkları görüntüleyebilir',
                    'Yöneticiler sadece personellerine ve kendilerine atanan varlıkları görebilir'
                ]
            }
        ],
        'zimmet_rules': [
            {
                'title': 'Değişim Tarihi',
                'items': [
                    'Değişim tarihi gelecekte veya geçmişte olabilir',
                    'Boş bırakılırsa, onay süreci tamamlandığında mevcut tarih ile doldurulur'
                ]
            },
            {
                'title': 'Red İşlemi',
                'items': [
                    'Kullanıcılardan biri reddederse süreç reddedilme ile sona erer',
                    'Varlık zimmet kayıtları değişmez ancak talep kayıtları tutulur'
                ]
            },
            {
                'title': 'Yetki Devri ile Onay',
                'items': [
                    'Bir yönetici çalışanı adına bir kaydı onaylarsa, personel tarafından PKI kart onayı alınmalı veya imzalı belge görüntüsü yüklenmelidir'
                ]
            },
            {
                'title': 'Otomatik Zimmet Oluşturma',
                'items': [
                    'İşten ayrılan (off-boarding) kullanıcının sahip olduğu her varlık için otomatik varlık zimmet oluşturulur',
                    'Alıcı kullanıcı ilk yöneticisi olur ancak değiştirilebilir'
                ]
            },
            {
                'title': 'PDF Doküman',
                'items': [
                    'Sürecin bilgilerini, yasal bilgileri ve onay bilgilerini içeren bir PDF doküman indirilebilmelidir',
                    'Kayıt tamamlandıktan sonra indirilebilir'
                ]
            }
        ]
    },
    'maintenance': {
        'title': '4. BAKIM YÖNETİMİ MODÜLÜ',
        'rules': [
            {
                'title': 'Görev ve Yorum Yönetimi',
                'items': [
                    'SL ve mühendisler, tamamlanmamış görevler sonucu girilen yorumları görebilmelidir',
                    'Yorumlar görevlere değil, görevlere (duties) yapılabilir',
                    'Ekler görevlere değil, görevlere (duties) eklenebilir'
                ]
            },
            {
                'title': 'Onay Süreci',
                'items': [
                    'SL ve Mühendisler görevleri (tasks) tek tek onaylamak yerine görevi (duty) onaylamalıdır'
                ]
            },
            {
                'title': 'Toplu Bakım',
                'items': [
                    'Aynı Periyodik Bakım ve Görevler tek girişle uygulanmalıdır',
                    'Görev ve diğer alanlarda yapılan güncellemeler toplu bakım gereksinimleri için uygulanmalıdır'
                ]
            },
            {
                'title': 'Görev Başlıkları',
                'items': [
                    'Görev başlıkları, toplu bakım görevi oluşturma için "Bakım gereksinimi adı" & "Bakım periyodu" & "Varlık Başlığı" ile otomatik oluşturulmalıdır'
                ]
            },
            {
                'title': 'Varlık Grupları',
                'items': [
                    'Varlık grupları birden fazla bakım gereksinimi için kullanılabilir',
                    'Listeye bir varlık eklenirse, geçmiş periyotlar için yeni görevler oluşturulmalıdır',
                    'Bir varlık çıkarılırsa, planlanan görevler (duties) silinmelidir'
                ]
            },
            {
                'title': 'Görev Durumları',
                'items': [
                    'Görevler "Planlandı" (Planned) durumu ile oluşturulur',
                    'Haftalık görevler kontrol edilir ve gerekli tarihe 5 haftadan az kalırsa durum "Aktif"e (Active) değişir'
                ]
            }
        ],
        'authorization': [
            {
                'title': 'Bakım Gereksinimleri ve Görevler',
                'items': [
                    'Bakım gereksinimleri, varlık grubu ve görevler (tasks) sadece SL ve Mühendisler tarafından eklenebilir'
                ]
            },
            {
                'title': 'Diğer İşlemler',
                'items': [
                    'Geri kalan tüm işlemler Bakım sorumlusu, SL ve mühendisler için erişilebilir olmalıdır'
                ]
            },
            {
                'title': 'Talep Erişimi',
                'items': [
                    'Her kullanıcı, sorumlu, onaylayıcı veya onaylanmış olduğu talebe erişebilir'
                ]
            }
        ]
    },
    'incident': {
        'title': '5. OLAY YÖNETİMİ MODÜLÜ',
        'rules': [
            {
                'title': 'Varlık Teslimi',
                'items': [
                    'Talep eden çözümü onayladığında ve varlık daha önce bakım departmanına verilmişse, talep edenin varlığı aldığı kabul edilir'
                ]
            },
            {
                'title': 'Alternatif Alıcı',
                'items': [
                    'Varlık almak için alternatif kullanıcı oluşturan veya bakım personeli tarafından değiştirilebilir',
                    'Değişiklik başka bir tabloda loglanmalıdır'
                ]
            },
            {
                'title': 'Talep Eden Onayı ve Varlık Teslimi',
                'items': [
                    'Talep eden onayı ve varlığın geri alınması aynı anda yapılmalıdır',
                    'Talep eden, talep edenin ilk yöneticisi ve varlık almak için alternatif kullanıcı tarafından yapılabilir'
                ]
            }
        ],
        'authorization': [
            {
                'title': 'Yetki Grupları',
                'items': [
                    'SL-TL yetkilendirme grubu oluşturulmalıdır',
                    'Standart kullanıcı yetkilendirme grubu oluşturulmalıdır'
                ]
            }
        ]
    }
}

print("\n" + "="*80)
print("ADDING BUSINESS RULES TO DOCUMENT")
print("="*80 + "\n")

# Load existing document
doc = Document(input_file)

def add_business_rules_section(doc, section_title, rules_data, authorization_data=None, extra_rules=None):
    """Add business rules section after module heading"""

    # Add Business Rules heading
    heading = doc.add_heading('İş Kuralları ve Yetkilendirme', 2)
    heading.runs[0].font.color.rgb = MAN_GREEN

    # Add Business Rules subheading
    doc.add_heading('📋 İş Kuralları', 3)

    for rule_group in rules_data:
        # Rule group title
        para = doc.add_paragraph()
        run = para.add_run(f"• {rule_group['title']}")
        run.font.bold = True
        run.font.size = Pt(11)

        # Rule items
        for item in rule_group['items']:
            para = doc.add_paragraph(f"  - {item}", style='List Bullet 2')
            para.paragraph_format.left_indent = Inches(0.5)

    # Add extra rules if provided (like zimmet rules for asset management)
    if extra_rules:
        doc.add_heading('📋 Varlık Zimmet İş Kuralları', 3)
        for rule_group in extra_rules:
            para = doc.add_paragraph()
            run = para.add_run(f"• {rule_group['title']}")
            run.font.bold = True
            run.font.size = Pt(11)

            for item in rule_group['items']:
                para = doc.add_paragraph(f"  - {item}", style='List Bullet 2')
                para.paragraph_format.left_indent = Inches(0.5)

    # Add Authorization Rules if provided
    if authorization_data:
        doc.add_heading('🔐 Yetkilendirme Kuralları', 3)

        for auth_group in authorization_data:
            # Auth group title
            para = doc.add_paragraph()
            run = para.add_run(f"• {auth_group['title']}")
            run.font.bold = True
            run.font.size = Pt(11)

            # Auth items
            for item in auth_group['items']:
                para = doc.add_paragraph(f"  - {item}", style='List Bullet 2')
                para.paragraph_format.left_indent = Inches(0.5)

    doc.add_paragraph()

# Find sections and add business rules
print("📋 Processing modules...")

# Track which sections to add business rules
sections_to_process = []
current_section = None

for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()

    # Job Request Module
    if '2. İŞ TALEPLERİ MODÜLÜ' in text and para.style.name == 'Heading 1':
        print(f"  ✓ Found Job Request module at paragraph {i}")
        sections_to_process.append({
            'index': i,
            'type': 'job_request',
            'title': text
        })

    # Asset Management Module
    elif '3. VARLIK YÖNETİMİ MODÜLÜ' in text and para.style.name == 'Heading 1':
        print(f"  ✓ Found Asset Management module at paragraph {i}")
        sections_to_process.append({
            'index': i,
            'type': 'asset_management',
            'title': text
        })

    # Maintenance Module
    elif '4. BAKIM YÖNETİMİ MODÜLÜ' in text and para.style.name == 'Heading 1':
        print(f"  ✓ Found Maintenance module at paragraph {i}")
        sections_to_process.append({
            'index': i,
            'type': 'maintenance',
            'title': text
        })

    # Incident Module
    elif '5. OLAY YÖNETİMİ MODÜLÜ' in text and para.style.name == 'Heading 1':
        print(f"  ✓ Found Incident module at paragraph {i}")
        sections_to_process.append({
            'index': i,
            'type': 'incident',
            'title': text
        })

print(f"\n📊 Found {len(sections_to_process)} sections to process")

# Process each section (in reverse order to maintain indices)
for section_info in reversed(sections_to_process):
    section_type = section_info['type']
    section_index = section_info['index']

    print(f"\n🔧 Adding business rules to {section_info['title']}...")

    # Find the position to insert (after the module heading and description)
    # Look for the first Heading 2 after this Heading 1
    insert_position = section_index + 1

    # Find where to insert (before first Heading 2)
    for j in range(section_index + 1, len(doc.paragraphs)):
        if doc.paragraphs[j].style.name == 'Heading 2':
            insert_position = j
            break

    # Create a new paragraph at the insert position
    new_para = doc.paragraphs[insert_position]._element
    parent = new_para.getparent()

    # Create business rules section
    temp_doc = Document()

    if section_type == 'job_request':
        add_business_rules_section(
            temp_doc,
            section_info['title'],
            business_rules['job_request']['rules'],
            business_rules['job_request']['authorization']
        )
    elif section_type == 'asset_management':
        add_business_rules_section(
            temp_doc,
            section_info['title'],
            business_rules['asset_management']['rules'],
            business_rules['asset_management']['authorization'],
            business_rules['asset_management']['zimmet_rules']
        )
    elif section_type == 'maintenance':
        add_business_rules_section(
            temp_doc,
            section_info['title'],
            business_rules['maintenance']['rules'],
            business_rules['maintenance']['authorization']
        )
    elif section_type == 'incident':
        add_business_rules_section(
            temp_doc,
            section_info['title'],
            business_rules['incident']['rules'],
            business_rules['incident']['authorization']
        )

    # Insert the content from temp_doc into main doc
    for element in temp_doc.element.body:
        parent.insert(parent.index(new_para), element)

print("\n💾 Saving updated document...")

# Save the updated document
doc.save(output_file)

print("\n" + "="*80)
print("✅ BUSINESS RULES ADDED SUCCESSFULLY!")
print(f"📁 Output file: {output_file}")
print("="*80 + "\n")
