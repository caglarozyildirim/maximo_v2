#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Desktop/new klasöründeki TÜM requirement dokümanlarını analiz et
Tamamen yeni bir HTML yapısı için eksiksiz analiz
"""

import pandas as pd
from docx import Document
import json
from pathlib import Path

# Paths
base_path = Path("/Users/caglarozyildirim/Desktop/new")

print("=" * 80)
print("BAKIM YÖNETİMİ SİSTEMİ - EKSİKSİZ REQUIREMENT ANALİZİ")
print("=" * 80)
print()

# 1. SCREEN DESIGNS - Ekran Tasarımları
print("1. EKRAN TASARIMLARI ANALİZİ (Screen Designs.xlsx)")
print("-" * 80)

try:
    screen_file = base_path / "Screen Designs.xlsx"
    xl = pd.ExcelFile(screen_file)

    print(f"Bulunan Sheet'ler: {xl.sheet_names}")
    print()

    all_screens = {}

    for sheet_name in xl.sheet_names:
        print(f"\n### {sheet_name} ###")
        df = pd.read_excel(screen_file, sheet_name=sheet_name)
        print(f"Kolonlar: {list(df.columns)}")
        print(f"Toplam satır: {len(df)}")

        # İlk 5 satırı göster
        print("\nİlk satırlar:")
        print(df.head().to_string())

        all_screens[sheet_name] = df.to_dict('records')

    # JSON'a kaydet
    with open('/Users/caglarozyildirim/WebstormProjects/Deneme/screen_designs_full.json', 'w', encoding='utf-8') as f:
        json.dump(all_screens, f, ensure_ascii=False, indent=2, default=str)

    print("\n✅ Screen Designs analizi tamamlandı")

except Exception as e:
    print(f"❌ Hata: {e}")

print("\n" + "=" * 80)
print("2. DATA STRUCTURE ANALİZİ (Data Structure.xlsx)")
print("-" * 80)

try:
    data_file = base_path / "Data Structure.xlsx"
    xl = pd.ExcelFile(data_file)

    print(f"Bulunan Sheet'ler: {xl.sheet_names}")
    print()

    all_data_structures = {}

    for sheet_name in xl.sheet_names:
        print(f"\n### {sheet_name} ###")
        df = pd.read_excel(data_file, sheet_name=sheet_name)
        print(f"Kolonlar: {list(df.columns)}")
        print(f"Toplam satır: {len(df)}")

        # İlk 5 satırı göster
        print("\nİlk satırlar:")
        print(df.head().to_string())

        all_data_structures[sheet_name] = df.to_dict('records')

    # JSON'a kaydet
    with open('/Users/caglarozyildirim/WebstormProjects/Deneme/data_structures_full.json', 'w', encoding='utf-8') as f:
        json.dump(all_data_structures, f, ensure_ascii=False, indent=2, default=str)

    print("\n✅ Data Structure analizi tamamlandı")

except Exception as e:
    print(f"❌ Hata: {e}")

print("\n" + "=" * 80)
print("3. LOCATIONS VE USER GROUPS ANALİZİ")
print("-" * 80)

try:
    loc_file = base_path / "Locations and user groups.xlsx"
    xl = pd.ExcelFile(loc_file)

    print(f"Bulunan Sheet'ler: {xl.sheet_names}")
    print()

    locations_data = {}

    for sheet_name in xl.sheet_names:
        print(f"\n### {sheet_name} ###")
        df = pd.read_excel(loc_file, sheet_name=sheet_name)
        print(f"Kolonlar: {list(df.columns)}")
        print(f"Toplam satır: {len(df)}")
        print("\nVeriler:")
        print(df.to_string())

        locations_data[sheet_name] = df.to_dict('records')

    # JSON'a kaydet
    with open('/Users/caglarozyildirim/WebstormProjects/Deneme/locations_users_full.json', 'w', encoding='utf-8') as f:
        json.dump(locations_data, f, ensure_ascii=False, indent=2, default=str)

    print("\n✅ Locations analizi tamamlandı")

except Exception as e:
    print(f"❌ Hata: {e}")

print("\n" + "=" * 80)
print("4. REQUIREMENT ANALYSIS DOKÜMANI")
print("-" * 80)

try:
    req_doc = Document(base_path / "Maintenance Management Application Requirement Analysis (Version1).docx")

    print(f"Toplam paragraf: {len(req_doc.paragraphs)}")
    print(f"Toplam tablo: {len(req_doc.tables)}")
    print()

    print("İLK 50 PARAGRAF:")
    for i, para in enumerate(req_doc.paragraphs[:50]):
        if para.text.strip():
            print(f"{i+1}. {para.text[:200]}")

    print("\n✅ Requirement document okundu")

except Exception as e:
    print(f"❌ Hata: {e}")

print("\n" + "=" * 80)
print("✅ TÜM ANALİZLER TAMAMLANDI")
print("=" * 80)
print("\nOluşturulan dosyalar:")
print("- screen_designs_full.json")
print("- data_structures_full.json")
print("- locations_users_full.json")
