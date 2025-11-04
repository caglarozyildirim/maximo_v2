#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import json
import zipfile
import xml.etree.ElementTree as ET
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import subprocess
import time

def read_visio_content(vsdx_path):
    """Extract detailed content from Visio file"""
    content = {
        'shapes': [],
        'text_content': [],
        'connections': []
    }

    try:
        with zipfile.ZipFile(vsdx_path, 'r') as zip_ref:
            # Read page XML files
            for file_name in zip_ref.namelist():
                if file_name.startswith('visio/pages/page') and file_name.endswith('.xml'):
                    xml_content = zip_ref.read(file_name)
                    try:
                        root = ET.fromstring(xml_content)
                        # Extract text content
                        for elem in root.iter():
                            if elem.text and elem.text.strip():
                                text = elem.text.strip()
                                if len(text) > 1 and text not in content['text_content']:
                                    content['text_content'].append(text)
                    except Exception as e:
                        print(f"Error parsing XML from {file_name}: {e}")
    except Exception as e:
        print(f"Error reading {vsdx_path}: {e}")

    return content

def take_html_screenshot(html_file, output_png, width=1400, height=900):
    """Take screenshot of HTML file using headless browser"""
    try:
        # Try using Chrome/Chromium
        chrome_paths = [
            '/Applications/Google Chrome.app/Contents/MacOS/Google Chrome',
            '/Applications/Chromium.app/Contents/MacOS/Chromium',
        ]

        chrome_path = None
        for path in chrome_paths:
            if Path(path).exists():
                chrome_path = path
                break

        if not chrome_path:
            print(f"Chrome not found, skipping screenshot for {html_file}")
            return False

        # Convert to absolute path
        html_abs = Path(html_file).resolve()
        output_abs = Path(output_png).resolve()

        cmd = [
            chrome_path,
            '--headless',
            '--disable-gpu',
            f'--screenshot={output_abs}',
            f'--window-size={width},{height}',
            f'file://{html_abs}'
        ]

        result = subprocess.run(cmd, capture_output=True, timeout=10)

        if result.returncode == 0 and output_abs.exists():
            print(f"✓ Screenshot created: {output_png}")
            return True
        else:
            print(f"✗ Failed to create screenshot for {html_file}")
            return False

    except Exception as e:
        print(f"Error taking screenshot: {e}")
        return False

def create_enhanced_document():
    """Create enhanced document with screenshots and visio content"""

    # Load existing JSON data
    json_path = '/Users/caglarozyildirim/WebstormProjects/Deneme/maintenance_docs_content.json'
    with open(json_path, 'r', encoding='utf-8') as f:
        data = json.load(f)

    # Paths
    workflows_path = Path("/Users/caglarozyildirim/Desktop/Şirketler/MAN Türkiye/Maintenance Management Application/Workflows")
    html_app_path = Path("/Users/caglarozyildirim/WebstormProjects/Deneme/bakim-yonetim-app")
    screenshots_path = Path("/Users/caglarozyildirim/WebstormProjects/Deneme/screenshots")
    screenshots_path.mkdir(exist_ok=True)

    # Workflow files
    workflow_files = {
        'job_request': 'Work Flow of Job Request.vsdx',
        'maintenance': 'Work Flow of Maintenance.vsdx',
        'asset_entry': 'Work Flow of Asset Entry.vsdx',
        'asset_assignment': 'Work flow of asset assigment.vsdx',
        'incident': 'Workflow of Incident Notification.vsdx',
        'retirement': 'Work Flow Asset Retirement.vsdx',
        'cost_center': 'Work Flow Cost Center Change.vsdx'
    }

    # Read Visio contents
    print("=== Visio Dosyaları Okunuyor ===")
    visio_contents = {}
    for key, filename in workflow_files.items():
        vsdx_path = workflows_path / filename
        if vsdx_path.exists():
            print(f"Reading: {filename}")
            visio_contents[key] = read_visio_content(vsdx_path)
        else:
            print(f"Not found: {filename}")

    # Take HTML screenshots
    print("\n=== HTML Ekran Görüntüleri Alınıyor ===")
    html_files = {
        'dashboard': 'index.html',
        'job_requests_list': 'pages/job-requests.html',
        'job_request_create': 'pages/job-request-create.html',
        'job_request_detail': 'pages/job-request-detail.html',
        'assets': 'pages/assets.html',
        'maintenance': 'pages/maintenance.html',
        'incidents': 'pages/incidents.html',
        'reports': 'pages/reports.html'
    }

    screenshots = {}
    for key, html_file in html_files.items():
        html_path = html_app_path / html_file
        output_png = screenshots_path / f"{key}.png"

        if html_path.exists():
            print(f"Processing: {html_file}")
            if take_html_screenshot(html_path, output_png):
                screenshots[key] = output_png
            time.sleep(1)  # Wait between screenshots
        else:
            print(f"HTML not found: {html_file}")

    # Create Word document
    print("\n=== Word Dokümanı Oluşturuluyor ===")
    doc = Document()

    # Set up styles
    section = doc.sections[0]
    section.page_height = Inches(11.69)
    section.page_width = Inches(8.27)

    # Title
    title = doc.add_heading('BAKIM YÖNETİMİ UYGULAMASI', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_heading('İŞ ANALİZİ DOKÜMANI (Görseller İle)', 0)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run('Versiyon: 2.0 (Görseller Dahil)\n').bold = True
    info.add_run('Tarih: Ekim 2025\n').bold = True
    info.add_run('Workflow Diyagramları ve Ekran Görüntüleri İçerir').bold = True

    doc.add_page_break()

    # İçindekiler
    doc.add_heading('İÇİNDEKİLER', 1)
    toc_items = [
        '1. YÖNETİCİ ÖZETİ',
        '2. İŞ SÜREÇLERİ VE WORKFLOW DİYAGRAMLARI',
        '3. UYGULAMA EKRAN GÖRÜNTÜLERİ',
        '4. FONKSİYONEL GEREKSİNİMLER',
        '5. VERİ MODELİ',
        '6. TEKNİK GEREKSİNİMLER'
    ]
    for item in toc_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

    # 1. YÖNETİCİ ÖZETİ
    doc.add_heading('1. YÖNETİCİ ÖZETİ', 1)
    doc.add_paragraph(
        'Bakım departmanı tüm operasyonlarını Maximo uygulaması üzerinde yönetmektedir. '
        'Bu proje, Maximo sisteminin yerine geçecek modern, kullanıcı dostu bir bakım yönetimi '
        'uygulaması geliştirmeyi hedeflemektedir. Proje 2027 yılında devreye girecek DIVA projesine '
        'kadar 2 yıl boyunca kullanılacaktır.'
    )

    doc.add_paragraph('Ana Fonksiyonlar:', style='Heading 3')
    functions = [
        'İş Talebi Yönetimi',
        'Sabit Varlık Yönetimi',
        'Bakım Yönetimi',
        'Olay Yönetimi',
        'Raporlama'
    ]
    for func in functions:
        doc.add_paragraph(func, style='List Bullet')

    doc.add_page_break()

    # 2. İŞ SÜREÇLERİ VE WORKFLOW DİYAGRAMLARI
    doc.add_heading('2. İŞ SÜREÇLERİ VE WORKFLOW DİYAGRAMLARI', 1)

    workflow_descriptions = {
        'job_request': {
            'title': '2.1 İş Talebi Süreci',
            'description': 'İş talepleri oluşturulur, detaylandırılır, onay süreçlerinden geçirilir ve çözüme kavuşturulur.',
            'file': 'Work Flow of Job Request.vsdx'
        },
        'maintenance': {
            'title': '2.2 Bakım Süreci',
            'description': 'Düzenli ve planlı bakım işlemlerinin yönetildiği süreç.',
            'file': 'Work Flow of Maintenance.vsdx'
        },
        'asset_entry': {
            'title': '2.3 Varlık Girişi Süreci',
            'description': 'Yeni varlıkların sisteme kaydedilmesi süreci.',
            'file': 'Work Flow of Asset Entry.vsdx'
        },
        'asset_assignment': {
            'title': '2.4 Varlık Atama Süreci',
            'description': 'Varlıkların çalışanlara veya departmanlara zimmetlenmesi süreci.',
            'file': 'Work flow of asset assigment.vsdx'
        },
        'incident': {
            'title': '2.5 Olay Bildirimi Süreci',
            'description': 'Acil arıza ve olayların bildirilmesi ve yönetilmesi süreci.',
            'file': 'Workflow of Incident Notification.vsdx'
        },
        'retirement': {
            'title': '2.6 Varlık Emekliliği Süreci',
            'description': 'Varlıkların hizmetten çıkarılması süreci.',
            'file': 'Work Flow Asset Retirement.vsdx'
        },
        'cost_center': {
            'title': '2.7 Maliyet Merkezi Değişikliği Süreci',
            'description': 'Varlıkların maliyet merkezleri arasında transfer edilmesi süreci.',
            'file': 'Work Flow Cost Center Change.vsdx'
        }
    }

    for key, info in workflow_descriptions.items():
        doc.add_heading(info['title'], 2)
        doc.add_paragraph(f"**Açıklama:** {info['description']}")
        doc.add_paragraph()

        # Add workflow diagram reference
        p = doc.add_paragraph()
        p.add_run(f"📊 Workflow Diyagramı: ").bold = True
        p.add_run(info['file'])

        # Add workflow content from Visio
        if key in visio_contents and visio_contents[key]['text_content']:
            doc.add_paragraph()
            doc.add_paragraph('Süreç Elemanları:', style='Heading 4')

            unique_items = []
            for item in visio_contents[key]['text_content']:
                if item and len(item) > 2 and item not in unique_items:
                    unique_items.append(item)

            # Limit to 25 items for readability
            for item in unique_items[:25]:
                doc.add_paragraph(f'• {item}', style='List Bullet')

        doc.add_paragraph()
        doc.add_paragraph('─' * 80)
        doc.add_paragraph()

    doc.add_page_break()

    # 3. UYGULAMA EKRAN GÖRÜNTÜLERİ
    doc.add_heading('3. UYGULAMA EKRAN GÖRÜNTÜLERİ', 1)
    doc.add_paragraph(
        'Bu bölümde, geliştirilen HTML prototip uygulamasının ekran görüntüleri yer almaktadır. '
        'Tüm ekranlar modern, responsive ve kullanıcı dostu tasarıma sahiptir.'
    )

    screenshot_descriptions = {
        'dashboard': {
            'title': '3.1 Ana Dashboard',
            'description': 'İstatistikler, bekleyen talepler, yaklaşan bakımlar ve hızlı erişim menüsü'
        },
        'job_requests_list': {
            'title': '3.2 İş Talepleri Listesi',
            'description': 'Filtreleme, arama ve sayfalama özellikleri ile talep listesi'
        },
        'job_request_create': {
            'title': '3.3 Yeni İş Talebi Oluşturma',
            'description': 'Detaylı form, dosya yükleme ve validation özellikleri'
        },
        'job_request_detail': {
            'title': '3.4 İş Talebi Detay',
            'description': 'Tam bilgi görüntüleme, onay süreci, yorum sistemi'
        },
        'assets': {
            'title': '3.5 Varlık Yönetimi',
            'description': 'Varlık listesi, durum filtreleme ve detay görüntüleme'
        },
        'maintenance': {
            'title': '3.6 Bakım Yönetimi',
            'description': 'Yaklaşan bakımlar timeline ve aktif işlemler'
        },
        'incidents': {
            'title': '3.7 Olay Yönetimi',
            'description': 'Acil olaylar, öncelik bazlı listeleme'
        },
        'reports': {
            'title': '3.8 Raporlar',
            'description': 'Rapor kategorileri ve hızlı erişim'
        }
    }

    for key, info in screenshot_descriptions.items():
        doc.add_heading(info['title'], 2)
        doc.add_paragraph(info['description'])
        doc.add_paragraph()

        if key in screenshots and screenshots[key].exists():
            try:
                # Add image with max width of 6.5 inches
                doc.add_picture(str(screenshots[key]), width=Inches(6.5))
                p = doc.paragraphs[-1]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_paragraph()
                print(f"✓ Added screenshot: {info['title']}")
            except Exception as e:
                doc.add_paragraph(f"[Ekran görüntüsü eklenemedi: {e}]")
                print(f"✗ Failed to add screenshot: {e}")
        else:
            p = doc.add_paragraph()
            p.add_run('[Ekran görüntüsü alınamadı - Chrome/Chromium gerekli]').italic = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            print(f"✗ Screenshot not available: {key}")

        doc.add_paragraph()
        doc.add_paragraph('─' * 80)
        doc.add_paragraph()

    doc.add_page_break()

    # 4. FONKSİYONEL GEREKSİNİMLER
    doc.add_heading('4. FONKSİYONEL GEREKSİNİMLER', 1)

    doc.add_heading('4.1 İş Talebi Yönetimi', 2)
    doc.add_paragraph('**Hedefler:**')
    objectives = [
        'Talepleri toplamak ve merkezi bir sistemde yönetmek',
        'Çok seviyeli onay sürecini otomatikleştirmek',
        'Çözüm sürecini takip etmek ve raporlamak',
        'Kullanılan malzemelerin tüketimini kaydetmek'
    ]
    for obj in objectives:
        doc.add_paragraph(obj, style='List Bullet')

    doc.add_paragraph()
    doc.add_heading('4.2 Varlık Yönetimi', 2)
    doc.add_paragraph('Sabit varlıkların yaşam döngüsü boyunca yönetimi ve takibi.')

    doc.add_paragraph()
    doc.add_heading('4.3 Bakım Yönetimi', 2)
    doc.add_paragraph('Düzenli, toplu ve önleyici bakım işlemlerinin planlanması ve yönetimi.')

    doc.add_paragraph()
    doc.add_heading('4.4 Olay Yönetimi', 2)
    doc.add_paragraph('Acil arıza ve olayların hızlı çözümü için kritik modül.')

    doc.add_page_break()

    # 5. VERİ MODELİ
    doc.add_heading('5. VERİ MODELİ', 1)

    data_structure = data.get('data_structure', {})
    doc.add_paragraph(f'Veri modeli **{len(data_structure)} adet tablo** içermektedir.')

    doc.add_paragraph('Ana Tablolar:', style='Heading 3')
    main_tables = [
        'JobRequest - İş talepleri',
        'Asset - Sabit varlıklar',
        'MaintenancePlan - Bakım planları',
        'Incident - Olay bildirimleri',
        'Approval - Onay işlemleri'
    ]
    for table in main_tables:
        doc.add_paragraph(table, style='List Bullet')

    doc.add_page_break()

    # 6. TEKNİK GEREKSİNİMLER
    doc.add_heading('6. TEKNİK GEREKSİNİMLER', 1)

    doc.add_heading('6.1 Teknoloji Stack', 2)
    tech = [
        'Frontend: HTML5, CSS3, JavaScript (ES6+)',
        'Backend: .NET Core veya Java Spring Boot (önerilen)',
        'Database: Microsoft SQL Server',
        'Authentication: Azure AD / OAuth 2.0',
        'Mobile: Progressive Web App (PWA)'
    ]
    for t in tech:
        doc.add_paragraph(t, style='List Bullet')

    doc.add_paragraph()
    doc.add_heading('6.2 Entegrasyonlar', 2)
    integrations = [
        'SAP: Varlık ana verileri, maliyet merkezi',
        'Active Directory: Kullanıcı kimlik doğrulama',
        'E-posta: Bildirimler ve onay süreçleri'
    ]
    for integ in integrations:
        doc.add_paragraph(integ, style='List Bullet')

    # Save document
    output_path = "/Users/caglarozyildirim/WebstormProjects/Deneme/Bakim_Yonetim_Uygulamasi_Is_Analizi_WITH_IMAGES.docx"
    doc.save(output_path)

    print(f"\n{'='*80}")
    print(f"✅ GÖRSELLER İLE DOKÜMAN OLUŞTURULDU!")
    print(f"{'='*80}")
    print(f"📄 Dosya: {output_path}")
    print(f"📊 Workflow Diyagramları: {len(visio_contents)} adet")
    print(f"🖼️  Ekran Görüntüleri: {len(screenshots)} adet")
    print(f"{'='*80}\n")

    return output_path

if __name__ == "__main__":
    create_enhanced_document()
