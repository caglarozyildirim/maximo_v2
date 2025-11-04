#!/usr/bin/env python3
"""
Update MAN_Turkiye_Bakim_Yonetimi_COMPLETE_ANALYSIS.docx:
1. Replace bullet point data structures with tables
2. Add asset-create screen
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import json

def set_cell_background(cell, color):
    """Set cell background color"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading_elm)

def add_data_structure_table(doc, fields, module_color="#E20714"):
    """Add data structure table in CleanShot format"""
    # Add table with 6 columns
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Light Grid Accent 1'

    # Header row
    header_cells = table.rows[0].cells
    headers = ['Alan Adı (EN)', 'Alan Adı (TR)', 'Veri Tipi', 'Zorunlu', 'Açıklama', 'SAP Mapping']

    for i, header_text in enumerate(headers):
        cell = header_cells[i]
        set_cell_background(cell, module_color)
        paragraph = cell.paragraphs[0]
        run = paragraph.runs[0] if paragraph.runs else paragraph.add_run(header_text)
        run.text = header_text
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(10)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    for field in fields:
        row_cells = table.add_row().cells
        row_cells[0].text = field['en']
        row_cells[1].text = field['tr']
        row_cells[2].text = field['type']
        row_cells[3].text = field['required']
        row_cells[4].text = field['desc']
        row_cells[5].text = field['sap']

        # Format cells
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)

def update_existing_document():
    """Update existing document"""
    # Load existing document
    doc = Document('/Users/caglarozyildirim/WebstormProjects/Deneme/MAN_Turkiye_Bakim_Yonetimi_COMPLETE_ANALYSIS.docx')

    # Load data structures
    with open('/Users/caglarozyildirim/WebstormProjects/Deneme/detailed_data_structures.json', 'r', encoding='utf-8') as f:
        data = json.load(f)

    # We need to iterate through paragraphs and find where to replace/add content
    # Strategy: Find "Veri Yapısı:" headings and replace the bullet points after them with tables

    i = 0
    while i < len(doc.paragraphs):
        para = doc.paragraphs[i]

        # Check if this is a "Veri Yapısı:" heading
        if para.text.strip().startswith('Veri Yapısı'):
            # Remove bullet points after this heading
            j = i + 1
            # Skip until we find bullet points or a heading
            while j < len(doc.paragraphs):
                next_para = doc.paragraphs[j]
                # If it's a bullet point or contains "Bu ekranda", remove it
                if (next_para.style.name.startswith('List') or
                    'Bu ekranda' in next_para.text or
                    'gösterilen alanlar' in next_para.text or
                    'Filtreler:' in next_para.text or
                    next_para.text.strip().startswith('•')):
                    # Delete this paragraph
                    p = next_para._element
                    p.getparent().remove(p)
                    # Don't increment j since we removed an element
                else:
                    # We've reached the end of bullet points
                    break

            # Now add table after "Veri Yapısı:" heading
            # Determine which module based on context (look at previous headings)
            module_name = None
            module_color = "#E20714"

            # Look backwards to find module
            for k in range(i-1, max(0, i-20), -1):
                prev_text = doc.paragraphs[k].text.lower()
                if 'iş talepleri' in prev_text or 'job request' in prev_text:
                    module_name = 'job_requests'
                    module_color = "#FFA726"
                    break
                elif 'varlık' in prev_text or 'asset' in prev_text:
                    module_name = 'assets'
                    module_color = "#42A5F5"
                    break
                elif 'bakım' in prev_text and 'sistem' not in prev_text:
                    module_name = 'maintenance'
                    module_color = "#66BB6A"
                    break
                elif 'olay' in prev_text or 'incident' in prev_text:
                    module_name = 'incidents'
                    module_color = "#EF5350"
                    break

            # Insert table if we found the module
            if module_name and module_name in data:
                # Add a blank paragraph first
                new_para = para._element.addnext(doc._element.body._new_p())

                # Now add the table
                add_data_structure_table(doc, data[module_name]['fields'], module_color)

                # Add another blank paragraph after table
                doc.add_paragraph()

        # Check for "3.2 Varlık Yönetimi" section to add asset-create
        if '3.2.2 Varlık - Detay Ekranı' in para.text:
            # Look for the section after this to insert before "3.3"
            # We need to find where 3.2.2 ends and insert 3.2.3 before 3.3

            # Find the end of 3.2.2 section (before next heading)
            k = i + 1
            while k < len(doc.paragraphs):
                next_para = doc.paragraphs[k]
                if (next_para.style.name.startswith('Heading') and
                    ('3.3' in next_para.text or '4.' in next_para.text)):
                    # Insert 3.2.3 section here

                    # Add heading
                    heading = doc.add_heading('3.2.3 Varlık Oluştur - Form Ekranı', level=3)

                    # Add description
                    desc = doc.add_paragraph(
                        'Sisteme yeni varlık kaydı oluşturma ekranıdır. Temel bilgiler, üretici bilgileri, '
                        'lokasyon, teknik özellikler ve ek belgeler girilir. SAP entegrasyonu için gerekli '
                        'alanlar mapping ile eşleştirilir.'
                    )

                    # Add screenshot heading
                    screenshot_heading = doc.add_heading('Ekran Görüntüsü:', level=4)

                    # Add screenshot
                    try:
                        img = doc.add_picture(
                            '/Users/caglarozyildirim/WebstormProjects/Deneme/screenshots/asset-create.png',
                            width=Inches(6)
                        )
                        last_paragraph = doc.paragraphs[-1]
                        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    except:
                        doc.add_paragraph('Ekran görüntüsü bulunamadı')

                    # Add data structure heading
                    data_heading = doc.add_heading('Veri Yapısı:', level=4)

                    # Add table
                    add_data_structure_table(doc, data['assets']['fields'], "#42A5F5")

                    # Add blank paragraph
                    doc.add_paragraph()

                    # Add page break
                    doc.add_page_break()

                    break
                k += 1

        i += 1

    # Save updated document
    output_path = '/Users/caglarozyildirim/WebstormProjects/Deneme/MAN_Turkiye_Bakim_Yonetimi_COMPLETE_ANALYSIS.docx'
    doc.save(output_path)
    print(f"✅ Doküman güncellendi: {output_path}")

    return output_path

if __name__ == '__main__':
    update_existing_document()
