#!/usr/bin/env python3
"""
MAN Turkey Maintenance Management System - Part 2
Adding remaining modules: Asset Management, Maintenance Management, Incident Management
Workflows, Data Structure, Test Scenarios, Technical Specs
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# MAN Corporate Colors
MAN_RED = RGBColor(226, 7, 20)
WHITE = RGBColor(255, 255, 255)
BLACK = RGBColor(0, 0, 0)
GRAY = RGBColor(128, 128, 128)

def set_cell_background(cell, color):
    """Set cell background color"""
    shading_elm = OxmlElement('w:shd')
    # RGBColor is a tuple-like object, access values by index
    shading_elm.set(qn('w:fill'), f'{color[0]:02X}{color[1]:02X}{color[2]:02X}')
    cell._element.get_or_add_tcPr().append(shading_elm)

def add_page_break(doc):
    """Add a page break"""
    doc.add_page_break()

def add_asset_management_module(doc, screenshot_list, screenshot_detail, screenshot_modal):
    """Add Asset Management module documentation"""
    heading = doc.add_heading("5. Modul 3: Varlik Yonetimi", level=1)
    heading.runs[0].font.color.rgb = MAN_RED

    doc.add_heading("5.1 Modul Aciklamasi / Module Description", level=2)

    p = doc.add_paragraph(
        "Varlik Yonetimi modulu, tum uretim ekipmanlarinin ve varliklarinin merkezi bir "
        "veritabaninda yonetilmesini saglar. Her varlik icin detayli bilgiler (SAP ID, "
        "kategori, lokasyon, satin alma tarihi, defter degeri vb.) kaydedilir. "
        "Varlik yasam dongusu, bakim gecmisi, mali bilgiler ve durumu bu modulde izlenir."
    )

    p = doc.add_paragraph(
        "The Asset Management module enables the management of all production equipment "
        "and assets in a centralized database. Detailed information (SAP ID, category, "
        "location, purchase date, book value, etc.) is recorded for each asset. "
        "Asset lifecycle, maintenance history, financial information, and status are tracked in this module."
    )

    doc.add_heading("5.2 Fonksiyonel Gereksinimler / Functional Requirements", level=2)

    requirements = [
        ("FR-AM-001", "Varlik kaydi olusturma",
         "Yeni varlik kaydı tum gerekli bilgilerle olusturulabilmelidir"),
        ("FR-AM-002", "SAP entegrasyonu",
         "SAP sisteminden varlik bilgileri senkronize edilebilmelidir"),
        ("FR-AM-003", "Varlik listesi goruntuleme",
         "Tum varliklar filtrelenebilir ve siralanabilir sekilde listelenmeli"),
        ("FR-AM-004", "Varlik detay sayfasi",
         "Her varlik icin detayli bilgi sayfasi gosterilmelidir"),
        ("FR-AM-005", "Varlik guncelleme",
         "Yetkili kullanicilar varlik bilgilerini guncelleyebilmelidir"),
        ("FR-AM-006", "Durum yonetimi",
         "Varlik durumu (AKTIF, BAKIMDA, ARIZALI) degistirilebilmelidir"),
        ("FR-AM-007", "Bakim gecmisi",
         "Varlik ile iliskili tum bakim kayitlari gosterilmelidir"),
        ("FR-AM-008", "Iliskili is talepleri",
         "Varliga bagli tum is talepleri listelenmelidir"),
        ("FR-AM-009", "QR kod olusturma",
         "Her varlik icin QR kod otomatik olusturulmalidir"),
        ("FR-AM-010", "Dokuman ekleme",
         "Varliga kullanim kilavuzu, sertifika vb. dokumanlar eklenebilmelidir"),
        ("FR-AM-011", "Mali bilgiler",
         "Satin alma degeri, defter degeri, amortisma bilgileri yonetilebilmelidir"),
        ("FR-AM-012", "Lokasyon takibi",
         "Varligin fiziksel konumu kayit altinda tutulmalidirr"),
        ("FR-AM-013", "Sonraki bakim tarihi",
         "Bir sonraki planli bakim tarihi otomatik hesaplanmalidir"),
        ("FR-AM-014", "Varlik raporlari",
         "Varlik envanteri, amortisman, bakim maliyeti raporlari uretilmelidir"),
        ("FR-AM-015", "Arama ve filtreleme",
         "Varliklar SAP ID, ad, kategori, lokasyon, duruma gore aranabilmelidir"),
    ]

    table = doc.add_table(rows=len(requirements)+1, cols=3)
    table.style = 'Light Grid Accent 1'

    header_cells = table.rows[0].cells
    header_cells[0].text = "ID"
    header_cells[1].text = "Gereksinim / Requirement"
    header_cells[2].text = "Aciklama / Description"

    for cell in header_cells:
        set_cell_background(cell, MAN_RED)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = WHITE
                run.font.bold = True

    for idx, (req_id, req, desc) in enumerate(requirements, 1):
        row = table.rows[idx]
        row.cells[0].text = req_id
        row.cells[1].text = req
        row.cells[2].text = desc

    add_page_break(doc)

    doc.add_heading("5.3 Veri Alanlari - 17 Alan / Data Fields - 17 Fields", level=2)

    data_fields = [
        ("assetId", "String", "Evet (Auto)", "Internal varlik ID",
         "AST-XXX format", "AST-001", "Auto-generated", "PK, Indexed"),
        ("sapId", "String", "Evet", "SAP sistem ID",
         "Numeric, 8-10 digits", "100234567", "SAP sync", "Unique, Indexed"),
        ("assetName", "String", "Evet", "Varlik adi",
         "5-200 karakter", "CNC Torna Makinesi - Sasi Isleme", "User input", "Indexed"),
        ("category", "Enum", "Evet", "Varlik kategorisi",
         "URETIM_MAKINESI, ROBOTIK_SISTEM, MONTAJ_EKIPMANI, TASIMA_EKIPMANI, ALTYAPI", "URETIM_MAKINESI", "Dropdown", "Indexed"),
        ("manufacturer", "String", "Evet", "Uretici firma",
         "Max 100 char", "HAAS", "User input", "-"),
        ("model", "String", "Evet", "Model bilgisi",
         "Max 100 char", "ST-30", "User input", "-"),
        ("serialNumber", "String", "Evet", "Seri numarasi",
         "Max 100 char", "HAAS-ST30-2021-1234", "User input", "Unique"),
        ("location", "String", "Evet", "Fiziksel lokasyon",
         "Max 200 char", "Sasi Uretim Hatti", "Dropdown/Text", "Indexed"),
        ("subLocation", "String", "Hayir", "Alt lokasyon",
         "Max 200 char", "Alan A1", "User input", "-"),
        ("costCenter", "String", "Hayir", "Maliyet merkezi",
         "SAP cost center code", "CC-PROD-001", "SAP sync", "-"),
        ("responsible", "String", "Evet", "Sorumlu kisi",
         "User ID reference", "Uretim Ekibi A", "Dropdown", "FK"),
        ("status", "Enum", "Evet", "Varlik durumu",
         "AKTIF, BAKIMDA, ARIZALI, DEVRE_DISI", "AKTIF", "System managed", "Indexed"),
        ("purchaseDate", "Date", "Evet", "Satin alma tarihi",
         "ISO 8601 date", "2021-03-15", "User input", "-"),
        ("purchaseValue", "Decimal", "Evet", "Satin alma degeri (TL)",
         ">= 0, 2 decimal", "850000.00", "User input", "-"),
        ("bookValue", "Decimal", "Hayir", "Guncel defter degeri (TL)",
         ">= 0, 2 decimal", "680000.00", "SAP sync/Calculated", "-"),
        ("lastMaintenanceDate", "Date", "Hayir", "Son bakim tarihi",
         "ISO 8601 date", "2025-09-15", "System calculated", "-"),
        ("nextMaintenanceDate", "Date", "Hayir", "Sonraki bakim tarihi",
         "ISO 8601 date", "2025-12-15", "System calculated", "Indexed"),
    ]

    table = doc.add_table(rows=len(data_fields)+1, cols=8)
    table.style = 'Light Grid Accent 1'

    header_cells = table.rows[0].cells
    headers = ["Alan Adi", "Veri Tipi", "Zorunlu", "Aciklama", "Validasyon",
               "Ornek", "Kaynak", "DB Index"]
    for idx, header in enumerate(headers):
        header_cells[idx].text = header
        set_cell_background(header_cells[idx], MAN_RED)
        for paragraph in header_cells[idx].paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = WHITE
                run.font.bold = True
                run.font.size = Pt(8)

    for idx, field_data in enumerate(data_fields, 1):
        row = table.rows[idx]
        for col_idx, value in enumerate(field_data):
            row.cells[col_idx].text = value
            for paragraph in row.cells[col_idx].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(7)

    add_page_break(doc)

    doc.add_heading("5.4 Ekran Goruntuleri / Screenshots", level=2)

    doc.add_heading("5.4.1 Varlik Yonetimi Listesi", level=3)
    if os.path.exists(screenshot_list):
        try:
            doc.add_picture(screenshot_list, width=Inches(6.5))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except:
            doc.add_paragraph(f"[Screenshot: {screenshot_list}]")

    add_page_break(doc)

    doc.add_heading("5.4.2 Varlik Detay Sayfasi", level=3)
    p = doc.add_paragraph(
        "Varlik detay sayfasinda tum 17 alan, bakim gecmisi ve iliskili is talepleri gosterilir."
    )
    if os.path.exists(screenshot_detail):
        try:
            doc.add_picture(screenshot_detail, width=Inches(6.5))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except:
            doc.add_paragraph(f"[Screenshot: {screenshot_detail}]")

    add_page_break(doc)

    doc.add_heading("5.4.3 Yeni Varlik Ekleme Modal", level=3)
    if os.path.exists(screenshot_modal):
        try:
            doc.add_picture(screenshot_modal, width=Inches(6.5))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except:
            doc.add_paragraph(f"[Screenshot: {screenshot_modal}]")

    add_page_break(doc)

    doc.add_heading("5.5 Is Kurallari / Business Rules", level=2)

    business_rules = [
        ("BR-AM-001", "SAP ID Benzersizligi", "Her SAP ID benzersiz olmalidir, duplicate kabul edilmez"),
        ("BR-AM-002", "Seri No Benzersizligi", "Seri numarasi benzersiz olmalidir"),
        ("BR-AM-003", "Bakim Tarihi Hesaplama", "nextMaintenanceDate, kategori ve bakim planina gore otomatik hesaplanir"),
        ("BR-AM-004", "Durum Gecis Kurallari", "AKTIF <-> BAKIMDA <-> ARIZALI <-> DEVRE_DISI (tum yonlerde gecis mumkun)"),
        ("BR-AM-005", "Arizali Varlik Engelleme", "ARIZALI durumdaki varlik uzerinde yeni is planlanamaaz"),
        ("BR-AM-006", "QR Kod Olusturma", "Varlik kaydedildiginde otomatik QR kod uretilir"),
        ("BR-AM-007", "Defter Degeri Guncelleme", "bookValue, SAP'dan gunluk senkronizasyonla guncellenir"),
        ("BR-AM-008", "Silme Kisitlamasi", "Aktif bakim plani olan varlik silinemez"),
    ]

    table = doc.add_table(rows=len(business_rules)+1, cols=3)
    table.style = 'Light Grid Accent 1'

    header_cells = table.rows[0].cells
    header_cells[0].text = "Kural ID"
    header_cells[1].text = "Kural Adi"
    header_cells[2].text = "Aciklama"

    for cell in header_cells:
        set_cell_background(cell, MAN_RED)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = WHITE
                run.font.bold = True

    for idx, (rule_id, rule_name, desc) in enumerate(business_rules, 1):
        row = table.rows[idx]
        row.cells[0].text = rule_id
        row.cells[1].text = rule_name
        row.cells[2].text = desc

    add_page_break(doc)

def add_maintenance_module(doc, screenshot_list, screenshot_modal):
    """Add Maintenance Management module"""
    heading = doc.add_heading("6. Modul 4: Bakim Yonetimi", level=1)
    heading.runs[0].font.color.rgb = MAN_RED

    doc.add_heading("6.1 Modul Aciklamasi / Module Description", level=2)
    p = doc.add_paragraph(
        "Bakim Yonetimi modulu, planli ve onleyici bakim faaliyetlerinin yonetimini saglar. "
        "Periyodik bakim planlari olusturulur, bakim gorevleri otomatik olusturulur ve "
        "teknisyenlere atanir. Bakim tipleri (onleyici, duzenli, kosiullu) yonetilebilir."
    )

    doc.add_heading("6.2 Veri Alanlari ve Bakim Tipleri / Data Fields and Maintenance Types", level=2)

    maintenance_types = [
        ("Onleyici Bakim", "Preventive Maintenance",
         "Zamanli bakim - Belirli periyotlarda (haftalik, aylik, yillik) yapilir",
         "Her 3 ayda bir yag degisimi"),
        ("Duzeltici Bakim", "Corrective Maintenance",
         "Ariza sonrasi yapilan onarim bakimi",
         "Motor arizasi sonrasi onarim"),
        ("Kosullu Bakim", "Condition-based Maintenance",
         "Sensor verileri veya durum kontrolune gore yapilan bakim",
         "Titresim olcumu esik degerini asti"),
        ("Ongorucu Bakim", "Predictive Maintenance",
         "AI/ML tabanli tahmine dayali bakim",
         "Makine ogrenmesi arizayi onceden tahmin etti"),
    ]

    table = doc.add_table(rows=len(maintenance_types)+1, cols=4)
    table.style = 'Light Grid Accent 1'

    header_cells = table.rows[0].cells
    header_cells[0].text = "Bakim Tipi (TR)"
    header_cells[1].text = "Maintenance Type (EN)"
    header_cells[2].text = "Aciklama / Description"
    header_cells[3].text = "Ornek / Example"

    for cell in header_cells:
        set_cell_background(cell, MAN_RED)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = WHITE
                run.font.bold = True

    for idx, (tr, en, desc, example) in enumerate(maintenance_types, 1):
        row = table.rows[idx]
        row.cells[0].text = tr
        row.cells[1].text = en
        row.cells[2].text = desc
        row.cells[3].text = example

    add_page_break(doc)

    doc.add_heading("6.3 Ekran Goruntuleri / Screenshots", level=2)

    doc.add_heading("6.3.1 Bakim Yonetimi Listesi", level=3)
    if os.path.exists(screenshot_list):
        try:
            doc.add_picture(screenshot_list, width=Inches(6.5))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except:
            pass

    add_page_break(doc)

    doc.add_heading("6.3.2 Yeni Bakim Plani Modal", level=3)
    if os.path.exists(screenshot_modal):
        try:
            doc.add_picture(screenshot_modal, width=Inches(6.5))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except:
            pass

    add_page_break(doc)

def add_incident_module(doc, screenshot_list, screenshot_modal):
    """Add Incident Management module"""
    heading = doc.add_heading("7. Modul 5: Olay Yonetimi", level=1)
    heading.runs[0].font.color.rgb = MAN_RED

    doc.add_heading("7.1 Modul Aciklamasi / Module Description", level=2)
    p = doc.add_paragraph(
        "Olay Yonetimi modulu, acil durumlar ve kritik olaylarin hizli bir sekilde "
        "rapor edilmesini ve yonetilmesini saglar. Oncelik seviyeleri, SLA sureleri "
        "ve eskalasyon kurallari tanimlanir."
    )

    doc.add_heading("7.2 Oncelik Seviyeleri ve SLA Sureleri", level=2)

    priority_levels = [
        ("KRITIK", "Critical", "1 saat", "15 dakika", "Uretim durduruldu, guvenlik riski"),
        ("YUKSEK", "High", "4 saat", "1 saat", "Onemli ekipman arizasi"),
        ("ORTA", "Medium", "24 saat", "4 saat", "Performans dusukluugu"),
        ("DUSUK", "Low", "72 saat", "24 saat", "Minor sorun"),
    ]

    table = doc.add_table(rows=len(priority_levels)+1, cols=5)
    table.style = 'Light Grid Accent 1'

    header_cells = table.rows[0].cells
    header_cells[0].text = "Oncelik (TR)"
    header_cells[1].text = "Priority (EN)"
    header_cells[2].text = "Cozum SLA"
    header_cells[3].text = "Yanit SLA"
    header_cells[4].text = "Ornek / Example"

    for cell in header_cells:
        set_cell_background(cell, MAN_RED)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = WHITE
                run.font.bold = True

    for idx, (tr, en, resolve, response, example) in enumerate(priority_levels, 1):
        row = table.rows[idx]
        row.cells[0].text = tr
        row.cells[1].text = en
        row.cells[2].text = resolve
        row.cells[3].text = response
        row.cells[4].text = example

    add_page_break(doc)

    doc.add_heading("7.3 Ekran Goruntuleri / Screenshots", level=2)

    doc.add_heading("7.3.1 Olay Yonetimi Listesi", level=3)
    if os.path.exists(screenshot_list):
        try:
            doc.add_picture(screenshot_list, width=Inches(6.5))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except:
            pass

    add_page_break(doc)

    doc.add_heading("7.3.2 Acil Olay Bildir Modal", level=3)
    if os.path.exists(screenshot_modal):
        try:
            doc.add_picture(screenshot_modal, width=Inches(6.5))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except:
            pass

    add_page_break(doc)

def add_workflows(doc, workflow_paths):
    """Add workflow diagrams section"""
    heading = doc.add_heading("8. Is Akis Diagramlari / Workflow Diagrams", level=1)
    heading.runs[0].font.color.rgb = MAN_RED

    workflows = [
        ("8.1 Is Talebi Akisi / Work Request Flow", workflow_paths[0],
         "Is talebinin olusturulmasindan tamamlanmasina kadar olan surec"),
        ("8.2 Bakim Planlama Akisi / Maintenance Planning Flow", workflow_paths[1],
         "Bakim planlama ve yurutme sureci"),
        ("8.3 Olay Yonetimi Akisi / Incident Management Flow", workflow_paths[2],
         "Acil olay bildirimi ve mudahale sureci"),
        ("8.4 Varlik Yonetimi Akisi / Asset Management Flow", workflow_paths[3],
         "Varlik yasam dongusu yonetimi"),
    ]

    for title, path, desc in workflows:
        doc.add_heading(title, level=2)
        doc.add_paragraph(desc)

        if os.path.exists(path):
            try:
                doc.add_picture(path, width=Inches(6.5))
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except:
                doc.add_paragraph(f"[Workflow diagram: {path}]")

        add_page_break(doc)

def add_data_structure(doc):
    """Add complete data structure section"""
    heading = doc.add_heading("9. Tam Veri Yapisi / Complete Data Structure", level=1)
    heading.runs[0].font.color.rgb = MAN_RED

    doc.add_heading("9.1 Veritabani Tablolari / Database Tables", level=2)

    tables_info = [
        ("users", "Kullanici bilgileri", "userId, username, email, role, department"),
        ("work_requests", "Is talepleri", "requestId, title, description, status, priority, ..."),
        ("assets", "Varliklar/Ekipmanlar", "assetId, sapId, assetName, category, ..."),
        ("maintenance_plans", "Bakim planlari", "planId, assetId, maintenanceType, frequency, ..."),
        ("incidents", "Olay kayitlari", "incidentId, title, priority, status, ..."),
        ("notifications", "Bildirimler", "notificationId, userId, type, message, ..."),
        ("attachments", "Dosya ekleri", "attachmentId, entityType, entityId, fileUrl, ..."),
        ("audit_logs", "Islem kayitlari", "logId, userId, action, entityType, timestamp, ..."),
    ]

    table = doc.add_table(rows=len(tables_info)+1, cols=3)
    table.style = 'Light Grid Accent 1'

    header_cells = table.rows[0].cells
    header_cells[0].text = "Tablo Adi"
    header_cells[1].text = "Aciklama"
    header_cells[2].text = "Anahtar Alanlar"

    for cell in header_cells:
        set_cell_background(cell, MAN_RED)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = WHITE
                run.font.bold = True

    for idx, (table_name, desc, fields) in enumerate(tables_info, 1):
        row = table.rows[idx]
        row.cells[0].text = table_name
        row.cells[1].text = desc
        row.cells[2].text = fields

    add_page_break(doc)

def add_test_scenarios(doc):
    """Add test scenarios for all modules"""
    heading = doc.add_heading("10. Test Senaryolari / Test Scenarios", level=1)
    heading.runs[0].font.color.rgb = MAN_RED

    doc.add_heading("10.1 Is Talepleri Test Senaryolari", level=2)

    test_scenarios = [
        ("TC-WR-001", "Yeni is talebi olusturma - Basarili",
         "1. Yeni is talebi butonuna tikla\n2. Tum zorunlu alanlari doldur\n3. Kaydet\n"
         "Beklenen: Is talebi basariyla olusturulur, liste sayfasinda gorulur"),
        ("TC-WR-002", "Zorunlu alan eksik - Hata",
         "1. Yeni is talebi formu ac\n2. Baslik alanini bos birak\n3. Kaydet\n"
         "Beklenen: 'Baslik zorunludur' hatasi gosterilir"),
        ("TC-WR-003", "Is talebi onaylama",
         "1. Yonetici olarak giris yap\n2. Bekleyen is talebini ac\n3. Onayla\n"
         "Beklenen: Durum ONAYLANDI olur, talep sahibine bildirim gider"),
    ]

    for test_id, test_name, steps in test_scenarios:
        p = doc.add_paragraph()
        run = p.add_run(f"{test_id}: {test_name}")
        run.font.bold = True
        doc.add_paragraph(steps)

    add_page_break(doc)

def add_technical_specs(doc):
    """Add technical specifications"""
    heading = doc.add_heading("11. Teknik Spesifikasyonlar / Technical Specifications", level=1)
    heading.runs[0].font.color.rgb = MAN_RED

    doc.add_heading("11.1 Performans Gereksinimleri", level=2)

    perf_reqs = [
        ("Sayfa Yukleme", "< 2 saniye"),
        ("API Yanit Suresi", "< 500ms (p95)"),
        ("Eszamanli Kullanici", "> 500"),
        ("Veritabani Sorgu", "< 100ms"),
        ("Dosya Yukleme", "< 10 saniye (10MB)"),
    ]

    for req, value in perf_reqs:
        p = doc.add_paragraph()
        run = p.add_run(f"{req}: ")
        run.font.bold = True
        p.add_run(value)

    doc.add_heading("11.2 Guvenlik Gereksinimleri", level=2)

    security_reqs = [
        "SSL/TLS sifreleme (minimum TLS 1.2)",
        "JWT token tabanli kimlik dogrulama",
        "Role-based access control (RBAC)",
        "SQL injection korunmasi",
        "XSS korunmasi",
        "CSRF token validation",
        "Rate limiting (100 req/min per user)",
        "Audit logging (tum islemler kaydedilir)",
    ]

    for req in security_reqs:
        doc.add_paragraph(req, style='List Bullet')

    add_page_break(doc)

def main():
    """Main function - Part 2"""
    print("Loading Part 1 document...")

    doc_path = "/Users/caglarozyildirim/WebstormProjects/Deneme/MAN_Turkiye_Bakim_Yonetimi_KAPSAMLI_Is_Analizi.docx"

    try:
        doc = Document(doc_path)
        print("Part 1 loaded successfully")
    except:
        print("ERROR: Could not load Part 1. Please run Part 1 first!")
        return

    # Add Asset Management Module
    print("Adding Asset Management module...")
    screenshot_asset_list = "/Users/caglarozyildirim/WebstormProjects/Deneme/screenshots/complete/03_Varlik_Yonetimi_Tam.png"
    screenshot_asset_detail = "/Users/caglarozyildirim/WebstormProjects/Deneme/screenshots/complete/04_Varlik_Detay_Tam.png"
    screenshot_asset_modal = "/Users/caglarozyildirim/WebstormProjects/Deneme/screenshots/complete/08_Modal_Varlik_Ekle.png"
    add_asset_management_module(doc, screenshot_asset_list, screenshot_asset_detail, screenshot_asset_modal)

    # Add Maintenance Module
    print("Adding Maintenance Management module...")
    screenshot_maint_list = "/Users/caglarozyildirim/WebstormProjects/Deneme/screenshots/complete/05_Bakim_Yonetimi_Tam.png"
    screenshot_maint_modal = "/Users/caglarozyildirim/WebstormProjects/Deneme/screenshots/complete/09_Modal_Bakim_Plani.png"
    add_maintenance_module(doc, screenshot_maint_list, screenshot_maint_modal)

    # Add Incident Module
    print("Adding Incident Management module...")
    screenshot_inc_list = "/Users/caglarozyildirim/WebstormProjects/Deneme/screenshots/complete/06_Olay_Yonetimi_Tam.png"
    screenshot_inc_modal = "/Users/caglarozyildirim/WebstormProjects/Deneme/screenshots/complete/10_Modal_Acil_Olay_Bildir.png"
    add_incident_module(doc, screenshot_inc_list, screenshot_inc_modal)

    # Add Workflows
    print("Adding workflow diagrams...")
    workflow_paths = [
        "/Users/caglarozyildirim/WebstormProjects/Deneme/diagrams/is_talebi_akisi.png",
        "/Users/caglarozyildirim/WebstormProjects/Deneme/diagrams/bakim_planlama_akisi.png",
        "/Users/caglarozyildirim/WebstormProjects/Deneme/diagrams/olay_yonetimi_akisi.png",
        "/Users/caglarozyildirim/WebstormProjects/Deneme/diagrams/varlik_yonetimi_akisi.png",
    ]
    add_workflows(doc, workflow_paths)

    # Add Data Structure
    print("Adding data structure...")
    add_data_structure(doc)

    # Add Test Scenarios
    print("Adding test scenarios...")
    add_test_scenarios(doc)

    # Add Technical Specs
    print("Adding technical specifications...")
    add_technical_specs(doc)

    # Save final document
    doc.save(doc_path)
    print(f"\nComplete document saved: {doc_path}")
    print("Document is now COMPLETE with all modules and sections!")

if __name__ == "__main__":
    main()
