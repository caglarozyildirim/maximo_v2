#!/usr/bin/env python3
"""
MAN Turkey Maintenance Management System - Comprehensive Business Analysis Document Generator
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# MAN Corporate Colors
MAN_RED = RGBColor(226, 7, 20)
WHITE = RGBColor(255, 255, 255)
BLACK = RGBColor(0, 0, 0)
GRAY = RGBColor(128, 128, 128)

def add_page_break(doc):
    """Add a page break"""
    doc.add_page_break()

def set_cell_background(cell, color):
    """Set cell background color"""
    shading_elm = OxmlElement('w:shd')
    # RGBColor is a tuple-like object, access values by index
    shading_elm.set(qn('w:fill'), f'{color[0]:02X}{color[1]:02X}{color[2]:02X}')
    cell._element.get_or_add_tcPr().append(shading_elm)

def create_title_page(doc):
    """Create title page with MAN branding"""
    # Add MAN logo text
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("MAN")
    run.font.size = Pt(48)
    run.font.bold = True
    run.font.color.rgb = MAN_RED

    # Add subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("MAN Turkiye")
    run.font.size = Pt(36)
    run.font.bold = True

    # Add spacing
    doc.add_paragraph()
    doc.add_paragraph()

    # Main title
    main_title = doc.add_paragraph()
    main_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = main_title.add_run("Bakim Yonetimi Sistemi")
    run.font.size = Pt(28)
    run.font.bold = True

    subtitle2 = doc.add_paragraph()
    subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle2.add_run("Maintenance Management System")
    run.font.size = Pt(24)
    run.font.color.rgb = GRAY

    doc.add_paragraph()
    doc.add_paragraph()

    # Document type
    doc_type = doc.add_paragraph()
    doc_type.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = doc_type.add_run("KAPSAMLI IS ANALIZI DOKUMANI")
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = MAN_RED

    subtitle3 = doc.add_paragraph()
    subtitle3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle3.add_run("Comprehensive Business Analysis Document")
    run.font.size = Pt(16)

    # Add spacing
    for _ in range(8):
        doc.add_paragraph()

    # Footer info
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Yazilim Gelistirme Ekibi Icin")
    run.font.size = Pt(12)
    run.font.italic = True

    footer2 = doc.add_paragraph()
    footer2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer2.add_run("For Software Development Team")
    run.font.size = Pt(12)
    run.font.italic = True

    doc.add_paragraph()

    date = doc.add_paragraph()
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date.add_run("Ekim 2025 / October 2025")
    run.font.size = Pt(12)

    add_page_break(doc)

def add_table_of_contents(doc):
    """Add table of contents"""
    heading = doc.add_heading("Icerik Tablosu / Table of Contents", level=1)
    heading.runs[0].font.color.rgb = MAN_RED

    toc_items = [
        ("1", "Proje Genel Bakis / Project Overview", "4"),
        ("2", "Sistem Mimarisi / System Architecture", "6"),
        ("3", "Modul 1: Ana Sayfa (Dashboard)", "8"),
        ("3.1", "Modul Aciklamasi", "8"),
        ("3.2", "Fonksiyonel Gereksinimler", "9"),
        ("3.3", "Veri Alanlari", "10"),
        ("3.4", "Ekran Goruntuleri", "11"),
        ("3.5", "Is Kurallari", "12"),
        ("4", "Modul 2: Is Talepleri Yonetimi", "13"),
        ("4.1", "Modul Aciklamasi", "13"),
        ("4.2", "Fonksiyonel Gereksinimler", "14"),
        ("4.3", "Veri Alanlari Detayli Tablo", "15"),
        ("4.4", "Ekran Goruntuleri", "17"),
        ("4.5", "Is Kurallari", "18"),
        ("4.6", "Kullanici Etkilesimleri", "19"),
        ("4.7", "Hata Yonetimi", "20"),
        ("5", "Modul 3: Varlik Yonetimi", "21"),
        ("5.1", "Modul Aciklamasi", "21"),
        ("5.2", "Fonksiyonel Gereksinimler", "22"),
        ("5.3", "Veri Alanlari (17 Alan)", "23"),
        ("5.4", "Ekran Goruntuleri", "25"),
        ("5.5", "Is Kurallari", "27"),
        ("5.6", "Kullanici Etkilesimleri", "28"),
        ("5.7", "Hata Yonetimi", "29"),
        ("6", "Modul 4: Bakim Yonetimi", "30"),
        ("6.1", "Modul Aciklamasi", "30"),
        ("6.2", "Fonksiyonel Gereksinimler", "31"),
        ("6.3", "Veri Alanlari ve Bakim Tipleri", "32"),
        ("6.4", "Ekran Goruntuleri", "34"),
        ("6.5", "Is Kurallari", "35"),
        ("6.6", "Kullanici Etkilesimleri", "36"),
        ("6.7", "Hata Yonetimi", "37"),
        ("7", "Modul 5: Olay Yonetimi", "38"),
        ("7.1", "Modul Aciklamasi", "38"),
        ("7.2", "Fonksiyonel Gereksinimler", "39"),
        ("7.3", "Veri Alanlari, Oncelik Seviyeleri, SLA", "40"),
        ("7.4", "Ekran Goruntuleri", "42"),
        ("7.5", "Is Kurallari", "43"),
        ("7.6", "Kullanici Etkilesimleri", "44"),
        ("7.7", "Hata Yonetimi", "45"),
        ("8", "Is Akis Diagramlari / Workflow Diagrams", "46"),
        ("8.1", "Is Talebi Akisi", "46"),
        ("8.2", "Bakim Planlama Akisi", "47"),
        ("8.3", "Olay Yonetimi Akisi", "48"),
        ("8.4", "Varlik Yonetimi Akisi", "49"),
        ("9", "Tam Veri Yapisi / Complete Data Structure", "50"),
        ("9.1", "Tum Veri Tabanlari", "50"),
        ("9.2", "Iliskiler", "55"),
        ("10", "Test Senaryolari / Test Scenarios", "56"),
        ("10.1", "Ana Sayfa Test Senaryolari", "56"),
        ("10.2", "Is Talepleri Test Senaryolari", "57"),
        ("10.3", "Varlik Yonetimi Test Senaryolari", "58"),
        ("10.4", "Bakim Yonetimi Test Senaryolari", "59"),
        ("10.5", "Olay Yonetimi Test Senaryolari", "60"),
        ("11", "Teknik Spesifikasyonlar / Technical Specifications", "61"),
        ("11.1", "Teknoloji Stack", "61"),
        ("11.2", "Performans Gereksinimleri", "62"),
        ("11.3", "Guvenlik Gereksinimleri", "63"),
        ("11.4", "Entegrasyon Noktalari", "64"),
    ]

    table = doc.add_table(rows=len(toc_items), cols=3)
    table.style = 'Light List Accent 1'

    for idx, (num, title, page) in enumerate(toc_items):
        row = table.rows[idx]
        row.cells[0].text = num
        row.cells[1].text = title
        row.cells[2].text = page

        # Make section numbers bold
        if "." not in num:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True

    add_page_break(doc)

def add_project_overview(doc):
    """Add project overview section"""
    heading = doc.add_heading("1. Proje Genel Bakis / Project Overview", level=1)
    heading.runs[0].font.color.rgb = MAN_RED

    # Turkish section
    doc.add_heading("1.1 Proje Tanimi", level=2)
    p = doc.add_paragraph(
        "MAN Turkiye Bakim Yonetimi Sistemi, uretim tesislerindeki tum bakim ve onarim "
        "faaliyetlerini merkezi bir platformda yoneten kapsamli bir yazilim cozumudur. "
        "Sistem, is taleplerinin olusturulmasindan baslayarak, varlik yonetimi, planlanan "
        "ve onleyici bakim faaliyetleri, acil olay mudahalesi ve raporlama sureclerini "
        "dijitallestirir."
    )

    # English section
    doc.add_heading("1.2 Project Definition", level=2)
    p = doc.add_paragraph(
        "The MAN Turkey Maintenance Management System is a comprehensive software solution "
        "that manages all maintenance and repair activities in production facilities through "
        "a centralized platform. The system digitalizes processes from work request creation "
        "to asset management, planned and preventive maintenance activities, emergency incident "
        "response, and reporting procedures."
    )

    # Key objectives
    doc.add_heading("1.3 Anahtar Hedefler / Key Objectives", level=2)

    objectives = [
        ("Dijital Donusum", "Digital Transformation",
         "Tum bakim surecleri kagitsiz ortamda yonetilecek"),
        ("Verimlilik Artisi", "Efficiency Increase",
         "Bakim faaliyetlerinde %30 verimlilik artisi hedeflenmektedir"),
        ("Varlik Yasam Dongusu", "Asset Lifecycle",
         "Ekipman yasam dongusunun tam takibi ve optimizasyonu"),
        ("Gerçek Zamanli Izleme", "Real-time Monitoring",
         "Anlık durum takibi ve hizli mudahale"),
        ("Veri Analizi", "Data Analytics",
         "Kapsamli raporlama ve analiz yetenekleri"),
        ("SLA Uyumu", "SLA Compliance",
         "Hizmet seviyesi anlasmalarinin tam uyumu"),
    ]

    table = doc.add_table(rows=len(objectives)+1, cols=3)
    table.style = 'Light Grid Accent 1'

    # Header
    header_cells = table.rows[0].cells
    header_cells[0].text = "Turkce"
    header_cells[1].text = "English"
    header_cells[2].text = "Aciklama / Description"

    for cell in header_cells:
        set_cell_background(cell, MAN_RED)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = WHITE
                run.font.bold = True

    # Data rows
    for idx, (tr, en, desc) in enumerate(objectives, 1):
        row = table.rows[idx]
        row.cells[0].text = tr
        row.cells[1].text = en
        row.cells[2].text = desc

    # Scope
    doc.add_heading("1.4 Kapsam / Scope", level=2)

    scope_items = [
        "Is Talepleri Yonetimi (Work Request Management)",
        "Varlik/Ekipman Yonetimi (Asset/Equipment Management)",
        "Planlanan Bakim (Planned Maintenance)",
        "Onleyici Bakim (Preventive Maintenance)",
        "Acil Olay Mudahalesi (Emergency Incident Response)",
        "Malzeme ve Yedek Parca Takibi (Material and Spare Parts Tracking)",
        "Personel Atama ve Takibi (Personnel Assignment and Tracking)",
        "Kapsamli Raporlama (Comprehensive Reporting)",
        "Bildirim Sistemi (Notification System)",
        "Mobil Erişim (Mobile Access)",
    ]

    for item in scope_items:
        doc.add_paragraph(item, style='List Bullet')

    # Stakeholders
    doc.add_heading("1.5 Paydaşlar / Stakeholders", level=2)

    stakeholders = [
        ("Bakim Yoneticisi", "Maintenance Manager", "Tum bakim faaliyetlerini yonetir"),
        ("Bakim Teknisyeni", "Maintenance Technician", "Bakim gorevlerini gerceklestirir"),
        ("Uretim Personeli", "Production Staff", "Is taleplerini olusturur"),
        ("Ekipman Sorumlusu", "Equipment Supervisor", "Varliklari yonetir"),
        ("Sistem Yoneticisi", "System Administrator", "Sistem ayarlarini yapar"),
        ("Satinalma Ekibi", "Procurement Team", "Yedek parca ve malzeme temini"),
        ("Ust Yonetim", "Top Management", "Stratejik kararlari alir"),
    ]

    table = doc.add_table(rows=len(stakeholders)+1, cols=3)
    table.style = 'Light Grid Accent 1'

    # Header
    header_cells = table.rows[0].cells
    header_cells[0].text = "Rol (TR)"
    header_cells[1].text = "Role (EN)"
    header_cells[2].text = "Sorumluluk / Responsibility"

    for cell in header_cells:
        set_cell_background(cell, MAN_RED)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = WHITE
                run.font.bold = True

    # Data rows
    for idx, (tr, en, resp) in enumerate(stakeholders, 1):
        row = table.rows[idx]
        row.cells[0].text = tr
        row.cells[1].text = en
        row.cells[2].text = resp

    add_page_break(doc)

def add_system_architecture(doc):
    """Add system architecture section"""
    heading = doc.add_heading("2. Sistem Mimarisi / System Architecture", level=1)
    heading.runs[0].font.color.rgb = MAN_RED

    doc.add_heading("2.1 Mimari Genel Bakis / Architecture Overview", level=2)

    p = doc.add_paragraph(
        "Sistem, modern web teknolojileri kullanilarak 3-katmanli mimari uzerine insa edilmistir:"
    )

    layers = [
        ("Sunum Katmani / Presentation Layer",
         "React.js tabanli responsive kullanici arayuzu, mobil uyumlu tasarim"),
        ("Is Mantigi Katmani / Business Logic Layer",
         "Node.js/Express.js tabanli RESTful API, is kurallari ve validasyon"),
        ("Veri Katmani / Data Layer",
         "PostgreSQL iliskisel veritabani, Redis cache katmani"),
    ]

    for title, desc in layers:
        p = doc.add_paragraph()
        run = p.add_run(f"{title}: ")
        run.font.bold = True
        p.add_run(desc)

    doc.add_heading("2.2 Teknoloji Stack", level=2)

    tech_stack = [
        ("Frontend", "React.js 18+, TypeScript, Material-UI, Redux Toolkit"),
        ("Backend", "Node.js 18+, Express.js 4.x, TypeScript"),
        ("Database", "PostgreSQL 15+, Redis 7+"),
        ("Authentication", "JWT, OAuth 2.0, LDAP Integration"),
        ("File Storage", "AWS S3 / Azure Blob Storage"),
        ("Notifications", "WebSocket, Email (SendGrid), SMS"),
        ("Reporting", "Chart.js, PDF Generation (PDFKit)"),
        ("Mobile", "Progressive Web App (PWA)"),
        ("DevOps", "Docker, Kubernetes, CI/CD (GitLab/Jenkins)"),
    ]

    table = doc.add_table(rows=len(tech_stack)+1, cols=2)
    table.style = 'Light Grid Accent 1'

    # Header
    header_cells = table.rows[0].cells
    header_cells[0].text = "Kategori / Category"
    header_cells[1].text = "Teknolojiler / Technologies"

    for cell in header_cells:
        set_cell_background(cell, MAN_RED)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = WHITE
                run.font.bold = True

    for idx, (cat, tech) in enumerate(tech_stack, 1):
        row = table.rows[idx]
        row.cells[0].text = cat
        row.cells[1].text = tech

    doc.add_heading("2.3 Sistem Bilesenleri / System Components", level=2)

    components = [
        "Web Arayuz / Web Interface",
        "API Gateway",
        "Kimlik Dogrulama Servisi / Authentication Service",
        "Is Talebi Servisi / Work Request Service",
        "Varlik Yonetimi Servisi / Asset Management Service",
        "Bakim Planlama Servisi / Maintenance Planning Service",
        "Olay Yonetimi Servisi / Incident Management Service",
        "Bildirim Servisi / Notification Service",
        "Raporlama Servisi / Reporting Service",
        "Dosya Yonetimi Servisi / File Management Service",
        "Cache Katmani / Cache Layer",
        "Message Queue (RabbitMQ/Kafka)",
        "Log Yonetimi (ELK Stack)",
    ]

    for comp in components:
        doc.add_paragraph(comp, style='List Bullet')

    doc.add_heading("2.4 Guvenlik Mimarisi / Security Architecture", level=2)

    security_features = [
        ("Kimlik Dogrulama", "JWT token tabanli, Multi-factor authentication (MFA)"),
        ("Yetkilendirme", "Role-based access control (RBAC), Fine-grained permissions"),
        ("Veri Sifreleme", "SSL/TLS transport layer, AES-256 data encryption"),
        ("Audit Trail", "Tum islemlerin kayit altina alinmasi, Degisiklik takibi"),
        ("Input Validation", "XSS, SQL injection, CSRF korumalari"),
        ("Session Management", "Secure session handling, Automatic timeout"),
    ]

    for title, desc in security_features:
        p = doc.add_paragraph()
        run = p.add_run(f"{title}: ")
        run.font.bold = True
        p.add_run(desc)

    doc.add_heading("2.5 Entegrasyon Noktalari / Integration Points", level=2)

    integrations = [
        ("SAP ERP", "Varlik bilgileri, malzeme yonetimi, finansal veriler"),
        ("Active Directory/LDAP", "Kullanici kimlik dogrulamasi"),
        ("Email Server", "Bildirim ve raporlar icin email gonderimi"),
        ("SMS Gateway", "Acil durum bildirimleri"),
        ("IoT Sensors", "Ekipman sensoru verilerinin entegrasyonu"),
        ("Document Management System", "Dokuman ve sertifika yonetimi"),
    ]

    table = doc.add_table(rows=len(integrations)+1, cols=2)
    table.style = 'Light Grid Accent 1'

    header_cells = table.rows[0].cells
    header_cells[0].text = "Sistem / System"
    header_cells[1].text = "Entegrasyon Amaci / Integration Purpose"

    for cell in header_cells:
        set_cell_background(cell, MAN_RED)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = WHITE
                run.font.bold = True

    for idx, (system, purpose) in enumerate(integrations, 1):
        row = table.rows[idx]
        row.cells[0].text = system
        row.cells[1].text = purpose

    add_page_break(doc)

def add_dashboard_module(doc, screenshot_path):
    """Add Dashboard module documentation"""
    heading = doc.add_heading("3. Modul 1: Ana Sayfa (Dashboard)", level=1)
    heading.runs[0].font.color.rgb = MAN_RED

    doc.add_heading("3.1 Modul Aciklamasi / Module Description", level=2)

    p = doc.add_paragraph(
        "Ana Sayfa modulu, kullanicilara sistemin genel durumunu tek bir ekranda gosteren "
        "merkezi bir kontrol paneli saglar. Real-time istatistikler, bekleyen is talepleri, "
        "yaklasan bakimlar ve aktif olaylar hakkinda ozet bilgiler sunar."
    )

    p = doc.add_paragraph(
        "The Dashboard module provides users with a centralized control panel that displays "
        "the overall system status on a single screen. It presents summary information about "
        "real-time statistics, pending work requests, upcoming maintenance, and active incidents."
    )

    doc.add_heading("3.2 Fonksiyonel Gereksinimler / Functional Requirements", level=2)

    requirements = [
        ("FR-DASH-001", "Toplam is talebi sayisini gosterme",
         "Sistem, toplam is talebi sayisini ve kategorize edilmis (bekleyen, onaylanan) sayilari gostermelidir"),
        ("FR-DASH-002", "Bekleyen onay sayisini gosterme",
         "Onay bekleyen is talebi sayisini gostermelidir"),
        ("FR-DASH-003", "Tamamlanan isler istatistigi",
         "Belirli donem icinde tamamlanan is sayisini gostermelidir"),
        ("FR-DASH-004", "Aktif olay sayisi",
         "Suanda aktif olan kritik ve yuksek oncelikli olaylari gostermelidir"),
        ("FR-DASH-005", "Uretim ekipmani istatistikleri",
         "Toplam ekipman sayisi, aktif, bakimda ve arizali ekipman sayilari"),
        ("FR-DASH-006", "Planli bakim takvimi",
         "Bu hafta ve gelecek hafta yapilacak planli bakimlari listelemelidir"),
        ("FR-DASH-007", "Yaklasan bakimlar listesi",
         "Son 7 gun icinde yaklasan bakimlari gostermelidir"),
        ("FR-DASH-008", "Aktif acil olaylar",
         "Kritik oncelikli aktif olaylari kirmizi renkte vurgulamalidirr"),
        ("FR-DASH-009", "Son aktiviteler akisi",
         "Son 10 sistem aktivitesini zaman damgali olarak listelemelidir"),
        ("FR-DASH-010", "Hizli erisim butonlari",
         "Yeni is talebi, yeni varlik, yeni bakim plani olusturma butonlari"),
        ("FR-DASH-011", "Son guncelleme zamani",
         "Dashboard verisinin son guncellenme zamani gosterilmelidir"),
        ("FR-DASH-012", "Otomatik yenileme",
         "Dashboard her 30 saniyede bir otomatik olarak yenilenmeli"),
    ]

    table = doc.add_table(rows=len(requirements)+1, cols=3)
    table.style = 'Light Grid Accent 1'

    header_cells = table.rows[0].cells
    header_cells[0].text = "ID"
    header_cells[1].text = "Gereksinim / Requirement"
    header_cells[2].text = "Aciklama / Description"

    for cell in header_cells:
        set_cell_background(cell, MAN_RED)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = WHITE
                run.font.bold = True

    for idx, (req_id, req, desc) in enumerate(requirements, 1):
        row = table.rows[idx]
        row.cells[0].text = req_id
        row.cells[1].text = req
        row.cells[2].text = desc

    add_page_break(doc)

    doc.add_heading("3.3 Veri Alanlari / Data Fields", level=2)

    data_fields = [
        ("toplamIsTalepleri", "Integer", "Evet", "Toplam is talebi sayisi",
         ">= 0", "87", "Dashboard API'den"),
        ("bekleyenOnaylar", "Integer", "Evet", "Bekleyen onay sayisi",
         ">= 0", "12", "Approval API'den"),
        ("tamamlananIsler", "Integer", "Evet", "Tamamlanan is sayisi",
         ">= 0", "68", "Completed API'den"),
        ("aktifOlaylar", "Integer", "Evet", "Aktif olay sayisi",
         ">= 0", "2", "Incident API'den"),
        ("uretimEkipmani", "Integer", "Evet", "Toplam ekipman sayisi",
         ">= 0", "324", "Asset API'den"),
        ("planlanmisİşSayısı", "Integer", "Evet", "Bu ay planlanan is sayisi",
         ">= 0", "18", "Maintenance API'den"),
        ("planliiBakimlar", "Integer", "Evet", "Bu hafta planli bakim sayisi",
         ">= 0", "8", "Maintenance API'den"),
        ("sonGuncelleme", "DateTime", "Evet", "Son guncelleme zamani",
         "ISO 8601", "2025-10-08 15:30", "System timestamp"),
        ("kullaniciAdi", "String", "Evet", "Oturum acik kullanici",
         "Max 100 char", "Mehmet Yilmaz", "Session data"),
    ]

    table = doc.add_table(rows=len(data_fields)+1, cols=7)
    table.style = 'Light Grid Accent 1'

    header_cells = table.rows[0].cells
    headers = ["Alan Adi", "Tip", "Zorunlu", "Aciklama", "Validasyon", "Ornek", "Kaynak"]
    for idx, header in enumerate(headers):
        header_cells[idx].text = header
        set_cell_background(header_cells[idx], MAN_RED)
        for paragraph in header_cells[idx].paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = WHITE
                run.font.bold = True
                run.font.size = Pt(9)

    for idx, field_data in enumerate(data_fields, 1):
        row = table.rows[idx]
        for col_idx, value in enumerate(field_data):
            row.cells[col_idx].text = value
            for paragraph in row.cells[col_idx].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)

    # Set column widths
    for row in table.rows:
        row.cells[0].width = Inches(1.5)
        row.cells[1].width = Inches(0.8)
        row.cells[2].width = Inches(0.7)
        row.cells[3].width = Inches(1.5)
        row.cells[4].width = Inches(1.0)
        row.cells[5].width = Inches(1.0)
        row.cells[6].width = Inches(1.0)

    add_page_break(doc)

    doc.add_heading("3.4 Ekran Goruntuleri / Screenshots", level=2)

    doc.add_heading("3.4.1 Ana Sayfa - Tam Ekran", level=3)
    p = doc.add_paragraph(
        "Asagidaki ekran goruntusunde Ana Sayfa'nin tam gorunumu yer almaktadir. "
        "Istatistik kartlari, bekleyen is talepleri, yaklasan bakimlar, aktif acil olaylar "
        "ve son aktiviteler gorulmektedir."
    )

    if os.path.exists(screenshot_path):
        try:
            doc.add_picture(screenshot_path, width=Inches(6.5))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except:
            doc.add_paragraph(f"[Ekran goruntusu: {screenshot_path}]")

    add_page_break(doc)

    doc.add_heading("3.5 Is Kurallari / Business Rules", level=2)

    business_rules = [
        ("BR-DASH-001", "Istatistik Guncelleme Siklig",
         "Dashboard istatistikleri her 30 saniyede bir otomatik olarak guncellenmelidir"),
        ("BR-DASH-002", "Kritik Olay Vurgulama",
         "Kritik oncelikli olaylar kirmizi renkte vurgulanmalidir"),
        ("BR-DASH-003", "Gecmis Is Talebi Renklendirme",
         "Suresi gecmis is talepleri kirmizi, yaklasan (24 saat icinde) turuncu renkte gosterilmelidir"),
        ("BR-DASH-004", "Kullanici Yetki Kontrolu",
         "Kullanici sadece yetkisi olan modullerin verilerini gorebilmelidir"),
        ("BR-DASH-005", "Yaklasan Bakim Uyarisi",
         "7 gun icinde bakim suresi dolacak ekipmanlar dashboard'da gosterilmelidir"),
        ("BR-DASH-006", "Performans Optimizasyonu",
         "Dashboard yukleme suresi 2 saniyeden kisa olmalidir"),
        ("BR-DASH-007", "Mobil Uyumluluk",
         "Dashboard mobil cihazlarda responsive olarak gorunmelidir"),
    ]

    table = doc.add_table(rows=len(business_rules)+1, cols=3)
    table.style = 'Light Grid Accent 1'

    header_cells = table.rows[0].cells
    header_cells[0].text = "Kural ID"
    header_cells[1].text = "Kural Adi"
    header_cells[2].text = "Aciklama"

    for cell in header_cells:
        set_cell_background(cell, MAN_RED)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = WHITE
                run.font.bold = True

    for idx, (rule_id, rule_name, desc) in enumerate(business_rules, 1):
        row = table.rows[idx]
        row.cells[0].text = rule_id
        row.cells[1].text = rule_name
        row.cells[2].text = desc

    add_page_break(doc)

def add_work_request_module(doc, screenshot_list, screenshot_modal):
    """Add Work Request module documentation"""
    heading = doc.add_heading("4. Modul 2: Is Talepleri Yonetimi", level=1)
    heading.runs[0].font.color.rgb = MAN_RED

    doc.add_heading("4.1 Modul Aciklamasi / Module Description", level=2)

    p = doc.add_paragraph(
        "Is Talepleri Yonetimi modulu, uretim personelinin bakim ve onarim taleplerini "
        "olusturmasini, takip etmesini ve yonetmesini saglar. Talepler kategorize edilir, "
        "onceliklendirilir ve uygun teknisyenlere atanir. Tum is talebi yasam dongusu "
        "(olusturma, onay, atama, gerceklestirme, tamamlama) bu modulde yonetilir."
    )

    p = doc.add_paragraph(
        "The Work Request Management module enables production staff to create, track, and "
        "manage maintenance and repair requests. Requests are categorized, prioritized, and "
        "assigned to appropriate technicians. The entire work request lifecycle (creation, "
        "approval, assignment, execution, completion) is managed in this module."
    )

    doc.add_heading("4.2 Fonksiyonel Gereksinimler / Functional Requirements", level=2)

    requirements = [
        ("FR-WR-001", "Is talebi olusturma",
         "Kullanicilar yeni is talebi olusturabilmelidir (baslik, aciklama, kategori, oncelik, lokasyon)"),
        ("FR-WR-002", "Is talebi listeleme",
         "Tum is talepleri filtrelenebilir ve siralanabilir sekilde listelenmeli"),
        ("FR-WR-003", "Is talebi detay goruntuleme",
         "Her is talebinin detayli bilgileri goruntulenbilmelidir"),
        ("FR-WR-004", "Is talebi guncelleme",
         "Yetkili kullanicilar is talebi bilgilerini guncelleyebilmelidir"),
        ("FR-WR-005", "Is talebi onaylama/reddetme",
         "Yonetici seviyesindeki kullanicilar talepleri onaylayabilir veya reddedebilmelidir"),
        ("FR-WR-006", "Teknisyen atama",
         "Onaylanan is talepleri uygun teknisyenlere atanabilmelidir"),
        ("FR-WR-007", "Durum degisikligi",
         "Is talebi durumu (Beklemede, Onaylandi, Atandi, Islemde, Tamamlandi, Reddedildi) degistirilebilmelidir"),
        ("FR-WR-008", "Yorum ekleme",
         "Kullanicilar is talebine yorum ekleyebilmelidir"),
        ("FR-WR-009", "Dosya ekleme",
         "Is talebine resim, dokuman vs. dosya eklenebilmelidir"),
        ("FR-WR-010", "Bildirim gonderme",
         "Durum degisikliklerinde ilgili kullanicilara bildirim gonderilmelidir"),
        ("FR-WR-011", "Arama ve filtreleme",
         "Is talepleri numara, baslik, kategori, durum, oncelik, lokasyon, tarih araligina gore filtrelenebilmelidir"),
        ("FR-WR-012", "Toplu islemler",
         "Birden fazla is talebi secilerek toplu durum degisikligi yapilabilmelidir"),
        ("FR-WR-013", "Is talebi gecmisi",
         "Tum degisiklik gecmisi (audit trail) goruntulenbilmelidir"),
        ("FR-WR-014", "Sure takibi",
         "Olusturulma ve tamamlanma sureleri otomatik hesaplanmalidir"),
        ("FR-WR-015", "SLA uyari sistemi",
         "SLA sureleri asildigi zaman sistem uyari vermelidir"),
    ]

    table = doc.add_table(rows=len(requirements)+1, cols=3)
    table.style = 'Light Grid Accent 1'

    header_cells = table.rows[0].cells
    header_cells[0].text = "ID"
    header_cells[1].text = "Gereksinim / Requirement"
    header_cells[2].text = "Aciklama / Description"

    for cell in header_cells:
        set_cell_background(cell, MAN_RED)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = WHITE
                run.font.bold = True

    for idx, (req_id, req, desc) in enumerate(requirements, 1):
        row = table.rows[idx]
        row.cells[0].text = req_id
        row.cells[1].text = req
        row.cells[2].text = desc

    add_page_break(doc)

    doc.add_heading("4.3 Veri Alanlari Detayli Tablo / Detailed Data Fields Table", level=2)

    data_fields = [
        ("requestId", "String (UUID)", "Evet (Auto)", "Benzersiz is talebi kimlik numarasi",
         "UUID format", "JR-2025-001", "Otomatik generate", "Primary key"),
        ("title", "String", "Evet", "Is talebi basligi",
         "5-200 karakter", "Boya Kabini Filtre Degisimi", "Kullanici girisi", "Indexed"),
        ("description", "Text", "Evet", "Detayli is aciklamasi",
         "10-2000 karakter", "Ana boya hatti boya kabini filtrelerinde tikanimalar...", "Kullanici girisi", "-"),
        ("category", "Enum", "Evet", "Is kategorisi",
         "HVAC, MEKANIK, ELEKTRIK, BINA, IT, DIGER", "HVAC", "Dropdown secimi", "Indexed"),
        ("priority", "Enum", "Evet", "Oncelik seviyesi",
         "DUSUK, NORMAL, YUKSEK, KRITIK", "YUKSEK", "Dropdown secimi", "Indexed"),
        ("location", "String", "Evet", "Is lokasyonu",
         "Max 200 karakter", "Sasi Uretim Hatti - Ana Boya Hatti", "Dropdown/Text", "Indexed"),
        ("status", "Enum", "Evet", "Is durumu",
         "BEKLEMEDE, ONAYLANDI, ATANDI, ISLEMDE, TAMAMLANDI, REDDEDILDI", "BEKLEMEDE", "System managed", "Indexed"),
        ("requestedBy", "String", "Evet (Auto)", "Talep eden kullanici",
         "User ID reference", "Mehmet Demir", "Session user", "Foreign key"),
        ("requestedDate", "DateTime", "Evet (Auto)", "Talep olusturma tarihi",
         "ISO 8601", "2025-10-08T08:00:00Z", "System timestamp", "Indexed"),
        ("assignedTo", "String", "Hayir", "Atanan teknisyen",
         "User ID reference", "Ali Yilmaz", "Admin secimi", "Foreign key"),
        ("assignedDate", "DateTime", "Hayir", "Atanma tarihi",
         "ISO 8601", "2025-10-08T09:30:00Z", "System timestamp", "-"),
        ("dueDate", "DateTime", "Hayir", "Bitis hedef tarihi",
         "ISO 8601, >= now", "2025-10-10T17:00:00Z", "Hesaplanir/Manuel", "Indexed"),
        ("completedDate", "DateTime", "Hayir", "Tamamlanma tarihi",
         "ISO 8601", "2025-10-09T16:45:00Z", "System timestamp", "-"),
        ("estimatedHours", "Decimal", "Hayir", "Tahmini sure (saat)",
         ">= 0, max 1000", "4.5", "Kullanici/Teknisyen", "-"),
        ("actualHours", "Decimal", "Hayir", "Gerceklesen sure (saat)",
         ">= 0, max 1000", "5.2", "Teknisyen girisi", "-"),
        ("costEstimate", "Decimal", "Hayir", "Tahmini maliyet (TL)",
         ">= 0", "5000.00", "Kullanici girisi", "-"),
        ("actualCost", "Decimal", "Hayir", "Gerceklesen maliyet (TL)",
         ">= 0", "5850.00", "System calculated", "-"),
        ("relatedAssetId", "String", "Hayir", "Iliskili ekipman",
         "Asset ID reference", "AST-001", "Dropdown secimi", "Foreign key"),
        ("attachments", "Array", "Hayir", "Ekli dosyalar",
         "File URLs", "[{name: 'foto1.jpg', url: '...'}]", "File upload", "-"),
        ("comments", "Array", "Hayir", "Yorumlar",
         "Comment objects", "[{user: '...', text: '...', date: '...'}]", "User input", "-"),
        ("approvedBy", "String", "Hayir", "Onaylayan kisi",
         "User ID reference", "Ayse Kaya", "System managed", "Foreign key"),
        ("approvedDate", "DateTime", "Hayir", "Onay tarihi",
         "ISO 8601", "2025-10-08T09:00:00Z", "System timestamp", "-"),
        ("rejectionReason", "Text", "Hayir", "Red sebebi",
         "Max 500 karakter", "Yedek parca bulunmuyor", "Admin girisi", "-"),
        ("createdAt", "DateTime", "Evet (Auto)", "Kayit olusturma",
         "ISO 8601", "2025-10-08T08:00:00Z", "System timestamp", "-"),
        ("updatedAt", "DateTime", "Evet (Auto)", "Son guncelleme",
         "ISO 8601", "2025-10-08T10:15:00Z", "System timestamp", "-"),
    ]

    # Create table with smaller font
    table = doc.add_table(rows=len(data_fields)+1, cols=8)
    table.style = 'Light Grid Accent 1'

    header_cells = table.rows[0].cells
    headers = ["Alan Adi", "Veri Tipi", "Zorunlu", "Aciklama", "Validasyon Kurallari",
               "Ornek Deger", "Veri Kaynagi", "DB Index"]
    for idx, header in enumerate(headers):
        header_cells[idx].text = header
        set_cell_background(header_cells[idx], MAN_RED)
        for paragraph in header_cells[idx].paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = WHITE
                run.font.bold = True
                run.font.size = Pt(8)

    for idx, field_data in enumerate(data_fields, 1):
        row = table.rows[idx]
        for col_idx, value in enumerate(field_data):
            row.cells[col_idx].text = value
            for paragraph in row.cells[col_idx].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(7)

    # Set column widths
    for row in table.rows:
        row.cells[0].width = Inches(1.2)
        row.cells[1].width = Inches(1.0)
        row.cells[2].width = Inches(0.6)
        row.cells[3].width = Inches(1.2)
        row.cells[4].width = Inches(1.3)
        row.cells[5].width = Inches(1.0)
        row.cells[6].width = Inches(0.9)
        row.cells[7].width = Inches(0.6)

    add_page_break(doc)

    doc.add_heading("4.4 Ekran Goruntuleri / Screenshots", level=2)

    doc.add_heading("4.4.1 Is Talepleri Listesi - Tam Ekran", level=3)
    p = doc.add_paragraph(
        "Asagida is talepleri liste ekraninin tam gorunumu bulunmaktadir. Arama/filtreleme "
        "alani, istatistik kartlari ve detayli tablo yer almaktadir."
    )

    if os.path.exists(screenshot_list):
        try:
            doc.add_picture(screenshot_list, width=Inches(6.5))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except:
            doc.add_paragraph(f"[Ekran goruntusu: {screenshot_list}]")

    add_page_break(doc)

    doc.add_heading("4.4.2 Yeni Is Talebi Olusturma Modal", level=3)
    p = doc.add_paragraph(
        "Asagida yeni is talebi olusturma modal ekraninin gorunumu bulunmaktadir. "
        "Form alanlari ve validasyon mesajlari gorulmektedir."
    )

    if os.path.exists(screenshot_modal):
        try:
            doc.add_picture(screenshot_modal, width=Inches(6.5))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except:
            doc.add_paragraph(f"[Ekran goruntusu: {screenshot_modal}]")

    add_page_break(doc)

    doc.add_heading("4.5 Is Kurallari / Business Rules", level=2)

    business_rules = [
        ("BR-WR-001", "Otomatik ID Uretimi",
         "Is talebi ID'si JR-YYYY-XXX formatinda otomatik olarak uretilir (JR = Job Request, YYYY = Yil, XXX = Sira no)"),
        ("BR-WR-002", "Oncelik Seviyesi Kurallari",
         "KRITIK: 4 saat icinde atanmali, YUKSEK: 24 saat, NORMAL: 72 saat, DUSUK: 1 hafta"),
        ("BR-WR-003", "Onay Sureci",
         "Is talepleri BEKLEMEDE durumunda olusturulur, yonetici onayiyla ONAYLANDI durumuna gecer"),
        ("BR-WR-004", "Atama Kurallari",
         "Sadece ONAYLANDI durumdaki is talepleri teknisyenlere atanabilir"),
        ("BR-WR-005", "Durum Gecis Kurallari",
         "BEKLEMEDE -> ONAYLANDI/REDDEDILDI -> ATANDI -> ISLEMDE -> TAMAMLANDI (yalnizca bu sirada degisilebilir)"),
        ("BR-WR-006", "Geri Donus Yasagi",
         "TAMAMLANDI veya REDDEDILDI durumundaki is talepleri geri alinmaaz"),
        ("BR-WR-007", "Zorunlu Alan Kontrolu",
         "Baslik, aciklama, kategori, oncelik ve lokasyon alanlari zorunludur"),
        ("BR-WR-008", "SLA Hesaplama",
         "Due date, oncelik seviyesine gore otomatik hesaplanir (KRITIK: +4h, YUKSEK: +24h, NORMAL: +72h, DUSUK: +7d)"),
        ("BR-WR-009", "Bildirim Kurallari",
         "Durum degisikliklerinde: talep eden, atanan teknisyen ve yoneticiye email bildirimi gonderilir"),
        ("BR-WR-010", "Dosya Yukleme Siniri",
         "Her is talebine maksimum 10 dosya (her biri max 10MB) eklenebilir"),
        ("BR-WR-011", "Yorum Kisitlamasi",
         "Yalnizca ilgili kullanicilar (talep eden, atanan, yonetici) yorum ekleyebilir"),
        ("BR-WR-012", "Maliyet Guncelleme",
         "actualCost, kullanilan malzeme ve calisan saatlerine gore otomatik hesaplanir"),
        ("BR-WR-013", "Sure Gecimi Uyarisi",
         "Due date'i gecmis is talepleri kirmizi renkte isaretlenir ve gunluk uyari maili gonderilir"),
        ("BR-WR-014", "Toplu Islem Yetkisi",
         "Toplu durum degisiklligi sadece yonetici ve supervior rolleriyle yapilabilir"),
    ]

    table = doc.add_table(rows=len(business_rules)+1, cols=3)
    table.style = 'Light Grid Accent 1'

    header_cells = table.rows[0].cells
    header_cells[0].text = "Kural ID"
    header_cells[1].text = "Kural Adi"
    header_cells[2].text = "Aciklama"

    for cell in header_cells:
        set_cell_background(cell, MAN_RED)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = WHITE
                run.font.bold = True

    for idx, (rule_id, rule_name, desc) in enumerate(business_rules, 1):
        row = table.rows[idx]
        row.cells[0].text = rule_id
        row.cells[1].text = rule_name
        row.cells[2].text = desc

    add_page_break(doc)

    doc.add_heading("4.6 Kullanici Etkilesimleri / User Interactions", level=2)

    interactions = [
        ("Yeni Is Talebi Olusturma",
         "1. Kullanici '+ Yeni Is Talebi' butonuna tiklar\n"
         "2. Modal formu acar\n"
         "3. Zorunlu alanlari doldurur\n"
         "4. Opsiyonel dosya ve resim ekler\n"
         "5. 'Talep Olustur' butonuna tiklar\n"
         "6. Sistem validasyon yapar\n"
         "7. Basarili ise listeye eklenir ve bildirim gonderilir"),
        ("Is Talebi Onaylama",
         "1. Yonetici dashboard'dan bekleyen onaylari gorur\n"
         "2. Is talebi detayini acar\n"
         "3. 'Onayla' veya 'Reddet' butonuna tiklar\n"
         "4. Reddetme durumunda sebep girer\n"
         "5. Sistem durumu gunceller\n"
         "6. Talep sahibine bildirim gonderir"),
        ("Teknisyen Atama",
         "1. Yonetici onaylanmis is talebini acar\n"
         "2. 'Teknisyen Ata' butonuna tiklar\n"
         "3. Musait teknisyenler listesinden secer\n"
         "4. Tahmini sure ve kaynak girer\n"
         "5. 'Ata' butonuna tiklar\n"
         "6. Sistem teknisyene bildirim gonderir"),
        ("Is Durumu Guncelleme",
         "1. Teknisyen kendisine atanan isi gorur\n"
         "2. 'Basla' butonuna tiklayarak ISLEMDE yapar\n"
         "3. Ilerme guncellemeleri ve yorumlar ekler\n"
         "4. Isi bitirdiginde 'Tamamla' butonuna tiklar\n"
         "5. Gerceklesen sure ve malzeme bilgisi girer\n"
         "6. Sistem durumu TAMAMLANDI yapar"),
        ("Filtreleme ve Arama",
         "1. Kullanici arama kutusuna metin girer\n"
         "2. Filtre dropdownlarindan secimler yapar\n"
         "3. Tarih araligi secer\n"
         "4. 'Ara' butonuna tiklar\n"
         "5. Sistem filtrelenmis sonuclari listeler\n"
         "6. 'Filtreleri Temizle' ile sifirlanabilir"),
    ]

    for title, steps in interactions:
        p = doc.add_paragraph()
        run = p.add_run(f"{title}:")
        run.font.bold = True
        run.font.size = Pt(11)

        p = doc.add_paragraph(steps)
        p.style = 'List Number'

    add_page_break(doc)

    doc.add_heading("4.7 Hata Yonetimi / Error Handling", level=2)

    error_scenarios = [
        ("ERR-WR-001", "Zorunlu Alan Eksik",
         "Zorunlu alanlardan biri bos birakilirsa -> 'Bu alan zorunludur' uyarisi goster"),
        ("ERR-WR-002", "Gecersiz Tarih",
         "Due date gecmis tarih secilirse -> 'Bitis tarihi gelecekte olmalidir' uyarisi"),
        ("ERR-WR-003", "Dosya Boyutu Asimi",
         "10MB'dan buyuk dosya yuklenirse -> 'Dosya boyutu 10MB'i asmamalidira' hatasi"),
        ("ERR-WR-004", "Yetkisiz Islem",
         "Kullanici yetkisi olmayan islemi yapmaya calissirsa -> '403 Forbidden' hatasi"),
        ("ERR-WR-005", "Network Hatasi",
         "API isteği basarisiz olursa -> 'Baglanti hatasi. Lutfen tekrar deneyin' mesaji"),
        ("ERR-WR-006", "Is Talebi Bulunamadi",
         "Silinmis/mevcut olmayan is talebi acilmaya calisilirsa -> '404 Not Found' hatasi"),
        ("ERR-WR-007", "Gecersiz Durum Gecisi",
         "Izin verilmeyen durum degisikligi yapilirsa -> 'Gecersiz durum gecisi' hatasi"),
        ("ERR-WR-008", "Teknisyen Atanamadi",
         "Teknisyen musait degilse -> 'Secili teknisyen bu tarihte musait degil' uyarisi"),
        ("ERR-WR-009", "Session Timeout",
         "Kullanici oturumu sona ermisse -> Login sayfasina yonlendir"),
        ("ERR-WR-010", "Dosya Tipi Desteklenmiyor",
         "Desteklenmeyen dosya tipi yuklenirse -> 'Sadece JPG, PNG, PDF dosyalari yuklenebilir' hatasi"),
    ]

    table = doc.add_table(rows=len(error_scenarios)+1, cols=3)
    table.style = 'Light Grid Accent 1'

    header_cells = table.rows[0].cells
    header_cells[0].text = "Hata Kodu"
    header_cells[1].text = "Hata Tipi"
    header_cells[2].text = "Cozum / Solution"

    for cell in header_cells:
        set_cell_background(cell, MAN_RED)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = WHITE
                run.font.bold = True

    for idx, (code, error_type, solution) in enumerate(error_scenarios, 1):
        row = table.rows[idx]
        row.cells[0].text = code
        row.cells[1].text = error_type
        row.cells[2].text = solution

    add_page_break(doc)

def main():
    """Main function to create the document"""
    print("Creating comprehensive MAN Turkey Maintenance Management System document...")

    # Create document
    doc = Document()

    # Set up document properties
    doc.core_properties.title = "MAN Turkiye Bakim Yonetimi Sistemi - Kapsamli Is Analizi"
    doc.core_properties.author = "MAN Turkiye"
    doc.core_properties.subject = "Comprehensive Business Analysis"

    # Create title page
    print("Creating title page...")
    create_title_page(doc)

    # Add table of contents
    print("Adding table of contents...")
    add_table_of_contents(doc)

    # Add project overview
    print("Adding project overview...")
    add_project_overview(doc)

    # Add system architecture
    print("Adding system architecture...")
    add_system_architecture(doc)

    # Add Dashboard module
    print("Adding Dashboard module...")
    screenshot_dashboard = "/Users/caglarozyildirim/WebstormProjects/Deneme/screenshots/complete/01_Ana_Sayfa_Tam.png"
    add_dashboard_module(doc, screenshot_dashboard)

    # Add Work Request module
    print("Adding Work Request module...")
    screenshot_list = "/Users/caglarozyildirim/WebstormProjects/Deneme/screenshots/complete/02_Is_Talepleri_Tam.png"
    screenshot_modal = "/Users/caglarozyildirim/WebstormProjects/Deneme/screenshots/complete/07_Modal_Is_Talebi_Olustur.png"
    add_work_request_module(doc, screenshot_list, screenshot_modal)

    # Note: Additional modules (Asset, Maintenance, Incident) will be added in continuation
    # This is Part 1 of the document generation

    # Save document
    output_path = "/Users/caglarozyildirim/WebstormProjects/Deneme/MAN_Turkiye_Bakim_Yonetimi_KAPSAMLI_Is_Analizi.docx"
    doc.save(output_path)
    print(f"\nDocument created successfully: {output_path}")
    print("Note: This is Part 1. Continuing with remaining modules...")

if __name__ == "__main__":
    main()
