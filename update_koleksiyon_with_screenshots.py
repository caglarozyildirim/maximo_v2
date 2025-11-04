#!/usr/bin/env python3
"""
Update Koleksiyon Mobilya Dealer Portal Document with Screenshots
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import os

# Koleksiyon Corporate Colors
KOLEKSIYON_BLUE = RGBColor(0, 51, 102)
KOLEKSIYON_ORANGE = RGBColor(255, 102, 0)
DARK_GRAY = RGBColor(51, 51, 51)
LIGHT_GRAY = RGBColor(128, 128, 128)
SUCCESS_GREEN = RGBColor(34, 139, 34)

# Screenshot paths
SCREENSHOTS_DIR = "/Users/caglarozyildirim/Desktop/Şirketler/Koleksiyon/görseller"

screenshots = [
    {
        "file": "CleanShot 2025-10-15 at 08.24.42@2x.png",
        "title": "SAP Ürünleri Listesi - Ana Tablo Görünümü",
        "description": "SAP sisteminden çekilen ürün listesi. MATNR (malzeme numarası), ürün adı, model kodu, tasarımcı, açıklama, tedarikçi ve malzeme açıklaması gibi alanları içerir."
    },
    {
        "file": "CleanShot 2025-10-15 at 08.35.38@2x.png",
        "title": "SAP Ürün Detay Formu - Genel Bilgiler",
        "description": "Ürün detay sayfası genel bilgiler bölümü. Müşteri ID, malzeme açıklaması (TR/EN), üretici parça numarası ve demo ürün bilgilerini içerir."
    },
    {
        "file": "CleanShot 2025-10-15 at 08.36.02@2x.png",
        "title": "PIM Yönetim Paneli - Ürün Kategori Bilgileri",
        "description": "Koleksiyon PIM yönetim panelinde ürün kategori bilgileri. E-Dükkan, Tekstil, Parfüm, Kanal, Tedarikçi, Ana Grup, Alt Grup, Grup ve Ürün alanlarını gösterir."
    },
    {
        "file": "CleanShot 2025-10-15 at 08.36.09@2x.png",
        "title": "PIM Yönetim Paneli - Malzeme ve Boyut Bilgileri",
        "description": "Ürün malzeme bilgisi (mermer), kumaş, renk ve tasarımcı bilgileri. Ayrıca model kodu, teklif kodu ve boyut bilgilerini içerir."
    },
    {
        "file": "CleanShot 2025-10-15 at 08.36.13@2x.png",
        "title": "PIM Yönetim Paneli - Ölçü ve Ağırlık Bilgileri",
        "description": "Ürün ölçü bilgileri: Genişlik, Genişlik ABD, Yükseklik, Yükseklik ABD, Çap, Çap ABD, Maksimum Beden, Maksimum Genişlik, Maksimum Yükseklik ve Paket Sayısı."
    },
    {
        "file": "CleanShot 2025-10-15 at 08.36.21@2x.png",
        "title": "PIM Yönetim Paneli - Ürün Hacmi ve Açıklamalar",
        "description": "Ürün hacmi, ürün hacmi ABD, ürün ağırlığı bilgileri. Montaj süresi, taşıma süresi, açıklama (TR/EN) ve TVM TR alanları."
    },
    {
        "file": "CleanShot 2025-10-15 at 08.36.27@2x.png",
        "title": "PIM Yönetim Paneli - Kategori ve Ürün Grupları",
        "description": "Z-Kanal (Home), Z-Tedarik (Production), Z-Ana Grup (Living), Z-Alt Grup (Coffee Table), Z-Grup (coffee table) gibi hiyerarşik kategori bilgileri."
    },
    {
        "file": "CleanShot 2025-10-15 at 08.36.33@2x.png",
        "title": "PIM Yönetim Paneli - Ürün Kumaş ve Kaydet",
        "description": "Ürün form son bölümü. Z-Murun ölçü bilgileri, kumaş metrajları ve KAYDET butonu görünümü."
    }
]

print("="*80)
print("UPDATING KOLEKSIYON MOBILYA DEALER PORTAL DOCUMENT")
print("WITH SCREENSHOTS")
print("="*80 + "\n")

# Load existing document
doc = Document()

# Set default font
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# ============================================================================
# TITLE PAGE
# ============================================================================
print("📄 Creating title page...")

title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title_para.add_run("KOLEKSİYON MOBİLYA\nBAYİ PORTALI PROJESİ")
title_run.font.size = Pt(24)
title_run.font.bold = True
title_run.font.color.rgb = KOLEKSIYON_BLUE

subtitle_para = doc.add_paragraph()
subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_run = subtitle_para.add_run("\nİş Analizi ve Kapsam Dokümanı\nEkran Görüntüleri ile Güncellenmiş Versiyon")
subtitle_run.font.size = Pt(16)
subtitle_run.font.color.rgb = KOLEKSIYON_ORANGE

doc.add_paragraph("\n" * 3)

# Document info table
info_para = doc.add_paragraph()
info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
info_run = info_para.add_run(f"Versiyon: 2.0 (Ekran Görüntüleri Eklendi)\n")
info_run.font.size = Pt(12)
info_run = info_para.add_run(f"Tarih: {datetime.now().strftime('%d.%m.%Y')}\n")
info_run.font.size = Pt(12)
info_run = info_para.add_run("Hazırlayan: İş Analisti")
info_run.font.size = Pt(12)

doc.add_page_break()

# ============================================================================
# EXECUTIVE SUMMARY
# ============================================================================
print("📝 Section: Executive Summary...")

heading = doc.add_heading('YÖNETİCİ ÖZETİ', 1)
heading.runs[0].font.color.rgb = KOLEKSIYON_BLUE

doc.add_paragraph(
    "Koleksiyon Mobilya, ana ERP sistemi olarak SAP kullanan ve katalog yönetimi için "
    "mevcut CMS (Butterfly/PIM) sistemini kullanan lider bir mobilya üreticisidir. "
    "Bu proje kapsamında, bayilerin SAP'de yer alan ürün fiyat listelerine erişebileceği, "
    "web tabanlı güvenli bir Bayi Portalı geliştirilecektir."
)

doc.add_paragraph(
    "Projenin temel amacı, bayilere ürün bilgileri ve fiyatlandırma kırılımlarını "
    "çevrimiçi olarak sunmak, manuel süreçleri azaltmak ve bayi memnuniyetini artırmaktır. "
    "Portal, mevcut CMS altyapısı üzerine entegre edilecek ve SAP'den gerçek zamanlı veri "
    "akışı sağlanacaktır."
)

doc.add_page_break()

# ============================================================================
# CURRENT SYSTEM SCREENSHOTS
# ============================================================================
print("📸 Adding screenshots section...")

heading = doc.add_heading('MEVCUT SİSTEM EKRAN GÖRÜNTÜLERİ', 1)
heading.runs[0].font.color.rgb = KOLEKSIYON_BLUE

doc.add_paragraph(
    "Bu bölümde, mevcut Koleksiyon PIM (Product Information Management) sisteminin "
    "ekran görüntüleri yer almaktadır. Bu görseller, SAP entegrasyonu, ürün yönetimi "
    "ve veri yapılarının mevcut durumunu göstermektedir. Bayi Portalı projesinde bu "
    "sistemlerden faydalanılacaktır."
)

doc.add_paragraph()

# Add each screenshot with description
for idx, screenshot in enumerate(screenshots, 1):
    screenshot_path = os.path.join(SCREENSHOTS_DIR, screenshot['file'])

    if os.path.exists(screenshot_path):
        print(f"  Adding screenshot {idx}/{len(screenshots)}: {screenshot['title']}")

        # Screenshot title
        screenshot_heading = doc.add_heading(f"Ekran {idx}: {screenshot['title']}", 2)
        screenshot_heading.runs[0].font.color.rgb = KOLEKSIYON_ORANGE
        screenshot_heading.runs[0].font.size = Pt(13)

        # Description
        desc_para = doc.add_paragraph()
        desc_run = desc_para.add_run(screenshot['description'])
        desc_run.font.italic = True
        desc_run.font.color.rgb = DARK_GRAY
        desc_run.font.size = Pt(10)

        doc.add_paragraph()

        # Add image
        try:
            # Add image with 6 inches width (fits nicely in page)
            doc.add_picture(screenshot_path, width=Inches(6.0))

            # Center the image
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            doc.add_paragraph()

            # Add figure caption
            caption_para = doc.add_paragraph()
            caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption_run = caption_para.add_run(f"Şekil {idx}: {screenshot['title']}")
            caption_run.font.size = Pt(9)
            caption_run.font.italic = True
            caption_run.font.color.rgb = LIGHT_GRAY

            doc.add_paragraph()
            doc.add_paragraph()

        except Exception as e:
            print(f"    ⚠️  Error adding image: {e}")
            error_para = doc.add_paragraph()
            error_run = error_para.add_run(f"[Görsel eklenemedi: {screenshot['file']}]")
            error_run.font.color.rgb = RGBColor(255, 0, 0)
    else:
        print(f"  ⚠️  Screenshot not found: {screenshot_path}")
        error_para = doc.add_paragraph()
        error_run = error_para.add_run(f"[Görsel bulunamadı: {screenshot['file']}]")
        error_run.font.color.rgb = RGBColor(255, 0, 0)

doc.add_page_break()

# ============================================================================
# KEY OBSERVATIONS FROM SCREENSHOTS
# ============================================================================
print("📊 Adding observations section...")

heading = doc.add_heading('EKRAN GÖRÜNTÜLERİNDEN ÇIKARIMLAR', 1)
heading.runs[0].font.color.rgb = KOLEKSIYON_BLUE

doc.add_heading('Mevcut Sistem Özellikleri', 2)

observations = {
    "Ürün Veri Yapısı": [
        "MATNR (Malzeme Numarası) - SAP standart ürün kodu",
        "Model Kodu (örn: MSY900H1) - Koleksiyon özel model kodları",
        "Teklif Kodu (36000010346) - Fiyatlandırma için kullanılan kod",
        "Çok dilli destek (TR/EN) - Açıklama ve malzeme açıklaması alanları",
        "Tasarımcı bilgisi (Studio Kairos, Faruk Malhan)",
        "Tedarikçi bilgisi (Üretim)"
    ],
    "Ürün Kategorileri ve Hiyerarşi": [
        "E-Dükkan: Ürünlerin satış kanalı bilgisi",
        "Z-Kanal: Home, Production gibi ana kategoriler",
        "Z-Ana Grup: Living, Production gibi üst kategori",
        "Z-Alt Grup: Coffee Table, Masa Sistemleri gibi alt kategori",
        "Z-Grup: Detay kategori (coffee table)",
        "Ürün: Milva gibi ürün adı"
    ],
    "Ürün Teknik Bilgileri": [
        "Malzeme: mermer, tekstil, parfüm",
        "Kumaş ve Renk: Renk tanımları",
        "Boyut Bilgileri: Genişlik, Yükseklik, Çap (hem metrik hem ABD)",
        "Ağırlık ve Hacim: Ürün ağırlığı, hacim bilgileri",
        "Beden Bilgileri: Maksimum beden, genişlik, yükseklik",
        "Paket Sayısı: Ambalaj bilgisi",
        "Montaj ve Taşıma Süresi: Lojistik bilgileri"
    ],
    "PIM Panel Özellikleri": [
        "Filtreleme: Liste görünümünde filtre özelliği",
        "Form bazlı veri girişi: Detaylı ürün bilgi formu",
        "Kaydet butonu: Veri kaydetme işlevi",
        "Log Tarih Zaman: Veri güncelleme takibi (14/10/2025 23:30)",
        "Müşteri ID: 050 gibi müşteri tanımları"
    ],
    "Bayi Portalı için Çıkarımlar": [
        "Fiyat bilgisi görsellerde görünmüyor - SAP'den ayrıca çekilecek",
        "Ürün kodları (MATNR, Teklif Kodu) fiyatlandırma için kullanılabilir",
        "Çok dilli destek mevcut - Portal TR/EN destekleyebilir",
        "Ürün kategorileri hiyerarşik - Portal'da filtreleme için kullanılabilir",
        "Teknik özellikler detaylı - PDF spec dosyası için veri hazır",
        "Görseller PIM'de yok - Butterfly CMS'den çekilmeli"
    ]
}

for category, items in observations.items():
    para = doc.add_paragraph()
    run = para.add_run(f"📌 {category}")
    run.font.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = KOLEKSIYON_ORANGE

    for item in items:
        doc.add_paragraph(f"  • {item}", style='List Bullet 2')

doc.add_paragraph()

doc.add_page_break()

# ============================================================================
# DATA MAPPING FOR DEALER PORTAL
# ============================================================================
print("🗺️  Adding data mapping section...")

heading = doc.add_heading('VERİ HARITALAMA: PIM → BAYİ PORTALI', 1)
heading.runs[0].font.color.rgb = KOLEKSIYON_BLUE

doc.add_paragraph(
    "Aşağıdaki tabloda, mevcut PIM sistemindeki alanların Bayi Portalı'nda "
    "nasıl kullanılacağı gösterilmektedir:"
)

doc.add_paragraph()

mapping_table = doc.add_table(rows=1, cols=4)
mapping_table.style = 'Light Grid Accent 1'

headers = mapping_table.rows[0].cells
headers[0].text = "PIM Alanı"
headers[1].text = "Örnek Değer"
headers[2].text = "Portal Kullanımı"
headers[3].text = "Öncelik"

for cell in headers:
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.bold = True
            run.font.size = Pt(9)

mapping_data = [
    ("MATNR", "36000010346", "Ürün benzersiz kimliği", "Yüksek"),
    ("Ürün", "Milva", "Ürün adı gösterimi", "Yüksek"),
    ("Model Kodu", "MSY900H1", "Model bilgisi", "Yüksek"),
    ("Tasarımcı", "Studio Kairos", "Tasarımcı bilgisi", "Orta"),
    ("Açıklama TR", "MILVA S0 MSY900H1 EPOKSI", "Türkçe ürün açıklaması", "Yüksek"),
    ("Açıklama EN", "MILVA S0 MSY900H1 EPOXY", "İngilizce ürün açıklaması", "Yüksek"),
    ("Tedarikçi", "Üretim", "Tedarikçi bilgisi", "Düşük"),
    ("Ana Grup", "Yaşam", "Kategori filtresi", "Yüksek"),
    ("Alt Grup", "Sehpa", "Alt kategori filtresi", "Yüksek"),
    ("Malzeme", "mermer", "Malzeme bilgisi", "Orta"),
    ("Boyut", "900", "Ölçü bilgisi", "Orta"),
    ("Genişlik ABD", "35,43", "ABD ölçü sistemi", "Düşük"),
    ("Ürün Hacmi", "0,20213", "Hacim bilgisi", "Düşük"),
    ("Ürün Ağırlığı", "12.326,84", "Ağırlık bilgisi", "Düşük"),
    ("TVM TR", "Detaylı ürün bilgisi", "Uzun açıklama", "Orta"),
    ("MAKET TR", "MARBLE EPOXY", "Teknik spec", "Orta"),
]

for pim_field, example, portal_use, priority in mapping_data:
    row_cells = mapping_table.add_row().cells
    row_cells[0].text = pim_field
    row_cells[1].text = example
    row_cells[2].text = portal_use
    row_cells[3].text = priority

    # Bold PIM field
    for para in row_cells[0].paragraphs:
        for run in para.runs:
            run.font.name = 'Courier New'
            run.font.size = Pt(9)

    # Color code priority
    if priority == "Yüksek":
        for para in row_cells[3].paragraphs:
            for run in para.runs:
                run.font.color.rgb = RGBColor(255, 0, 0)
                run.font.bold = True
    elif priority == "Orta":
        for para in row_cells[3].paragraphs:
            for run in para.runs:
                run.font.color.rgb = KOLEKSIYON_ORANGE

doc.add_paragraph()

# ============================================================================
# INTEGRATION REQUIREMENTS
# ============================================================================
print("🔗 Adding integration requirements...")

doc.add_page_break()
heading = doc.add_heading('ENTEGRASYON GEREKSİNİMLERİ', 1)
heading.runs[0].font.color.rgb = KOLEKSIYON_BLUE

doc.add_heading('SAP Entegrasyonu', 2)

doc.add_paragraph(
    "Ekran görüntülerinden anlaşıldığı üzere, SAP sisteminden aşağıdaki "
    "verilerin API ile çekilmesi gerekmektedir:"
)

sap_requirements = [
    "MATNR (Malzeme Numarası) - Ürün kodu",
    "Teklif Kodu - Fiyatlandırma için kritik",
    "Fiyat bilgileri (35/36 fiyat kodları, 3000'lu kırılımlar)",
    "Para birimi bazlı fiyatlar (TRY, USD, EUR)",
    "Ürün açıklamaları (TR/EN)",
    "Tedarikçi ve stok bilgileri (opsiyonel)"
]

for req in sap_requirements:
    doc.add_paragraph(f"  ✓ {req}", style='List Bullet')

doc.add_paragraph()

doc.add_heading('PIM/CMS Entegrasyonu', 2)

doc.add_paragraph(
    "Butterfly PIM sisteminden aşağıdaki veriler kullanılacaktır:"
)

pim_requirements = [
    "Ürün kategorileri (Ana Grup, Alt Grup, Grup)",
    "Ürün görselleri (PIM'de görünmüyor, CMS'den alınmalı)",
    "Teknik özellikler (Malzeme, Boyut, Ağırlık, Hacim)",
    "PDF spec dosyaları",
    "Tasarımcı bilgisi",
    "Çok dilli içerik (TR/EN)"
]

for req in pim_requirements:
    doc.add_paragraph(f"  ✓ {req}", style='List Bullet')

doc.add_paragraph()

# ============================================================================
# CONCLUSION
# ============================================================================
print("✅ Adding conclusion...")

doc.add_page_break()
heading = doc.add_heading('SONUÇ VE ÖNERİLER', 1)
heading.runs[0].font.color.rgb = KOLEKSIYON_BLUE

doc.add_paragraph(
    "Mevcut sistem ekran görüntüleri analiz edildiğinde, Koleksiyon'un güçlü bir "
    "ürün yönetim altyapısına sahip olduğu görülmektedir. PIM sistemi detaylı "
    "ürün bilgileri, kategori yapısı ve çok dilli destek sunmaktadır."
)

doc.add_paragraph(
    "Bayi Portalı projesi için öneriler:"
)

recommendations = [
    "SAP API bağlantısı için MATNR ve Teklif Kodu kullanılmalıdır",
    "Fiyat bilgileri SAP'den gerçek zamanlı çekilmeli, 15 dakika cache'lenmelidir",
    "Ürün kategorileri (Ana Grup, Alt Grup) filtreleme için kullanılmalıdır",
    "Çok dilli destek (TR/EN) mutlaka implement edilmelidir",
    "Ürün görselleri Butterfly CMS'den alınmalı, CDN üzerinden sunulmalıdır",
    "PDF spec dosyaları download özelliği eklenmelidir",
    "Tasarımcı bilgisi ürün detay sayfasında gösterilmelidir",
    "Teknik özellikler (Boyut, Ağırlık, Malzeme) opsiyonel olarak gösterilebilir",
    "Bayi bazlı para birimi yetkilendirmesi critical öneme sahiptir",
    "Müşteri ID bazlı erişim kontrolü yapılmalıdır"
]

for idx, rec in enumerate(recommendations, 1):
    para = doc.add_paragraph(f"{idx}. {rec}")
    para.style = 'List Number'

doc.add_paragraph()

# ============================================================================
# DOCUMENT CLOSE
# ============================================================================

doc.add_page_break()

# Final notes
final_para = doc.add_paragraph()
final_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
final_run = final_para.add_run("─" * 60)
final_run.font.color.rgb = LIGHT_GRAY

doc.add_paragraph()

closing_para = doc.add_paragraph()
closing_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
closing_run = closing_para.add_run(
    "Bu doküman, Koleksiyon Mobilya Bayi Portalı projesinin\n"
    "iş analizi ve kapsam tanımını içermektedir.\n"
    "Mevcut sistem ekran görüntüleri ile güncellenmiştir.\n\n"
    "Sorularınız ve geri bildirimleriniz için proje ekibi ile iletişime geçebilirsiniz."
)
closing_run.font.size = Pt(10)
closing_run.font.color.rgb = DARK_GRAY

doc.add_paragraph()

footer_para = doc.add_paragraph()
footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_run = footer_para.add_run(f"© {datetime.now().year} Koleksiyon Mobilya - Gizli ve Özel")
footer_run.font.size = Pt(9)
footer_run.font.color.rgb = LIGHT_GRAY

# Save document
output_file = "/Users/caglarozyildirim/WebstormProjects/Deneme/Koleksiyon_Mobilya_Bayi_Portali_Kapsam_Dokumani.docx"
doc.save(output_file)

print("\n" + "="*80)
print("✅ DOCUMENT UPDATED SUCCESSFULLY WITH SCREENSHOTS!")
print("="*80)
print(f"📁 File: {output_file}")
print(f"📸 Screenshots added: {len(screenshots)}")
print(f"📊 Data mapping entries: {len(mapping_data)}")
print(f"💡 Recommendations: {len(recommendations)}")
print("="*80 + "\n")