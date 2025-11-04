#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
MAN Turkey Maintenance Management System - Ultimate Comprehensive Analysis Document Generator
This script creates a complete, developer-ready analysis document with all data structures,
screenshots, workflow diagrams, and technical specifications.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
import os

# MAN Red Corporate Color
MAN_RED = RGBColor(226, 7, 20)
MAN_DARK_GRAY = RGBColor(26, 26, 26)
MAN_LIGHT_GRAY = RGBColor(245, 245, 245)

class MANDocumentGenerator:
    def __init__(self, output_path, screenshots_path, diagrams_path):
        self.doc = Document()
        self.output_path = output_path
        self.screenshots_path = screenshots_path
        self.diagrams_path = diagrams_path
        self.setup_document_properties()
        self.setup_styles()

    def setup_document_properties(self):
        """Set up document properties"""
        core_properties = self.doc.core_properties
        core_properties.title = "MAN Türkiye Bakım Yönetimi Sistemi - Kapsamlı İş Analizi"
        core_properties.subject = "Maintenance Management System Business Analysis"
        core_properties.author = "MAN Türkiye IT Department"
        core_properties.keywords = "MAN, Bakım Yönetimi, İş Analizi, CMS, SAP Integration"
        core_properties.comments = "Complete developer-ready analysis document with all data structures"
        core_properties.category = "Business Analysis"

        # Set page margins (narrow)
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(2)
            section.right_margin = Cm(2)

    def setup_styles(self):
        """Create custom styles for the document"""
        styles = self.doc.styles

        # Heading 1 - MAN Red
        heading1 = styles['Heading 1']
        heading1_font = heading1.font
        heading1_font.name = 'Arial'
        heading1_font.size = Pt(20)
        heading1_font.bold = True
        heading1_font.color.rgb = MAN_RED

        # Heading 2
        heading2 = styles['Heading 2']
        heading2_font = heading2.font
        heading2_font.name = 'Arial'
        heading2_font.size = Pt(16)
        heading2_font.bold = True
        heading2_font.color.rgb = MAN_DARK_GRAY

        # Heading 3
        heading3 = styles['Heading 3']
        heading3_font = heading3.font
        heading3_font.name = 'Arial'
        heading3_font.size = Pt(14)
        heading3_font.bold = True
        heading3_font.color.rgb = MAN_DARK_GRAY

        # Normal text
        normal = styles['Normal']
        normal_font = normal.font
        normal_font.name = 'Arial'
        normal_font.size = Pt(11)

    def add_page_break(self):
        """Add a page break"""
        self.doc.add_page_break()

    def add_title_page(self):
        """Create a professional title page with MAN branding"""
        # Add some vertical space
        for _ in range(3):
            self.doc.add_paragraph()

        # Main Title
        title = self.doc.add_paragraph()
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = title.add_run("MAN TÜRKİYE")
        run.font.name = 'Arial Black'
        run.font.size = Pt(36)
        run.font.bold = True
        run.font.color.rgb = MAN_RED

        # Subtitle
        subtitle = self.doc.add_paragraph()
        subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = subtitle.add_run("BAKIM YÖNETİMİ SİSTEMİ")
        run.font.name = 'Arial'
        run.font.size = Pt(28)
        run.font.bold = True
        run.font.color.rgb = MAN_DARK_GRAY

        # Document Type
        doc_type = self.doc.add_paragraph()
        doc_type.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = doc_type.add_run("KAPSAMLI İŞ ANALİZİ VE TEKNİK DOKÜMANTASYON")
        run.font.name = 'Arial'
        run.font.size = Pt(18)
        run.font.color.rgb = MAN_DARK_GRAY

        # Add space
        for _ in range(2):
            self.doc.add_paragraph()

        # Document Info Box
        info_table = self.doc.add_table(rows=5, cols=2)
        info_table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        labels = ["Doküman Tipi:", "Versiyon:", "Tarih:", "Hazırlayan:", "Durum:"]
        values = ["İş Analizi & Teknik Spesifikasyon",
                  "v3.0 FINAL",
                  datetime.now().strftime("%d.%m.%Y"),
                  "MAN Türkiye IT & Teknik Ekip",
                  "CMS Panel Geliştirme İçin Hazır"]

        for i, (label, value) in enumerate(zip(labels, values)):
            row = info_table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = value

            # Style cells
            for cell in row.cells:
                cell.paragraphs[0].runs[0].font.name = 'Arial'
                cell.paragraphs[0].runs[0].font.size = Pt(11)
                if cell == row.cells[0]:
                    cell.paragraphs[0].runs[0].font.bold = True

        # Add more space
        for _ in range(8):
            self.doc.add_paragraph()

        # Footer note
        footer = self.doc.add_paragraph()
        footer.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = footer.add_run("Bu doküman CMS panel geliştiricileri için hazırlanmıştır.\n")
        run.font.name = 'Arial'
        run.font.size = Pt(10)
        run.font.italic = True
        run.font.color.rgb = RGBColor(102, 102, 102)

        run = footer.add_run("Tüm veri yapıları, ekran görüntüleri ve akış diyagramları dahildir.")
        run.font.name = 'Arial'
        run.font.size = Pt(10)
        run.font.italic = True
        run.font.color.rgb = RGBColor(102, 102, 102)

        self.add_page_break()

    def add_table_of_contents(self):
        """Add table of contents"""
        heading = self.doc.add_heading('İÇİNDEKİLER', 1)

        toc_items = [
            ("1. YÖNETİCİ ÖZETİ", 3),
            ("2. AMAÇ VE HEDEFLER", 3),
            ("3. KAPSAM", 4),
            ("   3.1 Kapsam Dahili", 4),
            ("   3.2 Kapsam Dışı", 5),
            ("4. MODÜL DETAYLARI VE VERİ YAPILARI", 6),
            ("   4.1 Dashboard (Ana Sayfa)", 6),
            ("   4.2 İş Talepleri Modülü", 10),
            ("   4.3 Varlık Yönetimi Modülü", 25),
            ("   4.4 Bakım Yönetimi Modülü", 35),
            ("   4.5 Olay Yönetimi Modülü", 45),
            ("5. İŞ AKIŞI DİYAGRAMLARI", 52),
            ("6. SAP ENTEGRASYON GEREKSİNİMLERİ", 56),
            ("7. YETKİLENDİRME VE ROLLER", 60),
            ("8. TEST SENARYOLARI", 65),
            ("9. TEKNİK SPESİFİKASYONLAR", 70),
            ("10. EKLER", 75),
        ]

        for item, page in toc_items:
            p = self.doc.add_paragraph(style='List Number')
            p.add_run(item).font.name = 'Arial'
            p.add_run('\t' + str(page)).font.name = 'Arial'

        self.add_page_break()

    def add_management_summary(self):
        """Add management summary in Turkish"""
        self.doc.add_heading('1. YÖNETİCİ ÖZETİ', 1)

        content = """MAN Türkiye Bakım Yönetimi Sistemi, üretim tesislerindeki tüm bakım operasyonlarını dijitalleştirmek ve optimize etmek için tasarlanmış kapsamlı bir web tabanlı çözümdür. Bu sistem, iş taleplerinden varlık yönetimine, planlı bakımlardan acil olay yönetimine kadar tüm süreçleri tek bir platformda birleştirir.

Sistem, SAP ERP sistemi ile tam entegre çalışarak varlık verilerini, maliyet merkezlerini ve finansal bilgileri otomatik olarak senkronize eder. Kullanıcı dostu arayüzü ve MAN kurumsal kimliğine uygun tasarımı ile tüm çalışanlar tarafından kolayca benimsenebilir.

Temel Özellikler:
• İş Talebi Yönetimi: Çok aşamalı onay süreçleri ile talep yönetimi
• Varlık Yönetimi: SAP entegreli ekipman ve varlık takibi
• Bakım Yönetimi: Periyodik, önleyici ve düzeltici bakım planlaması
• Olay Yönetimi: Acil durumlar için SLA bazlı müdahale sistemi
• Raporlama: Detaylı analiz ve raporlama araçları

Hedef Kullanıcılar:
• Üretim Çalışanları: İş talebi oluşturma ve takip
• Bakım Teknisyenleri: İş emirlerini görüntüleme ve tamamlama
• SL Mühendisler: Teknik onay ve maliyet hesaplama
• Yöneticiler: İş onayı ve performans takibi
• IT Ekibi: Sistem yönetimi ve kullanıcı yönetimi

Beklenen Faydalar:
• %40 daha hızlı iş talebi süreçleri
• %30 azalma plansız duruş süresinde
• Tam dijital kayıt ve izlenebilirlik
• SAP ile gerçek zamanlı veri senkronizasyonu
• Mobil cihazlardan erişim imkanı"""

        p = self.doc.add_paragraph(content)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        self.add_page_break()

    def add_purpose_and_objectives(self):
        """Add purpose and objectives section"""
        self.doc.add_heading('2. AMAÇ VE HEDEFLER', 1)

        # Amaç
        self.doc.add_heading('2.1 Sistem Amacı', 2)
        purpose = """Bu sistem, MAN Türkiye üretim tesislerindeki bakım operasyonlarını dijitalleştirmek,
optimize etmek ve SAP ERP sistemi ile entegre bir şekilde yönetmek amacıyla geliştirilmektedir.
Sistem, manuel süreçleri otomatize ederek operasyonel verimliliği artırmayı ve maliyetleri
düşürmeyi hedeflemektedir."""
        self.doc.add_paragraph(purpose).alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        # Hedefler
        self.doc.add_heading('2.2 Temel Hedefler', 2)
        objectives = [
            "İş taleplerinin dijital ortamda oluşturulması ve çok aşamalı onay süreçlerinin yönetimi",
            "Tüm üretim varlıklarının SAP ile senkronize takibi ve dokümantasyonu",
            "Periyodik, önleyici ve düzeltici bakım planlarının sistematik yönetimi",
            "Acil olayların anında bildirimi ve SLA bazlı müdahale yönetimi",
            "Maliyet takibi ve raporlama ile finansal şeffaflık sağlama",
            "Mobil cihazlardan erişim ile saha çalışanlarının verimliliğini artırma",
            "Veri analitiği ile öngörülü bakım (predictive maintenance) altyapısı oluşturma",
            "Tam dijital kayıt ile audit ve compliance gereksinimlerini karşılama"
        ]

        for obj in objectives:
            p = self.doc.add_paragraph(obj, style='List Bullet')
            p.runs[0].font.size = Pt(11)

        self.add_page_break()

    def add_scope(self):
        """Add scope section"""
        self.doc.add_heading('3. KAPSAM', 1)

        # In Scope
        self.doc.add_heading('3.1 Kapsam Dahili', 2)
        in_scope = [
            "İş Talebi Yönetimi: Talep oluşturma, onay süreçleri, atama ve tamamlama",
            "Varlık Yönetimi: SAP entegrasyonu, varlık kayıt, takip ve dokümantasyon",
            "Bakım Yönetimi: Periyodik, önleyici, düzeltici, ölçüm bazlı ve toplu bakım planlaması",
            "Olay Yönetimi: Ekipman arızası, güvenlik, kalite ve çevre olayları yönetimi",
            "Kullanıcı Yönetimi: Rol bazlı yetkilendirme ve Active Directory entegrasyonu",
            "Raporlama: Dashboard, detaylı raporlar ve veri dışa aktarma",
            "SAP Entegrasyonu: Varlık verileri, maliyet merkezi ve finansal entegrasyon",
            "Bildirimler: E-posta, SMS ve sistem içi bildirimler",
            "Dokümantasyon: Dosya yükleme, versiyonlama ve arşivleme",
            "Mobil Uyumluluk: Responsive tasarım ve mobil erişim"
        ]

        for item in in_scope:
            p = self.doc.add_paragraph(item, style='List Bullet')
            p.runs[0].font.size = Pt(11)

        # Out of Scope
        self.doc.add_heading('3.2 Kapsam Dışı', 2)
        out_scope = [
            "Üretim Planlama (MES sistemi ile ilgili işlevler)",
            "Tedarik Zinciri Yönetimi (Procurement ve lojistik)",
            "İnsan Kaynakları Yönetimi (Personel yönetimi)",
            "Finansal Muhasebe (SAP FI/CO modülü işlevleri)",
            "Satış ve Müşteri İlişkileri Yönetimi (CRM)",
            "Yedek Parça Stoğu Yönetimi (SAP MM ile entegre ama bu sistemin dışında)",
            "IoT Sensör Entegrasyonu (Gelecek fazda planlanmaktadır)",
            "Yapay Zeka ve Makine Öğrenmesi (Predictive analytics için gelecek faz)"
        ]

        for item in out_scope:
            p = self.doc.add_paragraph(item, style='List Bullet')
            p.runs[0].font.size = Pt(11)

        self.add_page_break()

    def add_screenshot(self, filename, title, width=Inches(6.5)):
        """Add a screenshot with title"""
        self.doc.add_heading(title, 3)
        screenshot_path = os.path.join(self.screenshots_path, filename)
        if os.path.exists(screenshot_path):
            try:
                self.doc.add_picture(screenshot_path, width=width)
                last_paragraph = self.doc.paragraphs[-1]
                last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            except Exception as e:
                self.doc.add_paragraph(f"[Ekran görüntüsü yüklenemedi: {filename}]")
        else:
            self.doc.add_paragraph(f"[Ekran görüntüsü bulunamadı: {filename}]")

    def add_data_table(self, headers, data, column_widths=None):
        """Add a formatted data table"""
        table = self.doc.add_table(rows=1, cols=len(headers))
        table.style = 'Light Grid Accent 1'

        # Header row
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = header
            # Make header bold and colored
            for paragraph in header_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.color.rgb = MAN_RED
                    run.font.size = Pt(10)

        # Data rows
        for row_data in data:
            row_cells = table.add_row().cells
            for i, cell_data in enumerate(row_data):
                row_cells[i].text = str(cell_data)
                for paragraph in row_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)

        # Set column widths if provided
        if column_widths:
            for row in table.rows:
                for idx, width in enumerate(column_widths):
                    row.cells[idx].width = Inches(width)

        return table

    def add_dashboard_module(self):
        """Add Dashboard module documentation"""
        self.doc.add_heading('4. MODÜL DETAYLARI VE VERİ YAPILARI', 1)
        self.doc.add_heading('4.1 Dashboard (Ana Sayfa)', 2)

        # Module Description
        desc = """Dashboard, sistemin ana sayfasıdır ve kullanıcıya sistemin genel durumu hakkında özet bilgiler sunar.
Real-time veriler kullanılarak güncel durum görüntülenir. Rol bazlı widget'lar sayesinde her kullanıcı kendi
yetki seviyesine göre ilgili bilgileri görüntüler."""
        self.doc.add_paragraph(desc).alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        # Screenshot
        self.add_screenshot('01_Ana_Sayfa_Dashboard.png', 'Dashboard Ekran Görüntüsü')

        # Widgets
        self.doc.add_heading('Dashboard Widget\'ları', 3)

        widget_data = [
            ["Widget Adı", "Açıklama", "Veri Kaynağı"],
            ["Toplam İş Talepleri", "Bu aydaki toplam talep sayısı ve yeni talep sayısı", "JobRequests tablosu"],
            ["Bekleyen Onaylar", "Onay bekleyen iş talepleri, acil onay sayısı", "JobRequests (status=approval)"],
            ["Tamamlanan İşler", "Bu ay tamamlanan iş sayısı", "JobRequests (status=done)"],
            ["Aktif Olaylar", "Açık kritik ve yüksek öncelikli olaylar", "Incidents (status=open)"],
            ["Üretim Ekipmanı", "Toplam ekipman, aktif ve bakımdaki sayılar", "Assets tablosu"],
            ["Planlı Bakımlar", "Bu hafta yapılacak bakım sayısı", "MaintenancePlans (next 7 days)"],
        ]

        self.add_data_table(widget_data[0], widget_data[1:], [2.0, 3.0, 2.5])

        # Widget Data Fields
        self.doc.add_heading('Widget Veri Alanları Detayı', 3)

        widget_fields = [
            ["Widget", "Alan Adı", "Veri Tipi", "Hesaplama/Sorgu"],
            ["Toplam İş Talepleri", "totalRequests", "Integer", "COUNT(*) FROM JobRequests"],
            ["", "newThisMonth", "Integer", "COUNT(*) WHERE MONTH(createdDate) = current_month"],
            ["Bekleyen Onaylar", "pendingApprovals", "Integer", "COUNT(*) WHERE status='approval'"],
            ["", "urgentCount", "Integer", "COUNT(*) WHERE status='approval' AND priority='urgent'"],
            ["Tamamlanan İşler", "completedTotal", "Integer", "COUNT(*) WHERE status='done'"],
            ["", "completedThisMonth", "Integer", "COUNT(*) WHERE status='done' AND MONTH(completedDate)"],
            ["Aktif Olaylar", "activeIncidents", "Integer", "COUNT(*) WHERE status='open' OR status='investigating'"],
            ["", "criticalCount", "Integer", "COUNT(*) WHERE severity='critical'"],
            ["", "highCount", "Integer", "COUNT(*) WHERE severity='high'"],
            ["Üretim Ekipmanı", "totalAssets", "Integer", "COUNT(*) FROM Assets"],
            ["", "activeAssets", "Integer", "COUNT(*) WHERE status='active'"],
            ["", "inMaintenance", "Integer", "COUNT(*) WHERE status='maintenance'"],
            ["Planlı Bakımlar", "thisWeek", "Integer", "COUNT(*) WHERE scheduledDate BETWEEN today AND today+7"],
        ]

        self.add_data_table(widget_fields[0], widget_fields[1:], [1.5, 1.5, 1.0, 3.5])

        self.add_page_break()

    def add_job_requests_module(self):
        """Add Job Requests module documentation"""
        self.doc.add_heading('4.2 İş Talepleri (Job Requests) Modülü', 2)

        # Module Description
        desc = """İş Talepleri modülü, çalışanların bakım, onarım veya iyileştirme taleplerini sisteme girmelerini
ve bu taleplerin çok aşamalı onay süreçlerinden geçmesini sağlar. SAP ile entegre çalışarak maliyet
merkezleri ve varlık bilgilerini otomatik olarak getirir."""
        self.doc.add_paragraph(desc).alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        # List Page
        self.doc.add_heading('4.2.1 İş Talepleri Liste Sayfası', 3)
        self.add_screenshot('02_Is_Talepleri_Liste.png', 'Liste Sayfası Ekran Görüntüsü')

        # List Page Data Structure
        self.doc.add_heading('Liste Sayfası Veri Yapısı', 4)

        list_columns = [
            ["Sütun Adı", "Alan Adı (API)", "Veri Tipi", "Açıklama"],
            ["Talep No", "requestId", "String", "JR-YYYY-XXX formatında unique ID"],
            ["Başlık", "title", "String", "İş talebinin başlığı (max 200 karakter)"],
            ["Kategori", "category", "Enum", "HVAC, Elektrik, Mekanik, Bina, IT"],
            ["Lokasyon", "location", "String", "Ana lokasyon bilgisi"],
            ["Öncelik", "priority", "Enum", "Kritik, Yüksek, Orta, Düşük"],
            ["Durum", "status", "Enum", "Beklemede, Onayda, İşlemde, Tamamlandı, Reddedildi"],
            ["Oluşturan", "createdBy", "String", "Talebi oluşturan kullanıcı adı"],
            ["Tarih", "createdDate", "DateTime", "Talebin oluşturulma tarihi"],
        ]

        self.add_data_table(list_columns[0], list_columns[1:], [1.3, 1.3, 1.0, 3.9])

        # List Page Filters
        self.doc.add_heading('Filtre Alanları', 4)

        filter_data = [
            ["Filtre Adı", "Alan Adı", "Tip", "Değerler"],
            ["Arama", "search", "Text", "Talep No, Başlık, Açıklama içinde arama"],
            ["Durum Filtresi", "statusFilter", "Dropdown", "Tümü, Beklemede, Onayda, İşlemde, Tamamlandı, Reddedildi"],
            ["Öncelik Filtresi", "priorityFilter", "Dropdown", "Tümü, Kritik, Yüksek, Orta, Düşük"],
            ["Kategori Filtresi", "categoryFilter", "Dropdown", "Tümü, HVAC, Elektrik, Mekanik, Bina, IT"],
            ["Lokasyon Filtresi", "locationFilter", "Dropdown", "Tümü, Ana Tesis, Üretim, Ofis, Depo"],
            ["Tarih Aralığı", "dateRange", "DateRange", "Başlangıç ve bitiş tarihi seçimi"],
        ]

        self.add_data_table(filter_data[0], filter_data[1:], [1.5, 1.3, 1.0, 3.7])

        self.add_page_break()

        # Detail Page
        self.doc.add_heading('4.2.2 İş Talebi Detay Sayfası', 3)
        self.add_screenshot('03_Is_Talebi_Detay.png', 'Detay Sayfası Ekran Görüntüsü')

        # Detail Page Data Structure (36 fields)
        self.doc.add_heading('Detay Sayfası Tam Veri Yapısı (36 Alan)', 4)

        detail_fields_1 = [
            ["#", "Alan Adı", "Veri Tipi", "Zorunlu", "Açıklama"],
            ["1", "requestId", "String", "Evet", "Unique talep numarası (JR-YYYY-XXX)"],
            ["2", "title", "String(200)", "Evet", "İş talebi başlığı"],
            ["3", "description", "Text", "Evet", "Detaylı açıklama"],
            ["4", "category", "Enum", "Evet", "HVAC, Electrical, Mechanical, Plumbing, Building"],
            ["5", "priority", "Enum", "Evet", "Urgent, High, Normal, Low"],
            ["6", "status", "Enum", "Evet", "Waiting, Business-Approval, Technical-Approval, etc."],
            ["7", "requestReason", "Enum", "Evet", "OHS, Energy, Environment, Improvement, Investment"],
            ["8", "location", "String(100)", "Evet", "Ana lokasyon"],
            ["9", "subLocation1", "String(100)", "Hayır", "Alt lokasyon 1"],
            ["10", "subLocation2", "String(100)", "Hayır", "Alt lokasyon 2"],
            ["11", "assetSapId", "String(20)", "Hayır", "SAP varlık numarası"],
            ["12", "assetName", "String(200)", "Hayır", "Varlık adı"],
            ["13", "assetId", "Reference", "Hayır", "Assets tablosu referansı"],
            ["14", "requestedDate", "DateTime", "Evet", "Talep oluşturulma tarihi"],
            ["15", "desiredCompletionDate", "Date", "Hayır", "İstenen tamamlanma tarihi"],
        ]

        self.add_data_table(detail_fields_1[0], detail_fields_1[1:], [0.4, 1.5, 1.0, 0.8, 3.8])

        detail_fields_2 = [
            ["#", "Alan Adı", "Veri Tipi", "Zorunlu", "Açıklama"],
            ["16", "requestedBy", "String(100)", "Evet", "Talebi oluşturan kullanıcı"],
            ["17", "requestedByEmail", "String(100)", "Evet", "Oluşturan kullanıcı e-posta"],
            ["18", "requestedByManager", "String(100)", "Evet", "Talep edenin yöneticisi (iş onayı için)"],
            ["19", "department", "String(100)", "Evet", "Departman bilgisi"],
            ["20", "costCenter", "String(50)", "Evet", "SAP maliyet merkezi kodu"],
            ["21", "estimatedCost", "Decimal(10,2)", "Hayır", "Tahmini maliyet"],
            ["22", "actualCost", "Decimal(10,2)", "Hayır", "Gerçekleşen maliyet"],
            ["23", "currency", "Enum", "Evet", "TRY, EUR, USD (default: TRY)"],
            ["24", "slEngineer", "String(100)", "Hayır", "SL mühendis (teknik onay)"],
            ["25", "technicalApprovalDate", "DateTime", "Hayır", "Teknik onay tarihi"],
            ["26", "businessManager", "String(100)", "Hayır", "İş yöneticisi (maliyet onayı)"],
            ["27", "costApprovalDate", "DateTime", "Hayır", "Maliyet onay tarihi"],
            ["28", "solutionResponsible", "String(100)", "Hayır", "Çözüm sorumlusu"],
            ["29", "assignedTeam", "String(100)", "Hayır", "Atanan ekip"],
            ["30", "implementationStartDate", "DateTime", "Hayır", "Uygulama başlangıç tarihi"],
        ]

        self.add_data_table(detail_fields_2[0], detail_fields_2[1:], [0.4, 1.5, 1.0, 0.8, 3.8])

        detail_fields_3 = [
            ["#", "Alan Adı", "Veri Tipi", "Zorunlu", "Açıklama"],
            ["31", "completionDate", "DateTime", "Hayır", "Tamamlanma tarihi"],
            ["32", "rejectionReason", "Text", "Hayır", "Red nedeni (eğer reddedilmişse)"],
            ["33", "workOrder", "String(50)", "Hayır", "SAP iş emri numarası"],
            ["34", "attachments", "Array", "Hayır", "Ek dosyalar (JSON array)"],
            ["35", "comments", "Array", "Hayır", "Yorumlar ve notlar (JSON array)"],
            ["36", "activityLog", "Array", "Hayır", "Aktivite geçmişi (JSON array)"],
        ]

        self.add_data_table(detail_fields_3[0], detail_fields_3[1:], [0.4, 1.5, 1.0, 0.8, 3.8])

        self.add_page_break()

        # Create Form
        self.doc.add_heading('4.2.3 Yeni İş Talebi Oluşturma Formu', 3)
        self.add_screenshot('08_Modal_Yeni_Is_Talebi_Fixed.png', 'Yeni İş Talebi Formu')

        # Create Form Data Structure
        self.doc.add_heading('Form Alanları Veri Yapısı', 4)

        create_form_data = [
            ["Alan Adı", "Form Label (TR)", "Tip", "Zorunlu", "Validasyon", "Enum Değerleri"],
            ["title", "Talep Başlığı", "Text", "Evet", "Max 200 karakter", "-"],
            ["category", "Kategori", "Dropdown", "Evet", "Enum seçimi", "HVAC, Electrical, Mechanical, Plumbing, Building"],
            ["priority", "Öncelik", "Dropdown", "Evet", "Enum seçimi", "Urgent (Acil), High (Yüksek), Normal, Low (Düşük)"],
            ["requestReason", "Talep Nedeni", "Dropdown", "Evet", "Enum seçimi", "OHS (İSG), Energy (Enerji), Environment (Çevre), Improvement, Investment, Renovation"],
            ["location", "Lokasyon", "Dropdown", "Evet", "Master data", "Paint, Assembly, Chassis, Warehouse, Office"],
            ["subLocation1", "Alt Lokasyon 1", "Text", "Hayır", "Max 100 kar.", "-"],
            ["subLocation2", "Alt Lokasyon 2", "Text", "Hayır", "Max 100 kar.", "-"],
            ["description", "Açıklama", "Textarea", "Evet", "Max 2000 kar.", "-"],
            ["assetId", "İlişkili Varlık", "Autocomplete", "Hayır", "Asset seçimi", "Assets tablosundan"],
            ["desiredCompletionDate", "İstenen Tamamlanma Tarihi", "Date", "Hayır", "Bugünden sonra", "-"],
        ]

        self.add_data_table(create_form_data[0], create_form_data[1:], [1.2, 1.3, 0.8, 0.7, 1.0, 2.5])

        # Workflow States
        self.doc.add_heading('İş Talebi İş Akışı Durumları', 4)

        workflow_states = [
            ["Durum Kodu", "Durum Adı (TR)", "Açıklama", "Sonraki Durum"],
            ["waiting", "Beklemede", "Talep oluşturuldu, işlem bekliyor", "business-approval"],
            ["business-approval", "İş Onayı", "Yönetici onayı bekleniyor", "technical-approval / rejected"],
            ["technical-approval", "Teknik Onay", "SL mühendis teknik inceleme yapıyor", "cost-calculation / rejected"],
            ["cost-calculation", "Maliyet Hesaplama", "Maliyet birimi hesaplama yapıyor", "cost-approval"],
            ["cost-approval", "Maliyet Onayı", "İş yöneticisi maliyet onayı veriyor", "solution-assignment / rejected"],
            ["solution-assignment", "Çözüm Sorumlusu Atama", "Bakım ekibine atanıyor", "implementation"],
            ["implementation", "Uygulama", "İş uygulanıyor", "done"],
            ["done", "Tamamlandı", "İş başarıyla tamamlandı", "-"],
            ["rejected", "Reddedildi", "Talep reddedildi", "-"],
        ]

        self.add_data_table(workflow_states[0], workflow_states[1:], [1.5, 1.3, 2.5, 2.2])

        self.add_page_break()

    def add_assets_module(self):
        """Add Assets module documentation"""
        self.doc.add_heading('4.3 Varlık Yönetimi (Asset Management) Modülü', 2)

        # Module Description
        desc = """Varlık Yönetimi modülü, üretim tesislerindeki tüm ekipman ve varlıkların kaydını tutar.
SAP PM modülü ile entegre çalışarak varlık bilgilerini senkronize eder. Her varlık için bakım geçmişi,
maliyet bilgileri ve dokümantasyon saklanır."""
        self.doc.add_paragraph(desc).alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        # List Page
        self.doc.add_heading('4.3.1 Varlık Yönetimi Liste Sayfası', 3)
        self.add_screenshot('04_Varlik_Yonetimi_Liste.png', 'Liste Sayfası Ekran Görüntüsü')

        # List Page Data Structure
        self.doc.add_heading('Liste Sayfası Veri Yapısı', 4)

        asset_list_columns = [
            ["Sütun Adı", "Alan Adı (API)", "Veri Tipi", "Açıklama"],
            ["SAP ID", "sapAssetId", "String(20)", "SAP sistemindeki varlık numarası"],
            ["Varlık Adı", "assetName", "String(200)", "Varlığın adı"],
            ["Kategori", "category", "Enum", "Üretim Makinesi, Robotik, Boya, Taşıma, Altyapı"],
            ["Lokasyon", "location", "String(100)", "Ana lokasyon"],
            ["Durum", "status", "Enum", "Aktif, Pasif, Bakımda, Arızalı"],
            ["Zimmetli", "assignedTo", "String(100)", "Varlıktan sorumlu kişi/ekip"],
            ["Son Bakım", "lastMaintenanceDate", "Date", "Son yapılan bakım tarihi"],
        ]

        self.add_data_table(asset_list_columns[0], asset_list_columns[1:], [1.3, 1.3, 1.0, 3.9])

        # List Page Filters
        self.doc.add_heading('Filtre Alanları', 4)

        asset_filter_data = [
            ["Filtre Adı", "Alan Adı", "Tip", "Değerler"],
            ["Arama", "search", "Text", "SAP ID, Varlık Adı, Seri No içinde arama"],
            ["Durum Filtresi", "statusFilter", "Dropdown", "Tümü, Aktif, Pasif, Bakımda, Arızalı"],
            ["Kategori Filtresi", "categoryFilter", "Dropdown", "Tümü, Üretim Makinesi, Robotik, Boya, Taşıma, Altyapı"],
            ["Lokasyon Filtresi", "locationFilter", "Dropdown", "Master data'dan dinamik liste"],
        ]

        self.add_data_table(asset_filter_data[0], asset_filter_data[1:], [1.5, 1.3, 1.0, 3.7])

        self.add_page_break()

        # Detail Page
        self.doc.add_heading('4.3.2 Varlık Detay Sayfası', 3)
        self.add_screenshot('05_Varlik_Detay_Fixed.png', 'Detay Sayfası (Fixed - Olay Bildir butonu kaldırıldı)')

        # Detail Page Data Structure (28 fields)
        self.doc.add_heading('Detay Sayfası Tam Veri Yapısı (28 Alan)', 4)

        asset_detail_1 = [
            ["#", "Alan Adı", "Veri Tipi", "Zorunlu", "Açıklama"],
            ["1", "assetId", "String(50)", "Evet", "Internal unique ID (AST-XXXX)"],
            ["2", "sapAssetId", "String(20)", "Hayır", "SAP PM varlık numarası"],
            ["3", "assetName", "String(200)", "Evet", "Varlık adı"],
            ["4", "assetType", "Enum", "Evet", "Hand tools, Electric, Construction, Tool-Counter, Mechanic, Office, Meeting room, Other"],
            ["5", "category", "Enum", "Evet", "Üretim Ekipmanı, Bakım Ekipmanı, Kalite Kontrol, Güvenlik"],
            ["6", "status", "Enum", "Evet", "Active, Inactive, Maintenance, Faulty"],
            ["7", "manufacturer", "String(100)", "Hayır", "Üretici firma"],
            ["8", "model", "String(100)", "Hayır", "Model bilgisi"],
            ["9", "serialNumber", "String(100)", "Hayır", "Seri numarası"],
            ["10", "location", "String(100)", "Evet", "Ana lokasyon"],
            ["11", "subLocation1", "String(100)", "Hayır", "Alt lokasyon 1"],
            ["12", "subLocation2", "String(100)", "Hayır", "Alt lokasyon 2"],
            ["13", "purchaseDate", "Date", "Hayır", "Satın alma tarihi"],
            ["14", "purchaseValue", "Decimal(12,2)", "Hayır", "Satın alma değeri"],
        ]

        self.add_data_table(asset_detail_1[0], asset_detail_1[1:], [0.4, 1.3, 1.0, 0.8, 4.0])

        asset_detail_2 = [
            ["#", "Alan Adı", "Veri Tipi", "Zorunlu", "Açıklama"],
            ["15", "currency", "Enum", "Hayır", "TRY, EUR, USD"],
            ["16", "bookValue", "Decimal(12,2)", "Hayır", "Güncel defter değeri (SAP'den)"],
            ["17", "costCenter", "String(50)", "Evet", "Maliyet merkezi kodu"],
            ["18", "assignedTo", "String(100)", "Hayır", "Zimmetli kişi veya ekip"],
            ["19", "description", "Text", "Hayır", "Varlık açıklaması"],
            ["20", "technicalSpecifications", "JSON", "Hayır", "Teknik özellikler (JSON)"],
            ["21", "warrantyExpiry", "Date", "Hayır", "Garanti bitiş tarihi"],
            ["22", "lastMaintenanceDate", "Date", "Hayır", "Son bakım tarihi"],
            ["23", "nextMaintenanceDate", "Date", "Hayır", "Sonraki bakım tarihi"],
            ["24", "maintenanceHistory", "Array", "Hayır", "Bakım geçmişi (JSON array)"],
            ["25", "documents", "Array", "Hayır", "Dökümanlar (JSON array)"],
            ["26", "qrCode", "String(100)", "Hayır", "QR kod değeri"],
            ["27", "createdDate", "DateTime", "Evet", "Kayıt oluşturma tarihi"],
            ["28", "lastUpdated", "DateTime", "Evet", "Son güncelleme tarihi"],
        ]

        self.add_data_table(asset_detail_2[0], asset_detail_2[1:], [0.4, 1.3, 1.0, 0.8, 4.0])

        self.add_page_break()

        # Add Asset Form
        self.doc.add_heading('4.3.3 Varlık Ekleme Formu', 3)
        self.add_screenshot('09_Modal_Varlik_Ekle_Fixed.png', 'Yeni Varlık Ekleme Formu')

        # Add Asset Form Data Structure
        self.doc.add_heading('Form Alanları Veri Yapısı', 4)

        add_asset_form = [
            ["Alan Adı", "Form Label (TR)", "Tip", "Zorunlu", "Enum Değerleri / Notlar"],
            ["assetName", "Varlık Adı", "Text", "Evet", "Max 200 karakter"],
            ["sapAssetId", "SAP Varlık No", "Text", "Hayır", "SAP'den otomatik gelecek veya manuel girilebilir"],
            ["assetType", "Varlık Tipi", "Dropdown", "Evet", "Hand tools, Electric, Construction, Tool-Counter, Mechanic, Office, Meeting room, Other (8 değer)"],
            ["category", "Kategori", "Dropdown", "Evet", "Üretim Ekipmanı, Bakım Ekipmanı, Kalite Kontrol, Güvenlik"],
            ["status", "Durum", "Dropdown", "Evet", "Active, Inactive, Maintenance, Faulty"],
            ["manufacturer", "Üretici", "Text", "Hayır", "Max 100 karakter"],
            ["model", "Model", "Text", "Hayır", "Max 100 karakter"],
            ["serialNumber", "Seri No", "Text", "Hayır", "Max 100 karakter"],
            ["location", "Lokasyon", "Dropdown", "Evet", "Master data'dan"],
            ["subLocation1", "Alt Lokasyon 1", "Text", "Hayır", "Max 100 karakter"],
            ["subLocation2", "Alt Lokasyon 2", "Text", "Hayır", "Max 100 karakter"],
            ["purchaseDate", "Satın Alma Tarihi", "Date", "Hayır", "Date picker"],
            ["purchaseValue", "Satın Alma Değeri", "Number", "Hayır", "Decimal (12,2)"],
            ["currency", "Para Birimi", "Dropdown", "Hayır", "TRY (default), EUR, USD"],
            ["costCenter", "Maliyet Merkezi", "Text", "Evet", "SAP maliyet merkezi kodu"],
            ["assignedTo", "Zimmetli Kişi/Ekip", "Text", "Hayır", "Autocomplete kullanıcı/ekip listesinden"],
            ["description", "Açıklama", "Textarea", "Hayır", "Max 2000 karakter"],
        ]

        self.add_data_table(add_asset_form[0], add_asset_form[1:], [1.2, 1.3, 0.9, 0.7, 3.4])

        self.add_page_break()

    def add_maintenance_module(self):
        """Add Maintenance module documentation"""
        self.doc.add_heading('4.4 Bakım Yönetimi (Maintenance Management) Modülü', 2)

        # Module Description
        desc = """Bakım Yönetimi modülü, periyodik, önleyici, düzeltici ve ölçüm bazlı bakım planlarının
oluşturulmasını ve takibini sağlar. Varlıklarla ilişkili bakım planları oluşturulur, ekiplere atanır
ve tamamlanma durumu takip edilir. Bakım geçmişi raporlanabilir."""
        self.doc.add_paragraph(desc).alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        # List Page
        self.doc.add_heading('4.4.1 Bakım Yönetimi Liste Sayfası', 3)
        self.add_screenshot('06_Bakim_Yonetimi.png', 'Liste Sayfası Ekran Görüntüsü')

        # List Page Data Structure
        self.doc.add_heading('Liste Sayfası Veri Yapısı', 4)

        maintenance_list = [
            ["Sütun Adı", "Alan Adı (API)", "Veri Tipi", "Açıklama"],
            ["Bakım ID", "maintenanceId", "String", "MNT-YYYY-XXX formatında unique ID"],
            ["Varlık", "assetName", "String", "İlişkili varlığın adı"],
            ["Bakım Tipi", "maintenanceType", "Enum", "Periodic, Measured, Preventive, Corrective, Mass"],
            ["Planlanan Tarih", "scheduledDate", "Date", "Bakımın planlandığı tarih"],
            ["Durum", "status", "Enum", "Planned, In-Progress, Completed, Cancelled"],
            ["Atanan Ekip", "assignedTeam", "String", "Bakımdan sorumlu ekip"],
            ["Öncelik", "priority", "Enum", "Urgent, High, Normal, Low"],
        ]

        self.add_data_table(maintenance_list[0], maintenance_list[1:], [1.2, 1.2, 1.0, 4.1])

        # List Page Filters
        self.doc.add_heading('Filtre Alanları', 4)

        maintenance_filters = [
            ["Filtre Adı", "Alan Adı", "Tip", "Değerler"],
            ["Durum Filtresi", "statusFilter", "Dropdown", "Tümü, Planned, In-Progress, Completed, Cancelled"],
            ["Tip Filtresi", "typeFilter", "Dropdown", "Tümü, Periodic, Measured, Preventive, Corrective, Mass"],
            ["Tarih Aralığı", "dateRange", "DateRange", "Başlangıç ve bitiş tarihi seçimi"],
            ["Varlık Filtresi", "assetFilter", "Autocomplete", "Varlık adı veya SAP ID ile arama"],
        ]

        self.add_data_table(maintenance_filters[0], maintenance_filters[1:], [1.5, 1.3, 1.0, 3.7])

        self.add_page_break()

        # Create Maintenance Plan Form
        self.doc.add_heading('4.4.2 Yeni Bakım Planı Oluşturma Formu', 3)
        self.add_screenshot('10_Modal_Yeni_Bakim_Plani_Fixed.png', 'Yeni Bakım Planı Formu')

        # Form Data Structure
        self.doc.add_heading('Form Alanları Veri Yapısı', 4)

        maintenance_form = [
            ["Alan Adı", "Form Label (TR)", "Tip", "Zorunlu", "Enum Değerleri / Notlar"],
            ["title", "Bakım Planı Başlığı", "Text", "Evet", "Max 200 karakter"],
            ["maintenanceType", "Bakım Tipi", "Dropdown", "Evet", "Periodic (Periyodik), Measured (Ölçüm Bazlı), Preventive (Önleyici), Corrective (Düzeltici), Mass (Toplu) - 5 değer"],
            ["priority", "Öncelik", "Dropdown", "Evet", "Urgent, High, Normal, Low"],
            ["assetId", "Varlık Seçimi", "Autocomplete", "Evet", "Assets tablosundan, SAP ID veya isim ile arama"],
            ["scheduledStartDate", "Planlanan Başlangıç Tarihi", "Date", "Evet", "Date picker"],
            ["estimatedDuration", "Tahmini Süre (Saat)", "Number", "Hayır", "Saat cinsinden"],
            ["recurrencePeriod", "Tekrar Periyodu", "Dropdown", "Hayır", "Tek Seferlik, Haftalık, Aylık, 3 Aylık, 6 Aylık, Yıllık"],
            ["assignedTeam", "Sorumlu Ekip", "Dropdown", "Hayır", "Mekanik Bakım, Elektrik Bakım, Elektronik Bakım, Otomasyon"],
            ["description", "Bakım Açıklaması", "Textarea", "Evet", "Max 2000 karakter, yapılacak işlerin detaylı listesi"],
            ["requiredMaterials", "Gerekli Malzemeler", "Textarea", "Hayır", "Yedek parça ve malzeme listesi"],
            ["estimatedCost", "Tahmini Maliyet (TRY)", "Number", "Hayır", "Decimal (10,2)"],
            ["costCenter", "Maliyet Merkezi", "Text", "Hayır", "SAP maliyet merkezi kodu"],
        ]

        self.add_data_table(maintenance_form[0], maintenance_form[1:], [1.2, 1.5, 0.9, 0.7, 3.2])

        # Maintenance Types Detail
        self.doc.add_heading('Bakım Tipleri Detayı', 4)

        maintenance_types = [
            ["Bakım Tipi", "Kod", "Açıklama", "Kullanım Senaryosu"],
            ["Periyodik Bakım", "periodic", "Düzenli aralıklarla yapılan bakım", "Aylık, yıllık rutin bakımlar"],
            ["Ölçüm Bazlı Bakım", "measured", "Sayaç/ölçüm değerine göre", "Çalışma saati, üretim adedi bazlı"],
            ["Önleyici Bakım", "preventive", "Arıza olmadan önleyici", "Kritik ekipman koruyucu bakımları"],
            ["Düzeltici Bakım", "corrective", "Arıza sonrası onarım", "Reaktif bakım ve onarımlar"],
            ["Toplu Bakım", "mass", "Birden fazla varlık bir arada", "Duruş zamanlarında toplu bakım"],
        ]

        self.add_data_table(maintenance_types[0], maintenance_types[1:], [1.5, 0.9, 2.5, 2.6])

        self.add_page_break()

    def add_incidents_module(self):
        """Add Incidents module documentation"""
        self.doc.add_heading('4.5 Olay Yönetimi (Incident Management) Modülü', 2)

        # Module Description
        desc = """Olay Yönetimi modülü, ekipman arızaları, güvenlik olayları, kalite sorunları ve çevre olaylarının
hızlı bir şekilde bildirilmesini ve SLA (Service Level Agreement) bazlı müdahale edilmesini sağlar.
Kritik olaylar otomatik olarak ilgili kişilere SMS ve e-posta ile bildirilir."""
        self.doc.add_paragraph(desc).alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        # List Page
        self.doc.add_heading('4.5.1 Olay Yönetimi Liste Sayfası', 3)
        self.add_screenshot('07_Olay_Yonetimi.png', 'Liste Sayfası Ekran Görüntüsü')

        # List Page Data Structure
        self.doc.add_heading('Liste Sayfası Veri Yapısı', 4)

        incident_list = [
            ["Sütun Adı", "Alan Adı (API)", "Veri Tipi", "Açıklama"],
            ["Olay No", "incidentId", "String", "INC-YYYY-XXX formatında unique ID"],
            ["Başlık", "title", "String", "Olay başlığı"],
            ["Olay Tipi", "incidentType", "Enum", "Equipment Breakdown, Safety Incident, Quality Issue, Environmental"],
            ["Aciliyet", "severity", "Enum", "Critical, High, Medium, Low"],
            ["Lokasyon", "location", "String", "Olayın gerçekleştiği lokasyon"],
            ["Bildirim Tarihi", "reportedDate", "DateTime", "Olayın bildirilme tarihi"],
            ["Durum", "status", "Enum", "Open, Investigating, Resolved, Closed"],
            ["SLA Durumu", "slaStatus", "Calculated", "On-Time, At-Risk, Breached (SLA'ya göre hesaplanır)"],
        ]

        self.add_data_table(incident_list[0], incident_list[1:], [1.2, 1.2, 0.9, 4.2])

        # List Page Filters
        self.doc.add_heading('Filtre Alanları', 4)

        incident_filters = [
            ["Filtre Adı", "Alan Adı", "Tip", "Değerler"],
            ["Durum Filtresi", "statusFilter", "Dropdown", "Tümü, Open, Investigating, Resolved, Closed"],
            ["Tip Filtresi", "typeFilter", "Dropdown", "Tümü, Equipment Breakdown, Safety, Quality, Environmental"],
            ["Aciliyet Filtresi", "severityFilter", "Dropdown", "Tümü, Critical, High, Medium, Low"],
            ["Tarih Aralığı", "dateRange", "DateRange", "Başlangıç ve bitiş tarihi seçimi"],
        ]

        self.add_data_table(incident_filters[0], incident_filters[1:], [1.5, 1.3, 1.0, 3.7])

        self.add_page_break()

        # Report Incident Form
        self.doc.add_heading('4.5.2 Olay Bildirme Formu', 3)
        self.add_screenshot('11_Modal_Olay_Bildir_Fixed.png', 'Acil Olay Bildirme Formu')

        # Form Data Structure
        self.doc.add_heading('Form Alanları Veri Yapısı', 4)

        incident_form = [
            ["Alan Adı", "Form Label (TR)", "Tip", "Zorunlu", "Enum Değerleri / Notlar"],
            ["title", "Olay Başlığı", "Text", "Evet", "Max 200 karakter"],
            ["incidentType", "Olay Tipi", "Dropdown", "Evet", "Equipment Breakdown (Ekipman Arızası), Safety Incident (Güvenlik), Quality Issue (Kalite), Environmental Incident (Çevre) - 4 değer"],
            ["severity", "Aciliyet", "Dropdown", "Evet", "Critical (Kritik - Üretim Durdu), High (Yüksek), Medium (Orta), Low (Düşük)"],
            ["location", "Lokasyon", "Dropdown", "Evet", "Şasi Üretim, Boya Tesisi, Montaj Hattı, Depo, vb."],
            ["assetId", "İlişkili Varlık", "Autocomplete", "Hayır", "İlgili ekipman (opsiyonel)"],
            ["description", "Olay Açıklaması", "Textarea", "Evet", "Max 2000 karakter, olayın detaylı açıklaması"],
            ["immediateActions", "Alınan Acil Önlemler", "Textarea", "Hayır", "Şu ana kadar alınan önlemler"],
            ["productionStatus", "Üretim Durumu", "Dropdown", "Hayır", "Üretim Durdu, Üretim Yavaşladı, Normal Devam Ediyor"],
            ["injuryOccurred", "Yaralanma/Kaza Var mı?", "Dropdown", "Hayır", "Hayır (default), Evet - Detay Ekle"],
        ]

        self.add_data_table(incident_form[0], incident_form[1:], [1.2, 1.4, 0.9, 0.7, 3.3])

        # Incident Types and SLA
        self.doc.add_heading('Olay Tipleri ve SLA Tanımları', 4)

        incident_sla = [
            ["Olay Tipi", "Aciliyet", "İlk Müdahale SLA", "Çözüm SLA", "Bildirim"],
            ["Equipment Breakdown", "Critical", "15 dakika", "4 saat", "SMS + Email + Push"],
            ["Equipment Breakdown", "High", "30 dakika", "8 saat", "Email + Push"],
            ["Safety Incident", "Critical", "5 dakika", "2 saat", "SMS + Email + Push (EHS ekibi)"],
            ["Safety Incident", "High", "10 dakika", "4 saat", "Email + Push"],
            ["Quality Issue", "High", "1 saat", "24 saat", "Email + Push (Kalite ekibi)"],
            ["Environmental", "Critical", "10 dakika", "1 saat", "SMS + Email (EHS + Yönetim)"],
        ]

        self.add_data_table(incident_sla[0], incident_sla[1:], [1.5, 0.9, 1.2, 1.0, 2.9])

        # Warning Note
        note = self.doc.add_paragraph()
        note.add_run("ÖNEMLİ: ").font.bold = True
        note.add_run("Kritik olaylar otomatik olarak ilgili tüm yöneticilere, bakım ekibine ve güvenlik birimine bildirilir. ")
        note.add_run("SMS ve e-posta ile anında bildirim gönderilir. ")
        note.add_run("Güvenlik olaylarında EHS (Environment, Health & Safety) ekibi otomatik olarak devreye girer.")

        # Add styling to note
        for run in note.runs:
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(255, 102, 0)

        # Background color for note paragraph
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), 'FFF3CD')
        note._element.get_or_add_pPr().append(shading_elm)

        self.add_page_break()

    def add_workflow_diagrams(self):
        """Add workflow diagrams section"""
        self.doc.add_heading('5. İŞ AKIŞI DİYAGRAMLARI', 1)

        intro = """Bu bölümde sistemin ana iş akışları görsel olarak sunulmuştur. Her akış diyagramı,
ilgili modülün süreç adımlarını, karar noktalarını ve sistem davranışlarını göstermektedir."""
        self.doc.add_paragraph(intro).alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        # Workflow diagrams
        workflows = [
            ('is_talebi_akisi.png', '5.1 İş Talebi İş Akışı'),
            ('varlik_yonetimi_akisi.png', '5.2 Varlık Yönetimi İş Akışı'),
            ('bakim_planlama_akisi.png', '5.3 Bakım Planlama İş Akışı'),
            ('olay_yonetimi_akisi.png', '5.4 Olay Yönetimi İş Akışı'),
        ]

        for filename, title in workflows:
            self.doc.add_heading(title, 2)
            diagram_path = os.path.join(self.diagrams_path, filename)
            if os.path.exists(diagram_path):
                try:
                    self.doc.add_picture(diagram_path, width=Inches(6.5))
                    last_paragraph = self.doc.paragraphs[-1]
                    last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                except Exception as e:
                    self.doc.add_paragraph(f"[Diyagram yüklenemedi: {filename}]")
            else:
                self.doc.add_paragraph(f"[Diyagram bulunamadı: {filename}]")

            self.doc.add_paragraph()  # Add space

        self.add_page_break()

    def add_sap_integration(self):
        """Add SAP integration requirements"""
        self.doc.add_heading('6. SAP ENTEGRASYON GEREKSİNİMLERİ', 1)

        intro = """Sistem, SAP ERP sistemi ile gerçek zamanlı entegrasyon sağlayarak varlık bilgileri,
maliyet merkezi bilgileri ve iş emirlerini senkronize eder. Entegrasyon REST API veya RFC
(Remote Function Call) üzerinden gerçekleştirilecektir."""
        self.doc.add_paragraph(intro).alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        # SAP Modules
        self.doc.add_heading('6.1 Entegre Olunacak SAP Modülleri', 2)

        sap_modules = [
            ["SAP Modül", "Modül Adı", "Entegrasyon Amacı", "Veri Akışı"],
            ["SAP PM", "Plant Maintenance", "Varlık bilgileri, bakım planları", "SAP → Sistem (Read)"],
            ["SAP MM", "Materials Management", "Yedek parça stok bilgileri", "SAP → Sistem (Read)"],
            ["SAP FI/CO", "Finance/Controlling", "Maliyet merkezi, bütçe", "SAP ↔ Sistem (Read/Write)"],
            ["SAP HR", "Human Resources", "Çalışan bilgileri, organizasyon", "SAP → Sistem (Read)"],
        ]

        self.add_data_table(sap_modules[0], sap_modules[1:], [1.2, 1.5, 2.3, 2.5])

        # Integration Points
        self.doc.add_heading('6.2 Entegrasyon Noktaları', 2)

        integration_points = [
            ["Entegrasyon Noktası", "SAP Veri Kaynağı", "Senkronizasyon", "Açıklama"],
            ["Varlık Bilgileri", "SAP PM - Equipment Master", "Günlük (Batch)", "Ekipman listesi, teknik bilgiler"],
            ["Maliyet Merkezi", "SAP FI/CO - Cost Center", "Günlük (Batch)", "Aktif maliyet merkezi listesi"],
            ["Çalışan Bilgileri", "SAP HR - Employee Data", "Günlük (Batch)", "Aktif çalışanlar, departman, yönetici"],
            ["İş Emri Oluşturma", "SAP PM - Work Order", "Real-time (API)", "Onaylanan taleplerden iş emri oluştur"],
            ["Maliyet Aktarımı", "SAP FI/CO - Cost Posting", "Günlük (Batch)", "Gerçekleşen maliyetleri SAP'ye aktar"],
            ["Stok Sorgu", "SAP MM - Material Stock", "Real-time (API)", "Yedek parça stok durumu sorgusu"],
        ]

        self.add_data_table(integration_points[0], integration_points[1:], [1.5, 1.5, 1.2, 3.3])

        # Technical Requirements
        self.doc.add_heading('6.3 Teknik Gereksinimler', 2)

        tech_reqs = [
            "SAP NetWeaver Gateway REST API erişimi veya RFC bağlantısı",
            "SAP kullanıcı hesabı (Servis kullanıcısı) ve yetkilendirme",
            "SAP ABAP Custom Function Module geliştirme (gerekirse)",
            "SSL/TLS şifreli bağlantı güvenliği",
            "API rate limiting ve timeout yönetimi",
            "Hata durumunda retry mekanizması",
            "Entegrasyon log ve monitoring",
            "SAP veri değişikliklerinde Change Data Capture (CDC)"
        ]

        for req in tech_reqs:
            p = self.doc.add_paragraph(req, style='List Bullet')
            p.runs[0].font.size = Pt(11)

        self.add_page_break()

    def add_authorization_roles(self):
        """Add authorization and roles section"""
        self.doc.add_heading('7. YETKİLENDİRME VE ROLLER', 1)

        intro = """Sistem rol bazlı yetkilendirme (RBAC - Role Based Access Control) kullanır.
Her kullanıcı bir veya daha fazla role sahip olabilir ve roller belirli modüllere ve işlemlere
erişim hakkı tanımlar."""
        self.doc.add_paragraph(intro).alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        # Roles and Permissions
        self.doc.add_heading('7.1 Sistem Rolleri ve Yetkileri', 2)

        roles_data = [
            ["Rol Adı", "Açıklama", "İş Talebi", "Varlık", "Bakım", "Olay"],
            ["Çalışan (Employee)", "Standart kullanıcı", "Oluştur, Görüntüle", "Görüntüle", "Görüntüle", "Bildir"],
            ["Teknisyen (Technician)", "Bakım teknisyeni", "Görüntüle, Güncelle", "Görüntüle, Güncelle", "Görüntüle, Tamamla", "Görüntüle, Güncelle"],
            ["SL Mühendis (SL Engineer)", "Teknik onay", "Tümü (Teknik)", "Tümü", "Tümü", "Tümü"],
            ["Yönetici (Manager)", "Departman yöneticisi", "Onayla, Görüntüle", "Görüntüle", "Görüntüle", "Görüntüle"],
            ["Bakım Planlayıcı", "Bakım ekip lideri", "Görüntüle, Ata", "Görüntüle", "Tümü", "Görüntüle, Ata"],
            ["Maliyet Birimi", "Maliyet hesaplama", "Maliyet İşlemleri", "Görüntüle", "Maliyet İşlemleri", "Sadece Okuma"],
            ["Sistem Admin (Admin)", "IT yöneticisi", "Tümü", "Tümü", "Tümü", "Tümü"],
        ]

        self.add_data_table(roles_data[0], roles_data[1:], [1.3, 1.5, 1.2, 1.0, 1.0, 1.0])

        # Detailed Permissions
        self.doc.add_heading('7.2 Detaylı Yetki Matrisi', 2)

        permission_matrix = [
            ["İşlem", "Çalışan", "Teknisyen", "SL Müh.", "Yönetici", "Admin"],
            ["İş Talebi Oluştur", "✓", "✓", "✓", "✓", "✓"],
            ["İş Onayı Ver", "-", "-", "-", "✓", "✓"],
            ["Teknik Onay Ver", "-", "-", "✓", "-", "✓"],
            ["Maliyet Hesapla", "-", "-", "✓", "-", "✓"],
            ["Maliyet Onay", "-", "-", "-", "✓", "✓"],
            ["Teknisyen Ata", "-", "-", "✓", "-", "✓"],
            ["İş Tamamla", "-", "✓", "✓", "-", "✓"],
            ["Varlık Ekle", "-", "-", "✓", "-", "✓"],
            ["Varlık Düzenle", "-", "✓", "✓", "-", "✓"],
            ["Bakım Planı Oluştur", "-", "-", "✓", "-", "✓"],
            ["Bakım Tamamla", "-", "✓", "✓", "-", "✓"],
            ["Olay Bildir", "✓", "✓", "✓", "✓", "✓"],
            ["Olay Çöz", "-", "✓", "✓", "-", "✓"],
            ["Rapor Görüntüle", "Kendi", "Ekip", "Tümü", "Departman", "Tümü"],
            ["Kullanıcı Yönetimi", "-", "-", "-", "-", "✓"],
        ]

        self.add_data_table(permission_matrix[0], permission_matrix[1:], [1.8, 0.8, 0.9, 0.9, 0.9, 0.8])

        self.add_page_break()

    def add_test_scenarios(self):
        """Add test scenarios section"""
        self.doc.add_heading('8. TEST SENARYOLARI', 1)

        intro = """Bu bölümde, sistemin fonksiyonel testleri için temel test senaryoları listelenmiştir.
Her senaryo, beklenen davranış ve test adımlarını içermektedir."""
        self.doc.add_paragraph(intro).alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        # Test scenarios for each module
        test_scenarios = [
            ("8.1 İş Talebi Modülü Test Senaryoları", [
                "Yeni iş talebi oluşturma - başarılı",
                "Zorunlu alanlar boş bırakıldığında hata mesajı",
                "İş onayı verme - yönetici rolü",
                "Yetkisiz kullanıcı onay vermeye çalıştığında hata",
                "Teknik onay - SL mühendis",
                "Maliyet hesaplama ve onay süreci",
                "Teknisyen atama ve iş tamamlama",
                "İş talebi reddetme ve red nedeni girme",
                "İş talebi filtreleme - durum, öncelik, kategori",
                "Toplu export (Excel/PDF) işlemi"
            ]),
            ("8.2 Varlık Yönetimi Test Senaryoları", [
                "Yeni varlık ekleme - tüm alanlar dolu",
                "SAP ID ile varlık bilgilerini otomatik getirme",
                "Varlık detay sayfası görüntüleme",
                "Varlık düzenleme - yetkili kullanıcı",
                "Varlık durumu güncelleme (Aktif → Bakımda)",
                "QR kod oluşturma ve yazdırma",
                "Varlık bakım geçmişini görüntüleme",
                "Varlık listesi filtreleme ve arama",
                "Varlık dokümantasyon yükleme",
                "SAP ile senkronizasyon testi"
            ]),
            ("8.3 Bakım Yönetimi Test Senaryoları", [
                "Periyodik bakım planı oluşturma",
                "Bakım planına ekip atama",
                "Bakım planı tekrarlama ayarları (aylık, yıllık)",
                "Bakım tamamlama ve rapor girme",
                "Geciken bakım uyarıları",
                "Bakım maliyeti girişi",
                "Bakım takvimi görünümü",
                "Bakım geçmişi raporlama",
                "Varlığa bağlı bakım planlarını listeleme",
                "Bakım planı iptal etme"
            ]),
            ("8.4 Olay Yönetimi Test Senaryoları", [
                "Kritik olay bildirimi - otomatik bildirimler",
                "Olay detayları ve acil önlemler girme",
                "SLA süre hesaplama - kritik olay",
                "Olay durumu güncelleme (Open → Investigating → Resolved)",
                "Olaya teknisyen atama",
                "Olay raporlama ve kapama",
                "SLA ihlali uyarıları",
                "Güvenlik olayında EHS ekibine otomatik bildirim",
                "Olay istatistikleri ve trendler",
                "Olaya ilişkili varlık bağlama"
            ]),
        ]

        for heading, scenarios in test_scenarios:
            self.doc.add_heading(heading, 2)
            for scenario in scenarios:
                p = self.doc.add_paragraph(scenario, style='List Bullet')
                p.runs[0].font.size = Pt(11)

        self.add_page_break()

    def add_technical_specifications(self):
        """Add technical specifications"""
        self.doc.add_heading('9. TEKNİK SPESİFİKASYONLAR', 1)

        # Architecture
        self.doc.add_heading('9.1 Sistem Mimarisi', 2)

        arch_text = """Sistem modern web teknolojileri kullanılarak geliştirilecektir.
Frontend ve Backend ayrı katmanlar olarak tasarlanmış olup, RESTful API üzerinden iletişim kurar."""
        self.doc.add_paragraph(arch_text).alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        # Technology Stack
        self.doc.add_heading('9.2 Teknoloji Stack Önerisi', 2)

        tech_stack = [
            ["Katman", "Teknoloji", "Açıklama"],
            ["Frontend", "React.js / Angular / Vue.js", "Modern JavaScript framework"],
            ["", "Bootstrap / Material-UI", "UI component library"],
            ["", "Chart.js / D3.js", "Data visualization"],
            ["Backend", "Node.js + Express / .NET Core / Java Spring", "REST API geliştirme"],
            ["", "JWT Authentication", "Token-based authentication"],
            ["Database", "PostgreSQL / MySQL / SQL Server", "İlişkisel veritabanı"],
            ["", "Redis", "Caching ve session yönetimi"],
            ["Integration", "SAP REST API / RFC", "SAP entegrasyonu"],
            ["", "Active Directory / LDAP", "Kullanıcı yetkilendirme"],
            ["Notifications", "SendGrid / Twilio", "Email ve SMS gönderimi"],
            ["File Storage", "AWS S3 / Azure Blob", "Dosya depolama"],
            ["Deployment", "Docker + Kubernetes", "Container orchestration"],
            ["Monitoring", "Prometheus + Grafana", "System monitoring"],
        ]

        self.add_data_table(tech_stack[0], tech_stack[1:], [1.5, 2.5, 3.5])

        # Database Schema
        self.doc.add_heading('9.3 Veritabanı Tabloları', 2)

        db_tables = [
            "JobRequests - İş talepleri ana tablosu",
            "JobRequestComments - İş talebi yorumları",
            "JobRequestAttachments - İş talebi ekleri",
            "JobRequestHistory - İş talebi durum geçmişi",
            "Assets - Varlıklar ana tablosu",
            "AssetDocuments - Varlık dökümanları",
            "AssetMaintenanceHistory - Varlık bakım geçmişi",
            "MaintenancePlans - Bakım planları",
            "MaintenanceExecutions - Bakım uygulamaları",
            "Incidents - Olaylar ana tablosu",
            "IncidentActions - Olay aksiyon kayıtları",
            "Users - Kullanıcı bilgileri",
            "Roles - Roller",
            "UserRoles - Kullanıcı-rol ilişkisi",
            "Permissions - İzinler",
            "Locations - Lokasyon master data",
            "CostCenters - Maliyet merkezi master data",
            "Notifications - Bildirim kayıtları",
            "SystemSettings - Sistem ayarları",
            "AuditLog - Denetim kayıtları"
        ]

        for table in db_tables:
            p = self.doc.add_paragraph(table, style='List Bullet')
            p.runs[0].font.size = Pt(11)

        # Performance Requirements
        self.doc.add_heading('9.4 Performans Gereksinimleri', 2)

        perf_reqs = [
            "Sayfa yükleme süresi: < 2 saniye",
            "API response time: < 500ms (95th percentile)",
            "Eşzamanlı kullanıcı kapasitesi: 500+ kullanıcı",
            "Veritabanı sorgu optimizasyonu: indexleme stratejisi",
            "File upload: Max 50MB per file",
            "Dashboard widget'ları: Real-time data updates (WebSocket)",
            "Batch işlemler: Off-peak saatlerde (gece)",
            "Backup: Günlük otomatik yedekleme"
        ]

        for req in perf_reqs:
            p = self.doc.add_paragraph(req, style='List Bullet')
            p.runs[0].font.size = Pt(11)

        # Security
        self.doc.add_heading('9.5 Güvenlik Gereksinimleri', 2)

        security = [
            "HTTPS/TLS 1.2+ şifreli iletişim",
            "JWT token ile authentication (15 dakika expiry)",
            "Rol bazlı yetkilendirme (RBAC)",
            "SQL Injection koruması (Prepared statements)",
            "XSS (Cross-Site Scripting) koruması",
            "CSRF token validation",
            "Rate limiting ve DDoS koruması",
            "Dosya upload güvenlik kontrolü (virus scan)",
            "Şifre politikası: Min 8 karakter, büyük/küçük harf, rakam, özel karakter",
            "İki faktörlü kimlik doğrulama (2FA) - opsiyonel",
            "Audit log: Tüm kritik işlemler loglanır",
            "GDPR/KVKK uyumlu veri koruma"
        ]

        for item in security:
            p = self.doc.add_paragraph(item, style='List Bullet')
            p.runs[0].font.size = Pt(11)

        self.add_page_break()

    def add_appendices(self):
        """Add appendices section"""
        self.doc.add_heading('10. EKLER', 1)

        # Appendix A: Glossary
        self.doc.add_heading('10.1 Terimler Sözlüğü', 2)

        glossary = [
            ["Terim", "Açıklama"],
            ["SL Mühendis", "Service Level Engineer - Teknik onay veren mühendis"],
            ["İş Onayı", "Business Approval - Yöneticinin iş talebini onaylaması"],
            ["Teknik Onay", "Technical Approval - SL mühendisinin teknik fizibilite onayı"],
            ["Maliyet Onayı", "Cost Approval - İş yöneticisinin bütçe onayı"],
            ["SAP PM", "SAP Plant Maintenance - Bakım yönetimi modülü"],
            ["SAP MM", "SAP Materials Management - Malzeme yönetimi"],
            ["SAP FI/CO", "SAP Finance/Controlling - Finans ve kontrol"],
            ["RFC", "Remote Function Call - SAP uzaktan fonksiyon çağrısı"],
            ["SLA", "Service Level Agreement - Servis seviye sözleşmesi"],
            ["EHS", "Environment, Health & Safety - Çevre, sağlık ve güvenlik"],
            ["OHS / İSG", "Occupational Health & Safety - İş sağlığı ve güvenliği"],
            ["RBAC", "Role-Based Access Control - Rol bazlı erişim kontrolü"],
            ["JWT", "JSON Web Token - Token tabanlı kimlik doğrulama"],
            ["REST API", "RESTful Application Programming Interface"],
            ["CRUD", "Create, Read, Update, Delete - Temel veri işlemleri"],
        ]

        self.add_data_table(glossary[0], glossary[1:], [2.0, 5.5])

        # Appendix B: Enum Values
        self.doc.add_heading('10.2 Enum Değerleri Referans Tablosu', 2)

        self.doc.add_heading('İş Talebi - Kategori (Category)', 3)
        categories = [
            ["Değer", "Label (TR)", "Label (EN)"],
            ["hvac", "HVAC (İklimlendirme)", "HVAC"],
            ["electrical", "Elektrik Sistemleri", "Electrical Systems"],
            ["mechanical", "Mekanik Ekipman", "Mechanical Equipment"],
            ["plumbing", "Sıhhi Tesisat", "Plumbing"],
            ["building", "Bina Onarımı", "Building Repair"],
        ]
        self.add_data_table(categories[0], categories[1:], [1.5, 2.5, 2.5])

        self.doc.add_heading('İş Talebi - Talep Nedeni (Request Reason)', 3)
        reasons = [
            ["Değer", "Label (TR)", "Label (EN)"],
            ["ohs", "İSG (İş Sağlığı ve Güvenliği)", "OHS (Occupational Health & Safety)"],
            ["energy", "Enerji Tasarrufu", "Energy Saving"],
            ["environment", "Çevre", "Environment"],
            ["improvement", "Süreç İyileştirme", "Process Improvement"],
            ["investment", "Yatırım", "Investment"],
            ["renovation", "Yenileme", "Renovation"],
        ]
        self.add_data_table(reasons[0], reasons[1:], [1.5, 2.5, 2.5])

        self.doc.add_heading('Varlık - Varlık Tipi (Asset Type)', 3)
        asset_types = [
            ["Değer", "Label (TR)", "Label (EN)"],
            ["hand-tools", "El Aletleri", "Hand tools"],
            ["electric", "Elektrikli Ekipman", "Electric"],
            ["construction", "İnşaat Ekipmanı", "Construction"],
            ["tool-counter", "Alet - Sayaç", "Tool – Counter"],
            ["mechanic", "Mekanik", "Mechanic"],
            ["office", "Ofis Ekipmanı", "Office"],
            ["meeting", "Toplantı Odası", "Meeting room"],
            ["other", "Diğer", "Other"],
        ]
        self.add_data_table(asset_types[0], asset_types[1:], [1.5, 2.5, 2.5])

        self.doc.add_heading('Bakım - Bakım Tipi (Maintenance Type)', 3)
        maint_types = [
            ["Değer", "Label (TR)", "Label (EN)"],
            ["periodic", "Periyodik Bakım", "Periodic Maintenance"],
            ["measured", "Ölçüm Bazlı Bakım", "Measured Maintenance"],
            ["preventive", "Önleyici Bakım", "Preventive Maintenance"],
            ["corrective", "Düzeltici Bakım", "Corrective Maintenance"],
            ["mass", "Toplu Bakım", "Mass Maintenance"],
        ]
        self.add_data_table(maint_types[0], maint_types[1:], [1.5, 2.5, 2.5])

        self.doc.add_heading('Olay - Olay Tipi (Incident Type)', 3)
        incident_types = [
            ["Değer", "Label (TR)", "Label (EN)"],
            ["breakdown", "Ekipman Arızası", "Equipment Breakdown"],
            ["safety", "Güvenlik Olayı", "Safety Incident"],
            ["quality", "Kalite Sorunu", "Quality Issue"],
            ["environmental", "Çevre Olayı", "Environmental Incident"],
        ]
        self.add_data_table(incident_types[0], incident_types[1:], [1.5, 2.5, 2.5])

        # Appendix C: API Endpoints
        self.doc.add_heading('10.3 API Endpoints Referansı', 2)

        api_endpoints = [
            ["Method", "Endpoint", "Açıklama"],
            ["GET", "/api/job-requests", "İş talepleri listesi (filtreleme destekli)"],
            ["POST", "/api/job-requests", "Yeni iş talebi oluştur"],
            ["GET", "/api/job-requests/{id}", "İş talebi detay"],
            ["PUT", "/api/job-requests/{id}", "İş talebi güncelle"],
            ["POST", "/api/job-requests/{id}/approve", "İş talebi onayla"],
            ["POST", "/api/job-requests/{id}/reject", "İş talebi reddet"],
            ["GET", "/api/assets", "Varlık listesi"],
            ["POST", "/api/assets", "Yeni varlık ekle"],
            ["GET", "/api/assets/{id}", "Varlık detay"],
            ["PUT", "/api/assets/{id}", "Varlık güncelle"],
            ["GET", "/api/maintenance-plans", "Bakım planları listesi"],
            ["POST", "/api/maintenance-plans", "Yeni bakım planı oluştur"],
            ["GET", "/api/incidents", "Olay listesi"],
            ["POST", "/api/incidents", "Yeni olay bildir"],
            ["GET", "/api/dashboard/stats", "Dashboard istatistikleri"],
            ["POST", "/api/sap/sync-assets", "SAP varlık senkronizasyonu"],
        ]

        self.add_data_table(api_endpoints[0], api_endpoints[1:], [1.0, 2.5, 3.5])

        # Document End
        self.doc.add_paragraph()
        self.doc.add_paragraph()
        end_note = self.doc.add_paragraph()
        end_note.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = end_note.add_run("─── Doküman Sonu ───")
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(153, 153, 153)

    def generate(self):
        """Generate the complete document"""
        print("MAN Türkiye Bakım Yönetimi Sistemi - Kapsamlı Doküman Oluşturuluyor...")

        # Add all sections
        self.add_title_page()
        print("✓ Başlık sayfası eklendi")

        self.add_table_of_contents()
        print("✓ İçindekiler eklendi")

        self.add_management_summary()
        print("✓ Yönetici özeti eklendi")

        self.add_purpose_and_objectives()
        print("✓ Amaç ve hedefler eklendi")

        self.add_scope()
        print("✓ Kapsam eklendi")

        self.add_dashboard_module()
        print("✓ Dashboard modülü eklendi")

        self.add_job_requests_module()
        print("✓ İş Talepleri modülü eklendi")

        self.add_assets_module()
        print("✓ Varlık Yönetimi modülü eklendi")

        self.add_maintenance_module()
        print("✓ Bakım Yönetimi modülü eklendi")

        self.add_incidents_module()
        print("✓ Olay Yönetimi modülü eklendi")

        self.add_workflow_diagrams()
        print("✓ İş akışı diyagramları eklendi")

        self.add_sap_integration()
        print("✓ SAP entegrasyonu eklendi")

        self.add_authorization_roles()
        print("✓ Yetkilendirme ve roller eklendi")

        self.add_test_scenarios()
        print("✓ Test senaryoları eklendi")

        self.add_technical_specifications()
        print("✓ Teknik spesifikasyonlar eklendi")

        self.add_appendices()
        print("✓ Ekler eklendi")

        # Save document
        self.doc.save(self.output_path)
        file_size = os.path.getsize(self.output_path) / (1024 * 1024)  # MB
        print(f"\n✓ Doküman başarıyla oluşturuldu!")
        print(f"  Dosya: {self.output_path}")
        print(f"  Boyut: {file_size:.2f} MB")
        print(f"\nDoküman CMS panel geliştiricileri için hazırdır.")


if __name__ == "__main__":
    # Define paths
    output_path = "/Users/caglarozyildirim/WebstormProjects/Deneme/MAN_Turkiye_Bakim_Yonetimi_FINAL_COMPLETE.docx"
    screenshots_path = "/Users/caglarozyildirim/WebstormProjects/Deneme/screenshots/fixed/"
    diagrams_path = "/Users/caglarozyildirim/WebstormProjects/Deneme/diagrams/"

    # Create generator and generate document
    generator = MANDocumentGenerator(output_path, screenshots_path, diagrams_path)
    generator.generate()
