#!/usr/bin/env python3
"""
Create final document with correct screenshot paths
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import json
import os

def set_cell_background(cell, color):
    """Set cell background color"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading_elm)

def add_data_structure_table(doc, fields, module_color="#E20714"):
    """Add data structure table in CleanShot format"""
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Light Grid Accent 1'

    # Header row
    header_cells = table.rows[0].cells
    headers = ['Alan Adı (EN)', 'Alan Adı (TR)', 'Veri Tipi', 'Zorunlu', 'Açıklama', 'SAP Mapping']

    for i, header_text in enumerate(headers):
        cell = header_cells[i]
        set_cell_background(cell, module_color)
        paragraph = cell.paragraphs[0]
        paragraph.clear()
        run = paragraph.add_run(header_text)
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(10)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    for field in fields:
        row_cells = table.add_row().cells
        row_cells[0].text = field['en']
        row_cells[1].text = field['tr']
        row_cells[2].text = field['type']
        row_cells[3].text = field['required']
        row_cells[4].text = field['desc']
        row_cells[5].text = field.get('sap', '')

        for cell in row_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)

def add_screenshot(doc, path):
    """Add screenshot if exists"""
    if os.path.exists(path):
        try:
            doc.add_picture(path, width=Inches(6))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            return True
        except Exception as e:
            print(f"⚠️  Resim yüklenemedi {path}: {e}")
            return False
    else:
        print(f"❌ Dosya bulunamadı: {path}")
        return False

# Screenshot mapping
SCREENSHOTS = {
    'job-requests': '/Users/caglarozyildirim/WebstormProjects/Deneme/screenshots/02_is_talepleri_liste.png',
    'job-request-detail': '/Users/caglarozyildirim/WebstormProjects/Deneme/screenshots/03_is_talepleri_detay.png',
    'job-request-create': '/Users/caglarozyildirim/WebstormProjects/Deneme/screenshots/04_is_talepleri_olustur.png',
    'assets': '/Users/caglarozyildirim/WebstormProjects/Deneme/screenshots/05_varliklar_liste.png',
    'asset-detail': '/Users/caglarozyildirim/WebstormProjects/Deneme/screenshots/06_varliklar_detay.png',
    'asset-create': '/Users/caglarozyildirim/WebstormProjects/Deneme/screenshots/asset-create.png',
    'maintenance': '/Users/caglarozyildirim/WebstormProjects/Deneme/screenshots/07_bakim_liste.png',
    'maintenance-detail': '/Users/caglarozyildirim/WebstormProjects/Deneme/screenshots/08_bakim_detay.png',
    'maintenance-create': '/Users/caglarozyildirim/WebstormProjects/Deneme/screenshots/09_bakim_olustur.png',
    'incidents': '/Users/caglarozyildirim/WebstormProjects/Deneme/screenshots/10_olaylar_liste.png',
    'incident-detail': '/Users/caglarozyildirim/WebstormProjects/Deneme/screenshots/11_olaylar_detay.png',
    'incident-create': '/Users/caglarozyildirim/WebstormProjects/Deneme/screenshots/12_olaylar_olustur.png',
}

def create_document():
    """Create the complete document"""
    doc = Document()
    
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    with open('/Users/caglarozyildirim/WebstormProjects/Deneme/detailed_data_structures.json', 'r', encoding='utf-8') as f:
        data = json.load(f)

    print("📄 Doküman oluşturuluyor...")

    # Cover page
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('MAN TÜRKİYE')
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(226, 7, 20)
    
    doc.add_paragraph()
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('BAKIM YÖNETİMİ SİSTEMİ')
    run.bold = True
    run.font.size = Pt(24)
    
    doc.add_paragraph()
    subtitle2 = doc.add_paragraph()
    subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle2.add_run('İŞ ANALİZİ DOKÜMANI')
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(100, 100, 100)
    
    doc.add_page_break()

    # 1. Project Purpose
    print("✍️  1. Proje Amacı...")
    heading = doc.add_heading('1. PROJE AMACI VE HEDEFLER', level=1)
    heading.runs[0].font.color.rgb = RGBColor(226, 7, 20)

    doc.add_paragraph('MAN Türkiye Bakım Yönetimi Sistemi, üretim tesislerindeki tüm bakım süreçlerinin dijitalleştirilmesi ve SAP entegrasyonu ile merkezi yönetimini sağlamak amacıyla geliştirilmektedir.')

    doc.add_heading('1.1 Temel Hedefler', level=2)
    for goal in [
        'İş taleplerinin dijital ortamda oluşturulması ve onay sürecinin otomasyonu',
        'Varlık yönetimi ve SAP entegrasyonu ile merkezi kayıt sistemi',
        'Periyodik bakım planlaması ve takibi',
        'Olay yönetimi ile acil müdahalelerin hızlandırılması',
        'SLA takibi ve performans ölçümü',
        'Malzeme tüketimi ve maliyet takibi'
    ]:
        doc.add_paragraph(goal, style='List Bullet')

    doc.add_page_break()

    # 2. Project Scope
    print("✍️  2. Proje Kapsamı...")
    heading = doc.add_heading('2. PROJE KAPSAMI', level=1)
    heading.runs[0].font.color.rgb = RGBColor(226, 7, 20)

    doc.add_heading('2.1 Kapsam İçinde', level=2)
    for item in [
        'İş Talepleri Yönetimi: Talep oluşturma, onay süreci, çözüm takibi',
        'Varlık Yönetimi: Varlık kaydı, SAP entegrasyonu, transfer işlemleri',
        'Bakım Yönetimi: Periyodik bakım planlaması, görev atama, tamamlama',
        'Olay Yönetimi: Acil olay bildirimi, önceliklendirme, müdahale',
        'Mobil Uygulama: Teknisyen tarafı için mobil erişim',
        'Web Uygulaması: Planlama ve yönetim için masaüstü arayüz'
    ]:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('2.2 Kapsam Dışında', level=2)
    for item in [
        'Finans modülü yönetimi',
        'İnsan kaynakları entegrasyonu',
        'Satın alma süreçleri',
        'Üretim planlama modülü'
    ]:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

    # 3. Requirements
    heading = doc.add_heading('3. GEREKSINIMLER', level=1)
    heading.runs[0].font.color.rgb = RGBColor(226, 7, 20)

    # 3.1 Job Requests
    print("✍️  3.1 İş Talepleri...")
    doc.add_heading('3.1 İş Talepleri Yönetimi', level=2)
    doc.add_paragraph('İş talepleri modülü, kullanıcıların bakım, onarım veya yeni kurulum taleplerini sisteme girmesini, onay sürecinden geçirmesini ve çözüm aşamasına kadar takip edilmesini sağlar.')
    doc.add_heading('Müşteri Gereksinimi:', level=3)
    doc.add_paragraph(data['job_requests']['requirement'])
    
    doc.add_heading('3.1.1 İş Talepleri - Liste Ekranı', level=3)
    doc.add_paragraph('Bu ekran tüm iş taleplerinin listelendiği ana ekrandır.')
    doc.add_heading('Ekran Görüntüsü:', level=4)
    add_screenshot(doc, SCREENSHOTS['job-requests'])
    doc.add_heading('Veri Yapısı:', level=4)
    add_data_structure_table(doc, data['job_requests']['fields'], "#FFA726")
    doc.add_paragraph()

    doc.add_heading('3.1.2 İş Talebi - Detay Ekranı', level=3)
    doc.add_paragraph('Seçilen iş talebinin tüm detaylarının görüntülendiği ekrandır.')
    doc.add_heading('Ekran Görüntüsü:', level=4)
    add_screenshot(doc, SCREENSHOTS['job-request-detail'])
    doc.add_heading('Veri Yapısı:', level=4)
    add_data_structure_table(doc, data['job_requests']['fields'], "#FFA726")
    doc.add_paragraph()

    doc.add_heading('3.1.3 İş Talebi Oluştur - Form Ekranı', level=3)
    doc.add_paragraph('Yeni iş talebi oluşturma ekranıdır.')
    doc.add_heading('Ekran Görüntüsü:', level=4)
    add_screenshot(doc, SCREENSHOTS['job-request-create'])
    doc.add_heading('Veri Yapısı:', level=4)
    add_data_structure_table(doc, data['job_requests']['fields'], "#FFA726")
    doc.add_paragraph()
    doc.add_page_break()

    # 3.2 Assets
    print("✍️  3.2 Varlık Yönetimi...")
    doc.add_heading('3.2 Varlık Yönetimi', level=2)
    doc.add_paragraph('Varlık yönetimi modülü, tüm sabit varlıkların sisteme kaydedilmesini sağlar.')
    doc.add_heading('Müşteri Gereksinimi:', level=3)
    doc.add_paragraph(data['assets']['requirement'])
    
    doc.add_heading('3.2.1 Varlık Yönetimi - Liste Ekranı', level=3)
    doc.add_paragraph('Sistemde kayıtlı tüm varlıkların listelendiği ekrandır.')
    doc.add_heading('Ekran Görüntüsü:', level=4)
    add_screenshot(doc, SCREENSHOTS['assets'])
    doc.add_heading('Veri Yapısı:', level=4)
    add_data_structure_table(doc, data['assets']['fields'], "#42A5F5")
    doc.add_paragraph()

    doc.add_heading('3.2.2 Varlık - Detay Ekranı', level=3)
    doc.add_paragraph('Seçilen varlığın tüm bilgilerinin görüntülendiği ekrandır.')
    doc.add_heading('Ekran Görüntüsü:', level=4)
    add_screenshot(doc, SCREENSHOTS['asset-detail'])
    doc.add_heading('Veri Yapısı:', level=4)
    add_data_structure_table(doc, data['assets']['fields'], "#42A5F5")
    doc.add_paragraph()

    print("✍️  3.2.3 Varlık Oluştur...")
    doc.add_heading('3.2.3 Varlık Oluştur - Form Ekranı', level=3)
    doc.add_paragraph('Sisteme yeni varlık kaydı oluşturma ekranıdır.')
    doc.add_heading('Ekran Görüntüsü:', level=4)
    add_screenshot(doc, SCREENSHOTS['asset-create'])
    doc.add_heading('Veri Yapısı:', level=4)
    add_data_structure_table(doc, data['assets']['fields'], "#42A5F5")
    doc.add_paragraph()
    doc.add_page_break()

    # 3.3 Maintenance
    print("✍️  3.3 Bakım Yönetimi...")
    doc.add_heading('3.3 Bakım Yönetimi', level=2)
    doc.add_paragraph('Bakım yönetimi modülü, periyodik ve planlı bakım işlemlerinin planlanmasını sağlar.')
    doc.add_heading('Müşteri Gereksinimi:', level=3)
    doc.add_paragraph(data['maintenance']['requirement'])
    
    doc.add_heading('3.3.1 Bakım Yönetimi - Liste Ekranı', level=3)
    doc.add_paragraph('Planlanan ve devam eden bakım görevlerinin listelendiği ekrandır.')
    doc.add_heading('Ekran Görüntüsü:', level=4)
    add_screenshot(doc, SCREENSHOTS['maintenance'])
    doc.add_heading('Veri Yapısı:', level=4)
    add_data_structure_table(doc, data['maintenance']['fields'], "#66BB6A")
    doc.add_paragraph()

    doc.add_heading('3.3.2 Bakım Görevi - Detay Ekranı', level=3)
    doc.add_paragraph('Bakım görevinin detaylarının görüntülendiği ekrandır.')
    doc.add_heading('Ekran Görüntüsü:', level=4)
    add_screenshot(doc, SCREENSHOTS['maintenance-detail'])
    doc.add_heading('Veri Yapısı:', level=4)
    add_data_structure_table(doc, data['maintenance']['fields'], "#66BB6A")
    doc.add_paragraph()

    doc.add_heading('3.3.3 Bakım Görevi Oluştur - Form Ekranı', level=3)
    doc.add_paragraph('Yeni bakım görevi oluşturma ekranıdır.')
    doc.add_heading('Ekran Görüntüsü:', level=4)
    add_screenshot(doc, SCREENSHOTS['maintenance-create'])
    doc.add_heading('Veri Yapısı:', level=4)
    add_data_structure_table(doc, data['maintenance']['fields'], "#66BB6A")
    doc.add_paragraph()
    doc.add_page_break()

    # 3.4 Incidents
    print("✍️  3.4 Olay Yönetimi...")
    doc.add_heading('3.4 Olay Yönetimi', level=2)
    doc.add_paragraph('Olay yönetimi modülü, acil arızaların hızlı bildirilmesini sağlar.')
    doc.add_heading('Müşteri Gereksinimi:', level=3)
    doc.add_paragraph(data['incidents']['requirement'])
    
    doc.add_heading('3.4.1 Olay Yönetimi - Liste Ekranı', level=3)
    doc.add_paragraph('Bildirilen tüm olayların listelendiği ekrandır.')
    doc.add_heading('Ekran Görüntüsü:', level=4)
    add_screenshot(doc, SCREENSHOTS['incidents'])
    doc.add_heading('Veri Yapısı:', level=4)
    add_data_structure_table(doc, data['incidents']['fields'], "#EF5350")
    doc.add_paragraph()

    doc.add_heading('3.4.2 Olay - Detay Ekranı', level=3)
    doc.add_paragraph('Olayın detaylarının görüntülendiği ekrandır.')
    doc.add_heading('Ekran Görüntüsü:', level=4)
    add_screenshot(doc, SCREENSHOTS['incident-detail'])
    doc.add_heading('Veri Yapısı:', level=4)
    add_data_structure_table(doc, data['incidents']['fields'], "#EF5350")
    doc.add_paragraph()

    doc.add_heading('3.4.3 Olay Bildir - Form Ekranı', level=3)
    doc.add_paragraph('Yeni olay bildirimi oluşturma ekranıdır.')
    doc.add_heading('Ekran Görüntüsü:', level=4)
    add_screenshot(doc, SCREENSHOTS['incident-create'])
    doc.add_heading('Veri Yapısı:', level=4)
    add_data_structure_table(doc, data['incidents']['fields'], "#EF5350")
    doc.add_paragraph()

    # Save
    output_path = '/Users/caglarozyildirim/WebstormProjects/Deneme/MAN_Turkiye_Bakim_Yonetimi_COMPLETE_ANALYSIS.docx'
    doc.save(output_path)

    file_size = os.path.getsize(output_path)
    print(f"\n✅ Doküman oluşturuldu!")
    print(f"📄 Dosya adı: MAN_Turkiye_Bakim_Yonetimi_COMPLETE_ANALYSIS.docx")
    print(f"📊 Dosya boyutu: {file_size / (1024*1024):.2f} MB")

if __name__ == '__main__':
    create_document()
