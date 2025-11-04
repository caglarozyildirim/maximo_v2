#!/usr/bin/env python3
"""
Create Koleksiyon Mobilya Dealer Portal - Scope Document
Based on Patronlar Dünyası format
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime
import os

print("="*80)
print("CREATING KOLEKSIYON MOBILYA DEALER PORTAL - SCOPE DOCUMENT")
print("="*80 + "\n")

doc = Document()

# Set default font
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# Define custom styles
def setup_styles(doc):
    # Heading 1 style
    heading1 = doc.styles['Heading 1']
    heading1.font.name = 'Calibri'
    heading1.font.size = Pt(16)
    heading1.font.bold = True
    heading1.font.color.rgb = RGBColor(0, 0, 0)

    # Heading 3 style
    heading3 = doc.styles['Heading 3']
    heading3.font.name = 'Calibri'
    heading3.font.size = Pt(12)
    heading3.font.bold = True
    heading3.font.color.rgb = RGBColor(0, 0, 0)

    # Create whitespace-normal style
    try:
        ws_style = doc.styles.add_style('whitespace-normal', WD_STYLE_TYPE.PARAGRAPH)
        ws_style.font.name = 'Calibri'
        ws_style.font.size = Pt(11)
        ws_style.paragraph_format.line_spacing = 1.15
        ws_style.paragraph_format.space_after = Pt(8)
    except:
        pass  # Style already exists

setup_styles(doc)

# ============================================================================
# VERSION TABLE
# ============================================================================
print("📋 Creating version table...")

version_table = doc.add_table(rows=2, cols=4)
version_table.style = 'Light Grid Accent 1'

# Headers
headers = version_table.rows[0].cells
headers[0].text = "Version"
headers[1].text = "Date"
headers[2].text = "Author"
headers[3].text = "Description"

for cell in headers:
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.bold = True

# Version row
version_row = version_table.rows[1].cells
version_row[0].text = "1.0"
version_row[1].text = datetime.now().strftime('%d.%m.%Y')
version_row[2].text = "İş Analisti"
version_row[3].text = "Initial version - Koleksiyon Bayi Portalı Kapsam Dokümanı"

doc.add_paragraph()
doc.add_page_break()

# ============================================================================
# CONTENTS
# ============================================================================
print("📑 Creating table of contents...")

contents_para = doc.add_paragraph()
contents_run = contents_para.add_run("Contents")
contents_run.font.bold = True
contents_run.font.size = Pt(14)

doc.add_paragraph()

# TOC entries
toc_entries = [
    ("1", "PROJE AMACI", "2"),
    ("2", "PROJE KAPSAMI", "2"),
    ("2.1.1", "Bayi Kayıt ve Onay Sistemi", "2"),
    ("2.1.2", "SAP Fiyat Listesi Entegrasyonu", "3"),
    ("2.1.3", "PIM/CMS Ürün Bilgileri Entegrasyonu", "3"),
    ("2.1.4", "Para Birimi Bazlı Fiyat Gösterimi", "4"),
    ("2.1.5", "Çok Dilli Destek (TR/EN)", "4"),
]

for num, title, page in toc_entries:
    toc_para = doc.add_paragraph()

    if num.count('.') == 0:  # Level 1
        toc_para.paragraph_format.left_indent = Inches(0)
    else:  # Level 3
        toc_para.paragraph_format.left_indent = Inches(0.5)

    # Add number and title
    toc_para.add_run(f"{num}\t{title}\t{page}")

doc.add_page_break()

# ============================================================================
# 1. PROJE AMACI
# ============================================================================
print("📝 Section 1: PROJE AMACI...")

heading = doc.add_heading('PROJE AMACI', level=1)

para1 = doc.add_paragraph(style='whitespace-normal')
para1.add_run(
    "Koleksiyon Mobilya Bayi Portalı projesi kapsamında, mevcut SAP ERP sistemindeki "
    "ürün fiyat listelerine bayilerin güvenli ve kontrollü bir şekilde erişebileceği "
    "web tabanlı bir portal geliştirilecektir."
)

para2 = doc.add_paragraph(style='whitespace-normal')
para2.add_run(
    "Bu amaçla mevcut Butterfly DXP/PIM altyapısı kullanılarak SAP'den gerçek zamanlı "
    "fiyat verileri çekilerek bayilere sunulacak, manuel Excel gönderimi süreçleri "
    "ortadan kaldırılacak ve bayi deneyimi iyileştirilecektir."
)

para3 = doc.add_paragraph(style='whitespace-normal')
para3.add_run(
    "Butterfly DXP'nin esnek ve modüler yapısı kullanılarak bayi kayıt, onay, "
    "para birimi yetkilendirme ve fiyat listesi görüntüleme süreçleri otomatikleştirilecektir."
)

# ============================================================================
# 2. PROJE KAPSAMI
# ============================================================================
print("📝 Section 2: PROJE KAPSAMI...")

heading = doc.add_heading('PROJE KAPSAMI', level=1)

intro_para = doc.add_paragraph(style='whitespace-normal')
intro_para.add_run(
    "Butterfly DXP'nin güçlü teknolojisiyle mevcut manuel bayi fiyat paylaşım "
    "süreçlerini dijitalleştirerek bayilere self-servis erişim imkanı sağlanacaktır."
)

goal_para = doc.add_paragraph(style='whitespace-normal')
goal_para.add_run(
    "Sistemin temel amacı, SAP ERP sistemindeki fiyat listelerinin Butterfly PIM "
    "üzerinden bayilere güvenli bir şekilde sunulması, para birimi bazlı yetkilendirme "
    "ve çok dilli destek ile küresel bayi ağının etkin yönetimidir."
)

# ============================================================================
# 2.1.1 Bayi Kayıt ve Onay Sistemi
# ============================================================================
print("📝 Section 2.1.1: Bayi Kayıt ve Onay Sistemi...")

sub_heading = doc.add_heading('Bayi Kayıt ve Onay Sistemi', level=3)

para = doc.add_paragraph(style='whitespace-normal')
para.add_run(
    "Bayilerin portala erişimi için kayıt ve onay süreçleri Butterfly DXP'nin "
    "form oluşturma ve workflow yönetim kabiliyetleri kullanılarak geliştirilecektir."
)

# Registration system
reg_para = doc.add_paragraph(style='whitespace-normal')
reg_para.add_run(
    "Bayi Kayıt Formu: Butterfly DXP'nin form yapısı kullanılarak bayi kayıt formu "
    "oluşturulacaktır. Form alanları:"
)

registration_fields = [
    "Şirket Adı",
    "Vergi Numarası",
    "Yetkili Kişi Bilgileri",
    "E-posta Adresi",
    "Telefon",
    "Adres Bilgileri",
    "Ülke",
    "Dil Tercihi (TR/EN)"
]

for field in registration_fields:
    field_para = doc.add_paragraph(style='whitespace-normal')
    field_para.add_run(f"• {field}")
    field_para.paragraph_format.left_indent = Inches(0.5)

# Approval process
approval_para = doc.add_paragraph(style='whitespace-normal')
approval_para.add_run(
    "Onay Süreci: Yeni bayi kayıtları admin paneline düşecek ve Koleksiyon yetkilileri "
    "tarafından onaylanacaktır. Onay sonrası:"
)

approval_steps = [
    "Para birimi ataması yapılır (TRY, USD, EUR, vb.)",
    "E-posta ile bayi bilgilendirilir",
    "Bayi portala giriş yapabilir hale gelir",
    "Atanan para biriminde fiyat listelerine erişim sağlanır"
]

for step in approval_steps:
    step_para = doc.add_paragraph(style='whitespace-normal')
    step_para.add_run(f"• {step}")
    step_para.paragraph_format.left_indent = Inches(0.5)

# ============================================================================
# 2.1.2 SAP Fiyat Listesi Entegrasyonu
# ============================================================================
print("📝 Section 2.1.2: SAP Fiyat Listesi Entegrasyonu...")

sub_heading = doc.add_heading('SAP Fiyat Listesi Entegrasyonu', level=3)

para = doc.add_paragraph(style='whitespace-normal')
para.add_run(
    "SAP ERP sisteminden ürün fiyat listelerinin API üzerinden çekilmesi ve "
    "Butterfly PIM sistemi üzerinden bayilere sunulması sağlanacaktır."
)

# SAP Integration
sap_para = doc.add_paragraph(style='whitespace-normal')
sap_para.add_run(
    "SAP API Entegrasyonu: SAP OData/RFC protokolü kullanılarak gerçek zamanlı "
    "veri akışı sağlanacaktır. Çekilecek veriler:"
)

sap_data = [
    "MATNR (Malzeme Numarası): Ürün benzersiz kodu",
    "Teklif Kodu (35/36 fiyatlandırma kodları): Fiyat kırılım kodları",
    "3000'lu son kırılım kodları: Detaylı fiyat gösterimi",
    "Para birimi bazlı fiyatlar: TRY, USD, EUR, GBP, vb.",
    "Ürün açıklamaları: Türkçe ve İngilizce",
    "Model kodu: Ürün model bilgisi",
    "Tedarikçi bilgisi: Üretim bilgileri"
]

for data in sap_data:
    data_para = doc.add_paragraph(style='whitespace-normal')
    data_para.add_run(f"• {data}")
    data_para.paragraph_format.left_indent = Inches(0.5)

# Caching
cache_para = doc.add_paragraph(style='whitespace-normal')
cache_para.add_run(
    "Önbellekleme Stratejisi: Fiyat listeleri 15 dakika süre ile Redis cache'de "
    "tutulacak, performans optimize edilecektir."
)

# ============================================================================
# 2.1.3 PIM/CMS Ürün Bilgileri Entegrasyonu
# ============================================================================
print("📝 Section 2.1.3: PIM/CMS Ürün Bilgileri Entegrasyonu...")

sub_heading = doc.add_heading('PIM/CMS Ürün Bilgileri Entegrasyonu', level=3)

para = doc.add_paragraph(style='whitespace-normal')
para.add_run(
    "Butterfly PIM sisteminden ürün teknik bilgileri ve Butterfly CMS'den ürün "
    "görselleri çekilerek fiyat listesi ile birleştirilecektir."
)

# PIM Data
pim_para = doc.add_paragraph(style='whitespace-normal')
pim_para.add_run(
    "PIM'den Alınacak Veriler:"
)

pim_data = [
    "Ürün Kategorileri: Ana Grup, Alt Grup, Grup (örn: Yaşam > Sehpa > coffee table)",
    "Malzeme Bilgisi: mermer, tekstil, ahşap, vb.",
    "Boyut Bilgileri: Genişlik, Yükseklik, Çap (metrik ve ABD ölçü sistemleri)",
    "Ağırlık ve Hacim: Ürün fiziksel özellikleri",
    "Tasarımcı Bilgisi: Studio Kairos, Faruk Malhan, vb.",
    "Teknik Özellikler: Renk, kumaş, malzeme detayları"
]

for data in pim_data:
    data_para = doc.add_paragraph(style='whitespace-normal')
    data_para.add_run(f"• {data}")
    data_para.paragraph_format.left_indent = Inches(0.5)

# CMS Data
cms_para = doc.add_paragraph(style='whitespace-normal')
cms_para.add_run(
    "CMS'den Alınacak Veriler:"
)

cms_data = [
    "Ürün Görselleri: Yüksek çözünürlüklü ürün fotoğrafları",
    "PDF Teknik Dokümanları: Ürün spec dosyaları",
    "Katalog Bilgileri: Ürün kataloğu linkleri"
]

for data in cms_data:
    data_para = doc.add_paragraph(style='whitespace-normal')
    data_para.add_run(f"• {data}")
    data_para.paragraph_format.left_indent = Inches(0.5)

# Future Features
future_para = doc.add_paragraph(style='whitespace-normal')
future_para.add_run(
    "Gelecek Faz İçin Planlanan (Şu an beslenmeyecek):"
)

future_data = [
    "B (Büyük Metin): Detaylı ürün açıklamaları",
    "T (Teknik Resim): Teknik çizimler",
    "U (Uzun Metin): Genişletilmiş ürün bilgisi"
]

for data in future_data:
    data_para = doc.add_paragraph(style='whitespace-normal')
    data_para.add_run(f"• {data}")
    data_para.paragraph_format.left_indent = Inches(0.5)

# ============================================================================
# 2.1.4 Para Birimi Bazlı Fiyat Gösterimi
# ============================================================================
print("📝 Section 2.1.4: Para Birimi Bazlı Fiyat Gösterimi...")

sub_heading = doc.add_heading('Para Birimi Bazlı Fiyat Gösterimi', level=3)

para = doc.add_paragraph(style='whitespace-normal')
para.add_run(
    "Butterfly'ın kullanıcı yönetim sistemi ile entegre çalışarak her bayi için "
    "atanan para biriminde fiyat gösterimi yapılacaktır."
)

# Currency Management
currency_para = doc.add_paragraph(style='whitespace-normal')
currency_para.add_run(
    "Para Birimi Yönetimi:"
)

currency_features = [
    "Admin Tarafından Atama: Her bayi için admin panelinden para birimi seçilecek",
    "Desteklenen Para Birimleri: TRY (Türk Lirası), USD (Amerikan Doları), EUR (Euro), GBP (İngiliz Sterlini), vb.",
    "Otomatik Filtreleme: Bayi sadece kendi para birimindeki fiyatları görecek",
    "Para Birimi Güncelleme: Admin tarafından değiştirilebilir",
    "Format Uyumu: Sayısal formatlar para birimine göre ayarlanacak (. vs , kullanımı)"
]

for feature in currency_features:
    feature_para = doc.add_paragraph(style='whitespace-normal')
    feature_para.add_run(f"• {feature}")
    feature_para.paragraph_format.left_indent = Inches(0.5)

# Restrictions
restriction_para = doc.add_paragraph(style='whitespace-normal')
restriction_para.add_run(
    "Kısıtlamalar: Bayi başka para birimlerindeki fiyatları göremeyecek, sadece "
    "kendisine atanan para birimi ile çalışacaktır."
)

# ============================================================================
# 2.1.5 Çok Dilli Destek (TR/EN)
# ============================================================================
print("📝 Section 2.1.5: Çok Dilli Destek...")

sub_heading = doc.add_heading('Çok Dilli Destek (TR/EN)', level=3)

para = doc.add_paragraph(style='whitespace-normal')
para.add_run(
    "PIM sisteminde mevcut çok dilli altyapı kullanılarak Türkçe ve İngilizce "
    "arayüz desteği sağlanacaktır."
)

# Language Features
lang_para = doc.add_paragraph(style='whitespace-normal')
lang_para.add_run(
    "Dil Özellikleri:"
)

lang_features = [
    "Kullanıcı Dil Seçimi: Login sırasında veya profil ayarlarından dil değiştirilebilir",
    "Ürün Açıklamaları: Açıklama TR ve Açıklama EN alanları kullanılacak",
    "Arayüz Metinleri: Tüm butonlar, menüler ve mesajlar çok dilli olacak",
    "E-posta Bildirimleri: Kullanıcının diline göre e-posta gönderilecek",
    "PDF Dokümanları: Dil bazlı PDF spec dosyaları sunulacak"
]

for feature in lang_features:
    feature_para = doc.add_paragraph(style='whitespace-normal')
    feature_para.add_run(f"• {feature}")
    feature_para.paragraph_format.left_indent = Inches(0.5)

# Future Languages
future_lang_para = doc.add_paragraph(style='whitespace-normal')
future_lang_para.add_run(
    "Gelecek Genişleme: Sistem mimarisi diğer dillerin (FR, DE, IT, ES, vb.) "
    "kolayca eklenmesine imkan verecek şekilde tasarlanacaktır."
)

doc.add_page_break()

# ============================================================================
# SCREENSHOTS SECTION
# ============================================================================
print("📸 Adding screenshots section...")

screenshots_heading = doc.add_heading('EK: MEVCUT SİSTEM EKRAN GÖRÜNTÜLERİ', level=1)

screenshots_intro = doc.add_paragraph(style='whitespace-normal')
screenshots_intro.add_run(
    "Aşağıda mevcut Koleksiyon PIM sisteminin ekran görüntüleri yer almaktadır. "
    "Bu görseller SAP entegrasyonu ve ürün yönetimi için referans olarak kullanılacaktır."
)

doc.add_paragraph()

# Screenshot paths
SCREENSHOTS_DIR = "/Users/caglarozyildirim/Desktop/Şirketler/Koleksiyon/görseller"

screenshots = [
    {
        "file": "CleanShot 2025-10-15 at 08.24.42@2x.png",
        "title": "SAP Ürünleri Listesi",
        "description": "MATNR, Model Kodu, Tasarımcı, Malzeme Açıklaması alanlarını içeren ana liste görünümü"
    },
    {
        "file": "CleanShot 2025-10-15 at 08.35.38@2x.png",
        "title": "SAP Ürün Detay Formu",
        "description": "Müşteri ID, Malzeme Açıklaması (TR/EN), Üretici Parça Numarası bilgilerini içeren detay form"
    },
    {
        "file": "CleanShot 2025-10-15 at 08.36.02@2x.png",
        "title": "PIM Kategori Bilgileri",
        "description": "E-Dükkan, Kanal, Tedarikçi, Ana Grup, Alt Grup, Grup ve Ürün hiyerarşisi"
    },
    {
        "file": "CleanShot 2025-10-15 at 08.36.09@2x.png",
        "title": "PIM Malzeme ve Boyut",
        "description": "Malzeme, Kumaş, Tasarımcı, Model Kodu, Teklif Kodu, Boyut bilgileri"
    },
    {
        "file": "CleanShot 2025-10-15 at 08.36.13@2x.png",
        "title": "PIM Ölçü Bilgileri",
        "description": "Genişlik, Yükseklik, Çap (metrik ve ABD), Maksimum Beden bilgileri"
    },
    {
        "file": "CleanShot 2025-10-15 at 08.36.21@2x.png",
        "title": "PIM Hacim ve Açıklamalar",
        "description": "Ürün Hacmi, Ürün Ağırlığı, Montaj Süresi, Açıklama TR/EN alanları"
    },
    {
        "file": "CleanShot 2025-10-15 at 08.36.27@2x.png",
        "title": "PIM Z-Kategori Alanları",
        "description": "Z-Kanal (Home), Z-Tedarik (Production), Z-Ana Grup, Z-Alt Grup detayları"
    },
    {
        "file": "CleanShot 2025-10-15 at 08.36.33@2x.png",
        "title": "PIM Kumaş Metrajları",
        "description": "Z-Murun ölçü bilgileri ve KAYDET butonu"
    }
]

# Add each screenshot
for idx, screenshot in enumerate(screenshots, 1):
    screenshot_path = os.path.join(SCREENSHOTS_DIR, screenshot['file'])

    if os.path.exists(screenshot_path):
        print(f"  Adding screenshot {idx}/{len(screenshots)}: {screenshot['title']}")

        # Screenshot title
        screenshot_heading = doc.add_heading(f"{idx}. {screenshot['title']}", level=3)

        # Description
        desc_para = doc.add_paragraph(style='whitespace-normal')
        desc_para.add_run(screenshot['description'])

        doc.add_paragraph()

        # Add image
        try:
            doc.add_picture(screenshot_path, width=Inches(6.0))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()
        except Exception as e:
            print(f"    ⚠️  Error adding image: {e}")

# ============================================================================
# SAVE DOCUMENT
# ============================================================================

output_file = "/Users/caglarozyildirim/WebstormProjects/Deneme/Koleksiyon_Mobilya_Bayi_Portali_Kapsam_Dokumani.docx"
doc.save(output_file)

print("\n" + "="*80)
print("✅ DOCUMENT CREATED SUCCESSFULLY!")
print("="*80)
print(f"📁 File: {output_file}")
print(f"📸 Screenshots: {len(screenshots)}")
print(f"📄 Format: Patronlar Dünyası style")
print("="*80 + "\n")
