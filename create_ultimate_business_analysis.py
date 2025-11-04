#!/usr/bin/env python3
"""
MAN Türkiye Bakım Yönetimi - Kapsamlı İş Analizi Dokümanı Oluşturucu
"""

import os
import json
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime

# Dosya yolları
BASE_DIR = "/Users/caglarozyildirim/WebstormProjects/Deneme"
SCREENSHOTS_DIR = f"{BASE_DIR}/screenshots"
DATA_STRUCTURES_FILE = f"{BASE_DIR}/html_data_structures.json"
REQUIREMENTS_FILE = f"{BASE_DIR}/original_requirements.txt"
ANALYSIS_FILE = f"{BASE_DIR}/Bakim_Yonetim_Uygulamasi_Is_Analizi.md"
OUTPUT_FILE = f"{BASE_DIR}/MAN_Turkiye_Bakim_Yonetimi_ULTIMATE_BUSINESS_ANALYSIS.docx"

# Ekran görüntüleri mapping
SCREENSHOTS = {
    "index": {"file": "01_ana_sayfa.png", "title": "Ana Sayfa (Dashboard)", "path": "index.html"},
    "job-requests": {"file": "02_is_talepleri_liste.png", "title": "İş Talepleri - Liste", "path": "pages/job-requests.html"},
    "job-request-detail": {"file": "03_is_talepleri_detay.png", "title": "İş Talebi - Detay", "path": "pages/job-request-detail.html"},
    "job-request-create": {"file": "04_is_talepleri_olustur.png", "title": "İş Talebi - Oluştur", "path": "pages/job-request-create.html"},
    "assets": {"file": "05_varliklar_liste.png", "title": "Varlıklar - Liste", "path": "pages/assets.html"},
    "asset-detail": {"file": "06_varliklar_detay.png", "title": "Varlık - Detay", "path": "pages/asset-detail.html"},
    "maintenance": {"file": "07_bakim_liste.png", "title": "Bakım - Liste", "path": "pages/maintenance.html"},
    "maintenance-detail": {"file": "08_bakim_detay.png", "title": "Bakım - Detay", "path": "pages/maintenance-detail.html"},
    "maintenance-create": {"file": "09_bakim_olustur.png", "title": "Bakım - Oluştur", "path": "pages/maintenance-create.html"},
    "incidents": {"file": "10_olaylar_liste.png", "title": "Olaylar - Liste (GÜNCEL - Popup Kaldırıldı)", "path": "pages/incidents.html"},
    "incident-detail": {"file": "11_olaylar_detay.png", "title": "Olay - Detay", "path": "pages/incident-detail.html"},
    "incident-create": {"file": "12_olaylar_olustur.png", "title": "Olay - Oluştur (YENİ SAYFA!)", "path": "pages/incident-create.html"},
    "reports": {"file": "13_raporlar.png", "title": "Raporlar", "path": "pages/reports.html"},
}

def add_heading_with_style(doc, text, level=1):
    """Özel başlık stili ekle"""
    heading = doc.add_heading(text, level=level)
    if level == 1:
        heading.runs[0].font.color.rgb = RGBColor(226, 7, 20)  # MAN Kırmızı
        heading.runs[0].font.size = Pt(20)
        heading.runs[0].font.bold = True
    elif level == 2:
        heading.runs[0].font.color.rgb = RGBColor(226, 7, 20)
        heading.runs[0].font.size = Pt(16)
        heading.runs[0].font.bold = True
    elif level == 3:
        heading.runs[0].font.color.rgb = RGBColor(100, 100, 100)
        heading.runs[0].font.size = Pt(14)
    return heading

def add_page_break(doc):
    """Sayfa sonu ekle"""
    doc.add_page_break()

def add_screenshot_section(doc, key, data_structures):
    """Ekran görüntüsü ve veri yapısı bölümü ekle"""
    screenshot_info = SCREENSHOTS[key]
    screenshot_path = f"{SCREENSHOTS_DIR}/{screenshot_info['file']}"

    if not os.path.exists(screenshot_path):
        print(f"Uyarı: {screenshot_path} bulunamadı!")
        return

    # Başlık
    add_heading_with_style(doc, screenshot_info['title'], level=3)

    # Dosya yolu
    p = doc.add_paragraph()
    p.add_run("Dosya Yolu: ").bold = True
    p.add_run(f"/bakim-yonetim-app/{screenshot_info['path']}")

    # Ekran görüntüsü
    try:
        doc.add_picture(screenshot_path, width=Inches(6.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception as e:
        doc.add_paragraph(f"Ekran görüntüsü eklenemedi: {str(e)}")

    # Veri yapısı
    if key in data_structures and data_structures[key]:
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run("Veri Yapısı ve Sayfa Bileşenleri:").bold = True
        p.add_run().add_break()

        ds = data_structures[key]

        # Özet İstatistikler
        if 'summary' in ds:
            summary = ds['summary']
            doc.add_paragraph(f"Özet İstatistikler:", style='Heading 4')
            stats_text = (
                f"• Tablo Sayısı: {summary.get('table_count', 0)}\n"
                f"• Toplam Kolon: {summary.get('total_columns', 0)}\n"
                f"• Form Alanı: {summary.get('form_field_count', 0)}\n"
                f"• Badge Sayısı: {summary.get('badge_count', 0)}\n"
                f"• Kart Sayısı: {summary.get('card_count', 0)}\n"
                f"• Buton Sayısı: {summary.get('button_count', 0)}"
            )
            for line in stats_text.split('\n'):
                doc.add_paragraph(line, style='List Bullet')

        # Tablolar
        if 'tables' in ds and ds['tables'] and len(ds['tables']) > 0:
            doc.add_paragraph(f"Tablo Kolonları ({len(ds['tables'])} tablo):", style='Heading 4')
            for idx, table in enumerate(ds['tables'], 1):
                if 'columns' in table and table['columns']:
                    if len(ds['tables']) > 1:
                        doc.add_paragraph(f"Tablo {idx}:")
                    for col in table['columns']:
                        doc.add_paragraph(f"• {col}", style='List Bullet')

        # Form alanları
        if 'form_fields' in ds and ds['form_fields'] and len(ds['form_fields']) > 0:
            doc.add_paragraph(f"Form Alanları ({len(ds['form_fields'])} alan):", style='Heading 4')
            for form_field in ds['form_fields']:
                field_type = form_field.get('type', 'unknown')
                field_name = form_field.get('name', 'unnamed')
                field_label = form_field.get('label', '')

                if field_type == 'select' and 'options' in form_field:
                    option_count = form_field.get('option_count', 0)
                    doc.add_paragraph(f"• {field_name} (select, {option_count} seçenek)", style='List Bullet')
                    # İlk 5 seçeneği göster
                    if form_field.get('options'):
                        options_to_show = form_field['options'][:5]
                        for opt in options_to_show:
                            if opt.strip():
                                doc.add_paragraph(f"  - {opt}", style='List Bullet 2')
                        if len(form_field['options']) > 5:
                            doc.add_paragraph(f"  - ... ({len(form_field['options']) - 5} seçenek daha)", style='List Bullet 2')
                else:
                    label_text = f" - {field_label}" if field_label else ""
                    doc.add_paragraph(f"• {field_name} ({field_type}){label_text}", style='List Bullet')

        # Badge'ler
        if 'badges' in ds and ds['badges'] and len(ds['badges']) > 0:
            doc.add_paragraph(f"Durum Badge'leri ({len(ds['badges'])} badge):", style='Heading 4')
            badge_texts = list(set([b.get('text', '') for b in ds['badges'] if b.get('text')]))
            for badge_text in badge_texts:
                doc.add_paragraph(f"• {badge_text}", style='List Bullet')

        # Kartlar
        if 'stat_cards' in ds and ds['stat_cards'] and len(ds['stat_cards']) > 0:
            doc.add_paragraph(f"İstatistik Kartları ({len(ds['stat_cards'])} kart):", style='Heading 4')
            for card in ds['stat_cards']:
                title = card.get('title', '')
                value = card.get('value', '')
                if title:
                    card_text = f"• {title}"
                    if value:
                        card_text += f": {value}"
                    doc.add_paragraph(card_text, style='List Bullet')

        # Butonlar
        if 'action_buttons' in ds and ds['action_buttons'] and len(ds['action_buttons']) > 0:
            doc.add_paragraph(f"İşlem Butonları ({len(ds['action_buttons'])} buton):", style='Heading 4')
            for button in ds['action_buttons']:
                if button:
                    doc.add_paragraph(f"• {button}", style='List Bullet')

    doc.add_paragraph()

def create_document():
    """Ana doküman oluşturma fonksiyonu"""
    print("Doküman oluşturuluyor...")

    doc = Document()

    # Sayfa ayarları
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Veri yapılarını yükle
    data_structures = {}
    if os.path.exists(DATA_STRUCTURES_FILE):
        with open(DATA_STRUCTURES_FILE, 'r', encoding='utf-8') as f:
            data_structures = json.load(f)
        print(f"Veri yapıları yüklendi: {len(data_structures)} sayfa")

    # Gereksinim dokümanını oku
    requirements_text = ""
    if os.path.exists(REQUIREMENTS_FILE):
        with open(REQUIREMENTS_FILE, 'r', encoding='utf-8') as f:
            requirements_text = f.read()
        print(f"Gereksinim dokümanı yüklendi: {len(requirements_text)} karakter")

    # Analiz dokümanını oku
    analysis_text = ""
    if os.path.exists(ANALYSIS_FILE):
        with open(ANALYSIS_FILE, 'r', encoding='utf-8') as f:
            analysis_text = f.read()
        print(f"Analiz dokümanı yüklendi: {len(analysis_text)} karakter")

    # ==================== KAPAK SAYFASI ====================
    title = doc.add_heading('MAN TÜRKİYE', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.color.rgb = RGBColor(226, 7, 20)
    title.runs[0].font.size = Pt(36)

    subtitle = doc.add_heading('BAKIM YÖNETİMİ UYGULAMASI', 1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.color.rgb = RGBColor(0, 0, 0)
    subtitle.runs[0].font.size = Pt(24)

    doc.add_paragraph()
    doc.add_paragraph()

    subtitle2 = doc.add_heading('KAPSAMLI İŞ ANALİZİ DOKÜMANI', 2)
    subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle2.runs[0].font.size = Pt(18)

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    info_table = doc.add_table(rows=4, cols=2)
    info_table.style = 'Light Grid Accent 1'

    info_table.rows[0].cells[0].text = 'Versiyon'
    info_table.rows[0].cells[1].text = '2.0'
    info_table.rows[1].cells[0].text = 'Tarih'
    info_table.rows[1].cells[1].text = datetime.now().strftime('%d.%m.%Y')
    info_table.rows[2].cells[0].text = 'Proje Adı'
    info_table.rows[2].cells[1].text = 'Maintenance Management Application'
    info_table.rows[3].cells[0].text = 'Durum'
    info_table.rows[3].cells[1].text = 'GÜNCEL - Tüm Ekranlar ve Veri Yapıları Dahil'

    add_page_break(doc)

    # ==================== İÇİNDEKİLER ====================
    add_heading_with_style(doc, 'İÇİNDEKİLER', 1)

    toc_items = [
        "1. YÖNETİCİ ÖZETİ",
        "2. İŞ SÜREÇLERİ VE AKIŞLARI",
        "3. FONKSİYONEL GEREKSİNİMLER",
        "   3.1 İş Talebi Yönetimi",
        "   3.2 Varlık Yönetimi",
        "   3.3 Bakım Yönetimi",
        "   3.4 Olay Yönetimi",
        "   3.5 Diğer Modüller",
        "4. EKRAN TASARIMLARI VE VERİ YAPILARI",
        "   4.1 Dashboard",
        "   4.2 İş Talepleri Modülü",
        "   4.3 Varlık Yönetimi Modülü",
        "   4.4 Bakım Yönetimi Modülü",
        "   4.5 Olay Yönetimi Modülü",
        "   4.6 Raporlama Modülü",
        "5. VERİ MODELİ VE VERİ YAPISI",
        "6. ROLLER VE YETKİLENDİRME",
        "7. TEKNİK GEREKSİNİMLER",
        "8. OPERASYONEL GEREKSİNİMLER",
        "9. SONUÇ VE ÖNERİLER"
    ]

    for item in toc_items:
        doc.add_paragraph(item, style='List Number' if not item.startswith('   ') else 'List Bullet')

    add_page_break(doc)

    # ==================== 1. YÖNETİCİ ÖZETİ ====================
    add_heading_with_style(doc, '1. YÖNETİCİ ÖZETİ', 1)

    add_heading_with_style(doc, '1.1 Projenin Amacı ve Hedefleri', 2)
    doc.add_paragraph(
        "Bakım departmanı, tüm operasyonlarını Maximo uygulaması üzerinde yönetmektedir. "
        "IT departmanı, birden fazla uygulama teknolojisine sahip olup, lisans maliyetlerini "
        "azaltmak ve destek yeteneklerini artırmak amacıyla uygulamaları ortak bir teknolojiye "
        "taşıma çalışması yürütmektedir. Bu kapsamda, Maximo'nun artan lisans ve bakım maliyetleri "
        "nedeniyle başka bir platforma geçiş kararı alınmıştır."
    )

    doc.add_paragraph(
        "DIVA projesi 2027'de devreye girecek olup, bu tarihe kadar 2 yıllık geçici bir çözüme "
        "ihtiyaç vardır. Projenin basitliği ve maliyeti optimize etmek için bazı fonksiyonlar "
        "kapsam dışı bırakılırken, kullanıcı deneyimini artıracak ve geliştirme maliyeti düşük "
        "olan bazı yeni özellikler eklenmiştir."
    )

    doc.add_paragraph("Ana Hedefler:", style='Heading 3')
    goals = [
        "Bakım departmanının Maximo uygulamasındaki tüm operasyonlarını yeni platforma taşımak",
        "Lisans maliyetlerini azaltmak ve destek yeteneklerini artırmak",
        "2027'de DIVA projesinin devreye girmesine kadar 2 yıllık geçici çözüm sağlamak",
        "Süreçleri ve sorumlulukları takip etmek",
        "Hassas kararlar ve bilgiler için kayıt tutmak (onaylar ve maliyetler)",
        "Geliştirme maliyetlerini minimum düzeyde tutmak"
    ]
    for goal in goals:
        doc.add_paragraph(goal, style='List Bullet')

    add_heading_with_style(doc, '1.2 Kapsam', 2)
    doc.add_paragraph(
        "Bu proje, bakım departmanı ve ilgili departmanlar (maliyet kontrolü, muhasebe, "
        "lojistik ve depo) tarafından kullanılan Maximo uygulamasındaki fonksiyonları kapsamaktadır."
    )

    doc.add_paragraph("Ana Modüller:", style='Heading 3')
    modules = [
        "İş Talebi Yönetimi (Request Management)",
        "Sabit Varlık Yönetimi (Fixed Asset Management)",
        "Bakım Yönetimi - Düzenli ve Toplu Bakım",
        "Olay Yönetimi (Incident Management)",
        "Maliyet Merkezi Değişiklik Süreci",
        "Varlık Emekliliği (Asset Retirement)",
        "Raporlama (Reports)"
    ]
    for module in modules:
        doc.add_paragraph(module, style='List Bullet')

    add_page_break(doc)

    # ==================== 2. İŞ SÜREÇLERİ ====================
    add_heading_with_style(doc, '2. İŞ SÜREÇLERİ VE AKIŞLARI', 1)

    processes = [
        {
            "title": "2.1 İş Talebi Süreci (Job Request Workflow)",
            "desc": "İş talepleri oluşturulur, detaylandırılır, onay süreçlerinden geçirilir ve çözüme kavuşturulur.",
            "steps": [
                "Talep oluşturma",
                "Teknik onay (SL/Mühendis)",
                "İş yöneticisi onayı",
                "Maliyet onayı",
                "Çözüm süreci",
                "Çözüm onayı"
            ]
        },
        {
            "title": "2.2 Bakım Süreci (Maintenance Workflow)",
            "desc": "Düzenli ve planlı bakım işlemlerinin yönetildiği süreçtir.",
            "steps": [
                "Preventif bakım planlaması",
                "Bakım takvimi oluşturma",
                "Bakım ekiplerinin atanması",
                "Bakım görevlerinin yerine getirilmesi",
                "Bakım onayı"
            ]
        },
        {
            "title": "2.3 Varlık Girişi Süreci (Asset Entry Workflow)",
            "desc": "Yeni varlıkların sisteme kaydedilmesi sürecidir.",
            "steps": [
                "Varlık bilgilerinin girilmesi",
                "SAP entegrasyonu",
                "Varlık etiketleme",
                "Lokasyon atama"
            ]
        },
        {
            "title": "2.4 Olay Bildirimi Süreci (Incident Notification Workflow)",
            "desc": "Acil arıza ve olayların bildirilmesi ve yönetilmesi sürecidir.",
            "steps": [
                "Olay bildirimi",
                "Önceliklendirme (Kritik/Yüksek/Orta/Düşük)",
                "Müdahale ekibi atama",
                "Çözüm süreci",
                "Çözüm onayı"
            ]
        }
    ]

    for process in processes:
        add_heading_with_style(doc, process['title'], 2)
        doc.add_paragraph(process['desc'])
        doc.add_paragraph("İş Akışı Adımları:", style='Heading 3')
        for step in process['steps']:
            doc.add_paragraph(step, style='List Number')
        doc.add_paragraph()

    add_page_break(doc)

    # ==================== 3. FONKSİYONEL GEREKSİNİMLER ====================
    add_heading_with_style(doc, '3. FONKSİYONEL GEREKSİNİMLER', 1)

    # 3.1 İş Talebi Yönetimi
    add_heading_with_style(doc, '3.1 İş Talebi Yönetimi (Job Request)', 2)

    doc.add_paragraph(
        "Amaç: Talepleri toplamak, onay sürecini yönetmek ve tüm süreci takip etmek."
    )

    doc.add_paragraph("Hedefler:", style='Heading 3')
    jr_goals = [
        "Talepleri toplamak",
        "Onay sürecini yönetmek",
        "Çözüm sürecini yönetmek",
        "Kullanılan dolaylı malzemelerin tüketimini kaydetmek",
        "Raporlama"
    ]
    for goal in jr_goals:
        doc.add_paragraph(goal, style='List Bullet')

    doc.add_paragraph("Temel Veri Alanları:", style='Heading 3')
    jr_fields = [
        "Request Id - Talep benzersiz kimliği",
        "Request Title - Talep başlığı",
        "Request Description - Talep detaylı açıklaması",
        "Asset Id - Varlık kimliği",
        "Asset SAP Id - SAP sistemindeki varlık kimliği",
        "Location - Varlık lokasyonu",
        "Priority - Öncelik seviyesi (Urgent/High/Normal/Low)",
        "Creation Date Time - Oluşturulma tarihi ve saati",
        "Current Assignee - Mevcut atanan kişi",
        "Cost Value - Maliyet değeri",
        "Cost Currency - Maliyet para birimi",
        "Approval Status - Onay durumu",
        "Request Reason - Talep nedeni (OHS/Energy Saving/Environment/etc.)"
    ]
    for field in jr_fields:
        doc.add_paragraph(field, style='List Bullet')

    doc.add_paragraph("Durum Geçişleri:", style='Heading 3')
    jr_statuses = [
        "Business Approval → SL or Engineer Takeover",
        "Technical Approval → Cost Calculation",
        "Business Cost Approval → Solution Responsible Assignment",
        "Implementation → Solution approval",
        "Done / Rejected / Cancelled"
    ]
    for status in jr_statuses:
        doc.add_paragraph(status, style='List Bullet')

    add_page_break(doc)

    # 3.2 Varlık Yönetimi
    add_heading_with_style(doc, '3.2 Varlık Yönetimi (Asset Management)', 2)

    doc.add_paragraph(
        "Sabit varlıkların yaşam döngüsü boyunca yönetimi: varlık girişi, atama, "
        "transfer, ve emeklilik işlemlerini kapsar."
    )

    asset_modules = [
        {
            "title": "Varlık Girişi (Asset Entry)",
            "fields": [
                "Asset Id, Asset Title, Asset Description",
                "Asset Type, Asset Status",
                "SRM shopping bucket number, SRM number",
                "Location, Location sub unit 1, Location sub unit 2",
                "Producer name, Producer model name, Producer serial number",
                "Asset attachments",
                "Cost center",
                "SAP asset number (entegrasyon ile)"
            ]
        },
        {
            "title": "Varlık Atama (Asset Assignment)",
            "fields": [
                "Asset bilgileri (SAP ve Maintenance numaraları)",
                "Current Asset assigned user info",
                "New assignee user info",
                "Change Reason, Exchange date",
                "Approval fields (current holder, new holder manager, etc.)",
                "Rejection reason"
            ]
        },
        {
            "title": "Maliyet Merkezi Değişikliği (Cost Center Change)",
            "fields": [
                "Current cost center, New cost center",
                "Cost center responsible info",
                "Exchange date",
                "Approval statuses"
            ]
        },
        {
            "title": "Varlık Emekliliği (Asset Retirement)",
            "fields": [
                "Retirement Method (Scrapping/Donation/Selling/Missing/Stolen)",
                "Asset accounting info (IFRS ve Local)",
                "Selling Prize (if applicable)",
                "Physical retirement approval",
                "Scrapping manual attachment"
            ]
        }
    ]

    for module in asset_modules:
        doc.add_paragraph(module['title'], style='Heading 3')
        for field in module['fields']:
            doc.add_paragraph(field, style='List Bullet')
        doc.add_paragraph()

    add_page_break(doc)

    # 3.3 Bakım Yönetimi
    add_heading_with_style(doc, '3.3 Bakım Yönetimi (Maintenance)', 2)

    doc.add_paragraph(
        "Düzenli ve toplu bakım işlemlerinin planlanması ve yürütülmesi."
    )

    doc.add_paragraph("Bakım Planlama:", style='Heading 3')
    maintenance_planning = [
        "Periodic Maintenance Requirement Planning - Periyodik bakım ihtiyaç planlaması",
        "Period Planning - Dönemsel planlama (weekly/monthly/yearly)",
        "Task Planning - Görev listesi oluşturma",
        "Mass Maintenance Planning - Varlık grupları için toplu bakım",
        "Auto creation of duties - Ayın sonunda otomatik görev oluşturma"
    ]
    for item in maintenance_planning:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph("Bakım Yürütme:", style='Heading 3')
    maintenance_execution = [
        "Duty Assignment - Görev atama",
        "Visit Management - Ziyaret yönetimi",
        "Task Completion - Görev tamamlama",
        "Material Consumption - Malzeme tüketimi kaydı",
        "Approval/Rejection - Onay/Red"
    ]
    for item in maintenance_execution:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph("Temel Veri Alanları:", style='Heading 3')
    maintenance_fields = [
        "Duty id, Duty title, Duty description",
        "Maintenance date, SAP Asset number, Maintenance asset number",
        "Location, Task list number, Maintenance documents",
        "Maintenance responsible user info",
        "Visit start/finish date time",
        "Task checked, Task delay explanation",
        "Consumed materials (number, quantity, unit)"
    ]
    for field in maintenance_fields:
        doc.add_paragraph(field, style='List Bullet')

    add_page_break(doc)

    # 3.4 Olay Yönetimi
    add_heading_with_style(doc, '3.4 Olay Yönetimi (Incident Management)', 2)

    doc.add_paragraph(
        "Acil arıza ve olayların toplanması, çözümlenmesi ve gelecek analizler için "
        "veri kaydedilmesi."
    )

    doc.add_paragraph("Olay Toplama:", style='Heading 3')
    incident_collection = [
        "Create incident - Olay oluşturma",
        "SAP Asset bilgileri",
        "Incident Title ve Description",
        "Attachments - Ek dosyalar (fotoğraf vb.)",
        "Location bilgisi",
        "Alternative user to receive asset"
    ]
    for item in incident_collection:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph("Çözüm Süreci:", style='Heading 3')
    incident_solution = [
        "Asset handover to maintenance - Varlığın bakıma teslimi",
        "Solution responsible assignment - Çözüm sorumlusu atama",
        "Outsource service - Dış hizmet gönderme",
        "Delay explanation - Gecikme açıklaması",
        "Solution explanation - Çözüm açıklaması",
        "Consumed materials - Kullanılan malzemeler",
        "Solution approval from SL-TL - Teknik onay",
        "Solution approval from creator - Talep sahibi onayı",
        "Manager approval (if needed)"
    ]
    for item in incident_solution:
        doc.add_paragraph(item, style='List Bullet')

    add_page_break(doc)

    # ==================== 4. EKRAN TASARIMLARI ====================
    add_heading_with_style(doc, '4. EKRAN TASARIMLARI VE VERİ YAPILARI', 1)

    doc.add_paragraph(
        "Bu bölümde, uygulamanın tüm ekranlarının güncel görüntüleri ve detaylı veri yapıları "
        "yer almaktadır. Her ekran için sayfa yolu, görsel ve veri modeli sunulmuştur."
    )

    doc.add_paragraph()

    # 4.1 Dashboard
    add_heading_with_style(doc, '4.1 Dashboard (Ana Sayfa)', 2)
    add_screenshot_section(doc, "index", data_structures)
    add_page_break(doc)

    # 4.2 İş Talepleri Modülü
    add_heading_with_style(doc, '4.2 İş Talepleri Modülü', 2)
    doc.add_paragraph(
        "İş talepleri modülü, kullanıcıların bakım ve onarım taleplerini oluşturmasını, "
        "takip etmesini ve yönetmesini sağlar."
    )
    doc.add_paragraph()

    add_screenshot_section(doc, "job-requests", data_structures)
    add_page_break(doc)
    add_screenshot_section(doc, "job-request-detail", data_structures)
    add_page_break(doc)
    add_screenshot_section(doc, "job-request-create", data_structures)
    add_page_break(doc)

    # 4.3 Varlık Yönetimi Modülü
    add_heading_with_style(doc, '4.3 Varlık Yönetimi Modülü', 2)
    doc.add_paragraph(
        "Varlık yönetimi modülü, sabit varlıkların kaydını, takibini ve yaşam döngüsü "
        "yönetimini sağlar."
    )
    doc.add_paragraph()

    add_screenshot_section(doc, "assets", data_structures)
    add_page_break(doc)
    add_screenshot_section(doc, "asset-detail", data_structures)
    add_page_break(doc)

    # 4.4 Bakım Yönetimi Modülü
    add_heading_with_style(doc, '4.4 Bakım Yönetimi Modülü', 2)
    doc.add_paragraph(
        "Bakım yönetimi modülü, periyodik ve planlı bakım işlemlerinin yönetimini sağlar."
    )
    doc.add_paragraph()

    add_screenshot_section(doc, "maintenance", data_structures)
    add_page_break(doc)
    add_screenshot_section(doc, "maintenance-detail", data_structures)
    add_page_break(doc)
    add_screenshot_section(doc, "maintenance-create", data_structures)
    add_page_break(doc)

    # 4.5 Olay Yönetimi Modülü
    add_heading_with_style(doc, '4.5 Olay Yönetimi Modülü (GÜNCEL - ÖNEMLİ DEĞİŞİKLİKLER)', 2)

    p = doc.add_paragraph()
    p.add_run("ÖNEMLİ: ").bold = True
    p.add_run(
        "Bu modülde önemli güncellemeler yapılmıştır. Popup penceresi kaldırılmış ve "
        "yeni olay bildirimi için ayrı bir sayfa oluşturulmuştur. Bu değişiklik, kullanıcı "
        "deneyimini iyileştirmek ve daha kapsamlı form alanları sunmak amacıyla yapılmıştır."
    )
    doc.add_paragraph()

    add_screenshot_section(doc, "incidents", data_structures)
    add_page_break(doc)
    add_screenshot_section(doc, "incident-detail", data_structures)
    add_page_break(doc)
    add_screenshot_section(doc, "incident-create", data_structures)
    add_page_break(doc)

    # 4.6 Raporlama Modülü
    add_heading_with_style(doc, '4.6 Raporlama Modülü', 2)
    doc.add_paragraph(
        "Raporlama modülü, sistem genelindeki verilerin analiz edilmesi ve raporlanması "
        "için kullanılır."
    )
    doc.add_paragraph()

    add_screenshot_section(doc, "reports", data_structures)
    add_page_break(doc)

    # ==================== 5. VERİ MODELİ ====================
    add_heading_with_style(doc, '5. VERİ MODELİ VE VERİ YAPISI', 1)

    doc.add_paragraph(
        "Veri yapısı 37 adet tablo/sheet içermektedir. Bu tablolar, uygulamanın tüm "
        "fonksiyonel gereksinimlerini karşılamak üzere tasarlanmıştır."
    )

    doc.add_paragraph("Ana Veri Tabloları:", style='Heading 3')

    data_tables = [
        "Asset - Varlık ana tablosu",
        "Asset Group header - Varlık grup başlığı",
        "Asset Group item - Varlık grup öğeleri",
        "Asset Retirement - Varlık emekliliği",
        "Asset Status - Varlık durumları",
        "Asset Type - Varlık tipleri",
        "Assignment - Varlık atamaları",
        "Auth Group - Yetkilendirme grupları",
        "Comment - Yorumlar",
        "Consumed Materials - Tüketilen malzemeler",
        "Cost Center - Maliyet merkezleri",
        "Cost Center Change - Maliyet merkezi değişiklikleri",
        "Department - Departmanlar",
        "Document - Dokümanlar",
        "Incident - Olaylar",
        "Job Req. - İş talepleri",
        "Location - Lokasyonlar",
        "M. Duty - Bakım görevleri",
        "M. Req. - Bakım gereksinimleri",
        "M. Task - Bakım görevleri",
        "Priority - Öncelik seviyeleri",
        "User - Kullanıcılar",
        "User Group - Kullanıcı grupları",
        "Visit - Ziyaretler",
        "on behalf - Vekalet işlemleri"
    ]

    for table in data_tables:
        doc.add_paragraph(table, style='List Bullet')

    add_page_break(doc)

    # ==================== 6. ROLLER VE YETKİLENDİRME ====================
    add_heading_with_style(doc, '6. ROLLER VE YETKİLENDİRME', 1)

    doc.add_paragraph("Kullanıcı Rolleri:", style='Heading 3')

    roles = [
        {
            "role": "Talep Sahibi (Requester)",
            "permissions": "İş talebi oluşturabilir, kendi taleplerini görüntüleyebilir"
        },
        {
            "role": "Shift Leader / Mühendis",
            "permissions": "Teknik onay verebilir, talep atayabilir, çözüm sürecini yönetebilir"
        },
        {
            "role": "İş Yöneticisi (Business Manager)",
            "permissions": "Talep ve maliyet onayı verebilir"
        },
        {
            "role": "Bakım Teknisyeni",
            "permissions": "Bakım işlemlerini gerçekleştirebilir, malzeme kullanımını kaydedebilir"
        },
        {
            "role": "Varlık Yöneticisi",
            "permissions": "Varlık girişi, atama ve emeklilik işlemlerini yapabilir"
        },
        {
            "role": "Maliyet Kontrolörü",
            "permissions": "Maliyet raporlarına erişebilir, maliyet merkezi değişikliklerini onaylayabilir"
        },
        {
            "role": "Sistem Yöneticisi",
            "permissions": "Tüm yetkilere sahip, sistem yapılandırmasını yönetebilir"
        }
    ]

    for role in roles:
        p = doc.add_paragraph()
        p.add_run(f"{role['role']}: ").bold = True
        p.add_run(role['permissions'])

    add_page_break(doc)

    # ==================== 7. TEKNİK GEREKSİNİMLER ====================
    add_heading_with_style(doc, '7. TEKNİK GEREKSİNİMLER', 1)

    add_heading_with_style(doc, '7.1 Entegrasyonlar', 2)
    integrations = [
        "SAP Entegrasyonu: Varlık bilgileri, maliyet merkezi, muhasebe kayıtları",
        "Active Directory: Kullanıcı kimlik doğrulama ve yetkilendirme",
        "E-posta Sistemi: Bildirimler ve onay süreçleri",
        "IAM Integration: Gelişmiş yetkilendirme (opsiyonel)"
    ]
    for integration in integrations:
        doc.add_paragraph(integration, style='List Bullet')

    add_heading_with_style(doc, '7.2 Güvenlik Gereksinimleri', 2)
    security = [
        "Rol tabanlı erişim kontrolü (RBAC)",
        "Windows domain authentication",
        "PKI card approval support",
        "Audit logging (tüm işlemler kayıt altında)",
        "Veri şifreleme (transit ve rest)"
    ]
    for item in security:
        doc.add_paragraph(item, style='List Bullet')

    add_heading_with_style(doc, '7.3 Performans Gereksinimleri', 2)
    doc.add_paragraph("Hedef Kullanım Metrikleri:")
    performance = [
        "5000 kullanıcı (1500 bilgisayar erişimi)",
        "Günde 20 yeni iş talebi",
        "Günde 30 bakım görevi",
        "57,000 varlık yönetimi",
        "Aylık 150-200 yeni varlık eklenmesi"
    ]
    for item in performance:
        doc.add_paragraph(item, style='List Bullet')

    add_page_break(doc)

    # ==================== 8. OPERASYONEL GEREKSİNİMLER ====================
    add_heading_with_style(doc, '8. OPERASYONEL GEREKSİNİMLER', 1)

    operational_items = [
        {
            "title": "Kullanıcı Yönetimi",
            "items": [
                "User creation and deactivation",
                "User data update",
                "User authorization management",
                "Common user management (departman bazlı)",
                "On-behalf işlemleri"
            ]
        },
        {
            "title": "PKI Kart Onayları",
            "items": [
                "Bilgisayar erişimi olmayan kullanıcılar için PKI kart onayı",
                "On-behalf işlemlerinde PKI doğrulama",
                "Tüm PKI işlemlerinin loglanması"
            ]
        },
        {
            "title": "Yorum Yönetimi",
            "items": [
                "Her kayda yorum eklenebilir",
                "Yorumlar inactive yapılabilir (silinmez)",
                "Yorum alanları: Text, User info, Creation date time"
            ]
        },
        {
            "title": "Doküman Yönetimi",
            "items": [
                "Export: Excel, PDF, Word, Image, DWG",
                "Import: Excel toplu veri yükleme",
                "Sayfa bazlı özel dokümanlar (kullanım kılavuzları vb.)"
            ]
        },
        {
            "title": "Dil Desteği",
            "items": [
                "Türkçe",
                "İngilizce"
            ]
        }
    ]

    for item in operational_items:
        add_heading_with_style(doc, item['title'], 2)
        for sub_item in item['items']:
            doc.add_paragraph(sub_item, style='List Bullet')

    add_page_break(doc)

    # ==================== 9. SONUÇ VE ÖNERİLER ====================
    add_heading_with_style(doc, '9. SONUÇ VE ÖNERİLER', 1)

    add_heading_with_style(doc, '9.1 Kritik Başarı Faktörleri', 2)
    success_factors = [
        "Kullanıcı eğitiminin eksiksiz verilmesi",
        "SAP entegrasyonunun sorunsuz çalışması",
        "Mevcut Maximo verilerinin başarılı migrasyonu",
        "Süreç sahiplerinin aktif katılımı",
        "Düzenli geri bildirim ve iyileştirme döngüsü"
    ]
    for factor in success_factors:
        doc.add_paragraph(factor, style='List Bullet')

    add_heading_with_style(doc, '9.2 Riskler ve Öneriler', 2)

    risk_table = doc.add_table(rows=5, cols=3)
    risk_table.style = 'Light Grid Accent 1'

    risk_table.rows[0].cells[0].text = 'Risk'
    risk_table.rows[0].cells[1].text = 'Etki'
    risk_table.rows[0].cells[2].text = 'Öneri'

    risks = [
        ("Veri migrasyonu hataları", "Yüksek", "Pilot çalışma ve aşamalı geçiş"),
        ("Kullanıcı adaptasyonu", "Orta", "Yoğun eğitim ve süper kullanıcı desteği"),
        ("SAP entegrasyon sorunları", "Yüksek", "Erken test ve fallback planı"),
        ("Proje gecikmeleri", "Orta", "Agile metodoloji ve iteratif geliştirme")
    ]

    for i, risk in enumerate(risks, start=1):
        risk_table.rows[i].cells[0].text = risk[0]
        risk_table.rows[i].cells[1].text = risk[1]
        risk_table.rows[i].cells[2].text = risk[2]

    doc.add_paragraph()

    add_heading_with_style(doc, '9.3 Sonraki Adımlar', 2)
    next_steps = [
        "Detaylı teknik tasarım dokümantasyonunun hazırlanması",
        "Geliştirme ekibinin oluşturulması",
        "Sprint planlamasının yapılması",
        "Test ortamının hazırlanması",
        "Pilot kullanıcı grubunun belirlenmesi",
        "UAT (User Acceptance Testing) planlaması"
    ]
    for step in next_steps:
        doc.add_paragraph(step, style='List Number')

    # ==================== SON SAYFA ====================
    add_page_break(doc)

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    final_note = doc.add_paragraph()
    final_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = final_note.add_run("DOKÜMAN SONU")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(226, 7, 20)

    doc.add_paragraph()

    footer_text = doc.add_paragraph()
    footer_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_text.add_run(
        "Bu doküman, MAN Türkiye Bakım Yönetimi Uygulamasi projesi için hazırlanmış "
        "kapsamlı iş analizi dokümanıdır. Tüm gereksinim dokümanları, ekran tasarımları, "
        "veri yapıları ve güncel ekran görüntüleri içermektedir."
    )

    doc.add_paragraph()

    metadata = doc.add_paragraph()
    metadata.alignment = WD_ALIGN_PARAGRAPH.CENTER
    metadata.add_run(f"Oluşturulma Tarihi: {datetime.now().strftime('%d.%m.%Y %H:%M')}")

    # Dokümanı kaydet
    print(f"\nDoküman kaydediliyor: {OUTPUT_FILE}")
    doc.save(OUTPUT_FILE)

    # İstatistikler
    file_size = os.path.getsize(OUTPUT_FILE)
    print(f"\n{'='*60}")
    print("DOKÜMAN BAŞARIYLA OLUŞTURULDU!")
    print(f"{'='*60}")
    print(f"Dosya Adı: MAN_Turkiye_Bakim_Yonetimi_ULTIMATE_BUSINESS_ANALYSIS.docx")
    print(f"Dosya Konumu: {OUTPUT_FILE}")
    print(f"Dosya Boyutu: {file_size / (1024*1024):.2f} MB")
    print(f"Ekran Görüntüsü Sayısı: 13")
    print(f"Ana Bölüm Sayısı: 9")
    print(f"Modül Sayısı: 7")
    print(f"Oluşturulma Zamanı: {datetime.now().strftime('%d.%m.%Y %H:%M:%S')}")
    print(f"{'='*60}")

if __name__ == "__main__":
    create_document()
