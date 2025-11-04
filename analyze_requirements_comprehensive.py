#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Kapsamlı Bakım Yönetimi Sistemi Analizi
Desktop/new klasöründeki requirement dökümanlarını ve
mevcut HTML sayfalarını karşılaştırır
"""

import json
import os
from pathlib import Path
from docx import Document
import openpyxl

# Paths
NEW_DOCS_PATH = Path("/Users/caglarozyildirim/Desktop/new")
HTML_PATH = Path("/Users/caglarozyildirim/WebstormProjects/Deneme/bakim-yonetim-app/pages")

def read_docx(file_path):
    """DOCX dosyasını oku"""
    try:
        doc = Document(file_path)
        content = {
            'paragraphs': [],
            'tables': []
        }

        for para in doc.paragraphs:
            if para.text.strip():
                content['paragraphs'].append({
                    'text': para.text,
                    'style': para.style.name
                })

        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            content['tables'].append(table_data)

        return content
    except Exception as e:
        return {'error': str(e)}

def read_xlsx(file_path):
    """XLSX dosyasını oku"""
    try:
        wb = openpyxl.load_workbook(file_path)
        sheets_data = {}

        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            sheet_data = []

            for row in sheet.iter_rows(values_only=True):
                if any(cell for cell in row):
                    sheet_data.append(list(row))

            sheets_data[sheet_name] = sheet_data

        return sheets_data
    except Exception as e:
        return {'error': str(e)}

def count_html_files():
    """HTML dosyalarını say ve listele"""
    html_files = list(HTML_PATH.glob("*.html"))
    return {
        'count': len(html_files),
        'files': [f.name for f in html_files]
    }

def analyze_requirements():
    """Tüm requirement dökümanlarını analiz et"""

    analysis = {
        'requirement_doc': None,
        'screen_designs': None,
        'data_structure': None,
        'asset_retirement': None,
        'asset_assignment': None,
        'locations': None,
        'html_files': None
    }

    print("📄 Dökümanlar analiz ediliyor...")

    # 1. Main Requirement Document
    req_path = NEW_DOCS_PATH / "Maintenance Management Application Requirement Analysis (Version1).docx"
    if req_path.exists():
        print(f"  ✓ {req_path.name}")
        analysis['requirement_doc'] = read_docx(req_path)

    # 2. Screen Designs
    screen_path = NEW_DOCS_PATH / "Screen Designs.xlsx"
    if screen_path.exists():
        print(f"  ✓ {screen_path.name}")
        analysis['screen_designs'] = read_xlsx(screen_path)

    # 3. Data Structure
    data_path = NEW_DOCS_PATH / "Data Structure.xlsx"
    if data_path.exists():
        print(f"  ✓ {data_path.name}")
        analysis['data_structure'] = read_xlsx(data_path)

    # 4. Asset Retirement
    retire_path = NEW_DOCS_PATH / "Asset Retirement Printout.docx"
    if retire_path.exists():
        print(f"  ✓ {retire_path.name}")
        analysis['asset_retirement'] = read_docx(retire_path)

    # 5. Asset Assignment
    assign_path = NEW_DOCS_PATH / "Asset Assignment Form.docx"
    if assign_path.exists():
        print(f"  ✓ {assign_path.name}")
        analysis['asset_assignment'] = read_docx(assign_path)

    # 6. Locations
    loc_path = NEW_DOCS_PATH / "Locations and user groups.xlsx"
    if loc_path.exists():
        print(f"  ✓ {loc_path.name}")
        analysis['locations'] = read_xlsx(loc_path)

    # 7. HTML Files
    print(f"\n📱 HTML Dosyaları kontrol ediliyor...")
    analysis['html_files'] = count_html_files()
    print(f"  ✓ {analysis['html_files']['count']} HTML dosyası bulundu")

    return analysis

def create_comparison_report(analysis):
    """Karşılaştırma raporu oluştur"""

    report = []
    report.append("# BAKIM YÖNETİMİ SİSTEMİ - KAPSAMLI ANALİZ RAPORU\n")
    report.append("## Requirement Dökümanları vs Mevcut HTML Implementasyonu\n")
    report.append(f"**Analiz Tarihi:** {Path.cwd()}\n")
    report.append("---\n\n")

    # 1. Mevcut HTML Dosyaları
    report.append("## 1. MEVCUT HTML SAYFALAR\n")
    if analysis['html_files']:
        report.append(f"**Toplam:** {analysis['html_files']['count']} sayfa\n\n")
        for idx, file in enumerate(analysis['html_files']['files'], 1):
            report.append(f"{idx}. {file}\n")
    report.append("\n")

    # 2. Screen Designs Analizi
    if analysis['screen_designs']:
        report.append("## 2. EKRAN TASARIMLARI (Screen Designs.xlsx)\n\n")
        for sheet_name, data in analysis['screen_designs'].items():
            if 'error' not in analysis['screen_designs']:
                report.append(f"### {sheet_name}\n")
                report.append(f"**Satır Sayısı:** {len(data)}\n\n")

                if data and len(data) > 0:
                    # İlk 10 satırı göster
                    for row in data[:10]:
                        report.append(f"- {' | '.join([str(cell) if cell else '' for cell in row])}\n")
                report.append("\n")

    # 3. Data Structure Analizi
    if analysis['data_structure']:
        report.append("## 3. VERİ YAPILARI (Data Structure.xlsx)\n\n")
        for sheet_name, data in analysis['data_structure'].items():
            if 'error' not in analysis['data_structure']:
                report.append(f"### {sheet_name}\n")
                report.append(f"**Toplam Alan:** {len(data) if data else 0}\n\n")

                if data and len(data) > 1:
                    headers = data[0] if data else []
                    report.append("**Alanlar:**\n")
                    for row in data[1:11]:  # İlk 10 alan
                        report.append(f"- {' | '.join([str(cell) if cell else '' for cell in row])}\n")
                report.append("\n")

    # 4. Main Requirement Document
    if analysis['requirement_doc']:
        report.append("## 4. ANA REQUIREMENT DÖKÜMAN ANALİZİ\n\n")

        if 'paragraphs' in analysis['requirement_doc']:
            # Başlıkları bul
            report.append("### Ana Başlıklar:\n")
            for para in analysis['requirement_doc']['paragraphs'][:50]:
                if 'Heading' in para['style'] or len(para['text']) < 100:
                    report.append(f"- **{para['style']}**: {para['text'][:200]}\n")
            report.append("\n")

        if 'tables' in analysis['requirement_doc']:
            report.append(f"### Tablolar: {len(analysis['requirement_doc']['tables'])} adet\n\n")

    # 5. Asset Retirement
    if analysis['asset_retirement']:
        report.append("## 5. VARLIK HURDAYA ÇIKARMA (Asset Retirement)\n\n")
        if 'paragraphs' in analysis['asset_retirement']:
            report.append("**İçerik Özeti:**\n")
            for para in analysis['asset_retirement']['paragraphs'][:20]:
                if para['text']:
                    report.append(f"- {para['text'][:150]}\n")
        report.append("\n")

    # 6. Asset Assignment
    if analysis['asset_assignment']:
        report.append("## 6. VARLIK ZİMMETİ (Asset Assignment)\n\n")
        if 'paragraphs' in analysis['asset_assignment']:
            report.append("**İçerik Özeti:**\n")
            for para in analysis['asset_assignment']['paragraphs'][:20]:
                if para['text']:
                    report.append(f"- {para['text'][:150]}\n")
        report.append("\n")

    # 7. Locations
    if analysis['locations']:
        report.append("## 7. LOKASYONLAR VE KULLANICI GRUPLARI\n\n")
        for sheet_name, data in analysis['locations'].items():
            if 'error' not in analysis['locations']:
                report.append(f"### {sheet_name}\n")
                if data:
                    for row in data[:15]:
                        report.append(f"- {' | '.join([str(cell) if cell else '' for cell in row])}\n")
                report.append("\n")

    return ''.join(report)

def main():
    print("=" * 80)
    print("BAKIM YÖNETİMİ SİSTEMİ - KAPSAMLI ANALİZ")
    print("=" * 80)
    print()

    # Analiz yap
    analysis = analyze_requirements()

    # Rapor oluştur
    print("\n📊 Karşılaştırma raporu oluşturuluyor...")
    report = create_comparison_report(analysis)

    # Dosyaya kaydet
    output_path = Path("/Users/caglarozyildirim/WebstormProjects/Deneme/KAPSAMLI_ANALIZ_RAPORU.md")
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(report)

    print(f"✅ Rapor oluşturuldu: {output_path}")

    # JSON olarak da kaydet
    json_path = Path("/Users/caglarozyildirim/WebstormProjects/Deneme/analysis_data.json")

    # JSON serialization için temizle
    clean_analysis = {}
    for key, value in analysis.items():
        if value and isinstance(value, dict) and 'error' not in value:
            clean_analysis[key] = value
        elif value:
            clean_analysis[key] = str(value)[:500]  # Sadece özet

    with open(json_path, 'w', encoding='utf-8') as f:
        json.dump(clean_analysis, f, ensure_ascii=False, indent=2)

    print(f"✅ JSON data kaydedildi: {json_path}")
    print("\n✨ Analiz tamamlandı!")

if __name__ == "__main__":
    main()
