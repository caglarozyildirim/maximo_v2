#!/usr/bin/env python3
"""
Add Data Structure Tables to ALL Screen Sections
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json
import os

# Paths
base_path = "/Users/caglarozyildirim/WebstormProjects/Deneme"
input_file = f"{base_path}/MAN_Turkiye_Bakim_Yonetimi_FINAL_WITH_BUSINESS_RULES.docx"
output_file = f"{base_path}/MAN_Turkiye_Bakim_Yonetimi_COMPLETE_FINAL.docx"
data_structures_file = f"{base_path}/extracted_data_structures.json"

# Colors
MAN_RED = RGBColor(226, 7, 20)
MAN_GREEN = RGBColor(0, 168, 89)
MAN_DARK_GRAY = RGBColor(26, 26, 26)

# Load data structures
with open(data_structures_file, 'r', encoding='utf-8') as f:
    all_data_structures = json.load(f)

# Map data structures to modules
module_data = {
    'job_request': all_data_structures[0],      # 29 fields
    'asset': all_data_structures[3],            # 18 fields
    'maintenance': all_data_structures[6],      # 16 fields
    'incident': all_data_structures[9],         # 15 fields
}

def add_data_structure_table(doc, table_data, screen_type, module_name):
    """Add data structure table for a specific screen"""

    headers = table_data['headers']
    all_rows = table_data['rows']

    # Filter fields based on screen type
    if screen_type == 'list':
        # For list views: show summary fields only
        key_fields = ['id', 'title', 'status', 'priority', 'location', 'requestedBy', 'createdAt']
        rows = [row for row in all_rows if row[0] in key_fields]
        title = f"📊 {module_name} - Liste Görünümü Veri Alanları"
    elif screen_type == 'modal':
        # For modals: show creation fields (required fields)
        rows = [row for row in all_rows if row[3] == 'Evet']  # Only required fields
        title = f"📝 {module_name} - Oluşturma Formu Veri Alanları (Zorunlu)"
    elif screen_type == 'detail':
        # For detail views: show all fields
        rows = all_rows
        title = f"📋 {module_name} - Detay Sayfası Veri Alanları (Tüm Alanlar)"
    elif screen_type == 'create_form':
        # For full page create forms: show all input fields
        rows = all_rows
        title = f"📝 {module_name} - Tam Sayfa Form Veri Alanları (Tüm Alanlar)"
    elif screen_type == 'approval':
        # For approval dialogs: show approval related fields
        approval_fields = ['id', 'title', 'status', 'approvedBy', 'approvedAt', 'rejectedBy', 'rejectionReason']
        rows = [row for row in all_rows if row[0] in approval_fields]
        title = f"✅ {module_name} - Onay İşlemi Veri Alanları"
    elif screen_type == 'rejection':
        # For rejection dialogs: show rejection related fields
        rejection_fields = ['id', 'title', 'status', 'rejectedBy', 'rejectionReason', 'rejectionDate']
        rows = [row for row in all_rows if row[0] in rejection_fields]
        title = f"❌ {module_name} - Red İşlemi Veri Alanları"
    elif screen_type == 'rejected_state':
        # For rejected state views: show all fields in read-only
        rows = all_rows[:10]  # First 10 fields for brevity
        title = f"🔒 {module_name} - Reddedilmiş Durum Veri Alanları (Salt Okunur)"
    else:
        rows = all_rows
        title = f"📊 {module_name} - Veri Alanları"

    if not rows:
        return  # Skip if no rows to show

    # Add heading
    para = doc.add_paragraph()
    run = para.add_run(title)
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = MAN_GREEN

    # Add description
    if screen_type == 'list':
        doc.add_paragraph("Bu tablo, liste görünümünde gösterilen özet alanları içerir.")
    elif screen_type == 'modal':
        doc.add_paragraph("Bu tablo, hızlı oluşturma modalında zorunlu olan alanları içerir.")
    elif screen_type == 'detail':
        doc.add_paragraph("Bu tablo, detay sayfasında gösterilen tüm veri alanlarını içerir.")
    elif screen_type == 'create_form':
        doc.add_paragraph("Bu tablo, tam sayfa oluşturma formunda yer alan tüm veri alanlarını içerir.")
    elif screen_type == 'approval':
        doc.add_paragraph("Bu tablo, onay işlemi sırasında kullanılan veri alanlarını içerir.")
    elif screen_type == 'rejection':
        doc.add_paragraph("Bu tablo, red işlemi sırasında kullanılan veri alanlarını içerir.")
    elif screen_type == 'rejected_state':
        doc.add_paragraph("Bu tablo, reddedilmiş durumda gösterilen ana veri alanlarını içerir. Tüm alanlar salt okunurdur ve işlem butonları devre dışıdır.")

    # Create table
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'

    # Header row
    header_cells = table.rows[0].cells
    for i, header_text in enumerate(headers):
        header_cells[i].text = header_text
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = MAN_DARK_GRAY

    # Data rows
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, cell_text in enumerate(row_data):
            row_cells[i].text = str(cell_text)
            for paragraph in row_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)

    doc.add_paragraph()


print("\n" + "="*80)
print("ADDING DATA STRUCTURES TO ALL SCREENS")
print("="*80 + "\n")

# Load document
doc = Document(input_file)

# Define which sections need data structures
sections_config = [
    # Job Request
    {'title': '2.3. Liste Görünümü', 'type': 'list', 'module': 'job_request', 'name': 'İş Talepleri'},
    {'title': '2.4. Yeni Talep Oluşturma (Modal)', 'type': 'modal', 'module': 'job_request', 'name': 'İş Talebi'},
    {'title': '2.5. Detay Sayfası', 'type': 'detail', 'module': 'job_request', 'name': 'İş Talebi'},
    {'title': '2.6. ONAY İŞLEMİ', 'type': 'approval', 'module': 'job_request', 'name': 'İş Talebi'},
    {'title': '2.7. RED İŞLEMİ', 'type': 'rejection', 'module': 'job_request', 'name': 'İş Talebi'},
    {'title': '2.8. Reddedilmiş Durum', 'type': 'rejected_state', 'module': 'job_request', 'name': 'İş Talebi'},
    {'title': '2.9. Tam Sayfa Oluşturma Formu', 'type': 'create_form', 'module': 'job_request', 'name': 'İş Talebi'},

    # Asset Management
    {'title': '3.3. Liste Görünümü', 'type': 'list', 'module': 'asset', 'name': 'Varlık'},
    {'title': '3.4. Varlık Ekleme Modal', 'type': 'modal', 'module': 'asset', 'name': 'Varlık'},
    {'title': '3.5. Detay Sayfası - Onay Bekliyor Durumu', 'type': 'detail', 'module': 'asset', 'name': 'Varlık'},
    {'title': '3.6. Varlık Onay İşlemi', 'type': 'approval', 'module': 'asset', 'name': 'Varlık'},
    {'title': '3.7. Varlık Reddedilmiş Durumu', 'type': 'rejected_state', 'module': 'asset', 'name': 'Varlık'},
    {'title': '3.8. Tam Sayfa Oluşturma Formu', 'type': 'create_form', 'module': 'asset', 'name': 'Varlık'},

    # Maintenance
    {'title': '4.3. Liste Görünümü', 'type': 'list', 'module': 'maintenance', 'name': 'Bakım'},
    {'title': '4.4. Bakım Planı Modal', 'type': 'modal', 'module': 'maintenance', 'name': 'Bakım Planı'},
    {'title': '4.5. Detay Sayfası', 'type': 'detail', 'module': 'maintenance', 'name': 'Bakım'},
    {'title': '4.6. Bakım Planı Reddedilmiş Durumu', 'type': 'rejected_state', 'module': 'maintenance', 'name': 'Bakım'},
    {'title': '4.7. Tam Sayfa Oluşturma Formu', 'type': 'create_form', 'module': 'maintenance', 'name': 'Bakım Planı'},

    # Incident
    {'title': '5.3. Liste Görünümü', 'type': 'list', 'module': 'incident', 'name': 'Olay'},
    {'title': '5.4. Olay Bildirme Modal', 'type': 'modal', 'module': 'incident', 'name': 'Olay'},
    {'title': '5.5. Detay Sayfası', 'type': 'detail', 'module': 'incident', 'name': 'Olay'},
    {'title': '5.6. Çözüm Onayı', 'type': 'approval', 'module': 'incident', 'name': 'Olay'},
    {'title': '5.7. Olay Reddedilmiş Durumu', 'type': 'rejected_state', 'module': 'incident', 'name': 'Olay'},
    {'title': '5.8. Tam Sayfa Oluşturma Formu', 'type': 'create_form', 'module': 'incident', 'name': 'Olay'},
]

print(f"📋 Processing {len(sections_config)} sections...\n")

# Find all sections and mark insert positions
sections_found = []
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    for config in sections_config:
        if text == config['title'] and para.style.name == 'Heading 2':
            print(f"  ✓ Found: {text}")
            sections_found.append({
                'index': i,
                'config': config
            })

print(f"\n📊 Found {len(sections_found)} sections to process")

# Process in reverse order to maintain indices
for section_info in reversed(sections_found):
    para_index = section_info['index']
    config = section_info['config']

    print(f"\n🔧 Adding data structure to: {config['title']}")

    # Find next paragraph (after screenshot caption)
    # Look for the paragraph after the screenshot
    insert_position = para_index + 3  # After heading, screenshot, caption

    # Create temp document with data structure
    temp_doc = Document()
    add_data_structure_table(
        temp_doc,
        module_data[config['module']],
        config['type'],
        config['name']
    )

    # Insert into main document
    if insert_position < len(doc.paragraphs):
        insert_para = doc.paragraphs[insert_position]._element
        parent = insert_para.getparent()

        for element in temp_doc.element.body:
            parent.insert(parent.index(insert_para), element)

print("\n💾 Saving updated document...")
doc.save(output_file)

print("\n" + "="*80)
print("✅ DATA STRUCTURES ADDED TO ALL SCREENS!")
print(f"📁 Output: {output_file}")
print(f"📊 Total sections processed: {len(sections_found)}")
print("="*80 + "\n")
