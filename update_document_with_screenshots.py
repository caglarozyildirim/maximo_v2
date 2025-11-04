#!/usr/bin/env python3
"""
Word dosyasini yeni ekran goruntuleri ve veri yapilari ile guncelle
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json
import os
from datetime import datetime

# Paths
BASE_DIR = '/Users/caglarozyildirim/WebstormProjects/Deneme'
SOURCE_DOC = f'{BASE_DIR}/MAN_Turkiye_Bakim_Yonetimi_WITH_SCREENSHOTS.docx'
SCREENSHOTS_DIR = f'{BASE_DIR}/screenshots_updated'
DATA_FILE = f'{SCREENSHOTS_DIR}/data_structures.json'
OUTPUT_DOC = f'{BASE_DIR}/MAN_Turkiye_Bakim_Yonetimi_UPDATED_{datetime.now().strftime("%Y%m%d")}.docx'

# Sayfa bilgileri
PAGES = [
    {
        'name': '01_ana_sayfa',
        'title': 'Ana Sayfa - Dashboard',
        'description': 'Sistemin genel durumunu gosteren ozet ekran'
    },
    {
        'name': '02_is_talepleri_liste',
        'title': 'Is Talepleri - Liste',
        'description': 'Tum is taleplerinin listelenip filtrelendigi ekran'
    },
    {
        'name': '03_is_talepleri_detay',
        'title': 'Is Talepleri - Detay',
        'description': 'Is talebi detay bilgileri ve islem gecmisi'
    },
    {
        'name': '04_is_talepleri_olustur',
        'title': 'Is Talepleri - Yeni Talep Olustur',
        'description': 'Yeni is talebi olusturma formu'
    },
    {
        'name': '05_varliklar_liste',
        'title': 'Varlik Yonetimi - Liste',
        'description': 'Tum varliklarin listelenip filtrelendigi ekran'
    },
    {
        'name': '06_varliklar_detay',
        'title': 'Varlik Yonetimi - Detay',
        'description': 'Varlik detay bilgileri ve bakim gecmisi'
    },
    {
        'name': '07_bakim_liste',
        'title': 'Bakim Yonetimi - Liste',
        'description': 'Tum bakim planlarinin listelenip filtrelendigi ekran'
    },
    {
        'name': '08_bakim_detay',
        'title': 'Bakim Yonetimi - Detay',
        'description': 'Bakim plani detay bilgileri ve islem gecmisi'
    },
    {
        'name': '09_bakim_olustur',
        'title': 'Bakim Yonetimi - Plan Olustur',
        'description': 'Yeni bakim plani olusturma formu'
    },
    {
        'name': '10_olaylar_liste',
        'title': 'Olay Yonetimi - Liste (GUNCEL)',
        'description': 'Tum olaylarin listelenip filtrelendigi ekran - Popup kaldirildi, yeni olustur sayfasi eklendi'
    },
    {
        'name': '11_olaylar_detay',
        'title': 'Olay Yonetimi - Detay',
        'description': 'Olay detay bilgileri ve cozum adimlar'
    },
    {
        'name': '12_olaylar_olustur',
        'title': 'Olay Yonetimi - Yeni Olay Olustur (YENI!)',
        'description': 'Yeni olay kaydedilmesi icin tam sayfa form - Bu sayfa yeni eklendi'
    },
    {
        'name': '13_raporlar',
        'title': 'Raporlama',
        'description': 'Grafik ve tablolarla raporlama ekrani'
    }
]

def add_heading(doc, text, level=1):
    """Baslik ekle"""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading

def add_paragraph(doc, text, bold=False):
    """Paragraf ekle"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    return p

def add_screenshot(doc, screenshot_path, page_title):
    """Ekran goruntusu ekle"""
    if not os.path.exists(screenshot_path):
        print(f"  UYARI: Ekran goruntusu bulunamadi: {screenshot_path}")
        return False

    # Baslik
    p = doc.add_paragraph()
    run = p.add_run(page_title)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 51, 102)

    # Gorsel
    doc.add_picture(screenshot_path, width=Inches(6.5))

    # Bosluk
    doc.add_paragraph()

    return True

def format_data_structure(data):
    """Veri yapisini formatli metin olarak dondur"""
    lines = []

    # Tablolar
    if data.get('tables'):
        lines.append("TABLOLAR:")
        for idx, table in enumerate(data['tables'], 1):
            lines.append(f"  Tablo {idx}:")
            if table['headers']:
                lines.append(f"    Kolonlar: {', '.join(table['headers'])}")
            lines.append(f"    Satir Sayisi: {table['row_count']}")
        lines.append("")

    # Form Alanlari
    if data.get('form_fields'):
        lines.append("FORM ALANLARI:")
        for field in data['form_fields']:
            if field['type'] == 'input':
                lines.append(f"  - {field['label'] or field['name']}: {field['input_type']}")
            elif field['type'] == 'select':
                lines.append(f"  - {field['label'] or field['name']}: Secim Kutusu")
                if field['options']:
                    lines.append(f"    Secenekler: {', '.join(field['options'][:5])}")
            elif field['type'] == 'textarea':
                lines.append(f"  - {field['label'] or field['name']}: Metin Alani")
        lines.append("")

    # Badge'ler
    if data.get('badges'):
        lines.append("DURUM ETIKETLERI (BADGES):")
        badge_texts = [b['text'] for b in data['badges']]
        lines.append(f"  {', '.join(badge_texts)}")
        lines.append("")

    # Istatistik Kartlari
    if data.get('stats_cards'):
        lines.append("ISTATISTIK KARTLARI:")
        for stat in data['stats_cards']:
            if stat['title']:
                lines.append(f"  - {stat['title']}: {stat['value']}")
        lines.append("")

    return '\n'.join(lines)

def create_updated_document():
    """Guncellenmis dokumani olustur"""
    print("=" * 80)
    print("DOKUMAN GUNCELLENIYOR...")
    print("=" * 80)

    # Veri yapilarini yukle
    print("\n1. Veri yapilari yukleniyor...")
    with open(DATA_FILE, 'r', encoding='utf-8') as f:
        all_data = json.load(f)

    # Veri yapilarini dict'e donustur (name ile erisim icin)
    data_dict = {item['name']: item for item in all_data}

    # Yeni dokuman olustur
    print("\n2. Yeni dokuman olusturuluyor...")
    doc = Document(SOURCE_DOC)

    # Ekran goruntuleri bolumunu bul ve sil
    print("\n3. Eski ekran goruntuleri bolumu siliniyor...")
    screenshot_section_found = False
    paragraphs_to_remove = []

    for i, para in enumerate(doc.paragraphs):
        if 'EKRAN GORUNTULER' in para.text.upper() and 'VER' in para.text.upper():
            screenshot_section_found = True
            print(f"   Ekran goruntuleri bolumu bulundu: paragraf {i}")

            # Bu bolumden sonraki tum paragraflari isaretle
            # Bir sonraki ana bolum basligina kadar
            for j in range(i, len(doc.paragraphs)):
                # Eger yeni bir ana bolum basligina geldiysek dur
                if j > i and doc.paragraphs[j].style.name.startswith('Heading 1'):
                    break
                paragraphs_to_remove.append(j)

            break

    if not screenshot_section_found:
        print("   UYARI: Ekran goruntuleri bolumu bulunamadi, dokumana eklenecek")

    # Isaretlenen paragraflari sil (tersten)
    for idx in reversed(paragraphs_to_remove):
        p = doc.paragraphs[idx]._element
        p.getparent().remove(p)

    # Yeni ekran goruntuleri bolumunu ekle
    print("\n4. Yeni ekran goruntuleri ve veri yapilari ekleniyor...")

    # Ana baslik
    add_heading(doc, "EKRAN GORUNTULER VE VER YAPILARI", level=1)

    # Guncellenme bilgisi
    update_info = f"Bu bolum {datetime.now().strftime('%d.%m.%Y')} tarihinde guncellenmistir."
    add_paragraph(doc, update_info, bold=True)
    add_paragraph(doc, "Asagida sistemin tum ekranlari ve veri yapilari detayli olarak gosterilmektedir.")
    doc.add_paragraph()

    # Her sayfa icin ekran goruntusu ve veri yapisi ekle
    for idx, page in enumerate(PAGES, 1):
        print(f"\n   [{idx}/13] {page['title']}")

        # Alt baslik
        add_heading(doc, f"{idx}. {page['title']}", level=2)

        # Aciklama
        add_paragraph(doc, page['description'])
        doc.add_paragraph()

        # Ekran goruntusu
        screenshot_path = f"{SCREENSHOTS_DIR}/{page['name']}.png"
        if add_screenshot(doc, screenshot_path, "Ekran Goruntusu:"):
            print(f"      ✓ Ekran goruntusu eklendi")
        else:
            print(f"      ✗ Ekran goruntusu eklenemedi")

        # Veri yapisi
        if page['name'] in data_dict:
            data = data_dict[page['name']]
            formatted_data = format_data_structure(data)

            if formatted_data.strip():
                add_paragraph(doc, "Veri Yapisi:", bold=True)
                doc.add_paragraph(formatted_data)
                print(f"      ✓ Veri yapisi eklendi")
        else:
            print(f"      ✗ Veri yapisi bulunamadi")

        # Ayirici
        doc.add_paragraph()
        doc.add_paragraph("_" * 80)
        doc.add_paragraph()

    # Dokumani kaydet
    print(f"\n5. Dokuman kaydediliyor: {OUTPUT_DOC}")
    doc.save(OUTPUT_DOC)

    print("\n" + "=" * 80)
    print("ISLEM TAMAMLANDI!")
    print("=" * 80)
    print(f"\nGuncellemis dokuman: {OUTPUT_DOC}")
    print(f"Toplam sayfa sayisi: {len(PAGES)}")

    return OUTPUT_DOC

if __name__ == '__main__':
    create_updated_document()
