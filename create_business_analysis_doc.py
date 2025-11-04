#!/usr/bin/env python3
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def add_page_break(doc):
    doc.add_page_break()

def add_heading_with_style(doc, text, level):
    heading = doc.add_heading(text, level=level)
    if level == 1:
        heading.runs[0].font.color.rgb = RGBColor(0, 51, 153)
        heading.runs[0].font.size = Pt(20)
        heading.runs[0].font.bold = True
    elif level == 2:
        heading.runs[0].font.color.rgb = RGBColor(0, 102, 204)
        heading.runs[0].font.size = Pt(16)
        heading.runs[0].font.bold = True
    return heading

def add_table_with_border(doc, rows, cols):
    table = doc.add_table(rows=rows, cols=cols)
    table.style = 'Light Grid Accent 1'
    return table

def create_document():
    doc = Document()

    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # ====================================
    # BAŞLIK SAYFASI
    # ====================================
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('MAN TÜRKİYE\n')
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 153)

    run2 = title.add_run('BAKIM YÖNETİMİ SİSTEMİ\n\n')
    run2.font.size = Pt(24)
    run2.font.bold = True
    run2.font.color.rgb = RGBColor(0, 102, 204)

    run3 = title.add_run('İş Analizi ve Teknik Dokümantasyon')
    run3.font.size = Pt(18)
    run3.font.color.rgb = RGBColor(102, 102, 102)

    doc.add_paragraph('\n\n\n')

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run('Proje: ').font.bold = True
    info.add_run('Otobüs Üretim Tesisi Bakım Yönetimi\n')
    info.add_run('Versiyon: ').font.bold = True
    info.add_run('1.0\n')
    info.add_run('Tarih: ').font.bold = True
    info.add_run('Ekim 2025\n')

    add_page_break(doc)

    # ====================================
    # İÇİNDEKİLER
    # ====================================
    add_heading_with_style(doc, 'İçindekiler', 1)

    toc_items = [
        '1. Proje Özeti',
        '2. Sistem Genel Bakış',
        '3. Modül Detayları',
        '   3.1. Ana Sayfa Dashboard',
        '   3.2. İş Talepleri Yönetimi',
        '   3.3. Varlık Yönetimi',
        '   3.4. Bakım Yönetimi',
        '   3.5. Olay Yönetimi',
        '4. Veri Yapıları',
        '5. İş Akışları (Workflows)',
        '6. Ekran Görüntüleri',
        '7. Teknik Özellikler'
    ]

    for item in toc_items:
        doc.add_paragraph(item, style='List Number')

    add_page_break(doc)

    # ====================================
    # 1. PROJE ÖZETİ
    # ====================================
    add_heading_with_style(doc, '1. Proje Özeti', 1)

    doc.add_paragraph(
        'MAN Türkiye Bakım Yönetimi Sistemi, otobüs üretim tesisindeki tüm ekipman ve '
        'makinelerin bakım süreçlerini dijital olarak yöneten, iş taleplerini takip eden ve '
        'olayları (incidents) raporlayan entegre bir web uygulamasıdır.'
    )

    add_heading_with_style(doc, '1.1. Sistem Hedefleri', 2)
    goals = [
        'Üretim ekipmanlarının bakım süreçlerinin dijitalleştirilmesi',
        'İş taleplerinin merkezi takibi ve yönetimi',
        'Varlık (asset) yaşam döngüsü yönetimi',
        'Acil durum ve olay (incident) yönetimi',
        'Planlı ve reaktif bakım operasyonlarının koordinasyonu',
        'SAP ERP entegrasyonu için hazır veri yapıları'
    ]
    for goal in goals:
        doc.add_paragraph(goal, style='List Bullet')

    add_heading_with_style(doc, '1.2. Kullanıcı Rolleri', 2)

    table = add_table_with_border(doc, 5, 2)
    table.rows[0].cells[0].text = 'Rol'
    table.rows[0].cells[1].text = 'Yetkiler'

    roles = [
        ('Bakım Müdürü', 'Tüm modüllere tam erişim, onay yetkisi, raporlama'),
        ('Bakım Teknisyeni', 'İş talepleri ve bakım planlarını görüntüleme/güncelleme'),
        ('Üretim Sorumlusu', 'Olay bildirme, iş talebi oluşturma, varlık görüntüleme'),
        ('Sistem Yöneticisi', 'Kullanıcı yönetimi, sistem konfigürasyonu, veri yönetimi')
    ]

    for i, (role, permissions) in enumerate(roles, start=1):
        table.rows[i].cells[0].text = role
        table.rows[i].cells[1].text = permissions

    add_page_break(doc)

    # ====================================
    # 2. SİSTEM GENEL BAKIŞ
    # ====================================
    add_heading_with_style(doc, '2. Sistem Genel Bakış', 1)

    doc.add_paragraph(
        'Sistem 5 ana modülden oluşmaktadır:'
    )

    modules = [
        ('Ana Sayfa (Dashboard)', 'Sistemin özet görünümü ve hızlı erişim menüsü'),
        ('İş Talepleri Yönetimi', 'Bakım ve onarım taleplerinin yaratılması ve takibi'),
        ('Varlık Yönetimi', 'Üretim ekipmanlarının kayıt ve yaşam döngüsü yönetimi'),
        ('Bakım Yönetimi', 'Planlı bakım operasyonlarının programlanması'),
        ('Olay Yönetimi', 'Acil durumların ve olayların kaydı ve takibi')
    ]

    for module, desc in modules:
        p = doc.add_paragraph()
        p.add_run(f'{module}: ').font.bold = True
        p.add_run(desc)

    add_heading_with_style(doc, '2.2. Teknik Mimari', 2)

    doc.add_paragraph(
        'Sistem 3-katmanlı (3-tier) web uygulaması olarak tasarlanmıştır:'
    )

    tech_layers = [
        'Sunum Katmanı (Frontend): HTML5, CSS3, JavaScript (Vanilla)',
        'İş Mantığı Katmanı (Business Logic): JavaScript modülleri',
        'Veri Katmanı (Data Layer): JSON veri yapıları (Gelecekte SAP ERP entegrasyonu)'
    ]

    for layer in tech_layers:
        doc.add_paragraph(layer, style='List Bullet')

    add_page_break(doc)

    # ====================================
    # 3. MODÜL DETAYLARI
    # ====================================
    add_heading_with_style(doc, '3. Modül Detayları', 1)

    # ====================================
    # 3.1. ANA SAYFA DASHBOARD
    # ====================================
    add_heading_with_style(doc, '3.1. Ana Sayfa (Dashboard)', 2)

    add_heading_with_style(doc, 'Modül Açıklaması', 3)
    doc.add_paragraph(
        'Ana sayfa, sistemin özet görünümünü sunar. Kullanıcılar aktif iş taleplerini, '
        'yaklaşan bakımları, acil olayları ve genel istatistikleri görüntüleyebilir. '
        'Hızlı eylem butonları ile yeni talep/olay oluşturma ve varlık ekleme '
        'işlemleri tek tıkla yapılabilir.'
    )

    add_heading_with_style(doc, 'Fonksiyonel Gereksinimler', 3)

    dashboard_reqs = [
        'İstatistik kartları: Toplam iş talepleri, bekleyen onaylar, tamamlanan işler, aktif olaylar, üretim ekipmanı sayısı, planlı bakımlar',
        'Bekleyen İş Talepleri listesi (son 5 talep)',
        'Yaklaşan Bakımlar timeline görünümü',
        'Aktif Acil Olaylar listesi',
        'Son Aktiviteler akışı',
        'Hızlı İşlemler butonları: Yeni İş Talebi, Olay Bildir, Varlık Ekle'
    ]

    for req in dashboard_reqs:
        doc.add_paragraph(req, style='List Bullet')

    add_heading_with_style(doc, 'Ekran Elemanları', 3)

    table = add_table_with_border(doc, 7, 3)
    table.rows[0].cells[0].text = 'Eleman Adı'
    table.rows[0].cells[1].text = 'Tipi'
    table.rows[0].cells[2].text = 'Açıklama'

    dashboard_elements = [
        ('İstatistik Kartları', 'Widget', '6 adet özet kart (iş talepleri, onaylar, ekipman vb.)'),
        ('Bekleyen Talepler Tablosu', 'Tablo', 'Talep No, Başlık, Öncelik, Durum gösterir'),
        ('Yaklaşan Bakımlar', 'Timeline', 'Tarih, bakım adı, lokasyon, sorumlu ekip'),
        ('Aktif Olaylar', 'Alert List', 'Kritik olayların listesi, renk kodlu öncelik'),
        ('Hızlı İşlem Butonları', 'Button Group', '3 modal açma butonu'),
        ('Son Aktiviteler', 'Activity Feed', 'Zaman damgalı sistem olayları')
    ]

    for i, (name, type, desc) in enumerate(dashboard_elements, start=1):
        table.rows[i].cells[0].text = name
        table.rows[i].cells[1].text = type
        table.rows[i].cells[2].text = desc

    doc.add_paragraph('\n')
    doc.add_paragraph('Ekran Görüntüsü:', style='Intense Quote')

    screenshot_path = '/Users/caglarozyildirim/WebstormProjects/Deneme/screenshots_full/01_ana_sayfa.png'
    if os.path.exists(screenshot_path):
        doc.add_picture(screenshot_path, width=Inches(6.5))

    add_page_break(doc)

    # ====================================
    # 3.2. İŞ TALEPLERİ YÖNETİMİ
    # ====================================
    add_heading_with_style(doc, '3.2. İş Talepleri Yönetimi', 2)

    add_heading_with_style(doc, 'Modül Açıklaması', 3)
    doc.add_paragraph(
        'İş Talepleri modülü, bakım ve onarım taleplerinin oluşturulması, takip edilmesi '
        've yönetilmesi için kullanılır. Kullanıcılar talepleri filtreleyebilir, '
        'arayabilir ve detaylarını görüntüleyebilir.'
    )

    add_heading_with_style(doc, 'Fonksiyonel Gereksinimler', 3)

    job_reqs = [
        'İş talebi listesi görüntüleme (tablo formatında)',
        'Çoklu filtreleme: Durum, Öncelik, Kategori, Lokasyon',
        'Arama fonksiyonu: Talep No, Başlık, Açıklama bazlı',
        'Filtreleri temizleme butonu',
        'İş talebi detay görüntüleme',
        'Yeni iş talebi oluşturma (modal form)',
        'İstatistik özeti: Toplam, Bekleyen, Onay Bekleyen, İşlemde, Tamamlanan, Reddedilen'
    ]

    for req in job_reqs:
        doc.add_paragraph(req, style='List Bullet')

    add_heading_with_style(doc, 'Veri Alanları', 3)

    table = add_table_with_border(doc, 10, 4)
    table.rows[0].cells[0].text = 'Alan Adı'
    table.rows[0].cells[1].text = 'Veri Tipi'
    table.rows[0].cells[2].text = 'Zorunlu'
    table.rows[0].cells[3].text = 'Açıklama'

    job_fields = [
        ('Talep No', 'String', 'Evet', 'JR-2025-XXX formatında otomatik üretilir'),
        ('Başlık', 'String', 'Evet', 'İş talebinin kısa açıklaması'),
        ('Açıklama', 'Text', 'Evet', 'Detaylı açıklama'),
        ('Kategori', 'Enum', 'Evet', 'HVAC, Elektrik, Mekanik, Bina, IT'),
        ('Lokasyon', 'Enum', 'Evet', 'Ana Tesis, Üretim, Ofis, Depo'),
        ('Öncelik', 'Enum', 'Evet', 'Kritik, Yüksek, Orta, Düşük'),
        ('Durum', 'Enum', 'Evet', 'Beklemede, Onayda, İşlemde, Tamamlandı, Reddedildi'),
        ('Oluşturan', 'String', 'Evet', 'Kullanıcı adı'),
        ('Tarih', 'Date', 'Evet', 'Oluşturulma tarihi')
    ]

    for i, (field, type, req, desc) in enumerate(job_fields, start=1):
        table.rows[i].cells[0].text = field
        table.rows[i].cells[1].text = type
        table.rows[i].cells[2].text = req
        table.rows[i].cells[3].text = desc

    add_heading_with_style(doc, 'Filtreleme Kriterleri', 3)

    filter_table = add_table_with_border(doc, 5, 2)
    filter_table.rows[0].cells[0].text = 'Filtre Adı'
    filter_table.rows[0].cells[1].text = 'Değerler'

    filters = [
        ('Durum', 'Tümü, Beklemede, Onayda, İşlemde, Tamamlandı, Reddedildi'),
        ('Öncelik', 'Tümü, Kritik, Yüksek, Orta, Düşük'),
        ('Kategori', 'Tümü, Elektrik, Mekanik, Bina, IT, HVAC'),
        ('Lokasyon', 'Tümü, Ana Tesis, Üretim, Ofis, Depo')
    ]

    for i, (fname, values) in enumerate(filters, start=1):
        filter_table.rows[i].cells[0].text = fname
        filter_table.rows[i].cells[1].text = values

    doc.add_paragraph('\n')
    doc.add_paragraph('Ekran Görüntüsü:', style='Intense Quote')

    screenshot_path = '/Users/caglarozyildirim/WebstormProjects/Deneme/screenshots_full/02_is_talepleri_liste.png'
    if os.path.exists(screenshot_path):
        doc.add_picture(screenshot_path, width=Inches(6.5))

    add_page_break(doc)

    # ====================================
    # 3.3. VARLIK YÖNETİMİ
    # ====================================
    add_heading_with_style(doc, '3.3. Varlık Yönetimi', 2)

    add_heading_with_style(doc, 'Modül Açıklaması', 3)
    doc.add_paragraph(
        'Varlık Yönetimi modülü, otobüs üretim tesisindeki tüm ekipman ve makinelerin '
        'kaydını tutar. Her varlığın teknik özellikleri, lokasyonu, bakım geçmişi ve '
        'mali bilgileri saklanır. SAP ID entegrasyonu ile ERP sistemine bağlanabilir.'
    )

    add_heading_with_style(doc, 'Fonksiyonel Gereksinimler', 3)

    asset_reqs = [
        'Varlık listesi görüntüleme (tablo formatında)',
        'Varlık arama: SAP ID, Ad, Seri No ile',
        'Varlık detay sayfası',
        'Yeni varlık ekleme (modal form)',
        'İstatistik özeti: Toplam Ekipman, Aktif, Bakımda, Arızalı',
        'Her varlık için: Temel bilgiler, Lokasyon, Mali bilgiler, Bakım geçmişi, İlgili iş talepleri'
    ]

    for req in asset_reqs:
        doc.add_paragraph(req, style='List Bullet')

    add_heading_with_style(doc, 'Veri Alanları', 3)

    table = add_table_with_border(doc, 18, 4)
    table.rows[0].cells[0].text = 'Alan Adı'
    table.rows[0].cells[1].text = 'Veri Tipi'
    table.rows[0].cells[2].text = 'Zorunlu'
    table.rows[0].cells[3].text = 'Açıklama'

    asset_fields = [
        ('Varlık ID', 'String', 'Evet', 'AST-XXX formatında otomatik üretilir'),
        ('SAP ID', 'String', 'Evet', 'SAP sisteminden gelen ID (10 haneli)'),
        ('Varlık Adı', 'String', 'Evet', 'Ekipman adı'),
        ('Kategori', 'Enum', 'Evet', 'Üretim Makinesi, Robotik Sistem, Boya Sistemi, Taşıma, Altyapı'),
        ('Üretici', 'String', 'Evet', 'Üretici firma adı (örn: HAAS, ABB, DÜRR)'),
        ('Model', 'String', 'Evet', 'Model numarası'),
        ('Seri No', 'String', 'Evet', 'Seri numarası'),
        ('Lokasyon', 'Enum', 'Evet', 'Şasi Üretim, Gövde Montaj, Boya Tesisi, Final Montaj, Depo'),
        ('Alt Lokasyon', 'String', 'Hayır', 'Detaylı lokasyon bilgisi'),
        ('Maliyet Merkezi', 'String', 'Evet', 'SAP maliyet merkezi kodu'),
        ('Durum', 'Enum', 'Evet', 'Aktif, Bakımda, Arızalı, Emekli'),
        ('Zimmetli', 'String', 'Hayır', 'Sorumlu kişi/ekip'),
        ('Satın Alma Tarihi', 'Date', 'Evet', 'Alım tarihi'),
        ('Defter Değeri', 'Decimal', 'Evet', 'TRY cinsinden değer'),
        ('Son Bakım', 'Date', 'Hayır', 'Son bakım tarihi'),
        ('Sonraki Bakım', 'Date', 'Hayır', 'Planlanan bakım tarihi'),
        ('Açıklama', 'Text', 'Hayır', 'Varlık hakkında detaylı bilgi')
    ]

    for i, (field, type, req, desc) in enumerate(asset_fields, start=1):
        table.rows[i].cells[0].text = field
        table.rows[i].cells[1].text = type
        table.rows[i].cells[2].text = req
        table.rows[i].cells[3].text = desc

    doc.add_paragraph('\n')
    doc.add_paragraph('Ekran Görüntüsü - Liste:', style='Intense Quote')

    screenshot_path = '/Users/caglarozyildirim/WebstormProjects/Deneme/screenshots_full/03_varlik_yonetimi_liste.png'
    if os.path.exists(screenshot_path):
        doc.add_picture(screenshot_path, width=Inches(6.5))

    doc.add_paragraph('\n')
    doc.add_paragraph('Ekran Görüntüsü - Detay:', style='Intense Quote')

    screenshot_path = '/Users/caglarozyildirim/WebstormProjects/Deneme/screenshots_full/04_varlik_detay.png'
    if os.path.exists(screenshot_path):
        doc.add_picture(screenshot_path, width=Inches(6.5))

    add_page_break(doc)

    # ====================================
    # 3.4. BAKIM YÖNETİMİ
    # ====================================
    add_heading_with_style(doc, '3.4. Bakım Yönetimi', 2)

    add_heading_with_style(doc, 'Modül Açıklaması', 3)
    doc.add_paragraph(
        'Bakım Yönetimi modülü, planlı bakım operasyonlarının programlanması ve takip '
        'edilmesi için kullanılır. Önleyici bakım, rutin kontrol, düzeltici bakım ve '
        'tahmine dayalı bakım planları oluşturulabilir.'
    )

    add_heading_with_style(doc, 'Fonksiyonel Gereksinimler', 3)

    maint_reqs = [
        'Yaklaşan bakımlar listesi (timeline görünümü)',
        'Aktif bakım işlemleri',
        'Yeni bakım planı oluşturma (modal form)',
        'Bakım tipi seçimi: Önleyici, Rutin, Düzeltici, Tahmine Dayalı',
        'Varlık bazlı planlama',
        'Ekip atama',
        'İstatistik özeti: Toplam Plan, Bu Hafta, Tamamlanan, Geciken'
    ]

    for req in maint_reqs:
        doc.add_paragraph(req, style='List Bullet')

    add_heading_with_style(doc, 'Veri Alanları', 3)

    table = add_table_with_border(doc, 9, 4)
    table.rows[0].cells[0].text = 'Alan Adı'
    table.rows[0].cells[1].text = 'Veri Tipi'
    table.rows[0].cells[2].text = 'Zorunlu'
    table.rows[0].cells[3].text = 'Açıklama'

    maint_fields = [
        ('Plan No', 'String', 'Evet', 'MNT-2025-XXX formatında otomatik'),
        ('Bakım Başlığı', 'String', 'Evet', 'Bakım açıklaması'),
        ('Varlık ID', 'String', 'Evet', 'İlgili varlık referansı'),
        ('Bakım Tipi', 'Enum', 'Evet', 'Önleyici, Rutin, Düzeltici, Tahmine Dayalı'),
        ('Planlanan Tarih', 'Date', 'Evet', 'Bakım tarihi'),
        ('Tahmini Süre', 'String', 'Evet', 'Örn: "4 saat"'),
        ('Sorumlu Ekip', 'Enum', 'Evet', 'Bakım Ekibi A/B, Robotik Ekip, Boya Ekibi, Teknik Ekip'),
        ('Açıklama', 'Text', 'Hayır', 'Detaylı bakım açıklaması')
    ]

    for i, (field, type, req, desc) in enumerate(maint_fields, start=1):
        table.rows[i].cells[0].text = field
        table.rows[i].cells[1].text = type
        table.rows[i].cells[2].text = req
        table.rows[i].cells[3].text = desc

    doc.add_paragraph('\n')
    doc.add_paragraph('Ekran Görüntüsü:', style='Intense Quote')

    screenshot_path = '/Users/caglarozyildirim/WebstormProjects/Deneme/screenshots_full/05_bakim_yonetimi.png'
    if os.path.exists(screenshot_path):
        doc.add_picture(screenshot_path, width=Inches(6.5))

    add_page_break(doc)

    # ====================================
    # 3.5. OLAY YÖNETİMİ
    # ====================================
    add_heading_with_style(doc, '3.5. Olay Yönetimi (Incident Management)', 2)

    add_heading_with_style(doc, 'Modül Açıklaması', 3)
    doc.add_paragraph(
        'Olay Yönetimi modülü, acil durumların ve beklenmedik olayların kaydedilmesi '
        've takip edilmesi için kullanılır. Kritik üretim durmaları, ekipman arızaları '
        've güvenlik olayları bu modül üzerinden yönetilir.'
    )

    add_heading_with_style(doc, 'Fonksiyonel Gereksinimler', 3)

    incident_reqs = [
        'Aktif olaylar listesi',
        'Acil olay bildirme (modal form)',
        'Öncelik bazlı renk kodlama (Kritik: Kırmızı, Yüksek: Turuncu)',
        'Geçen süre gösterimi',
        'Varlık ilişkilendirme (opsiyonel)',
        'İstatistik özeti: Aktif Olaylar, Kritik, Yüksek, Bu Ay Kapatılan',
        'Olay numarası otomatik üretimi (INC-2025-XXX)'
    ]

    for req in incident_reqs:
        doc.add_paragraph(req, style='List Bullet')

    add_heading_with_style(doc, 'Veri Alanları', 3)

    table = add_table_with_border(doc, 8, 4)
    table.rows[0].cells[0].text = 'Alan Adı'
    table.rows[0].cells[1].text = 'Veri Tipi'
    table.rows[0].cells[2].text = 'Zorunlu'
    table.rows[0].cells[3].text = 'Açıklama'

    incident_fields = [
        ('Olay No', 'String', 'Evet', 'INC-2025-XXX formatında otomatik'),
        ('Başlık', 'String', 'Evet', 'Olay başlığı'),
        ('Açıklama', 'Text', 'Evet', 'Olay detayları'),
        ('Öncelik', 'Enum', 'Evet', 'Düşük, Orta, Yüksek, Kritik'),
        ('Lokasyon', 'Enum', 'Evet', 'Şasi Üretim, Gövde Montaj, Boya Tesisi, vb.'),
        ('İlgili Varlık', 'String', 'Hayır', 'Varlık ID referansı'),
        ('Başlangıç Zamanı', 'DateTime', 'Evet', 'Olay başlangıç zamanı')
    ]

    for i, (field, type, req, desc) in enumerate(incident_fields, start=1):
        table.rows[i].cells[0].text = field
        table.rows[i].cells[1].text = type
        table.rows[i].cells[2].text = req
        table.rows[i].cells[3].text = desc

    doc.add_paragraph('\n')
    doc.add_paragraph('Ekran Görüntüsü:', style='Intense Quote')

    screenshot_path = '/Users/caglarozyildirim/WebstormProjects/Deneme/screenshots_full/06_olay_yonetimi.png'
    if os.path.exists(screenshot_path):
        doc.add_picture(screenshot_path, width=Inches(6.5))

    add_page_break(doc)

    # ====================================
    # 4. VERİ YAPILARI
    # ====================================
    add_heading_with_style(doc, '4. Veri Yapıları (Data Structures)', 1)

    doc.add_paragraph(
        'Sistemde kullanılan ana veri yapıları aşağıdaki gibidir. Tüm veriler JSON formatında '
        'saklanmakta ve gelecekte SAP ERP sistemine entegre edilecektir.'
    )

    add_heading_with_style(doc, '4.1. Varlık (Asset) Veri Yapısı', 2)

    asset_json = '''
{
  "id": "AST-001",
  "sapId": "100234567",
  "name": "CNC Torna Makinesi - Şasi İşleme",
  "category": "Üretim Makinesi",
  "manufacturer": "HAAS",
  "model": "ST-30",
  "serialNo": "HAAS-ST30-2021-1234",
  "location": "Şasi Üretim Hattı",
  "subLocation": "Alan A1",
  "costCenter": "CC-PROD-001",
  "status": "active",
  "statusText": "Aktif",
  "assignedTo": "Üretim Ekibi A",
  "purchaseDate": "2021-03-15",
  "bookValue": 850000,
  "lastMaintenance": "2025-09-15",
  "nextMaintenance": "2025-12-15",
  "description": "Otobüs şasi parçalarının hassas işlenmesi için kullanılıyor"
}
'''

    p = doc.add_paragraph(asset_json)
    p.style = 'Intense Quote'
    for run in p.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)

    add_heading_with_style(doc, '4.2. İş Talebi (Job Request) Veri Yapısı', 2)

    job_json = '''
{
  "id": "JR-2025-001",
  "title": "Boya Kabini Havalandırma Arızası",
  "description": "Ana boya kabininde havalandırma sistemi yetersiz çalışıyor",
  "category": "HVAC",
  "priority": "high",
  "priorityText": "Yüksek",
  "status": "waiting",
  "statusText": "Beklemede",
  "location": "Boya Tesisi",
  "assetId": "AST-002",
  "assetName": "Boya Kabini - Ana Hat",
  "createdBy": "Mehmet Kaya",
  "createdDate": "2025-10-08 09:30",
  "estimatedCost": 15000,
  "currency": "TRY"
}
'''

    p = doc.add_paragraph(job_json)
    p.style = 'Intense Quote'
    for run in p.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)

    add_heading_with_style(doc, '4.3. Bakım Planı (Maintenance Plan) Veri Yapısı', 2)

    maint_json = '''
{
  "id": "MNT-2025-034",
  "title": "CNC Torna Yıllık Bakım",
  "assetId": "AST-001",
  "assetName": "CNC Torna Makinesi - Şasi İşleme",
  "type": "preventive",
  "scheduledDate": "2025-12-15",
  "status": "scheduled",
  "assignedTo": "Bakım Ekibi A",
  "estimatedDuration": "8 saat"
}
'''

    p = doc.add_paragraph(maint_json)
    p.style = 'Intense Quote'
    for run in p.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)

    add_heading_with_style(doc, '4.4. Olay (Incident) Veri Yapısı', 2)

    incident_json = '''
{
  "id": "INC-2025-012",
  "title": "Montaj Platformu Elektrik Kesintisi",
  "description": "Platform aniden durdu, elektrik sistemi arızası",
  "priority": "critical",
  "priorityText": "Kritik",
  "location": "Final Montaj Hattı",
  "assetId": "AST-006",
  "startTime": "2025-10-08 14:20",
  "status": "active",
  "assignedTo": "Elektrik Ekibi"
}
'''

    p = doc.add_paragraph(incident_json)
    p.style = 'Intense Quote'
    for run in p.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)

    add_page_break(doc)

    # ====================================
    # 5. İŞ AKIŞLARI
    # ====================================
    add_heading_with_style(doc, '5. İş Akışları (Workflows)', 1)

    doc.add_paragraph(
        'Sistemin ana iş akışları aşağıdaki diyagramlar ile gösterilmiştir. '
        'Her akış, Visio dokümanlarında detaylandırılan süreçlerin '
        'uygulamaya yansımasıdır.'
    )

    add_heading_with_style(doc, '5.1. İş Talebi Oluşturma İş Akışı', 2)

    workflow_steps_job = [
        '1. Kullanıcı "Yeni İş Talebi" butonuna tıklar',
        '2. İş talebi formu modal olarak açılır',
        '3. Kullanıcı zorunlu alanları doldurur:',
        '   - Başlık',
        '   - Açıklama',
        '   - Kategori',
        '   - Öncelik',
        '   - Lokasyon',
        '4. Form validasyonu yapılır',
        '5. Sistem otomatik talep numarası üretir (JR-2025-XXX)',
        '6. Talep veritabanına kaydedilir',
        '7. Başarı mesajı gösterilir',
        '8. Kullanıcı iş talepleri listesine yönlendirilir',
        '9. E-posta bildirimi gönderilir (gelecek versiyon)'
    ]

    for step in workflow_steps_job:
        doc.add_paragraph(step, style='List Number')

    doc.add_paragraph('\n')
    doc.add_paragraph('İş Akışı Diyagramı:', style='Intense Quote')
    doc.add_paragraph(
        'Başla → Form Aç → Veri Girişi → Validasyon → Talep No Üret → '
        'Veritabanına Kaydet → Bildirim Gönder → Listeye Yönlendir → Bitir'
    )

    add_page_break(doc)

    add_heading_with_style(doc, '5.2. Varlık Detay Görüntüleme İş Akışı', 2)

    workflow_steps_asset = [
        '1. Kullanıcı varlık listesinde "Detay" butonuna tıklar',
        '2. Sistem URL parametresinden Varlık ID\'yi alır (?id=AST-XXX)',
        '3. Varlık detay sayfası yüklenir',
        '4. JavaScript data.js dosyasından varlık verisini bulur',
        '5. Sayfa elemanları dinamik olarak doldurulur:',
        '   - Varlık bilgileri',
        '   - Lokasyon bilgileri',
        '   - Mali bilgiler',
        '   - Bakım geçmişi',
        '   - İlgili iş talepleri',
        '6. Durum badge\'i renklendirilir (Aktif: Yeşil, Bakımda: Sarı, Arızalı: Kırmızı)',
        '7. Detay sayfası gösterilir'
    ]

    for step in workflow_steps_asset:
        doc.add_paragraph(step, style='List Number')

    doc.add_paragraph('\n')
    doc.add_paragraph('İş Akışı Diyagramı:', style='Intense Quote')
    doc.add_paragraph(
        'Başla → Detay Tıkla → URL Oku → Veri Bul → Sayfayı Doldur → '
        'Durum Renklendir → Göster → Bitir'
    )

    add_page_break(doc)

    add_heading_with_style(doc, '5.3. Acil Olay Bildirme İş Akışı', 2)

    workflow_steps_incident = [
        '1. Kullanıcı "Acil Olay Bildir" butonuna tıklar',
        '2. Olay bildirme formu modal olarak açılır',
        '3. Kullanıcı zorunlu alanları doldurur:',
        '   - Olay Başlığı',
        '   - Açıklama',
        '   - Öncelik (Kritik seçilirse kırmızı uyarı)',
        '   - Lokasyon',
        '   - İlgili Varlık (opsiyonel)',
        '4. Form validasyonu yapılır',
        '5. Sistem otomatik olay numarası üretir (INC-2025-XXX)',
        '6. Başlangıç zamanı otomatik eklenir',
        '7. Olay veritabanına kaydedilir',
        '8. Önceliğe göre bildirim gönderilir:',
        '   - Kritik: SMS + E-posta + Push Notification',
        '   - Yüksek: E-posta + Push Notification',
        '   - Orta/Düşük: E-posta',
        '9. İlgili ekip otomatik atanır',
        '10. Kullanıcı olay yönetimi sayfasına yönlendirilir'
    ]

    for step in workflow_steps_incident:
        doc.add_paragraph(step, style='List Number')

    doc.add_paragraph('\n')
    doc.add_paragraph('İş Akışı Diyagramı:', style='Intense Quote')
    doc.add_paragraph(
        'Başla → Acil Olay Butonu → Form Aç → Veri Girişi → Validasyon → '
        'Olay No Üret → Zaman Damgası Ekle → Öncelik Kontrolü → '
        'Bildirim Gönder (SMS/Email/Push) → Ekip Ata → Listeye Yönlendir → Bitir'
    )

    add_page_break(doc)

    add_heading_with_style(doc, '5.4. Filtreleme İş Akışı (İş Talepleri)', 2)

    workflow_steps_filter = [
        '1. Kullanıcı filtre değişikliği yapar veya arama yapar',
        '2. JavaScript applyFilters() fonksiyonu tetiklenir',
        '3. Tüm filtre değerleri okunur:',
        '   - Arama metni',
        '   - Durum filtresi',
        '   - Öncelik filtresi',
        '   - Kategori filtresi',
        '   - Lokasyon filtresi',
        '4. Tablo satırları döngüye alınır',
        '5. Her satır için filtreleme kontrolü yapılır',
        '6. Eşleşmeyen satırlar gizlenir (display: none)',
        '7. Eşleşen satırlar gösterilir',
        '8. Görünen satır sayısı hesaplanır',
        '9. Eğer sonuç yoksa "Filtreye uygun sonuç bulunamadı" mesajı gösterilir',
        '10. Filtreleme tamamlanır'
    ]

    for step in workflow_steps_filter:
        doc.add_paragraph(step, style='List Number')

    doc.add_paragraph('\n')
    doc.add_paragraph('İş Akışı Diyagramı:', style='Intense Quote')
    doc.add_paragraph(
        'Başla → Filtre Değişikliği → Değerleri Oku → Satırları Döngüye Al → '
        'Koşul Kontrolü → [Eşleşiyor mu?] → Evet: Göster / Hayır: Gizle → '
        'Sonuç Sayısı Hesapla → [Sonuç Var mı?] → Hayır: Mesaj Göster → Bitir'
    )

    add_page_break(doc)

    # ====================================
    # 6. TEKNİK ÖZELLİKLER
    # ====================================
    add_heading_with_style(doc, '6. Teknik Özellikler ve Gereksinimler', 1)

    add_heading_with_style(doc, '6.1. Frontend Teknolojileri', 2)

    tech_frontend = [
        'HTML5',
        'CSS3 (Custom styling, Grid, Flexbox)',
        'JavaScript (ES6+, Vanilla JS)',
        'Responsive Design (Mobile-first approach)',
        'Modal yönetimi (Custom implementation)',
        'Form validasyonu (Client-side)'
    ]

    for tech in tech_frontend:
        doc.add_paragraph(tech, style='List Bullet')

    add_heading_with_style(doc, '6.2. Veri Yönetimi', 2)

    tech_data = [
        'JSON veri formatı',
        'Client-side data store (data.js)',
        'SAP ERP entegrasyonu hazır veri yapıları',
        'RESTful API hazırlığı (gelecek versiyon)'
    ]

    for tech in tech_data:
        doc.add_paragraph(tech, style='List Bullet')

    add_heading_with_style(doc, '6.3. Tarayıcı Desteği', 2)

    browsers = [
        'Chrome 90+ (Önerilen)',
        'Firefox 88+',
        'Safari 14+',
        'Edge 90+',
        'IE11 desteklenmez'
    ]

    for browser in browsers:
        doc.add_paragraph(browser, style='List Bullet')

    add_heading_with_style(doc, '6.4. Performans Gereksinimleri', 2)

    perf_table = add_table_with_border(doc, 6, 2)
    perf_table.rows[0].cells[0].text = 'Metrik'
    perf_table.rows[0].cells[1].text = 'Hedef Değer'

    perf_reqs = [
        ('Sayfa Yükleme Süresi', '< 2 saniye'),
        ('Modal Açma Süresi', '< 300 ms'),
        ('Filtreleme Yanıt Süresi', '< 100 ms'),
        ('Tablo Render Süresi (100 satır)', '< 500 ms'),
        ('API Yanıt Süresi (gelecek)', '< 1 saniye')
    ]

    for i, (metric, value) in enumerate(perf_reqs, start=1):
        perf_table.rows[i].cells[0].text = metric
        perf_table.rows[i].cells[1].text = value

    add_heading_with_style(doc, '6.5. Güvenlik Gereksinimleri', 2)

    security = [
        'XSS koruması (Input sanitization)',
        'CSRF token kullanımı (gelecek versiyon)',
        'Form validasyonu (client + server)',
        'Rol bazlı erişim kontrolü (RBAC) - gelecek versiyon',
        'HTTPS zorunluluğu',
        'Session yönetimi (gelecek versiyon)'
    ]

    for sec in security:
        doc.add_paragraph(sec, style='List Bullet')

    add_page_break(doc)

    # ====================================
    # 7. GELECEK VERSİYON ÖZELLİKLERİ
    # ====================================
    add_heading_with_style(doc, '7. Gelecek Versiyon Özellikleri (Roadmap)', 1)

    add_heading_with_style(doc, '7.1. Versiyon 2.0 (Q1 2026)', 2)

    v2_features = [
        'SAP ERP entegrasyonu (RFC bağlantısı)',
        'Gerçek zamanlı bildirimler (WebSocket)',
        'Mobil uygulama (iOS/Android)',
        'QR kod ile varlık etiketleme',
        'Raporlama modülü (Charts, Grafikler)',
        'Excel/PDF export fonksiyonları'
    ]

    for feature in v2_features:
        doc.add_paragraph(feature, style='List Bullet')

    add_heading_with_style(doc, '7.2. Versiyon 3.0 (Q3 2026)', 2)

    v3_features = [
        'IoT sensör entegrasyonu',
        'Tahmine dayalı bakım (Predictive Maintenance) AI modeli',
        'Çoklu dil desteği (İngilizce, Almanca)',
        'Gelişmiş analitik ve BI dashboard',
        'Workflow otomasyonu (BPMN)',
        'Doküman yönetimi (DMS entegrasyonu)'
    ]

    for feature in v3_features:
        doc.add_paragraph(feature, style='List Bullet')

    add_page_break(doc)

    # ====================================
    # 8. SONUÇ
    # ====================================
    add_heading_with_style(doc, '8. Sonuç ve Uygulama Notları', 1)

    doc.add_paragraph(
        'MAN Türkiye Bakım Yönetimi Sistemi, otobüs üretim tesisinin bakım süreçlerini '
        'dijitalleştiren, kullanıcı dostu ve ölçeklenebilir bir web uygulamasıdır. '
        'Modüler yapısı sayesinde gelecekteki genişlemelere açıktır ve SAP ERP '
        'entegrasyonu için hazır veri yapılarına sahiptir.'
    )

    add_heading_with_style(doc, '8.1. Geliştirme Notları', 2)

    dev_notes = [
        'Tüm JavaScript fonksiyonları ES6+ standartlarında yazılmıştır',
        'Modal yönetimi global fonksiyonlar (openModal, closeModal) ile yapılmaktadır',
        'Filtreleme işlemleri client-side olarak gerçekleştirilmektedir',
        'Form validasyonu hem required attribute hem de JavaScript ile yapılmaktadır',
        'Notification sistemi custom implementation ile geliştirilmiştir',
        'Tüm ID\'ler otomatik üretilmekte ve unique garantisi verilmektedir'
    ]

    for note in dev_notes:
        doc.add_paragraph(note, style='List Bullet')

    add_heading_with_style(doc, '8.2. Test Senaryoları', 2)

    test_table = add_table_with_border(doc, 6, 3)
    test_table.rows[0].cells[0].text = 'Modül'
    test_table.rows[0].cells[1].text = 'Test Senaryosu'
    test_table.rows[0].cells[2].text = 'Beklenen Sonuç'

    tests = [
        ('İş Talepleri', 'Durum filtresi: "Beklemede" seç', 'Sadece beklemedeki talepler gösterilir'),
        ('İş Talepleri', 'Arama: "JR-2025-001" yaz', 'Tek talep gösterilir'),
        ('Varlık Yönetimi', 'Detay butonuna tıkla', 'Varlık detay sayfası açılır, tüm bilgiler gösterilir'),
        ('Bakım Yönetimi', 'Yeni bakım planı oluştur', 'MNT-2025-XXX numarası üretilir, başarı mesajı görülür'),
        ('Olay Yönetimi', 'Kritik olay bildir', 'INC-2025-XXX numarası üretilir, uyarı mesajı görülür')
    ]

    for i, (module, scenario, expected) in enumerate(tests, start=1):
        test_table.rows[i].cells[0].text = module
        test_table.rows[i].cells[1].text = scenario
        test_table.rows[i].cells[2].text = expected

    doc.add_paragraph('\n\n')

    closing = doc.add_paragraph()
    closing.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = closing.add_run('\n\n--- DOKÜMAN SONU ---\n')
    run.font.bold = True
    run.font.color.rgb = RGBColor(102, 102, 102)

    run2 = closing.add_run('\nHazırlayan: MAN Türkiye IT Ekibi\n')
    run2.font.size = Pt(10)
    run2.font.color.rgb = RGBColor(102, 102, 102)

    run3 = closing.add_run('Ekim 2025')
    run3.font.size = Pt(10)
    run3.font.italic = True
    run3.font.color.rgb = RGBColor(102, 102, 102)

    return doc

if __name__ == '__main__':
    print('MAN Türkiye Bakım Yönetimi - İş Analizi Dokümanı Oluşturuluyor...\n')

    doc = create_document()

    output_path = '/Users/caglarozyildirim/WebstormProjects/Deneme/MAN_Turkiye_Bakim_Yonetimi_Is_Analizi_FINAL.docx'
    doc.save(output_path)

    file_size = os.path.getsize(output_path) / 1024
    print(f'✅ Doküman başarıyla oluşturuldu!')
    print(f'📁 Dosya: {output_path}')
    print(f'📊 Boyut: {file_size:.1f} KB')
