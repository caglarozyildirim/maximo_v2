#!/usr/bin/env python3
"""
MAN Türkiye Bakım Yönetimi - Ekran Görüntüsü ve Veri Yapısı Çıkarma
Tüm HTML sayfalarını işler, ekran görüntüsü alır ve dokümanı günceller
"""

import os
import json
import time
import asyncio
from pathlib import Path
from datetime import datetime
from playwright.sync_api import sync_playwright
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from bs4 import BeautifulSoup

# Yollar
BASE_DIR = Path('/Users/caglarozyildirim/WebstormProjects/Deneme')
APP_DIR = BASE_DIR / 'bakim-yonetim-app'
SCREENSHOTS_DIR = BASE_DIR / 'screenshots'
DOCX_FILE = BASE_DIR / 'MAN_Turkiye_Bakim_Yonetimi_ULTIMATE_DEVELOPER_READY.docx'

# Klasörü oluştur
SCREENSHOTS_DIR.mkdir(exist_ok=True)

# İşlenecek sayfa tanımları
PAGES = [
    # Ana Sayfalar
    {
        'name': 'Ana Sayfa (Dashboard)',
        'file': 'index.html',
        'url': 'index.html',
        'type': 'main',
        'description': 'Sistem ana sayfası. Dashboard görünümü, istatistikler, grafikler ve son aktiviteler.',
        'group': 'Dashboard'
    },
    {
        'name': 'İş Talepleri Listesi',
        'file': 'pages/job-requests.html',
        'url': 'pages/job-requests.html',
        'type': 'main',
        'description': 'Tüm iş taleplerinin listelendiği sayfa. Filtreleme, arama ve durum bazlı görüntüleme.',
        'group': 'İş Talepleri'
    },
    {
        'name': 'İş Talebi Detay',
        'file': 'pages/job-request-detail.html',
        'url': 'pages/job-request-detail.html?id=JR-2025-001',
        'type': 'detail',
        'description': 'İş talebi detay sayfası. İş akışı, yorumlar, ekler ve tüm talep bilgileri.',
        'group': 'İş Talepleri'
    },
    {
        'name': 'İş Talebi Oluşturma',
        'file': 'pages/job-request-create.html',
        'url': 'pages/job-request-create.html',
        'type': 'create',
        'description': 'Yeni iş talebi oluşturma formu. Tüm gerekli alanlar ve validasyonlar.',
        'group': 'İş Talepleri'
    },
    {
        'name': 'Varlık Yönetimi Listesi',
        'file': 'pages/assets.html',
        'url': 'pages/assets.html',
        'type': 'main',
        'description': 'Tüm varlıkların (asset) listelendiği sayfa. Durum, kategori ve lokasyon bazlı filtreleme.',
        'group': 'Varlık Yönetimi'
    },
    {
        'name': 'Varlık Detay',
        'file': 'pages/asset-detail.html',
        'url': 'pages/asset-detail.html?id=AST-8840',
        'type': 'detail',
        'description': 'Varlık detay sayfası. SAP bilgileri, bakım geçmişi, teknik özellikler.',
        'group': 'Varlık Yönetimi'
    },
    {
        'name': 'Bakım Yönetimi Listesi',
        'file': 'pages/maintenance.html',
        'url': 'pages/maintenance.html',
        'type': 'main',
        'description': 'Bakım planları ve kayıtlarının listelendiği sayfa. Periyodik ve koruyucu bakımlar.',
        'group': 'Bakım Yönetimi'
    },
    {
        'name': 'Bakım Detay',
        'file': 'pages/maintenance-detail.html',
        'url': 'pages/maintenance-detail.html?id=MNT-445',
        'type': 'detail',
        'description': 'Bakım planı/kaydı detay sayfası. Bakım adımları, kullanılan malzemeler, süre bilgileri.',
        'group': 'Bakım Yönetimi'
    },
    {
        'name': 'Bakım Planı Oluşturma',
        'file': 'pages/maintenance-create.html',
        'url': 'pages/maintenance-create.html',
        'type': 'create',
        'description': 'Yeni bakım planı oluşturma formu. Periyot tanımlama, sorumlu atama.',
        'group': 'Bakım Yönetimi'
    },
    {
        'name': 'Olay Yönetimi Listesi',
        'file': 'pages/incidents.html',
        'url': 'pages/incidents.html',
        'type': 'main',
        'description': 'Tüm olayların (incident) listelendiği sayfa. Kritiklik seviyesi ve durum bazlı görüntüleme.',
        'group': 'Olay Yönetimi'
    },
    {
        'name': 'Olay Detay',
        'file': 'pages/incident-detail.html',
        'url': 'pages/incident-detail.html?id=INC-2025-012',
        'type': 'detail',
        'description': 'Olay detay sayfası. Olay bilgileri, etkilenen varlıklar, çözüm adımları.',
        'group': 'Olay Yönetimi'
    },
    {
        'name': 'Olay Bildirimi Oluşturma',
        'file': 'pages/incident-create.html',
        'url': 'pages/incident-create.html',
        'type': 'create',
        'description': 'Yeni olay bildirimi formu. Acil durum raporlama.',
        'group': 'Olay Yönetimi'
    },
    {
        'name': 'Raporlar',
        'file': 'pages/reports.html',
        'url': 'pages/reports.html',
        'type': 'main',
        'description': 'Raporlama sayfası. Çeşitli analiz ve raporlar.',
        'group': 'Raporlama'
    }
]


def take_screenshot(browser, page_info):
    """Sayfa ekran görüntüsünü al"""
    url = f"file://{APP_DIR}/{page_info['url']}"
    print(f"  📸 Ekran görüntüsü alınıyor: {page_info['name']}")

    try:
        page = browser.new_page(viewport={'width': 1920, 'height': 1080})
        page.goto(url)
        page.wait_for_timeout(2000)  # 2 saniye bekle

        # Ekran görüntüsünü kaydet (tam sayfa)
        screenshot_filename = f"{page_info['file'].replace('/', '_').replace('.html', '')}.png"
        screenshot_path = SCREENSHOTS_DIR / screenshot_filename
        page.screenshot(path=str(screenshot_path), full_page=True)

        page.close()

        print(f"    ✓ Kaydedildi: {screenshot_filename}")
        return screenshot_path

    except Exception as e:
        print(f"    ✗ HATA: {e}")
        return None


def extract_data_structure(page_info):
    """HTML dosyasından veri yapısını çıkar"""
    html_path = APP_DIR / page_info['file']
    print(f"  📊 Veri yapısı çıkarılıyor: {page_info['name']}")

    try:
        with open(html_path, 'r', encoding='utf-8') as f:
            soup = BeautifulSoup(f.read(), 'html.parser')

        data_structure = {
            'page_name': page_info['name'],
            'page_type': page_info['type'],
            'tables': [],
            'forms': [],
            'cards': [],
            'badges': [],
            'charts': []
        }

        # Tablo kolonlarını bul
        tables = soup.find_all('table')
        for table in tables:
            thead = table.find('thead')
            if thead:
                headers = [th.get_text(strip=True) for th in thead.find_all('th')]
                data_structure['tables'].append({
                    'columns': headers
                })

        # Form alanlarını bul
        forms = soup.find_all(['input', 'select', 'textarea'])
        form_fields = []
        for field in forms:
            field_info = {
                'type': field.name,
                'id': field.get('id', ''),
                'name': field.get('name', ''),
                'placeholder': field.get('placeholder', '')
            }
            if field.name == 'select':
                options = [opt.get_text(strip=True) for opt in field.find_all('option')]
                field_info['options'] = options
            form_fields.append(field_info)

        if form_fields:
            data_structure['forms'] = form_fields

        # Badge'leri bul
        badges = soup.find_all(class_=lambda x: x and 'badge' in x)
        badge_types = list(set([b.get('class', [''])[0] for b in badges if b.get('class')]))
        data_structure['badges'] = badge_types

        # Kartları bul
        cards = soup.find_all(class_=lambda x: x and ('card' in x or 'summary-card' in x))
        data_structure['cards'] = [
            {'title': card.find('h3').get_text(strip=True) if card.find('h3') else
                     card.find('h4').get_text(strip=True) if card.find('h4') else 'Başlıksız'}
            for card in cards[:5]  # İlk 5 kart
        ]

        # Chart canvas'ları bul
        charts = soup.find_all('canvas')
        data_structure['charts'] = [chart.get('id', 'unnamed') for chart in charts]

        print(f"    ✓ Veri yapısı çıkarıldı")
        return data_structure

    except Exception as e:
        print(f"    ✗ HATA: {e}")
        return None


def add_heading_with_style(doc, text, level=1):
    """Başlık ekle"""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Stil ayarları
    if level == 1:
        run = heading.runs[0]
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(26, 26, 26)
        run.font.bold = True
    elif level == 2:
        run = heading.runs[0]
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(226, 7, 20)  # MAN Red
        run.font.bold = True
    elif level == 3:
        run = heading.runs[0]
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(51, 51, 51)
        run.font.bold = True

    return heading


def add_page_to_document(doc, page_info, screenshot_path, data_structure):
    """Doküman'a sayfa bilgisini ekle"""
    print(f"  📝 Dokümana ekleniyor: {page_info['name']}")

    try:
        # Sayfa başlığı
        add_heading_with_style(doc, f"{page_info['name']}", level=3)

        # Açıklama
        p = doc.add_paragraph()
        p.add_run('Açıklama: ').bold = True
        p.add_run(page_info['description'])

        # Sayfa tipi
        p = doc.add_paragraph()
        p.add_run('Sayfa Tipi: ').bold = True
        type_map = {'main': 'Ana Liste Sayfası', 'detail': 'Detay Sayfası', 'create': 'Oluşturma/Düzenleme Formu'}
        p.add_run(type_map.get(page_info['type'], page_info['type']))

        # Dosya yolu
        p = doc.add_paragraph()
        p.add_run('Dosya: ').bold = True
        run = p.add_run(page_info['file'])
        run.font.name = 'Courier New'
        run.font.size = Pt(9)

        # Ekran görüntüsü
        if screenshot_path and screenshot_path.exists():
            doc.add_paragraph()
            p = doc.add_paragraph()
            p.add_run('Ekran Görüntüsü:').bold = True

            try:
                # Resmi ekle (maksimum genişlik 6 inch)
                doc.add_picture(str(screenshot_path), width=Inches(6.5))
            except Exception as e:
                p = doc.add_paragraph(f"[Ekran görüntüsü eklenemedi: {e}]")
                p.runs[0].font.color.rgb = RGBColor(255, 0, 0)

        # Veri yapısı
        if data_structure:
            doc.add_paragraph()
            p = doc.add_paragraph()
            p.add_run('Veri Yapısı (JSON):').bold = True

            # JSON'u formatla
            json_text = json.dumps(data_structure, ensure_ascii=False, indent=2)

            # Kod bloğu olarak ekle
            p = doc.add_paragraph(json_text)
            p.style = 'No Spacing'
            for run in p.runs:
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(51, 51, 51)

            # Arka plan rengi (hafif gri)
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), 'F5F5F5')
            p._element.get_or_add_pPr().append(shading_elm)

        # Ayırıcı
        doc.add_paragraph()
        p = doc.add_paragraph('─' * 80)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

        print(f"    ✓ Eklendi")

    except Exception as e:
        print(f"    ✗ HATA: {e}")


def update_document():
    """Dokümanı güncelle"""
    print("\n📄 Doküman güncelleniyor...")

    try:
        # Dokümanı aç
        doc = Document(str(DOCX_FILE))

        # Yeni bölüm ekle
        doc.add_page_break()
        add_heading_with_style(doc, "EKRAN GÖRÜNTÜLERİ VE VERİ YAPILARI", level=1)

        p = doc.add_paragraph()
        p.add_run('Oluşturma Tarihi: ').bold = True
        p.add_run(datetime.now().strftime('%d.%m.%Y %H:%M'))

        p = doc.add_paragraph()
        p.add_run('Toplam Sayfa: ').bold = True
        p.add_run(str(len(PAGES)))

        doc.add_paragraph()

        # Her sayfa için işlem - Playwright kullan
        with sync_playwright() as p:
            browser = p.chromium.launch(headless=True)

            try:
                # Gruplara ayır
                groups = {}
                for page in PAGES:
                    group = page['group']
                    if group not in groups:
                        groups[group] = []
                    groups[group].append(page)

                # Her grup için
                for group_name, group_pages in groups.items():
                    print(f"\n{'='*60}")
                    print(f"📁 GRUP: {group_name}")
                    print(f"{'='*60}")

                    # Grup başlığı
                    add_heading_with_style(doc, group_name, level=2)

                    # Gruptaki her sayfa için
                    for page_info in group_pages:
                        print(f"\n▶ İşleniyor: {page_info['name']}")

                        # Ekran görüntüsü al
                        screenshot_path = take_screenshot(browser, page_info)

                        # Veri yapısını çıkar
                        data_structure = extract_data_structure(page_info)

                        # Dokümana ekle
                        add_page_to_document(doc, page_info, screenshot_path, data_structure)

                        print(f"  ✓ Tamamlandı\n")

            finally:
                browser.close()

        # Dokümanı kaydet
        output_file = BASE_DIR / 'MAN_Turkiye_Bakim_Yonetimi_WITH_SCREENSHOTS.docx'
        doc.save(str(output_file))

        print(f"\n{'='*60}")
        print(f"✅ İŞLEM TAMAMLANDI!")
        print(f"{'='*60}")
        print(f"📄 Güncellenen doküman: {output_file}")
        print(f"📸 Ekran görüntüleri: {SCREENSHOTS_DIR}")
        print(f"📊 İşlenen sayfa sayısı: {len(PAGES)}")

        return output_file

    except Exception as e:
        print(f"\n❌ HATA: {e}")
        import traceback
        traceback.print_exc()
        return None


def main():
    """Ana fonksiyon"""
    print("╔" + "═" * 58 + "╗")
    print("║" + " " * 58 + "║")
    print("║" + " MAN TÜRKİYE BAKIM YÖNETİMİ ".center(58) + "║")
    print("║" + " Ekran Görüntüsü ve Veri Yapısı Çıkarma ".center(58) + "║")
    print("║" + " " * 58 + "║")
    print("╚" + "═" * 58 + "╝")
    print()

    # İşlemi başlat
    output_file = update_document()

    if output_file:
        print("\n🎉 İşlem başarıyla tamamlandı!")
    else:
        print("\n❌ İşlem başarısız oldu!")


if __name__ == "__main__":
    main()
