#!/usr/bin/env python3
"""
Desktop/new klasöründeki dokümanları analiz etmek için script
"""
import json
from pathlib import Path

def analyze_desktop_docs():
    desktop_path = Path.home() / "Desktop" / "new"

    print("=== Desktop/new Klasörü Analizi ===\n")

    # Dosya listesi
    files = list(desktop_path.glob("*.*"))
    print("Dosyalar:")
    for f in files:
        print(f"  - {f.name} ({f.stat().st_size / 1024:.1f} KB)")

    print("\n")

    # Workflows klasörü
    workflows_path = desktop_path / "Workflows"
    if workflows_path.exists():
        print("Workflows Klasörü:")
        for f in workflows_path.glob("*.*"):
            print(f"  - {f.name}")

    print("\n")

    # Use Cases klasörü
    usecases_path = desktop_path / "Use Cases"
    if usecases_path.exists():
        print("Use Cases Klasörü:")
        for f in usecases_path.glob("*.*"):
            print(f"  - {f.name}")

    print("\n")

    # Docx dosyalarını okuma denemesi
    try:
        from docx import Document

        req_doc_path = desktop_path / "Maintenance Management Application Requirement Analysis (Version1).docx"
        if req_doc_path.exists():
            print("\n=== Requirement Analysis Dokümanı ===")
            doc = Document(str(req_doc_path))

            print(f"Toplam paragraf sayısı: {len(doc.paragraphs)}")

            # İlk 50 başlık ve önemli paragrafları topla
            headers = []
            for para in doc.paragraphs[:200]:
                if para.style.name.startswith('Heading'):
                    headers.append({
                        'level': para.style.name,
                        'text': para.text
                    })

            print(f"\nBaşlıklar ({len(headers)} adet):")
            for h in headers[:30]:
                print(f"  {h['level']}: {h['text']}")

            # Tabloları say
            print(f"\nToplam tablo sayısı: {len(doc.tables)}")

    except ImportError:
        print("python-docx kütüphanesi yüklü değil. pip install python-docx ile yükleyiniz.")
    except Exception as e:
        print(f"Docx okuma hatası: {e}")

    # Excel dosyalarını okuma
    try:
        import openpyxl

        # Screen Designs
        screen_designs_path = desktop_path / "Screen Designs.xlsx"
        if screen_designs_path.exists():
            print("\n\n=== Screen Designs Excel ===")
            wb = openpyxl.load_workbook(str(screen_designs_path), read_only=True, data_only=True)
            print(f"Sayfalar: {wb.sheetnames}")

            for sheet_name in wb.sheetnames[:5]:
                sheet = wb[sheet_name]
                print(f"\n{sheet_name} Sayfası:")
                print(f"  Boyut: {sheet.max_row} satır x {sheet.max_column} sütun")

                # İlk 5 satırı göster
                for row_idx, row in enumerate(sheet.iter_rows(max_row=10, values_only=True), 1):
                    if any(cell for cell in row):
                        print(f"  Satır {row_idx}: {[str(cell)[:50] if cell else '' for cell in row[:6]]}")

        # Data Structure
        data_struct_path = desktop_path / "Data Structure.xlsx"
        if data_struct_path.exists():
            print("\n\n=== Data Structure Excel ===")
            wb = openpyxl.load_workbook(str(data_struct_path), read_only=True, data_only=True)
            print(f"Sayfalar: {wb.sheetnames}")

            for sheet_name in wb.sheetnames[:5]:
                sheet = wb[sheet_name]
                print(f"\n{sheet_name} Sayfası:")
                print(f"  Boyut: {sheet.max_row} satır x {sheet.max_column} sütun")

                # İlk 5 satırı göster
                for row_idx, row in enumerate(sheet.iter_rows(max_row=10, values_only=True), 1):
                    if any(cell for cell in row):
                        print(f"  Satır {row_idx}: {[str(cell)[:50] if cell else '' for cell in row[:6]]}")

    except ImportError:
        print("openpyxl kütüphanesi yüklü değil. pip install openpyxl ile yükleyiniz.")
    except Exception as e:
        print(f"Excel okuma hatası: {e}")

if __name__ == "__main__":
    analyze_desktop_docs()