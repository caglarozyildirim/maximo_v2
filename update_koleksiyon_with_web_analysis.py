#!/usr/bin/env python3
"""
Update Koleksiyon Document with Web Analysis and 3-Tier Structure
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

print("="*80)
print("UPDATING KOLEKSIYON DOCUMENT WITH WEB ANALYSIS")
print("="*80 + "\n")

# Load existing document
doc = Document('/Users/caglarozyildirim/WebstormProjects/Deneme/Koleksiyon_Mobilya_Bayi_Portali_Kapsam_Dokumani.docx')

# Find where to insert new section (after 2.1.5)
insert_index = None
for idx, para in enumerate(doc.paragraphs):
    if "Çok Dilli Destek (TR/EN)" in para.text and para.style.name == 'Heading 3':
        # Find the end of this section
        for i in range(idx + 1, len(doc.paragraphs)):
            if doc.paragraphs[i].style.name.startswith('Heading'):
                insert_index = i
                break
        break

if not insert_index:
    insert_index = len(doc.paragraphs) - 20  # Before screenshots section

print(f"📍 Insert position found at paragraph {insert_index}")

# Helper function to insert paragraph at specific position
def insert_paragraph_after(doc, index, text='', style='whitespace-normal'):
    """Insert a new paragraph after the given index"""
    para = doc.paragraphs[index]._element
    new_para = doc.add_paragraph(text, style=style)._element
    para.addnext(new_para)
    return index + 1

# ============================================================================
# 2.1.6 Web Sitesi Entegrasyonu ve 3 Katmanlı Erişim Yapısı
# ============================================================================
print("📝 Adding Section 2.1.6: Web Sitesi Entegrasyonu...")

current_idx = insert_index

# Add heading
heading_para = doc.paragraphs[current_idx]._element
new_heading = doc.add_heading('Web Sitesi Entegrasyonu ve 3 Katmanlı Erişim Yapısı', level=3)._element
heading_para.addnext(new_heading)
current_idx += 1

# Intro paragraph
intro_text = (
    "Mevcut Koleksiyon web sitesinde yer alan bayilik başvuru sayfası "
    "(https://www.koleksiyondesign.com/tr/bayilik-basvurusu/) yeniden yapılandırılarak "
    "3 farklı erişim noktası oluşturulacaktır."
)
p1 = doc.paragraphs[current_idx]._element
new_p1 = doc.add_paragraph(intro_text, style='whitespace-normal')._element
p1.addnext(new_p1)
current_idx += 1

# Current analysis
analysis_text = (
    "Mevcut Durum Analizi: Koleksiyon web sitesinin bayilik başvuru sayfası, "
    "50 yıllık marka geçmişi, 33 ülkede faaliyet, 17 ulusal ve uluslararası tasarım "
    "ödülü gibi kurumsal bilgileri içermektedir. Sayfa, potansiyel bayileri çekmek "
    "için tasarlanmıştır."
)
p2 = doc.paragraphs[current_idx]._element
new_p2 = doc.add_paragraph(analysis_text, style='whitespace-normal')._element
p2.addnext(new_p2)
current_idx += 1

# 3-tier structure
structure_text = "Yeni 3 Katmanlı Yapı:"
p3 = doc.paragraphs[current_idx]._element
new_p3 = doc.add_paragraph(structure_text, style='whitespace-normal')._element
p3.addnext(new_p3)
current_idx += 1

# ============================================================================
# Layer 1: Bayilik Başvurusu
# ============================================================================
layer1_title = "1. Bayilik Başvurusu (Genel Başvuru - Mevcut Yapı)"
p4 = doc.paragraphs[current_idx]._element
new_p4 = doc.add_paragraph(layer1_title, style='whitespace-normal')._element
# Make it bold
for run in new_p4.iter():
    if hasattr(run, 'text'):
        run.set('w:b', '1')
p4.addnext(new_p4)
current_idx += 1

layer1_desc = (
    "Mevcut bayilik başvuru süreci korunacak ve geliştirilecektir. Bu form, "
    "fiziksel mağaza açmak isteyen veya genel bayilik talebinde bulunan firmaların "
    "başvurusu için kullanılacaktır."
)
p5 = doc.paragraphs[current_idx]._element
new_p5 = doc.add_paragraph(layer1_desc, style='whitespace-normal')._element
p5.addnext(new_p5)
current_idx += 1

layer1_fields_title = "Form Alanları:"
p6 = doc.paragraphs[current_idx]._element
new_p6 = doc.add_paragraph(layer1_fields_title, style='whitespace-normal')._element
p6.addnext(new_p6)
current_idx += 1

layer1_fields = [
    "Şirket Adı",
    "Yetkili Kişi Adı Soyadı",
    "E-posta Adresi",
    "Telefon Numarası",
    "İşletme Türü (Bayilik, Distribütörlük, vb.)",
    "Mağaza Alanı (m²)",
    "Bölge/Lokasyon",
    "İlgi Duyulan Ürün Grupları",
    "Mesaj/Açıklama"
]

for field in layer1_fields:
    p = doc.paragraphs[current_idx]._element
    new_p = doc.add_paragraph(f"• {field}", style='whitespace-normal')._element
    # Add left indent
    new_p.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr')
    p.addnext(new_p)
    current_idx += 1

layer1_flow = (
    "Akış: Başvuru → Koleksiyon Satış Ekibine Bildirim → Manuel Değerlendirme → "
    "İletişim (Fiziksel Mağaza Görüşmesi)"
)
p_flow = doc.paragraphs[current_idx]._element
new_p_flow = doc.add_paragraph(layer1_flow, style='whitespace-normal')._element
p_flow.addnext(new_p_flow)
current_idx += 1

# ============================================================================
# Layer 2: Panel Girişi
# ============================================================================
layer2_title = "2. Panel Girişi (Onaylı Bayiler İçin Login)"
p7 = doc.paragraphs[current_idx]._element
new_p7 = doc.add_paragraph(layer2_title, style='whitespace-normal')._element
p7.addnext(new_p7)
current_idx += 1

layer2_desc = (
    "Daha önce onaylanmış ve aktif olan bayilerin Bayi Portalı'na giriş yapması "
    "için kullanılacak login sayfasıdır. PIM panel (https://pim.koleksiyon.com.tr/admin) "
    "yapısından ilham alınarak tasarlanacaktır."
)
p8 = doc.paragraphs[current_idx]._element
new_p8 = doc.add_paragraph(layer2_desc, style='whitespace-normal')._element
p8.addnext(new_p8)
current_idx += 1

layer2_features_title = "Özellikler:"
p9 = doc.paragraphs[current_idx]._element
new_p9 = doc.add_paragraph(layer2_features_title, style='whitespace-normal')._element
p9.addnext(new_p9)
current_idx += 1

layer2_features = [
    "E-posta ve Şifre ile Giriş",
    "Beni Hatırla (Remember Me) Özelliği",
    "Şifremi Unuttum Linki",
    "Koleksiyon PIM Logosu",
    "Minimalist ve Kurumsal Tasarım",
    "Dil Seçeneği (TR/EN)",
    "SSL Güvenli Bağlantı"
]

for feature in layer2_features:
    p = doc.paragraphs[current_idx]._element
    new_p = doc.add_paragraph(f"• {feature}", style='whitespace-normal')._element
    p.addnext(new_p)
    current_idx += 1

layer2_access = (
    "Erişim: Sadece admin tarafından onaylanmış ve aktif durumda olan bayiler "
    "giriş yapabilir. Giriş sonrası para birimine göre fiyat listelerine erişim sağlanır."
)
p_access = doc.paragraphs[current_idx]._element
new_p_access = doc.add_paragraph(layer2_access, style='whitespace-normal')._element
p_access.addnext(new_p_access)
current_idx += 1

# ============================================================================
# Layer 3: Panele Üyelik Talebi
# ============================================================================
layer3_title = "3. Panele Üyelik Talebi (Yeni Bayi Kayıt - Self Registration)"
p10 = doc.paragraphs[current_idx]._element
new_p10 = doc.add_paragraph(layer3_title, style='whitespace-normal')._element
p10.addnext(new_p10)
current_idx += 1

layer3_desc = (
    "Henüz bayilik başvurusu yapmamış ancak fiyat listelerine erişim talep eden "
    "firmaların self-registration (kendi kendine kayıt) yapabileceği formdur. "
    "Bu form, online fiyat erişimi isteyen firmalar içindir."
)
p11 = doc.paragraphs[current_idx]._element
new_p11 = doc.add_paragraph(layer3_desc, style='whitespace-normal')._element
p11.addnext(new_p11)
current_idx += 1

layer3_fields_title = "Kayıt Formu Alanları:"
p12 = doc.paragraphs[current_idx]._element
new_p12 = doc.add_paragraph(layer3_fields_title, style='whitespace-normal')._element
p12.addnext(new_p12)
current_idx += 1

layer3_fields = [
    "Şirket Adı (Zorunlu)",
    "Vergi Numarası (Zorunlu)",
    "Yetkili Kişi Adı Soyadı (Zorunlu)",
    "E-posta Adresi (Login için kullanılacak - Zorunlu)",
    "Şifre ve Şifre Tekrar (Zorunlu - min 8 karakter)",
    "Telefon Numarası",
    "Ülke (Zorunlu)",
    "Şehir",
    "Adres",
    "Dil Tercihi (TR/EN)",
    "Firma Türü (Bayi, Distribütör, Müşteri, vb.)",
    "GDPR/KVKK Onay Checkbox"
]

for field in layer3_fields:
    p = doc.paragraphs[current_idx]._element
    new_p = doc.add_paragraph(f"• {field}", style='whitespace-normal')._element
    p.addnext(new_p)
    current_idx += 1

layer3_flow_title = "Onay Akışı:"
p13 = doc.paragraphs[current_idx]._element
new_p13 = doc.add_paragraph(layer3_flow_title, style='whitespace-normal')._element
p13.addnext(new_p13)
current_idx += 1

layer3_flow_steps = [
    "Kullanıcı kayıt formunu doldurur",
    "E-posta doğrulama linki gönderilir",
    "Kullanıcı e-posta doğrulaması yapar",
    "Kayıt 'Onay Bekliyor' durumuna geçer",
    "Koleksiyon admin paneline bildirim düşer",
    "Admin, bayi bilgilerini inceler",
    "Admin, para birimi atar (TRY, USD, EUR, vb.)",
    "Admin, kaydı onaylar veya reddeder",
    "Kullanıcıya onay/red e-postası gönderilir",
    "Onaylanan kullanıcı portal'a giriş yapabilir",
    "Kullanıcı sadece kendi para birimindeki fiyatları görür"
]

for step in layer3_flow_steps:
    p = doc.paragraphs[current_idx]._element
    new_p = doc.add_paragraph(f"{layer3_flow_steps.index(step) + 1}. {step}", style='whitespace-normal')._element
    p.addnext(new_p)
    current_idx += 1

# ============================================================================
# Web Sitesi Navigasyon Yapısı
# ============================================================================
nav_title = "Web Sitesi Navigasyon Yapısı:"
p14 = doc.paragraphs[current_idx]._element
new_p14 = doc.add_paragraph(nav_title, style='whitespace-normal')._element
p14.addnext(new_p14)
current_idx += 1

nav_desc = (
    "https://www.koleksiyondesign.com/tr/bayilik-basvurusu/ sayfası "
    "3 bölüme ayrılacaktır:"
)
p15 = doc.paragraphs[current_idx]._element
new_p15 = doc.add_paragraph(nav_desc, style='whitespace-normal')._element
p15.addnext(new_p15)
current_idx += 1

nav_structure = [
    "Üst Bölüm: Mevcut kurumsal içerik (50 yıllık geçmiş, ödüller, vb.) korunacak",
    "Orta Bölüm: 3 kart/buton yapısı:",
    "  - Kart 1: 'Bayilik Başvurusu' → Genel başvuru formu",
    "  - Kart 2: 'Panel Girişi' → Login sayfasına yönlendirme",
    "  - Kart 3: 'Panele Üye Ol' → Self-registration formu",
    "Alt Bölüm: İletişim bilgileri ve destek"
]

for item in nav_structure:
    p = doc.paragraphs[current_idx]._element
    new_p = doc.add_paragraph(f"• {item}", style='whitespace-normal')._element
    p.addnext(new_p)
    current_idx += 1

# ============================================================================
# Karşılaştırma Tablosu
# ============================================================================
comparison_title = "3 Yapının Karşılaştırması:"
p16 = doc.paragraphs[current_idx]._element
new_p16 = doc.add_paragraph(comparison_title, style='whitespace-normal')._element
p16.addnext(new_p16)
current_idx += 1

# Insert a table
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

table = doc.add_table(rows=1, cols=4)
table.style = 'Light Grid Accent 1'

# Headers
headers = table.rows[0].cells
headers[0].text = "Özellik"
headers[1].text = "Bayilik Başvurusu"
headers[2].text = "Panel Girişi"
headers[3].text = "Panele Üyelik"

for cell in headers:
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.bold = True
            run.font.size = Pt(9)

# Data rows
comparison_data = [
    ("Hedef Kitle", "Fiziksel mağaza açmak isteyenler", "Onaylı bayiler", "Online erişim isteyenler"),
    ("Kayıt Tipi", "Başvuru formu", "Var olan hesap", "Self-registration"),
    ("Onay Süreci", "Manuel (Satış ekibi)", "Önceden onaylı", "Admin onayı gerekli"),
    ("Fiyat Erişimi", "Yok", "Var (Para birimine göre)", "Onay sonrası"),
    ("Giriş Şifresi", "Yok", "Var", "Kayıt sırasında oluşturulur"),
    ("Kullanım Amacı", "Fiziksel bayilik", "Fiyat listesi erişimi", "Online fiyat erişimi"),
]

for feature, app, login, reg in comparison_data:
    row_cells = table.add_row().cells
    row_cells[0].text = feature
    row_cells[1].text = app
    row_cells[2].text = login
    row_cells[3].text = reg

    # Bold first column
    for para in row_cells[0].paragraphs:
        for run in para.runs:
            run.font.bold = True

# Insert table after current position
table_element = table._element
p_before_table = doc.paragraphs[current_idx]._element
p_before_table.addnext(table_element)
current_idx += 1

# Add space after table
p17 = doc.paragraphs[current_idx]._element
new_p17 = doc.add_paragraph('', style='whitespace-normal')._element
p17.addnext(new_p17)
current_idx += 1

# ============================================================================
# Teknik Gereksinimler
# ============================================================================
tech_title = "Teknik Gereksinimler:"
p18 = doc.paragraphs[current_idx]._element
new_p18 = doc.add_paragraph(tech_title, style='whitespace-normal')._element
p18.addnext(new_p18)
current_idx += 1

tech_reqs = [
    "Butterfly DXP form builder ile 3 ayrı form yapısı oluşturulacak",
    "Panel girişi için secure authentication sistemi (JWT/OAuth2)",
    "E-posta doğrulama servisi (SMTP/SendGrid)",
    "Admin panel'de bayi onay workflow'u",
    "Para birimi yönetim modülü",
    "Session yönetimi (Redis)",
    "GDPR/KVKK uyumlu veri saklama",
    "Responsive tasarım (mobil uyumlu)",
    "SSL sertifikası (HTTPS zorunlu)"
]

for req in tech_reqs:
    p = doc.paragraphs[current_idx]._element
    new_p = doc.add_paragraph(f"• {req}", style='whitespace-normal')._element
    p.addnext(new_p)
    current_idx += 1

# Save document
output_file = "/Users/caglarozyildirim/WebstormProjects/Deneme/Koleksiyon_Mobilya_Bayi_Portali_Kapsam_Dokumani.docx"
doc.save(output_file)

print("\n" + "="*80)
print("✅ DOCUMENT UPDATED SUCCESSFULLY!")
print("="*80)
print(f"📁 File: {output_file}")
print(f"✨ Added: Section 2.1.6 - Web Sitesi Entegrasyonu ve 3 Katmanlı Yapı")
print(f"📊 Comparison table created")
print(f"🔗 Web site analysis: https://www.koleksiyondesign.com/tr/bayilik-basvurusu/")
print(f"🔗 PIM panel analysis: https://pim.koleksiyon.com.tr/admin")
print("="*80 + "\n")
