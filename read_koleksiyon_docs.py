#!/usr/bin/env python3
"""
Read Koleksiyon Mobilya Documentation
"""

from docx import Document
import json

# Paths
docs_path = "/Users/caglarozyildirim/Desktop/Şirketler/Koleksiyon"
dealer_app_file = f"{docs_path}/koleksiyon_dealer_app.docx"
meeting_notes_file = f"{docs_path}/Koleksiyon Mobilya Toplantısı.docx"
suggestions_file = f"{docs_path}/koleksiyon_oneriler.docx"

def read_docx(file_path):
    """Read all text from a docx file"""
    doc = Document(file_path)
    content = []

    for para in doc.paragraphs:
        if para.text.strip():
            content.append({
                'text': para.text,
                'style': para.style.name
            })

    # Read tables
    tables_content = []
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_data.append(row_data)
        tables_content.append(table_data)

    return {
        'paragraphs': content,
        'tables': tables_content
    }

print("="*80)
print("READING KOLEKSIYON MOBILYA DOCUMENTATION")
print("="*80 + "\n")

# Read all documents
docs = {}

print("📄 Reading koleksiyon_dealer_app.docx...")
docs['dealer_app'] = read_docx(dealer_app_file)
print(f"   ✓ {len(docs['dealer_app']['paragraphs'])} paragraphs, {len(docs['dealer_app']['tables'])} tables\n")

print("📄 Reading Koleksiyon Mobilya Toplantısı.docx...")
docs['meeting_notes'] = read_docx(meeting_notes_file)
print(f"   ✓ {len(docs['meeting_notes']['paragraphs'])} paragraphs, {len(docs['meeting_notes']['tables'])} tables\n")

print("📄 Reading koleksiyon_oneriler.docx...")
docs['suggestions'] = read_docx(suggestions_file)
print(f"   ✓ {len(docs['suggestions']['paragraphs'])} paragraphs, {len(docs['suggestions']['tables'])} tables\n")

# Save to JSON for easier access
output_file = "koleksiyon_documentation.json"
with open(output_file, 'w', encoding='utf-8') as f:
    json.dump(docs, f, ensure_ascii=False, indent=2)

print(f"💾 Saved to: {output_file}\n")

# Print summary
print("="*80)
print("DEALER APP CONTENT SUMMARY")
print("="*80)
for para in docs['dealer_app']['paragraphs'][:50]:  # First 50 paragraphs
    if para['style'].startswith('Heading'):
        print(f"\n{'='*60}")
        print(f"{para['style']}: {para['text']}")
        print('='*60)
    else:
        print(para['text'])

print("\n\n" + "="*80)
print("MEETING NOTES SUMMARY")
print("="*80)
for para in docs['meeting_notes']['paragraphs'][:50]:
    if para['style'].startswith('Heading'):
        print(f"\n{'='*60}")
        print(f"{para['style']}: {para['text']}")
        print('='*60)
    else:
        print(para['text'])

print("\n\n" + "="*80)
print("✅ DOCUMENTATION READ COMPLETE")
print("="*80)
