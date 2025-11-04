#!/usr/bin/env python3
"""
Rebuild the document properly with table format for data structures
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import json

def set_cell_background(cell, color):
    """Set cell background color"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading_elm)

def add_data_structure_table(doc, fields, module_color="#E20714"):
    """Add data structure table in CleanShot format"""
    # Add table with 6 columns
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Light Grid Accent 1'

    # Set column widths
    for row in table.rows:
        row.cells[0].width = Inches(1.2)  # EN
        row.cells[1].width = Inches(1.2)  # TR
        row.cells[2].width = Inches(1.0)  # Type
        row.cells[3].width = Inches(0.8)  # Required
        row.cells[4].width = Inches(2.0)  # Description
        row.cells[5].width = Inches(1.0)  # SAP

    # Header row
    header_cells = table.rows[0].cells
    headers = ['Alan Adı (EN)', 'Alan Adı (TR)', 'Veri Tipi', 'Zorunlu', 'Açıklama', 'SAP Mapping']

    for i, header_text in enumerate(headers):
        cell = header_cells[i]
        set_cell_background(cell, module_color)
        paragraph = cell.paragraphs[0]
        paragraph.clear()
        run = paragraph.add_run(header_text)
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(10)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    for field in fields:
        row_cells = table.add_row().cells
        row_cells[0].text = field['en']
        row_cells[1].text = field['tr']
        row_cells[2].text = field['type']
        row_cells[3].text = field['required']
        row_cells[4].text = field['desc']
        row_cells[5].text = field.get('sap', '')

        # Format cells
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)

    return table

def create_document():
    """Create the complete document"""
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Load data structures
    with open('/Users/caglarozyildirim/WebstormProjects/Deneme/detailed_data_structures.json', 'r', encoding='utf-8') as f:
        data = json.load(f)

    # ============= COVER PAGE =============
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Title
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('MAN TÜRKİYE')
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(226, 7, 20)

    doc.add_paragraph()

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('BAKIM YÖNETİMİ SİSTEMİ')
    run.bold = True
    run.font.size = Pt(24)

    doc.add_paragraph()

    subtitle2 = doc.add_paragraph()
    subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle2.add_run('İŞ ANALİZİ DOKÜMANI')
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_page_break()

    # ============= 1. PROJE AMACI VE HEDEFLER =============
    heading = doc.add_heading('1. PROJE AMACI VE HEDEFLER', level=1)
    heading.runs[0].font.color.rgb = RGBColor(226, 7, 20)

    doc.add_paragraph(
        'MAN Türkiye Bakım Yönetimi Sistemi, üretim tesislerindeki tüm bakım süreçlerinin '
        'dijitalleştirilmesi ve SAP entegrasyonu ile merkezi yönetimini sağlamak amacıyla geliştirilmektedir.'
    )

    doc.add_heading('1.1 Temel Hedefler', level=2)
    goals = [
        'İş taleplerinin dijital ortamda oluşturulması ve onay sürecinin otomasyonu',
        'Varlık yönetimi ve SAP entegrasyonu ile merkezi kayıt sistemi',
        'Periyodik bakım planlaması ve takibi',
        'Olay yönetimi ile acil müdahalelerin hızlandırılması',
        'SLA takibi ve performans ölçümü',
        'Malzeme tüketimi ve maliyet takibi'
    ]
    for goal in goals:
        doc.add_paragraph(goal, style='List Bullet')

    doc.add_page_break()

    # ============= 2. PROJE KAPSAMI =============
    heading = doc.add_heading('2. PROJE KAPSAMI', level=1)
    heading.runs[0].font.color.rgb = RGBColor(226, 7, 20)

    doc.add_heading('2.1 Kapsam İçinde', level=2)
    in_scope = [
        'İş Talepleri Yönetimi: Talep oluşturma, onay süreci, çözüm takibi',
        'Varlık Yönetimi: Varlık kaydı, SAP entegrasyonu, transfer işlemleri',
        'Bakım Yönetimi: Periyodik bakım planlaması, görev atama, tamamlama',
        'Olay Yönetimi: Acil olay bildirimi, önceliklendirme, müdahale',
        'Mobil Uygulama: Teknisyen tarafı için mobil erişim',
        'Web Uygulaması: Planlama ve yönetim için masaüstü arayüz'
    ]
    for item in in_scope:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('2.2 Kapsam Dışında', level=2)
    out_scope = [
        'Finans modülü yönetimi',
        'İnsan kaynakları entegrasyonu',
        'Satın alma süreçleri',
        'Üretim planlama modülü'
    ]
    for item in out_scope:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

    # ============= 3. GEREKSINIMLER =============
    heading = doc.add_heading('3. GEREKSINIMLER', level=1)
    heading.runs[0].font.color.rgb = RGBColor(226, 7, 20)

    # ============= 3.1 İŞ TALEPLERİ YÖNETİMİ =============
    doc.add_heading('3.1 İş Talepleri Yönetimi', level=2)

    doc.add_paragraph(
        'İş talepleri modülü, kullanıcıların bakım, onarım veya yeni kurulum taleplerini '
        'sisteme girmesini, onay sürecinden geçirmesini ve çözüm aşamasına kadar takip '
        'edilmesini sağlar. SLA takibi ile zamanında tamamlanma garanti edilir.'
    )

    doc.add_heading('Müşteri Gereksinimi:', level=3)
    doc.add_paragraph(data['job_requests']['requirement'])

    doc.add_heading('İş Akışı:', level=3)
    try:
        doc.add_picture('/Users/caglarozyildirim/WebstormProjects/Deneme/diagrams/job-request-workflow.png', width=Inches(6))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except:
        doc.add_paragraph('İş akışı diyagramı: diagrams/job-request-workflow.png')

    # === Liste Ekranı ===
    doc.add_heading('3.1.1 İş Talepleri - Liste Ekranı', level=3)

    doc.add_paragraph(
        'Bu ekran tüm iş taleplerinin listelendiği ana ekrandır. Kullanıcılar duruma, '
        'önceliğe ve kategoriye göre filtreleme yapabilir, arama yapabilir ve yeni talep oluşturabilir.'
    )

    doc.add_heading('Ekran Görüntüsü:', level=4)
    try:
        doc.add_picture('/Users/caglarozyildirim/WebstormProjects/Deneme/screenshots/job-requests.png', width=Inches(6))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except:
        doc.add_paragraph('Ekran görüntüsü bulunamadı')

    doc.add_heading('Veri Yapısı:', level=4)
    add_data_structure_table(doc, data['job_requests']['fields'], "#FFA726")
    doc.add_paragraph()

    # === Detay Ekranı ===
    doc.add_heading('3.1.2 İş Talebi - Detay Ekranı', level=3)

    doc.add_paragraph(
        'Seçilen iş talebinin tüm detaylarının görüntülendiği, onay işlemlerinin yapıldığı, '
        'yorumların eklendiği ve durum değişikliklerinin takip edildiği ekrandır.'
    )

    doc.add_heading('Ekran Görüntüsü:', level=4)
    try:
        doc.add_picture('/Users/caglarozyildirim/WebstormProjects/Deneme/screenshots/job-request-detail.png', width=Inches(6))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except:
        doc.add_paragraph('Ekran görüntüsü bulunamadı')

    doc.add_heading('Veri Yapısı:', level=4)
    add_data_structure_table(doc, data['job_requests']['fields'], "#FFA726")
    doc.add_paragraph()

    # === Oluştur Ekranı ===
    doc.add_heading('3.1.3 İş Talebi Oluştur - Form Ekranı', level=3)

    doc.add_paragraph(
        'Yeni iş talebi oluşturma ekranıdır. Kullanıcı kategoriye göre uygun formu doldurur, '
        'varlık seçimi yapar, öncelik belirler ve talebi sisteme gönderir.'
    )

    doc.add_heading('Ekran Görüntüsü:', level=4)
    try:
        doc.add_picture('/Users/caglarozyildirim/WebstormProjects/Deneme/screenshots/job-request-create.png', width=Inches(6))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except:
        doc.add_paragraph('Ekran görüntüsü bulunamadı')

    doc.add_heading('Veri Yapısı:', level=4)
    add_data_structure_table(doc, data['job_requests']['fields'], "#FFA726")
    doc.add_paragraph()

    doc.add_page_break()

    # ============= 3.2 VARLIK YÖNETİMİ =============
    doc.add_heading('3.2 Varlık Yönetimi', level=2)

    doc.add_paragraph(
        'Varlık yönetimi modülü, tüm sabit varlıkların sisteme kaydedilmesini, '
        'SAP ile entegrasyonunu, lokasyon takibini, zimmet işlemlerini ve yaşam '
        'döngüsü yönetimini sağlar.'
    )

    doc.add_heading('Müşteri Gereksinimi:', level=3)
    doc.add_paragraph(data['assets']['requirement'])

    doc.add_heading('İş Akışı:', level=3)
    try:
        doc.add_picture('/Users/caglarozyildirim/WebstormProjects/Deneme/diagrams/asset-registration-workflow.png', width=Inches(6))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except:
        doc.add_paragraph('İş akışı diyagramı: diagrams/asset-registration-workflow.png')

    # === Liste Ekranı ===
    doc.add_heading('3.2.1 Varlık Yönetimi - Liste Ekranı', level=3)

    doc.add_paragraph(
        'Sistemde kayıtlı tüm varlıkların listelendiği ekrandır. Kategori, durum ve '
        'lokasyona göre filtreleme, arama ve yeni varlık ekleme işlemleri yapılabilir.'
    )

    doc.add_heading('Ekran Görüntüsü:', level=4)
    try:
        doc.add_picture('/Users/caglarozyildirim/WebstormProjects/Deneme/screenshots/assets.png', width=Inches(6))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except:
        doc.add_paragraph('Ekran görüntüsü bulunamadı')

    doc.add_heading('Veri Yapısı:', level=4)
    add_data_structure_table(doc, data['assets']['fields'], "#42A5F5")
    doc.add_paragraph()

    # === Detay Ekranı ===
    doc.add_heading('3.2.2 Varlık - Detay Ekranı', level=3)

    doc.add_paragraph(
        'Seçilen varlığın tüm bilgilerinin, teknik özelliklerinin, bakım geçmişinin, '
        'ilgili iş taleplerinin ve ek belgelerin görüntülendiği ekrandır. '
        'Onay işlemleri ve zimmet atamaları bu ekrandan yapılır.'
    )

    doc.add_heading('Ekran Görüntüsü:', level=4)
    try:
        doc.add_picture('/Users/caglarozyildirim/WebstormProjects/Deneme/screenshots/asset-detail.png', width=Inches(6))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except:
        doc.add_paragraph('Ekran görüntüsü bulunamadı')

    doc.add_heading('Veri Yapısı:', level=4)
    add_data_structure_table(doc, data['assets']['fields'], "#42A5F5")
    doc.add_paragraph()

    # === Oluştur Ekranı ===
    doc.add_heading('3.2.3 Varlık Oluştur - Form Ekranı', level=3)

    doc.add_paragraph(
        'Sisteme yeni varlık kaydı oluşturma ekranıdır. Temel bilgiler, üretici bilgileri, '
        'lokasyon, teknik özellikler ve ek belgeler girilir. SAP entegrasyonu için gerekli '
        'alanlar mapping ile eşleştirilir.'
    )

    doc.add_heading('Ekran Görüntüsü:', level=4)
    try:
        doc.add_picture('/Users/caglarozyildirim/WebstormProjects/Deneme/screenshots/asset-create.png', width=Inches(6))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except:
        doc.add_paragraph('Ekran görüntüsü bulunamadı')

    doc.add_heading('Veri Yapısı:', level=4)
    add_data_structure_table(doc, data['assets']['fields'], "#42A5F5")
    doc.add_paragraph()

    doc.add_page_break()

    # ============= 3.3 BAKIM YÖNETİMİ =============
    doc.add_heading('3.3 Bakım Yönetimi', level=2)

    doc.add_paragraph(
        'Bakım yönetimi modülü, periyodik ve planlı bakım işlemlerinin planlanmasını, '
        'görev listelerinin oluşturulmasını, teknisyen atamalarını ve bakım tamamlama '
        'sürecinin takibini sağlar.'
    )

    doc.add_heading('Müşteri Gereksinimi:', level=3)
    doc.add_paragraph(data['maintenance']['requirement'])

    doc.add_heading('İş Akışı:', level=3)
    try:
        doc.add_picture('/Users/caglarozyildirim/WebstormProjects/Deneme/diagrams/maintenance-duty-workflow.png', width=Inches(6))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except:
        doc.add_paragraph('İş akışı diyagramı: diagrams/maintenance-duty-workflow.png')

    # === Liste Ekranı ===
    doc.add_heading('3.3.1 Bakım Yönetimi - Liste Ekranı', level=3)

    doc.add_paragraph(
        'Planlanan ve devam eden bakım görevlerinin listelendiği ekrandır. Takvim görünümü, '
        'durum filtreleme ve yeni bakım görevi oluşturma işlemleri yapılabilir.'
    )

    doc.add_heading('Ekran Görüntüsü:', level=4)
    try:
        doc.add_picture('/Users/caglarozyildirim/WebstormProjects/Deneme/screenshots/maintenance.png', width=Inches(6))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except:
        doc.add_paragraph('Ekran görüntüsü bulunamadı')

    doc.add_heading('Veri Yapısı:', level=4)
    add_data_structure_table(doc, data['maintenance']['fields'], "#66BB6A")
    doc.add_paragraph()

    # === Detay Ekranı ===
    doc.add_heading('3.3.2 Bakım Görevi - Detay Ekranı', level=3)

    doc.add_paragraph(
        'Bakım görevinin detaylarının, görev listesinin, tamamlanan işlerin ve '
        'kullanılan malzemelerin görüntülendiği ekrandır. Teknisyen ataması ve '
        'onay işlemleri bu ekrandan yapılır.'
    )

    doc.add_heading('Ekran Görüntüsü:', level=4)
    try:
        doc.add_picture('/Users/caglarozyildirim/WebstormProjects/Deneme/screenshots/maintenance-detail.png', width=Inches(6))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except:
        doc.add_paragraph('Ekran görüntüsü bulunamadı')

    doc.add_heading('Veri Yapısı:', level=4)
    add_data_structure_table(doc, data['maintenance']['fields'], "#66BB6A")
    doc.add_paragraph()

    # === Oluştur Ekranı ===
    doc.add_heading('3.3.3 Bakım Görevi Oluştur - Form Ekranı', level=3)

    doc.add_paragraph(
        'Yeni bakım görevi oluşturma ekranıdır. Varlık seçimi, bakım tarihi belirleme, '
        'görev listesi atama ve sorumlu teknisyen seçimi yapılır.'
    )

    doc.add_heading('Ekran Görüntüsü:', level=4)
    try:
        doc.add_picture('/Users/caglarozyildirim/WebstormProjects/Deneme/screenshots/maintenance-create.png', width=Inches(6))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except:
        doc.add_paragraph('Ekran görüntüsü bulunamadı')

    doc.add_heading('Veri Yapısı:', level=4)
    add_data_structure_table(doc, data['maintenance']['fields'], "#66BB6A")
    doc.add_paragraph()

    doc.add_page_break()

    # ============= 3.4 OLAY YÖNETİMİ =============
    doc.add_heading('3.4 Olay Yönetimi', level=2)

    doc.add_paragraph(
        'Olay yönetimi modülü, acil arızaların ve beklenmedik olayların hızlı bir şekilde '
        'bildirilmesini, önceliklendirilmesini ve müdahale ekibinin atanmasını sağlar.'
    )

    doc.add_heading('Müşteri Gereksinimi:', level=3)
    doc.add_paragraph(data['incidents']['requirement'])

    doc.add_heading('İş Akışı:', level=3)
    try:
        doc.add_picture('/Users/caglarozyildirim/WebstormProjects/Deneme/diagrams/incident-workflow.png', width=Inches(6))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except:
        doc.add_paragraph('İş akışı diyagramı: diagrams/incident-workflow.png')

    # === Liste Ekranı ===
    doc.add_heading('3.4.1 Olay Yönetimi - Liste Ekranı', level=3)

    doc.add_paragraph(
        'Bildirilen tüm olayların listelendiği ekrandır. Öncelik ve duruma göre '
        'filtreleme, acil olayların vurgulanması ve yeni olay bildirimi yapılabilir.'
    )

    doc.add_heading('Ekran Görüntüsü:', level=4)
    try:
        doc.add_picture('/Users/caglarozyildirim/WebstormProjects/Deneme/screenshots/incidents.png', width=Inches(6))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except:
        doc.add_paragraph('Ekran görüntüsü bulunamadı')

    doc.add_heading('Veri Yapısı:', level=4)
    add_data_structure_table(doc, data['incidents']['fields'], "#EF5350")
    doc.add_paragraph()

    # === Detay Ekranı ===
    doc.add_heading('3.4.2 Olay - Detay Ekranı', level=3)

    doc.add_paragraph(
        'Olayın detaylarının, müdahale geçmişinin ve çözüm bilgilerinin '
        'görüntülendiği ekrandır. Sorumlu atama ve durum güncellemeleri yapılır.'
    )

    doc.add_heading('Ekran Görüntüsü:', level=4)
    try:
        doc.add_picture('/Users/caglarozyildirim/WebstormProjects/Deneme/screenshots/incident-detail.png', width=Inches(6))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except:
        doc.add_paragraph('Ekran görüntüsü bulunamadı')

    doc.add_heading('Veri Yapısı:', level=4)
    add_data_structure_table(doc, data['incidents']['fields'], "#EF5350")
    doc.add_paragraph()

    # === Oluştur Ekranı ===
    doc.add_heading('3.4.3 Olay Bildir - Form Ekranı', level=3)

    doc.add_paragraph(
        'Yeni olay bildirimi oluşturma ekranıdır. Olay detayları, öncelik seviyesi, '
        'lokasyon ve ilgili varlık bilgileri girilir.'
    )

    doc.add_heading('Ekran Görüntüsü:', level=4)
    try:
        doc.add_picture('/Users/caglarozyildirim/WebstormProjects/Deneme/screenshots/incident-create.png', width=Inches(6))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except:
        doc.add_paragraph('Ekran görüntüsü bulunamadı')

    doc.add_heading('Veri Yapısı:', level=4)
    add_data_structure_table(doc, data['incidents']['fields'], "#EF5350")
    doc.add_paragraph()

    # Save document
    output_path = '/Users/caglarozyildirim/WebstormProjects/Deneme/MAN_Turkiye_Bakim_Yonetimi_COMPLETE_ANALYSIS.docx'
    doc.save(output_path)
    print(f"✅ Doküman oluşturuldu: {output_path}")

    import os
    file_size = os.path.getsize(output_path)
    print(f"📊 Dosya boyutu: {file_size / (1024*1024):.2f} MB")

    return output_path

if __name__ == '__main__':
    create_document()
